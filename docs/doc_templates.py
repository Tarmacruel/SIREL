# -*- coding: utf-8 -*-
from pathlib import Path
from datetime import datetime, date, time
from decimal import Decimal
from io import BytesIO
from django.conf import settings
from django.apps import apps
from django.db import OperationalError, ProgrammingError
from docx import Document

Processo = apps.get_model('core', 'Processo')
OrgaoEntidade = apps.get_model('core', 'OrgaoEntidade')
BASE = Path(settings.BASE_DIR)

# --- Diretórios onde procurar modelos ---
CANDIDATE_DIRS = []
if getattr(settings, "DOCS_TEMPLATES_DIR", None):
    CANDIDATE_DIRS.append(Path(settings.DOCS_TEMPLATES_DIR))

CANDIDATE_DIRS += [
    BASE/"docs",
    BASE/"docs"/"templates",
    BASE/"docs"/"docx",
    BASE/"docs"/"modelos",
    BASE/"docs"/"templates"/"word",
    BASE/"docs"/"word",
    BASE/"docs"/"docx"/"templates",
    BASE/"docs"/"modelos"/"word",
]

def _normalize_names(filename: str):
    stem = Path(filename).stem
    ext = ".docx"
    bases = {stem, stem.replace("-", "_"), stem.replace("_", "-"),
             stem.replace(" ", "_"), stem.replace("_", " ")}
    return [b + ext for b in bases]

def _find_template(filename: str) -> Path:
    candidates = []
    for alt in _normalize_names(filename):
        for d in CANDIDATE_DIRS:
            candidates.append(d/alt)
    for p in candidates:
        if p.exists():
            return p
    tried = "\n - " + "\n - ".join(str(p) for p in candidates)
    raise FileNotFoundError(f"Modelo DOCX não encontrado: '{filename}'. Caminhos testados:{tried}")

MESES = ["janeiro","fevereiro","março","abril","maio","junho","julho","agosto","setembro","outubro","novembro","dezembro"]

# ---------------- utils ----------------
def _brl(value):
    try:
        if value is None or (isinstance(value, str) and value.strip() == ""):
            return ""
        val = Decimal(str(value))
    except Exception:
        return str(value) if value is not None else ""
    q = f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return q

def _data_br(d):
    if not d: return ""
    if isinstance(d, (datetime, date)): 
        return (d if isinstance(d, datetime) else datetime.combine(d, time.min)).strftime("%d/%m/%Y")
    try:
        return datetime.strptime(str(d), "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return str(d)

def _data_extenso(d):
    """16 de agosto de 2025"""
    if not d: return ""
    if not isinstance(d, (datetime, date)):
        try:
            d = datetime.strptime(str(d), "%Y-%m-%d").date()
        except Exception:
            return str(d)
    if isinstance(d, datetime):
        d = d.date()
    return f"{d.day} de {MESES[d.month-1]} de {d.year}"

def _text(x): return (str(x) if x is not None else "").strip()

def _pessoa_nome(p):
    if not p: return ""
    return _text(getattr(p, "nome", None) or getattr(p, "razao_social", None) or p)

def _pessoa_cargo(p):
    if not p: return ""
    for attr in ("cargo","funcao","funcao_cargo","cargo_funcao","cargo_pessoa","role","position","job_title","titulo"):
        v = getattr(p, attr, None)
        if v: return _text(v)
    prof = getattr(p, "perfil", None) or getattr(p, "profile", None)
    if prof:
        for attr in ("cargo","funcao","role","position","job_title"):
            v = getattr(prof, attr, None)
            if v: return _text(v)
    return ""

def _resolve_id(*args, **kwargs):
    for key in ("processo_id", "pk", "id"):
        if key in kwargs:
            try: return int(kwargs[key])
            except Exception: pass
    for a in reversed(args):
        if hasattr(a, "pk"):
            try: return int(a.pk)
            except Exception: pass
        if isinstance(a, (int,)) or (isinstance(a, str) and a.isdigit()):
            return int(a)
    return int(args[-1]) if args else None

# ---- número por extenso (pt-BR) ----
_UNIDADES = ["zero","um","dois","três","quatro","cinco","seis","sete","oito","nove"]
_DEZ_A_DEZENOVE = ["dez","onze","doze","treze","quatorze","quinze","dezesseis","dezessete","dezoito","dezenove"]
_DEZENAS = ["","dez","vinte","trinta","quarenta","cinquenta","sessenta","setenta","oitenta","noventa"]
_CENTENAS = ["","cento","duzentos","trezentos","quatrocentos","quinhentos","seiscentos","setecentos","oitocentos","novecentos"]

def _centena_por_extenso(n):
    assert 0 <= n < 1000
    if n == 0: return ""
    if n == 100: return "cem"
    c = n // 100
    r = n % 100
    partes = []
    if c: partes.append(_CENTENAS[c])
    if r:
        if r < 10:
            partes.append(_UNIDADES[r])
        elif r < 20:
            partes.append(_DEZ_A_DEZENOVE[r-10])
        else:
            d = r // 10
            u = r % 10
            if u:
                partes.append(f"{_DEZENAS[d]} e {_UNIDADES[u]}")
            else:
                partes.append(_DEZENAS[d])
    return " e ".join([p for p in partes if p])

def _numero_por_extenso(n):
    if n == 0: return "zero"
    partes = []
    bilhoes = n // 1_000_000_000
    n %= 1_000_000_000
    milhoes = n // 1_000_000
    n %= 1_000_000
    milhares = n // 1000
    resto = n % 1000

    if bilhoes:
        partes.append(("um bilhão" if bilhoes==1 else f"{_numero_por_extenso(bilhoes)} bilhões"))
    if milhoes:
        partes.append(("um milhão" if milhoes==1 else f"{_numero_por_extenso(milhoes)} milhões"))
    if milhares:
        if milhares == 1:
            partes.append("mil")
        else:
            partes.append(f"{_centena_por_extenso(milhares)} mil")
    if resto:
        partes.append(_centena_por_extenso(resto))

    txt = ""
    for i, p in enumerate(partes):
        if not txt:
            txt = p
        else:
            conj = " e " if i == len(partes)-1 else ", "
            txt += conj + p
    return txt

def _moeda_por_extenso(valor):
    try:
        val = Decimal(str(valor)).quantize(Decimal("0.01"))
    except Exception:
        return ""
    inteiro = int(val)
    centavos = int((val - Decimal(inteiro)) * 100)
    reais_txt = ""
    if inteiro == 0:
        reais_txt = "zero real"
    elif inteiro == 1:
        reais_txt = "um real"
    else:
        reais_txt = f"{_numero_por_extenso(inteiro)} reais"
    if centavos:
        if centavos == 1:
            reais_txt += f" e um centavo"
        else:
            reais_txt += f" e {_numero_por_extenso(centavos)} centavos"
    return reais_txt

def _data_extenso_por_extenso(d):
    """'dezesseis de agosto de dois mil e vinte e cinco'"""
    if not d: return ""
    if not isinstance(d, (datetime, date)):
        try:
            d = datetime.strptime(str(d), "%Y-%m-%d").date()
        except Exception:
            return str(d)
    if isinstance(d, datetime):
        d = d.date()
    dia = _numero_por_extenso(d.day)
    ano = _numero_por_extenso(d.year)
    return f"{dia} de {MESES[d.month-1]} de {ano}"

def _first_attr(obj, names):
    for n in names:
        if hasattr(obj, n):
            v = getattr(obj, n)
            if v not in (None, "", []):
                return v
    return None

def _fmt_hhmm(h):
    if not h: return ""
    if isinstance(h, time):
        return f"{h.hour:02d}:{h.minute:02d}"
    s = str(h)
    if len(s) >= 5 and s[2]==":":
        return s[:5]
    try:
        dt = datetime.strptime(s, "%H:%M:%S")
        return dt.strftime("%H:%M")
    except Exception:
        try:
            dt = datetime.strptime(s, "%H:%M")
            return dt.strftime("%H:%M")
        except Exception:
            return s

def _compose_data_hora(processo):
    # tenta pares data/hora comuns
    data = _first_attr(processo, ["data_abertura","dt_abertura","data_abertura_processo","data_abertura_sessao"])
    hora = _first_attr(processo, ["hora_abertura","hr_abertura","hora_abertura_processo","hora_abertura_sessao"])
    # se não houver pares, tenta datetimes conhecidos
    if not data:
        dt = _first_attr(processo, ["inicio_disputa","inicio_recebimento_propostas"])
        if isinstance(dt, datetime):
            data, hora = dt.date(), dt.time()
    if not data:
        return ""
    return f"{_data_br(data)} às {_fmt_hhmm(hora)}"

# ---------- contexto para merge ----------
def _ctx(processo):
    try:
        orgao = OrgaoEntidade.objects.order_by("-atualizado_em", "-id").first()
    except (OperationalError, ProgrammingError):
        orgao = None
    secretaria = getattr(processo, "secretaria", "")
    modalidade = getattr(processo, "modalidade", "")
    condutor = getattr(processo, "condutor_processo", "")
    try:
        autoridade = processo.autoridade_competente.first()
    except Exception:
        autoridade = getattr(processo, "autoridade_competente", None)

    criado_em = getattr(processo, "criado_em", None) or getattr(processo, "created_at", None)

    tipo_obj = ""
    try:
        tipo_obj = getattr(processo, "get_tipo_objeto_display")()
    except Exception:
        tipo_obj = getattr(processo, "tipo_objeto", "")

    criterio = ""
    try:
        criterio = getattr(processo, "get_criterio_julgamento_display")()
    except Exception:
        criterio = getattr(processo, "criterio_julgamento", "")

    valor_estimado_raw = getattr(processo, "valor_estimado", "")

    # tipo_contratacao especial: SRP se "Registro de Preço"
    tipo_contratacao = ""
    try:
        tipo_contratacao = getattr(processo, "get_tipo_contratacao_display")()
    except Exception:
        tipo_contratacao = getattr(processo, "tipo_contratacao", "") or ""
    tipo_contratacao_sigla = "SRP" if str(tipo_contratacao).strip().lower().startswith("registro de preço") or str(tipo_contratacao).strip().lower().startswith("registro de preco") or "registro de preço" in str(tipo_contratacao).strip().lower() or "registro de preco" in str(tipo_contratacao).strip().lower() or "srp" == str(tipo_contratacao).strip().lower() else " "

    base = {
        "ORGAO_NOME": _text(getattr(orgao, "nome_fantasia", "") or getattr(orgao, "razao_social", "")),
        "ORGAO_RAZAO_SOCIAL": _text(getattr(orgao, "razao_social", "")),
        "ORGAO_CNPJ": _text(getattr(orgao, "cnpj", "")),
        "ORGAO_ENDERECO": _text(getattr(orgao, "endereco_completo", "")),
        "ORGAO_TELEFONE": _text(getattr(orgao, "telefone", "")),
        "ORGAO_EMAIL": _text(getattr(orgao, "email", "")),
        # Identificação
        "NUMERO_PROCESSO_ADM": _text(getattr(processo, "numero_processo_adm", "")),
        "NUMERO_EDITAL": _text(getattr(processo, "numero_edital", "")),
        "ANO_REFERENCIA": _text(getattr(processo, "ano_referencia", "")),
        "SECRETARIA": _text(secretaria).upper(),
        "OBJETO": _text(getattr(processo, "objeto", "")).upper(),
        "MODALIDADE": _text(modalidade).upper(),

        # Pessoas
        "CONDUTOR_PROCESSO": _pessoa_nome(condutor).upper(),
        "CARGO_CONDUTOR_PROCESSO": _pessoa_cargo(condutor).upper(),
        "AUTORIDADE_COMPETENTE": _pessoa_nome(autoridade).upper(),
        "CARGO_AUTORIDADE_COMPETENTE": _pessoa_cargo(autoridade).upper(),
        "AUTORIDADE_CARGO": _pessoa_cargo(autoridade).upper(),

        # Campos do processo
        "TIPO_OBJETO": _text(tipo_obj).upper(),
        "CRITERIO_JULGAMENTO": _text(criterio),
        "DATA_HORA_ABERTURA": _compose_data_hora(processo),
        "TIPO_CONTRATACAO": tipo_contratacao_sigla,

        # Valores
        "VALOR_ESTIMADO": _brl(valor_estimado_raw),
        "VALOR_ESTIMADO_EXTENSO": _moeda_por_extenso(valor_estimado_raw),
        "VALOR_HOMOLOGADO": _brl(getattr(processo, "valor_homologado", "")),

        # Datas
        "CRIADO_EM": _data_extenso(criado_em),
        "CRIADO_EM_EXTENSO": _data_extenso_por_extenso(criado_em),
        "HOJE": _data_extenso(datetime.today()),
    }

    # aliases/minúsculas
    lower = {k.lower(): v for k, v in base.items()}
    base.update(lower)
    return base

# ---------- Merge no DOCX ----------
def _replace_runs(p, mapping, literal_map=None):
    text = "".join(r.text for r in p.runs) or p.text
    if not text: return
    new = text
    for k,v in mapping.items():
        new = new.replace(f"«{k}»", v).replace(f"<{k}>", v)
    if literal_map:
        for lk, lv in literal_map.items():
            new = new.replace(lk, lv)
    if new != text:
        for i in range(len(p.runs)-1, -1, -1):
            r = p.runs[i]._element
            r.getparent().remove(r)
        p.add_run(new)

def _apply(doc, mapping):
    literal_map = {
        "(Valor estimado por extenso)": f"({mapping.get('VALOR_ESTIMADO_EXTENSO','')})",
        "(valor estimado por extenso)": f"({mapping.get('valor_estimado_extenso','')})",
    }
    for p in doc.paragraphs: _replace_runs(p, mapping, literal_map)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _replace_runs(p, mapping, literal_map)

def _render(template_filename, processo_id):
    path = _find_template(template_filename)
    proc = Processo.objects.get(pk=int(processo_id))
    doc = Document(str(path))
    mapping = _ctx(proc)
    _apply(doc, mapping)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --------- Entradas públicas ---------
def gerar_termo_autuacao_docx(*args, **kwargs):               return _render("termo_autuacao.docx", _resolve_id(*args, **kwargs))
def gerar_termo_autorizacao_docx(*args, **kwargs):            return _render("ato_autorizacao.docx", _resolve_id(*args, **kwargs))
def gerar_aviso_licitacao_docx(*args, **kwargs):              return _render("aviso_licitacao.docx", _resolve_id(*args, **kwargs))
def gerar_ci_procuradoria(*args, **kwargs):                   return _render("ci_procuradoria.docx", _resolve_id(*args, **kwargs))
def gerar_ci_controladoria(*args, **kwargs):                  return _render("ci_controladoria.docx", _resolve_id(*args, **kwargs))
def gerar_ci_contabilidade(*args, **kwargs):                  return _render("ci_contabilidade.docx", _resolve_id(*args, **kwargs))
def gerar_declaracao_nao_fracionamento_docx(*args, **kwargs): return _render("declaracao_nao_fracionamento.docx", _resolve_id(*args, **kwargs))

# Aliases para retrocompatibilidade
gerar_termo_autuacao               = gerar_termo_autuacao_docx
gerar_termo_autorizacao            = gerar_termo_autorizacao_docx
gerar_ato_autorizacao              = gerar_termo_autorizacao_docx
gerar_aviso_licitacao              = gerar_aviso_licitacao_docx
ci_procuradoria_docx               = gerar_ci_procuradoria
ci_controladoria_docx              = gerar_ci_controladoria
ci_contabilidade_docx              = gerar_ci_contabilidade
gerar_declaracao_nao_fracionamento = gerar_declaracao_nao_fracionamento_docx
declaracao_nao_fracionamento_docx  = gerar_declaracao_nao_fracionamento_docx
