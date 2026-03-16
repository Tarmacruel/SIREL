# -*- coding: utf-8 -*-
import csv
from django.contrib import admin, messages
from django.http import HttpResponseRedirect, HttpResponse
from django.urls import reverse, path
from django.utils.html import format_html
from import_export.admin import ImportExportModelAdmin
from simple_history.admin import SimpleHistoryAdmin
from django.contrib.admin.sites import NotRegistered
from django.db.models import Count, F, Sum
from .utils.formatters import fmt_brl
from .models import (
    Secretaria, Modalidade, StatusProcesso,
    Fornecedor, FornecedorDocumentoExterno, ItemCatalogo, ProcessoItem, ProcessoItemResultado, ProcessoLoteItem,
    Pessoa, OrgaoEntidade, Processo, FornecimentoItem, Observacao,
    Lote, ItemResultado, Contrato, ContratoItem, Aditivo
)
from decimal import Decimal
from django.template.response import TemplateResponse
from django.contrib.admin.views.main import ChangeList
from django.http import HttpResponse
from . import admin_process_dashboards as proc_dash  # views dos dashboards

# --- Pessoa precisa estar registrada para autocomplete_fields ---
@admin.register(Pessoa)
class PessoaAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('nome','cpf','cargo','secretaria')
    search_fields = ('nome','cpf','cargo')

# ======= Inlines / util =======
class LoteInline(admin.TabularInline):
    model = Lote
    extra = 0

class ContratoItemInline(admin.TabularInline):
    model = ContratoItem
    extra = 0

# ======= PROCESSO =======
class ProcessoAdmin(SimpleHistoryAdmin, ImportExportModelAdmin, admin.ModelAdmin):
    change_form_template = "admin/core/processo/change_form.html"
    readonly_fields = ('id',)
    fields = (
        'id',
        'protocolo',
        ('numero_processo_sirel','numero_processo_adm','numero_edital','ano_referencia'),
        ('identificador_bll','promotor_bll','link_bll'),
        'objeto',
        ('secretaria','modalidade','status'),
        ('criterio_julgamento','modo_disputa','tipo_objeto','tipo_contratacao'),
        ('autoridade_competente','condutor_processo'),
        ('data_publicacao',),
        ('data_hora_abertura','inicio_recolhimento_propostas','fim_recolhimento_propostas','fim_impugnacao_esclarecimentos'),
        ('valor_estimado','valor_homologado'),
    )
    list_display = (
        'id',
        'numero_processo_sirel',
        'numero_processo_adm',
        'numero_edital',
        'identificador_bll',
        'ano_referencia',
        'modalidade',
        'status',
        'valor_estimado_brl',
        'valor_homologado_brl',
        'atualizado_em',
    )
    inlines = [LoteInline]
    autocomplete_fields = ('autoridade_competente','condutor_processo')
    raw_id_fields = ('autoridade_competente','condutor_processo')
    search_fields = ('numero_processo_sirel','numero_processo_adm','numero_edital','identificador_bll','objeto')
    list_filter = ('modalidade','status','ano_referencia','secretaria')

    def valor_estimado_brl(self, obj): return fmt_brl(obj.valor_estimado)
    valor_estimado_brl.short_description = "Valor estimado"
    def valor_homologado_brl(self, obj): return fmt_brl(obj.valor_homologado)
    valor_homologado_brl.short_description = "Valor homologado"

    def change_view(self, request, object_id, form_url='', extra_context=None):
        extra_context = extra_context or {}
        def safe(name):
            try:
                return reverse(name, args=[object_id])
            except Exception:
                return '#'
        imp_url = safe('admin:core_processo_importar_bll_arquivo')
        exp_xlsx = safe('admin:core_processo_exportar_bll_xlsx')
        try:
            itens_url = reverse('admin:core_fornecimentoitem_changelist') + f'?processo__id__exact={object_id}'
        except Exception:
            itens_url = '#'
        try:
            gen_ctt = reverse('admin:core_processo_gerar_contratos', args=[object_id])
        except Exception:
            gen_ctt = '#'
        try:
            contratos_list = reverse('admin:core_contrato_changelist') + f'?processo__id__exact={object_id}'
        except Exception:
            contratos_list = '#'

        extra_context['additional_object_tools_items'] = format_html(
            '<li><a class="button" href="{}">Importar (XLSX/CSV)</a></li>'
            '<li><a class="button" href="{}" target="_blank">Exportar XLSX</a></li>'
            '<li><a class="button" href="{}" target="_blank">Itens</a></li>'
            '<li><a class="button" href="{}">Gerar contratos</a></li>'
            '<li><a class="button" href="{}" target="_blank">Ver contratos</a></li>',
            imp_url, exp_xlsx, itens_url, gen_ctt, contratos_list
        )
        return super().change_view(request, object_id, form_url, extra_context=extra_context)

    def get_urls(self):
        urls = super().get_urls()
        my = [
            path('<int:pk>/importar-arquivo-bll/', self.admin_site.admin_view(self.importar_bll_view), name='core_processo_importar_bll_arquivo'),
            path('<int:pk>/exportar-bll-xlsx/', self.admin_site.admin_view(self.exportar_bll_xlsx_view), name='core_processo_exportar_bll_xlsx'),
            path('<int:pk>/exportar-padrao-csv/', self.admin_site.admin_view(self.exportar_padrao_csv_view), name='core_processo_exportar_padrao_csv'),
            path("<path:object_id>/dashboard-itens/", self.admin_site.admin_view(proc_dash.processo_dashboard_itens), name="core_processo_dashboard_itens"),
            path("<path:object_id>/dashboard-contratos/", self.admin_site.admin_view(proc_dash.processo_dashboard_contratos), name="core_processo_dashboard_contratos"),
            path("<path:object_id>/dashboard-mal-sucedidos/", self.admin_site.admin_view(proc_dash.processo_dashboard_malsucedidos), name="core_processo_dashboard_malsucedidos"),
            path("<path:object_id>/dashboard-lotes-mal-sucedidos/", self.admin_site.admin_view(proc_dash.processo_dashboard_malsucedidos), name="core_processo_dashboard_malsucedidos_alias"),
            path("<path:object_id>/dashboard-geral/", self.admin_site.admin_view(proc_dash.processo_dashboard_geral), name="core_processo_dashboard_geral"),
            path("dashboard-filtro/", self.admin_site.admin_view(self.dashboard_filtro_view), name="core_processo_dashboard_filtro"),
            path("exportar-filtro-csv/", self.admin_site.admin_view(self.exportar_filtro_csv_view), name="core_processo_exportar_filtro_csv"),
        ]
        return my + urls

    # --- helper: reconstrói a queryset com base no filtro atual do changelist ---
    def _get_filtered_queryset(self, request):
        """
        Retorna (qs, cl) usando exatamente os filtros/buscas/ordenação atuais.
        """
        cl = ChangeList(
            request,
            self.model,
            self.get_list_display(request),
            self.get_list_display_links(request),
            self.get_list_filter(request),
            self.date_hierarchy,
            self.get_search_fields(request),
            self.get_list_select_related(request),
            self.list_per_page,
            self.list_max_show_all,
            self.list_editable,
            self,
            self.get_sortable_by(request),
            getattr(self, "search_help_text", None),
        )
        qs = cl.get_queryset(request)
        return qs, cl

    # --- DASHBOARD DO FILTRO ---
    def dashboard_filtro_view(self, request):
        if not self.has_view_permission(request):
            return self._get_obj_does_not_exist_redirect(request, self.model._meta, None)

        qs, cl = self._get_filtered_queryset(request)

        soma_estimado = qs.aggregate(v=Sum("valor_estimado"))["v"] or Decimal("0")
        soma_homolog = qs.aggregate(v=Sum("valor_homologado"))["v"] or Decimal("0")

        por_status = (
            qs.values("status__nome")
            .annotate(qtd=Count("id"), estimado=Sum("valor_estimado"), homologado=Sum("valor_homologado"))
            .order_by("-qtd", "status__nome")
        )

        por_modalidade = (
            qs.values("modalidade__nome")
            .annotate(qtd=Count("id"), estimado=Sum("valor_estimado"), homologado=Sum("valor_homologado"))
            .order_by("-qtd", "modalidade__nome")
        )

        context = {
            **self.admin_site.each_context(request),
            "title": "Dashboard — filtro atual",
            "opts": self.model._meta,
            "total": qs.count(),
            "soma_estimado": soma_estimado,
            "soma_homolog": soma_homolog,
            "por_status": por_status,
            "por_modalidade": por_modalidade,
            "qs_preview": qs.select_related("modalidade", "status")[:50],  # amostra
            "back_query": request.META.get("QUERY_STRING", ""),
        }
        return TemplateResponse(request, "admin/core/processo/dashboard_filtro.html", context)

    # --- EXPORTAÇÃO CSV DO FILTRO ---
    def exportar_filtro_csv_view(self, request):
        if not self.has_view_permission(request):
            return self._get_obj_does_not_exist_redirect(request, self.model._meta, None)

        qs, cl = self._get_filtered_queryset(request)

        headers = [
            "ID",
            "Nº processo adm",
            "Nº edital",
            "Ano ref.",
            "Modalidade",
            "Status",
            "Valor estimado",
            "Valor homologado",
            "Atualizado em",
        ]

        resp = HttpResponse(content_type="text/csv; charset=utf-8")
        resp["Content-Disposition"] = 'attachment; filename="processos_filtro.csv"'

        import csv
        w = csv.writer(resp)
        w.writerow(headers)

        for p in qs.select_related("modalidade", "status"):
            w.writerow([
                p.id,
                p.numero_processo_adm,
                p.numero_edital,
                p.ano_referencia,
                getattr(p.modalidade, "nome", "") or "",
                getattr(p.status, "nome", "") or "",
                p.valor_estimado or "",
                p.valor_homologado or "",
                getattr(p, "updated_at", "") or "",
            ])

        return resp
    # ------ Importar/Exportar BLL ------
    def importar_bll_view(self, request, pk):
        from django.template.response import TemplateResponse
        processo = Processo.objects.get(pk=pk)
        if request.method == 'POST':
            up = request.FILES.get('arquivo')
            if not up:
                messages.error(request, "Selecione um arquivo XLSX ou CSV.")
                return HttpResponseRedirect(reverse('admin:core_processo_importar_bll_arquivo', args=[pk]))
            try:
                from core.utils.bll_import import import_bll_file
                res = import_bll_file(processo, up)
                messages.success(request, f"Importação concluída. Ofertas criadas: {res.get('created',0)}; atualizadas: {res.get('updated',0)}.")
                return HttpResponseRedirect(reverse('admin:core_processo_change', args=[pk]))
            except Exception as e:
                messages.error(request, f"Erro ao importar: {e}")
                return HttpResponseRedirect(reverse('admin:core_processo_importar_bll_arquivo', args=[pk]))
        contexto = dict(self.admin_site.each_context(request), pk=pk, title="Importar arquivo BLL (XLSX/CSV)")
        return TemplateResponse(request, "admin/core/processo/importar_bll.html", contexto)

    def exportar_bll_xlsx_view(self, request, pk):
        try:
            from io import BytesIO
            from core.utils.bll_export import export_bll_xlsx
            processo = Processo.objects.get(pk=pk)
            buf = BytesIO()
            export_bll_xlsx(processo, buf)
            buf.seek(0)
            resp = HttpResponse(buf.getvalue(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            resp['Content-Disposition'] = f'attachment; filename="processo_{pk}.xlsx"'
            return resp
        except Exception as e:
            messages.error(request, f"Erro ao exportar XLSX: {e}")
            return HttpResponseRedirect(reverse('admin:core_processo_change', args=[pk]))

    def exportar_bll_csv_view(self, request, pk):
        messages.error(request, "Exportação BLL em CSV está desabilitada. Utilize apenas XLSX.")
        return HttpResponseRedirect(reverse('admin:core_processo_change', args=[pk]))
        
    def exportar_padrao_csv_view(self, request, pk):
        import csv
        from django.http import HttpResponse
        from .models import Processo

        p = Processo.objects.select_related(
            'secretaria', 'modalidade', 'status',
            'autoridade_competente', 'condutor_processo'
        ).get(pk=pk)

        # Cabeçalhos e valores (ajuste os campos que quiser aqui)
        headers = [
            'id', 'numero_processo_adm', 'numero_edital', 'ano_referencia',
            'secretaria', 'modalidade', 'status',
            'criterio_julgamento', 'modo_disputa', 'tipo_objeto', 'tipo_contratacao',
            'autoridade_competente', 'cargo_autoridade',
            'condutor_processo', 'cargo_condutor',
            'data_publicacao', 'data_hora_abertura',
            'valor_estimado', 'valor_homologado',
            'objeto',
        ]
        row = [
            p.id, p.numero_processo_adm, p.numero_edital, p.ano_referencia,
            getattr(p.secretaria, 'sigla', '') or str(getattr(p, 'secretaria', '') or ''),
            getattr(p.modalidade, 'nome', '') or '',
            getattr(p.status, 'nome', '') or '',
            p.criterio_julgamento, p.modo_disputa, p.tipo_objeto, p.tipo_contratacao,
            getattr(p.autoridade_competente, 'nome', '') or '',
            getattr(p.autoridade_competente, 'cargo', '') or '',
            getattr(p.condutor_processo, 'nome', '') or '',
            getattr(p.condutor_processo, 'cargo', '') or '',
            p.data_publicacao, p.data_hora_abertura,
            p.valor_estimado, p.valor_homologado,
            p.objeto,
        ]

        resp = HttpResponse(content_type='text/csv; charset=utf-8')
        resp['Content-Disposition'] = f'attachment; filename="processo_{pk}_padrao.csv"'
        writer = csv.writer(resp)
        writer.writerow(headers)
        writer.writerow(row)
        return resp

# --- Registro idempotente (evita AlreadyRegistered no autoreload) ---
try:
    admin.site.unregister(Processo)
except NotRegistered:
    pass
admin.site.register(Processo, ProcessoAdmin)

# ======= LOTE =======
@admin.register(Lote)
class LoteAdmin(SimpleHistoryAdmin, admin.ModelAdmin):
    change_form_template = "admin/core/lote/change_form.html"
    list_display = ('processo', 'numero', 'titulo', 'status', 'tipo_lance', 'qtd_itens', 'valor_referencia', 'melhor_oferta', 'exclusivo_me')
    list_filter = ('status', 'tipo_lance', 'exclusivo_me', 'processo')
    search_fields = ('processo__numero_processo_adm', 'processo__numero_edital', 'titulo')

# ======= ITENS =======
try:
    from ofertas.models import ItemOferta
except Exception:
    ItemOferta = None

class ItemOfertaInline(admin.TabularInline):
    model = ItemOferta
    extra = 0
    fields = ('fornecedor','classificacao','valor_unitario','valor_total','proposta_inicial','proposta_final','status')
    readonly_fields = ('valor_total',)
    autocomplete_fields = ('fornecedor',)

@admin.register(FornecimentoItem)
class FornecimentoItemAdmin(SimpleHistoryAdmin, admin.ModelAdmin):
    change_form_template = "admin/core/fornecimentoitem/change_form.html"
    list_display = ('processo','lote','numero_item','descricao','fornecedor_vencedor','valor_unitario_vencedor_brl','proposta_final_vencedor_brl','status_item')
    list_filter = ('processo','lote','status_item','fornecedor')
    search_fields = ('descricao','processo__numero_edital','processo__numero_processo_adm')
    inlines = [ItemOfertaInline] if ItemOferta is not None else []
    exclude = ('fornecedor','proposta_inicial','proposta_final','valor_unitario','valor_total')

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.prefetch_related('ofertas','ofertas__fornecedor')

    def _oferta_vencedora(self, obj):
        ov = None
        try:
            ov = next((o for o in obj.ofertas.all() if o.status == "VENCEDOR"), None)
            if ov is None and obj.ofertas.exists():
                ov = obj.ofertas.order_by('classificacao').first()
        except Exception:
            ov = None
        return ov

    def fornecedor_vencedor(self, obj):
        ov = self._oferta_vencedora(obj)
        return getattr(ov, 'fornecedor', None) if ov else None
    fornecedor_vencedor.short_description = "Fornecedor (vencedor)"

    def valor_unitario_vencedor_brl(self, obj):
        from .utils.formatters import fmt_brl
        ov = self._oferta_vencedora(obj)
        return fmt_brl(getattr(ov, 'valor_unitario', None)) if ov else "-"
    valor_unitario_vencedor_brl.short_description = "Valor unitário"

    def proposta_final_vencedor_brl(self, obj):
        from .utils.formatters import fmt_brl
        ov = self._oferta_vencedora(obj)
        return fmt_brl(getattr(ov, 'proposta_final', None)) if ov else "-"
    proposta_final_vencedor_brl.short_description = "Proposta final"

# ======= DEMAIS CADASTROS =======
@admin.register(Fornecedor)
class FornecedorAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('razao_social','cnpj','cidade','estado')
    search_fields = ('razao_social','cnpj')


@admin.register(FornecedorDocumentoExterno)
class FornecedorDocumentoExternoAdmin(admin.ModelAdmin):
    list_display = ("origem", "documento_digits", "identificador_externo", "fornecedor", "atualizado_em")
    search_fields = ("documento_digits", "identificador_externo", "fornecedor__razao_social", "fornecedor__cnpj")
    list_filter = ("origem",)


@admin.register(ItemCatalogo)
class ItemCatalogoAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ("codigo", "descricao_padrao", "unidade_padrao", "ativo", "atualizado_em")
    search_fields = ("codigo", "descricao_padrao")
    list_filter = ("ativo",)


@admin.register(ProcessoItem)
class ProcessoItemAdmin(admin.ModelAdmin):
    list_display = (
        "processo",
        "numero_item",
        "status_consolidado",
        "fornecedor_homologado",
        "valor_referencia_total",
        "valor_homologado_total",
        "conflito_lote",
    )
    search_fields = ("processo__numero_processo_adm", "processo__numero_edital", "descricao_snapshot")
    list_filter = ("status_consolidado", "conflito_lote")


@admin.register(ProcessoLoteItem)
class ProcessoLoteItemAdmin(admin.ModelAdmin):
    list_display = ("processo", "lote", "item", "ativo", "atualizado_em")
    search_fields = ("processo__numero_processo_adm", "processo__numero_edital", "item__descricao_snapshot")
    list_filter = ("ativo",)


@admin.register(ProcessoItemResultado)
class ProcessoItemResultadoAdmin(admin.ModelAdmin):
    list_display = (
        "processo",
        "processo_item",
        "origem",
        "status_resultado",
        "classificacao",
        "fornecedor",
        "valor_unitario",
        "valor_total",
        "ativo",
    )
    search_fields = (
        "processo__numero_processo_adm",
        "processo__numero_edital",
        "processo_item__descricao_snapshot",
        "fornecedor__razao_social",
    )
    list_filter = ("origem", "status_resultado", "ativo")


@admin.register(OrgaoEntidade)
class OrgaoEntidadeAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('razao_social', 'cnpj', 'cidade', 'uf', 'telefone')
    search_fields = ('razao_social', 'nome_fantasia', 'cnpj', 'cidade')

@admin.register(Observacao)
class ObservacaoAdmin(SimpleHistoryAdmin, admin.ModelAdmin):
    list_display = ('processo','texto','data_hora')

@admin.register(Secretaria)
class SecretariaAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('sigla','nome')
    search_fields = ('sigla','nome')

@admin.register(Modalidade)
class ModalidadeAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('nome',)
    search_fields = ('nome',)

@admin.register(StatusProcesso)
class StatusAdmin(ImportExportModelAdmin, admin.ModelAdmin):
    list_display = ('nome',)

# ====== CONTRATOS (acompanhamento) ======
@admin.register(Contrato)
class ContratoAdmin(SimpleHistoryAdmin, admin.ModelAdmin):
    change_form_template = "admin/core/contrato/change_form.html"
    list_display = ('processo','fornecedor','numero','valor_inicial','valor_atual','vigencia_inicio','vigencia_fim')

    def get_urls(self):
        urls = super().get_urls()
        my = [
            path('<int:pk>/doc/', self.admin_site.admin_view(self.gerar_doc_view), name='core_contrato_doc'),
        ]
        return my + urls

    def gerar_doc_view(self, request, pk):
        try:
            from io import BytesIO
            from docx import Document
            from docx.shared import Pt
            c = Contrato.objects.get(pk=pk)
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            doc.add_heading('Minuta de Contrato', 1)
            doc.add_paragraph(f'Processo: {getattr(c.processo, "numero_processo_adm", "")} — Edital {getattr(c.processo, "numero_edital", "")}/{getattr(c.processo, "ano_referencia", "")}')
            doc.add_paragraph(f'Contratante: {getattr(c.processo, "secretaria", "")}')
            doc.add_paragraph(f'Contratada: {getattr(c, "fornecedor", "")}')
            doc.add_paragraph(f'Número do contrato: {getattr(c, "numero", "")}')
            doc.add_paragraph(f'Vigência: {getattr(c, "vigencia_inicio", "")} a {getattr(c, "vigencia_fim", "")}')
            doc.add_paragraph(f'Valor inicial: {getattr(c, "valor_inicial", "")} — Valor atual: {getattr(c, "valor_atual", "")}')
            bio = BytesIO(); doc.save(bio); bio.seek(0)
            resp = HttpResponse(bio.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            resp['Content-Disposition'] = f'attachment; filename="contrato_{pk}.docx"'
            return resp
        except Exception as e:
            messages.error(request, f"Erro ao gerar DOCX: {e}")
            return HttpResponseRedirect(reverse('admin:core_contrato_change', args=[pk]))

@admin.register(ItemResultado)
class ItemResultadoAdmin(SimpleHistoryAdmin, admin.ModelAdmin):
    list_display = ('lote','posicao','fornecedor','valor_total_brl','microempresa','classificado','habilitado')
    def valor_total_brl(self, obj): return fmt_brl(obj.valor_total)

# Ativa patches auxiliares (se existirem), mas não registra Processo novamente
try:
    from . import admin_autocomplete_patch  # ativa autocomplete global
except Exception:
    pass
try:
    from . import admin_ofertas_filter_patch  # ativa filtro/pesquisa no inline de ofertas
except Exception:
    pass
