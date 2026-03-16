# -*- coding: utf-8 -*-
from __future__ import annotations
from io import BytesIO
from dataclasses import dataclass
from django.template.loader import render_to_string
from django.http import HttpResponse
from django.conf import settings
from core.utils.formatters import fmt_brl, valor_por_extenso
from core.models import Processo, Contrato, Aditivo
from decimal import Decimal

@dataclass
class DocResult:
    filename: str
    content_type: str
    data: bytes

# ---------- HTML render ----------
def render_html(template: str, ctx: dict) -> str:
    return render_to_string(template, ctx)

# ---------- PDF via xhtml2pdf ----------
def html_to_pdf_bytes(html: str) -> bytes:
    try:
        from xhtml2pdf import pisa
    except Exception as e:
        raise RuntimeError("Biblioteca 'xhtml2pdf' não instalada. Rode: pip install xhtml2pdf") from e
    out = BytesIO()
    pisa.CreatePDF(html, dest=out, encoding='utf-8')
    return out.getvalue()

# Helpers to prepare context for templates (avoid calling functions inside template)
def contrato_context(contrato: Contrato) -> dict:
    return {
        "municipio": getattr(settings, "MUNICIPIO_NOME", "Município"),
        "uf": getattr(settings, "MUNICIPIO_UF", "UF"),
        "contrato": contrato,
        "processo_str": f"{contrato.processo.numero_edital}-{contrato.processo.ano_referencia}",
        "valor_atual_brl": fmt_brl(contrato.valor_atual),
        "valor_atual_ext": valor_por_extenso(contrato.valor_atual),
        "vigencia_str": f"{contrato.vigencia_inicio:%d/%m/%Y} a {contrato.vigencia_fim:%d/%m/%Y}" if (contrato.vigencia_inicio and contrato.vigencia_fim) else "",
    }

def aditivo_context(aditivo: Aditivo) -> dict:
    efeito = (aditivo.valor_acrescimo or 0) - (aditivo.valor_supressao or 0)
    return {
        "municipio": getattr(settings, "MUNICIPIO_NOME", "Município"),
        "uf": getattr(settings, "MUNICIPIO_UF", "UF"),
        "aditivo": aditivo,
        "efeito_brl": fmt_brl(efeito),
        "efeito_ext": valor_por_extenso(efeito),
    }

# High-level helpers returning HttpResponse (for Admin buttons)
def export_contrato_pdf(contrato: Contrato, template: str) -> HttpResponse:
    html = render_html(template, contrato_context(contrato))
    pdf = html_to_pdf_bytes(html)
    resp = HttpResponse(pdf, content_type="application/pdf")
    resp["Content-Disposition"] = f'attachment; filename="Contrato_{contrato.id}.pdf"'
    return resp

def export_contrato_docx(contrato: Contrato, tipo: str) -> HttpResponse:
    data = build_contrato_docx(contrato, tipo=tipo)
    resp = HttpResponse(data, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    resp["Content-Disposition"] = f'attachment; filename="Contrato_{contrato.id}.docx"'
    return resp

def export_aditivo_pdf(aditivo: Aditivo, template: str) -> HttpResponse:
    html = render_html(template, aditivo_context(aditivo))
    pdf = html_to_pdf_bytes(html)
    resp = HttpResponse(pdf, content_type="application/pdf")
    resp["Content-Disposition"] = f'attachment; filename="Aditivo_{aditivo.id}.pdf"'
    return resp

def export_aditivo_docx(aditivo: Aditivo, tipo: str) -> HttpResponse:
    data = build_aditivo_docx(aditivo, tipo=tipo)
    resp = HttpResponse(data, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    resp["Content-Disposition"] = f'attachment; filename="Aditivo_{aditivo.id}.docx"'
    return resp

# ---------- DOCX via python-docx ----------
def build_contrato_docx(contrato: Contrato, tipo: str = "minuta") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("Biblioteca 'python-docx' não instalada. Rode: pip install python-docx") from e

    doc = Document()

    def h1(txt):
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def p(txt):
        doc.add_paragraph(txt)

    municipio = getattr(settings, 'MUNICIPIO_NOME', 'Município')
    uf = getattr(settings, 'MUNICIPIO_UF', 'UF')

    if tipo == "minuta":
        h1("MINUTA DE CONTRATO")
        p(f"{municipio}/{uf}")
        p(f"Processo: {contrato.processo.numero_edital}-{contrato.processo.ano_referencia}")
        p(f"Contratado: {contrato.fornecedor.razao_social} — CNPJ: {contrato.fornecedor.cnpj}")
        p(f"Objeto: {contrato.objeto}")
        p(f"Valor: {fmt_brl(contrato.valor_atual)} ({valor_por_extenso(contrato.valor_atual)})")
        if contrato.vigencia_inicio and contrato.vigencia_fim:
            p(f"Vigência: de {contrato.vigencia_inicio:%d/%m/%Y} a {contrato.vigencia_fim:%d/%m/%Y}")
        doc.add_page_break()
        h1("CLÁUSULAS PADRÃO (modelo baseado em diretrizes da AGU)")
        p("1. DO OBJETO — descreve o objeto de forma precisa e suficiente.")
        p("2. DA FUNDAMENTAÇÃO LEGAL — Lei 14.133/2021 e demais normas aplicáveis.")
        p("3. DO PREÇO E REAJUSTE — condições, periodicidade e índices, quando couber.")
        p("4. DA VIGÊNCIA — prazos de execução e vigência contratual.")
        p("5. DA EXECUÇÃO — condições de entrega/prestação, locais e prazos.")
        p("6. DA FISCALIZAÇÃO — designação de fiscal/titular e substituto.")
        p("7. DAS GARANTIAS — quando houver.")
        p("8. DAS SANÇÕES — penalidades e hipóteses de aplicação.")
        p("9. DA RESCISÃO — hipóteses e procedimentos.")
        p("10. DAS DISPOSIÇÕES FINAIS — foro, publicação e outras.")
    elif tipo == "extrato":
        h1("EXTRATO DE CONTRATO")
        p(f"Contratante: {municipio}/{uf}")
        p(f"Contratado: {contrato.fornecedor.razao_social} — CNPJ {contrato.fornecedor.cnpj}")
        p(f"Objeto: {contrato.objeto}")
        p(f"Valor: {fmt_brl(contrato.valor_atual)} ({valor_por_extenso(contrato.valor_atual)})")
        if contrato.vigencia_inicio and contrato.vigencia_fim:
            p(f"Vigência: {contrato.vigencia_inicio:%d/%m/%Y} a {contrato.vigencia_fim:%d/%m/%Y}")
        if contrato.link_publicacao:
            p(f"Publicação: {contrato.link_publicacao}")
    else:
        h1("DOCUMENTO CONTRATUAL")
        p("Tipo não reconhecido.")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

def build_aditivo_docx(aditivo: Aditivo, tipo: str = "minuta") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise RuntimeError("Biblioteca 'python-docx' não instalada. Rode: pip install python-docx") from e

    doc = Document()
    def h1(txt):
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def p(txt):
        doc.add_paragraph(txt)

    contrato = aditivo.contrato
    municipio = getattr(settings, 'MUNICIPIO_NOME', 'Município')
    uf = getattr(settings, 'MUNICIPIO_UF', 'UF')

    if tipo == "minuta":
        h1("MINUTA DE TERMO ADITIVO")
        p(f"Contrato: {contrato.numero} — Processo {contrato.processo.numero_edital}-{contrato.processo.ano_referencia}")
        p(f"Contratado: {contrato.fornecedor.razao_social} — CNPJ {contrato.fornecedor.cnpj}")
        p(f"Tipo: {aditivo.get_tipo_display()}")
        if aditivo.valor_acrescimo or aditivo.valor_supressao:
            efeito = (aditivo.valor_acrescimo or 0) - (aditivo.valor_supressao or 0)
            p(f"Ajuste de valor: {fmt_brl(efeito)} ({valor_por_extenso(efeito)})")
        if aditivo.novo_vigencia_fim:
            p(f"Nova vigência final: {aditivo.novo_vigencia_fim:%d/%m/%Y}")
        p(f"Justificativa: {aditivo.justificativa}")
    elif tipo == "extrato":
        h1("EXTRATO DE TERMO ADITIVO")
        p(f"Contrato: {contrato.numero}")
        p(f"Contratado: {contrato.fornecedor.razao_social} — CNPJ {contrato.fornecedor.cnpj}")
        p(f"Tipo: {aditivo.get_tipo_display()}")
        if aditivo.valor_acrescimo or aditivo.valor_supressao:
            efeito = (aditivo.valor_acrescimo or 0) - (aditivo.valor_supressao or 0)
            p(f"Ajuste de valor: {fmt_brl(efeito)} ({valor_por_extenso(efeito)})")
        if aditivo.novo_vigencia_fim:
            p(f"Nova vigência final: {aditivo.novo_vigencia_fim:%d/%m/%Y}")
    buf = BytesIO(); doc.save(buf); return buf.getvalue()
