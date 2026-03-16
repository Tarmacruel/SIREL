from io import BytesIO
from datetime import date

def _mk_doc():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    return doc

def gerar_termo_autuacao_docx(p):
    doc = _mk_doc()
    doc.add_heading('Termo de Autuação', 1)
    doc.add_paragraph(f'Processo Administrativo: {getattr(p,"numero_processo_adm","")}')
    doc.add_paragraph(f'Edital: {getattr(p,"numero_edital","")}/{getattr(p,"ano_referencia","")} — {getattr(p,"modalidade","")}')
    doc.add_paragraph(f'Objeto: {getattr(p,"objeto","")}')
    doc.add_paragraph(f'Data: {date.today().strftime("%d/%m/%Y")}')
    bio = BytesIO(); doc.save(bio)
    return (f"termo_autuacao_{p.id}.docx", bio.getvalue())

def gerar_termo_autorizacao_docx(p):
    doc = _mk_doc()
    doc.add_heading('Termo de Autorização', 1)
    doc.add_paragraph(f'Processo Administrativo: {getattr(p,"numero_processo_adm","")}')
    doc.add_paragraph(f'Edital: {getattr(p,"numero_edital","")}/{getattr(p,"ano_referencia","")} — {getattr(p,"modalidade","")}')
    doc.add_paragraph('Fica autorizada a abertura do certame conforme especificações.')
    bio = BytesIO(); doc.save(bio)
    return (f"termo_autorizacao_{p.id}.docx", bio.getvalue())

def gerar_aviso_licitacao_docx(p):
    doc = _mk_doc()
    doc.add_heading('Aviso de Licitação', 1)
    doc.add_paragraph(f'Órgão/Secretaria: {getattr(p,"secretaria","")}')
    doc.add_paragraph(f'Modalidade: {getattr(p,"modalidade","")} • Nº Edital: {getattr(p,"numero_edital","")}/{getattr(p,"ano_referencia","")}')
    doc.add_paragraph(f'Objeto: {getattr(p,"objeto","")}')
    bio = BytesIO(); doc.save(bio)
    return (f"aviso_licitacao_{p.id}.docx", bio.getvalue())
