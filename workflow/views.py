import json
import unicodedata
import mimetypes
from datetime import date, datetime, time
from decimal import Decimal, InvalidOperation
from functools import lru_cache
from html import unescape as html_unescape
from io import BytesIO
from pathlib import Path
import threading
from uuid import uuid4
import calendar
import re
import zipfile
from xml.sax.saxutils import escape

from PIL import Image as PILImage, ImageSequence
import numpy as np
from django.apps import apps
from django.conf import settings
from django.contrib.auth import get_user_model
from django.contrib import messages
from django.core.cache import cache
from django.db import OperationalError, transaction, IntegrityError, close_old_connections
from django.db.models import Count, Avg, Q, Max, Sum, Value, DecimalField, OuterRef, Subquery
from django.db.models.functions import Coalesce
from django.http import JsonResponse, Http404, HttpResponse, HttpResponseNotFound
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import NoReverseMatch, reverse
from django.utils import timezone
from django.views.decorators.clickjacking import xframe_options_exempt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from core.models import (
    ElementoDespesa, FonteRecurso, Fornecedor, FornecimentoItem, Lote, Modalidade, ProcessoItem,
    ProcessoItemResultado, ProcessoLoteItem,
    OrgaoEntidade, Pessoa, Processo, ProjetoAtividade, Secretaria, StatusProcesso,
    UnidadeOrcamentaria,
)
from core.utils.formatters import fmt_brl, valor_por_extenso
from .forms import (
    ComunicacaoInternaForm,
    ComprasComprovanteUploadForm,
    ComprasPesquisaForm,
    CadastroCatalogoItemForm, CadastroElementoDespesaForm, CadastroFonteRecursoForm,
    CadastroFornecedorForm, CadastroOrgaoForm,
    CadastroPessoaForm, CadastroProjetoAtividadeForm, CadastroSecretariaForm,
    CadastroUnidadeOrcamentariaForm, CadastroUsuarioForm, PerfilUsuarioForm,
    DocumentoAssinadoUploadForm,
    LicitacaoDocumentoUploadForm,
    LicitacaoProcessoExternoForm,
    DFDItemEdicaoForm, ProcessoPlanejamentoForm, DFDForm, DFDItemForm, DFDItemCatalogoForm,
    ETPForm, ETPCotacaoFonteForm, ETPCotacaoItemForm, ETPCotacaoEdicaoForm,
    TRForm, TRLoteForm, TRDotacaoForm, TRDistribuicaoSecretariaForm,
)
from .models import (
    IntegracaoProcesso, ModuloSistema, ProcessoMovimentacao, ProcessoWorkflow, SituacaoWorkflow,
    DocumentoProcessoWorkflow, PNCPContratacaoSnapshot, PNCPDetalhamentoFila,
    PlanejamentoDFD, DFDSecretaria, DFDItemCatalogo, DFDItem,
    ETPPlanejamento, ETPCotacaoFonte, ETPCotacaoItem,
    TRPlanejamento, TRLote, TRDotacao, TRDistribuicaoSecretaria, ComunicacaoInterna,
    FrequenciaRegistro,
)
from .planning_exports import (
    export_ci_pdf_preview,
    export_dfd_docx, export_dfd_pdf, export_dfd_xlsx,
    export_mapa_docx, export_mapa_pdf, export_mapa_xlsx,
    export_distribuicao_docx, export_distribuicao_pdf, export_distribuicao_xlsx,
    export_sd_docx, export_sd_pdf, export_sd_xlsx,
    export_frequencia_pdf, export_frequencia_xlsx,
    build_sd_payload,
    export_tr_docx,
)
from .services.item_registry import (
    ensure_fornecedor_documento_externo,
    sync_canonical_items_for_processo,
)
from .services.pncp import PNCPClient
from .services.pncp_publish import PNCPPublishClient
from .services.pncp_queue import enqueue_pncp_detalhamento, processar_fila_pncp as processar_fila_pncp_service

User = get_user_model()
EDITAL_TEMPLATES_DIR = Path(__file__).resolve().parent / 'edital_templates'


def _form_errors_text(form):
    erros = []
    for campo, msgs in form.errors.items():
        field = form.fields.get(campo)
        if campo == '__all__':
            label = 'Formulário'
        else:
            label = field.label if field else campo
        erros.append(f'{label}: {", ".join(msgs)}')
    return ' | '.join(erros)


def _is_ajax(request):
    return request.headers.get('x-requested-with') == 'XMLHttpRequest'


def _ajax_or_redirect(request, *, ok, message, redirect_name, redirect_args):
    if _is_ajax(request):
        return JsonResponse({'ok': ok, 'message': message})
    if ok:
        messages.success(request, message)
    else:
        messages.error(request, message)
    return redirect(redirect_name, *redirect_args)


def _set_pncp_error_popup(request, *, titulo: str, detalhes, contexto: dict | None = None):
    if isinstance(detalhes, (list, tuple)):
        detalhes_list = [str(x) for x in detalhes if str(x).strip()]
    else:
        detalhes_list = [str(detalhes)]
    contexto_out = {}
    for k, v in (contexto or {}).items():
        if v is None:
            continue
        contexto_out[str(k)] = str(v)
    request.session['pncp_erro_popup'] = {
        'titulo': str(titulo),
        'detalhes': detalhes_list[:50],
        'contexto': contexto_out,
        'criado_em': timezone.localtime().strftime('%d/%m/%Y %H:%M:%S'),
    }


def _set_pncp_success_popup(
    request,
    *,
    titulo: str,
    processos: list[dict] | None = None,
    contexto: dict | None = None,
):
    contexto_out = {}
    for k, v in (contexto or {}).items():
        if v is None:
            continue
        contexto_out[str(k)] = str(v)

    processos_out = []
    for item in (processos or [])[:80]:
        if not isinstance(item, dict):
            continue
        numero = str(item.get('numero') or '').strip()
        objeto = str(item.get('objeto') or '').strip()
        acao = str(item.get('acao') or '').strip()
        processos_out.append(
            {
                'numero': numero[:120],
                'objeto': objeto[:220],
                'acao': acao[:30],
            }
        )

    request.session['pncp_sucesso_popup'] = {
        'titulo': str(titulo),
        'processos': processos_out,
        'contexto': contexto_out,
        'criado_em': timezone.localtime().strftime('%d/%m/%Y %H:%M:%S'),
    }


def _pncp_progress_key(user_id: int, op_id: str) -> str:
    return f'pncp_import_progress:{int(user_id)}:{op_id}'


def _set_pncp_import_progress(request, op_id: str, **data):
    if not op_id or not getattr(request, 'user', None) or not request.user.is_authenticated:
        return
    key = _pncp_progress_key(request.user.id, op_id)
    current = cache.get(key) or {}
    current.update(data)
    current['updated_at'] = timezone.localtime().strftime('%d/%m/%Y %H:%M:%S')
    cache.set(key, current, timeout=PNCP_IMPORT_PROGRESS_TIMEOUT)


_PNCP_QUEUE_BG_LOCK = threading.Lock()
_PNCP_QUEUE_BG_RUNNING = False


def _worker_fila_pncp_background(limit: int):
    global _PNCP_QUEUE_BG_RUNNING
    try:
        close_old_connections()
        processar_fila_pncp_service(limit=max(1, int(limit or 20)))
    finally:
        close_old_connections()
        with _PNCP_QUEUE_BG_LOCK:
            _PNCP_QUEUE_BG_RUNNING = False


def _disparar_fila_pncp_background(limit: int = 20) -> bool:
    global _PNCP_QUEUE_BG_RUNNING
    if not bool(getattr(settings, 'PNCP_DETALHAMENTO_AUTOSTART', True)):
        return False
    with _PNCP_QUEUE_BG_LOCK:
        if _PNCP_QUEUE_BG_RUNNING:
            return False
        _PNCP_QUEUE_BG_RUNNING = True
    th = threading.Thread(
        target=_worker_fila_pncp_background,
        args=(max(1, int(limit or 20)),),
        daemon=True,
        name='pncp-detalhamento-bg',
    )
    th.start()
    return True


def _marcar_pendente_detalhamento_pncp(processo: Processo, numero_controle: str = ''):
    wf = _ensure_workflow(processo)
    wf.etapa_atual = 'PNCP - AGUARDANDO DETALHAMENTO'
    wf.save(update_fields=['etapa_atual', 'atualizado_em'])
    enqueue_pncp_detalhamento(
        processo,
        numero_controle=numero_controle or getattr(wf, 'pncp_numero_controle', ''),
        origem='IMPORTACAO_RAPIDA',
        prioridade=20,
    )


def _ajax_redirect_payload(request, redirect_name: str, ok=True):
    if _is_ajax(request):
        return JsonResponse({'ok': bool(ok), 'redirect': reverse(redirect_name)})
    return redirect(redirect_name)


def _pncp_auto_paginacao(dt_ini: date | None, dt_fim: date | None):
    if not dt_ini or not dt_fim:
        return 30, 200
    dias = max(1, (dt_fim - dt_ini).days + 1)
    if dias <= 7:
        return 50, 90
    if dias <= 31:
        return 50, 160
    if dias <= 90:
        return 40, 240
    if dias <= 180:
        return 30, 320
    return 20, 420

MODULOS = [
    {'slug': 'PLANEJAMENTO', 'titulo': 'Planejamento', 'descricao': 'DFD, ETP, TR, lotes, secretarias e dotações.', 'cor': '#1f5fa8'},
    {'slug': 'CADASTROS', 'titulo': 'Cadastros', 'descricao': 'Gestão de usuários, pessoas, órgão, fornecedores e referências base.', 'cor': '#315b9c'},
    {'slug': 'COMPRAS', 'titulo': 'Compras', 'descricao': 'Pesquisa de preços, SD, mapa comparativo e envio à licitação.', 'cor': '#2f7d4d'},
    {'slug': 'LICITACAO', 'titulo': 'Licitação', 'descricao': 'Autuação, publicidade, impugnações, sessão, recursos e homologação.', 'cor': '#8b5c1b'},
    {'slug': 'PROCURADORIA', 'titulo': 'Procuradoria', 'descricao': 'Parecer jurídico padronizado e devolução ao fluxo.', 'cor': '#7a3f8f'},
    {'slug': 'CONTROLADORIA', 'titulo': 'Controladoria', 'descricao': 'Parecer técnico de controle interno.', 'cor': '#3f7f85'},
    {'slug': 'CONTRATOS', 'titulo': 'Contratos', 'descricao': 'Contratos, atas, aditivos, apostilamentos e saldo.', 'cor': '#8d3345'},
    {'slug': 'DASHBOARDS', 'titulo': 'Dashboards', 'descricao': 'Indicadores de gestão processual, fornecedores e contratos.', 'cor': '#4a4d94'},
    {'slug': 'FREQUENCIA', 'titulo': 'Frequência', 'descricao': 'Registro diário de ponto, horas extras e relatório mensal.', 'cor': '#1f6d7a'},
    {'slug': 'DOCUMENTOS', 'titulo': 'Documentos', 'descricao': 'OCR, consolidação, paginação e geração padrão e-TCM.', 'cor': '#245a63'},
    {'slug': 'INTEGRACAO', 'titulo': 'Integrações', 'descricao': 'PNCP por API, BLL por CSV/XLSX e trilha de sincronização.', 'cor': '#4f5a6d'},
]


def _modulo_href(slug: str) -> str:
    key = (slug or '').upper().strip()
    if key == 'PLANEJAMENTO':
        return reverse('workflow:planejamento_dashboard')
    if key == 'CADASTROS':
        return reverse('workflow:cadastros_dashboard')
    if key == 'DASHBOARDS':
        return reverse('workflow:dashboards_geral')
    if key == 'FREQUENCIA':
        return reverse('workflow:frequencia')
    if key == 'INTEGRACAO':
        return reverse('workflow:integracoes')
    return reverse('workflow:modulo_detail', args=[key])


def _modulos_ui_data():
    cards = []
    for m in MODULOS:
        card = dict(m)
        card['href'] = _modulo_href(card.get('slug', ''))
        cards.append(card)
    return cards


FLUXO_MODULOS_PROCESSO = [
    ModuloSistema.PLANEJAMENTO,
    ModuloSistema.COMPRAS,
    ModuloSistema.LICITACAO,
    ModuloSistema.PROCURADORIA,
    ModuloSistema.CONTROLADORIA,
    ModuloSistema.CONTRATOS,
]

DOC_ASSINADO_CONFIG = {
    'dfd': {'label': 'DFD', 'modulo': ModuloSistema.PLANEJAMENTO},
    'mapa': {'label': 'Mapa comparativo (ETP)', 'modulo': ModuloSistema.PLANEJAMENTO},
    'distribuicao': {'label': 'Mapa de distribuição por secretaria', 'modulo': ModuloSistema.PLANEJAMENTO},
    'tr': {'label': 'Termo de Referência', 'modulo': ModuloSistema.PLANEJAMENTO},
    'ci': {'label': 'Comunicação Interna (C.I.)', 'modulo': None},
    'sd': {'label': 'Solicitação de Despesa', 'modulo': ModuloSistema.COMPRAS},
    'mapa_compras': {'label': 'Mapa comparativo (Compras)', 'modulo': ModuloSistema.COMPRAS},
    'declaracao_desconsideracao': {'label': 'Declaração de preços desconsiderados', 'modulo': ModuloSistema.COMPRAS},
}
DOC_ASSINADO_PREFIX = 'ASSINADO::'
COMPRAS_COMPROVANTE_TIPO = 'COMPRAS_COMPROVANTE_PESQUISA'
COMPRAS_SD_GERADA_DESCRICAO = 'SD gerada no módulo Compras.'
LICITACAO_DOC_PREFIX = 'LICITACAO_DOC::'
LICITACAO_ETAPA_FASE_INTERNA = 'LICITACAO - FASE INTERNA'

LICITACAO_CHECKLIST = [
    {
        'codigo': 'dfd',
        'label': 'DFD',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'ci_abertura_despacho',
        'label': 'CI de abertura ou despacho (I-GOV)',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'documentos_suporte',
        'label': 'Documentos de suporte',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'cotacoes_preliminares',
        'label': 'Cotações preliminares',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'termo_referencia',
        'label': 'Termo de Referência',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'justificativa_modalidade_simplificada',
        'label': 'Justificativa da modalidade simplificada',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'ci_planejamento',
        'label': 'Comunicação de encaminhamento para Planejamento',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'ci_compras',
        'label': 'Comunicação de encaminhamento para Compras',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'pesquisa_preco',
        'label': 'Pesquisa de preço',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'sd',
        'label': 'Solicitação de Despesa (SD)',
        'fase': 'ENTRADA DO PROCESSO',
    },
    {
        'codigo': 'decreto_comissao_permanente',
        'label': 'Decreto da Comissão Permanente',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
    {
        'codigo': 'decreto_equipe_apoio',
        'label': 'Decreto da equipe de apoio',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
    {
        'codigo': 'ci_orcamento_reserva',
        'label': 'CI para Orçamento (reserva orçamentária)',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:ci_contabilidade_docx', 'ci_contabilidade_docx'),
        'gerar_ci_html': True,
    },
    {
        'codigo': 'reserva_orcamentaria',
        'label': 'Reserva orçamentária',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
    {
        'codigo': 'declaracao_nao_fracionamento',
        'label': 'Declaração de não fracionamento de despesa',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:declaracao_nao_fracionamento_docx', 'declaracao_nao_fracionamento_docx'),
        'gerar_ci_html': True,
    },
    {
        'codigo': 'decreto_secretario_ordenador',
        'label': 'Decreto do secretário ordenador de despesas',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
    {
        'codigo': 'ato_autorizacao',
        'label': 'Ato de autorização',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:ato_autorizacao_docx', 'ato_autorizacao_docx'),
        'gerar_ci_html': True,
    },
    {
        'codigo': 'minuta_aviso',
        'label': 'Minuta do aviso',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:aviso_licitacao_docx', 'aviso_licitacao_docx'),
    },
    {
        'codigo': 'ci_pgm_parecer',
        'label': 'CI para PGM solicitando parecer',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:ci_procuradoria_docx', 'ci_procuradoria_docx'),
        'gerar_ci_html': True,
    },
    {
        'codigo': 'parecer_juridico',
        'label': 'Parecer jurídico',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
    {
        'codigo': 'aviso_licitacao_dispensa',
        'label': 'Aviso de Licitação de Dispensa',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:aviso_licitacao_docx', 'aviso_licitacao_docx'),
    },
    {
        'codigo': 'termo_autuacao_agente',
        'label': 'Termo de autuação do agente de contratação',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
        'gerar_doc_names': ('docs:termo_autuacao_docx', 'termo_autuacao_docx'),
        'gerar_ci_html': True,
    },
    {
        'codigo': 'decreto_agente_contratacao',
        'label': 'Decreto do agente de contratação',
        'fase': 'FASE INTERNA NA LICITAÇÃO',
    },
]
LICITACAO_CHECKLIST_MAP = {item['codigo']: item for item in LICITACAO_CHECKLIST}

LICITACAO_FASES_FUTURAS = [
    {
        'fase': 'FASE DE JULGAMENTO',
        'documentos': [
            'Ata da sessão de julgamento',
            'Mapa de propostas classificadas',
            'Documento de julgamento por item/lote',
            'Registro de diligências da fase de julgamento',
            'Resultado preliminar',
        ],
    },
    {
        'fase': 'FASE DE HABILITAÇÃO',
        'documentos': [
            'Ata de habilitação',
            'Checklist de habilitação por fornecedor',
            'Parecer técnico de habilitação',
            'Resultado da habilitação',
        ],
    },
    {
        'fase': 'FASES FINAIS',
        'documentos': [
            'Ata/termo de adjudicação',
            'Homologação',
            'Publicação do resultado',
            'Termo de encerramento da fase licitatória',
        ],
    },
]

PNCP_MODALIDADES_PADRAO = [
    {'codigo': 1, 'nome': 'Pregão - Lei 14.133/2021'},
    {'codigo': 2, 'nome': 'Concorrência - Lei 14.133/2021'},
    {'codigo': 3, 'nome': 'Concurso - Lei 14.133/2021'},
    {'codigo': 4, 'nome': 'Leilão - Lei 14.133/2021'},
    {'codigo': 5, 'nome': 'Diálogo competitivo - Lei 14.133/2021'},
    {'codigo': 6, 'nome': 'Dispensa de licitação - Lei 14.133/2021'},
    {'codigo': 7, 'nome': 'Inexigibilidade - Lei 14.133/2021'},
    {'codigo': 8, 'nome': 'Credenciamento - Lei 14.133/2021'},
]
PNCP_IMPORT_PROGRESS_TIMEOUT = 60 * 40


def _ensure_workflow(processo: Processo) -> ProcessoWorkflow:
    wf, _ = ProcessoWorkflow.objects.get_or_create(processo=processo)
    return wf


def _doc_assinado_tipo(key: str) -> str:
    return f'{DOC_ASSINADO_PREFIX}{key}'


def _listar_documentos_assinados(processo: Processo):
    docs = []
    qs = DocumentoProcessoWorkflow.objects.filter(
        processo=processo,
        tipo_documento__startswith=DOC_ASSINADO_PREFIX,
    ).order_by('-criado_em', '-id')
    for doc in qs:
        key = (doc.tipo_documento or '').replace(DOC_ASSINADO_PREFIX, '', 1)
        conf = DOC_ASSINADO_CONFIG.get(key, {})
        docs.append({
            'id': doc.id,
            'key': key,
            'label': conf.get('label', key.upper() or 'Documento'),
            'arquivo': doc.arquivo,
            'criado_em': doc.criado_em,
        })
    return docs


def _compras_sd_status(processo: Processo):
    sd_assinada = DocumentoProcessoWorkflow.objects.filter(
        processo=processo,
        tipo_documento=_doc_assinado_tipo('sd'),
    ).exists()
    sd_gerada = ProcessoMovimentacao.objects.filter(
        processo=processo,
        descricao=COMPRAS_SD_GERADA_DESCRICAO,
    ).exists()
    return {
        'gerada': sd_gerada,
        'assinada': sd_assinada,
        'concluida': bool(sd_assinada or sd_gerada),
    }


def _atualizar_etapa_compras_sd(processo: Processo):
    wf = _ensure_workflow(processo)
    if wf.modulo_atual != ModuloSistema.COMPRAS:
        return
    sd_status = _compras_sd_status(processo)
    etapa = 'COMPRAS - SD CONCLUIDA' if sd_status['concluida'] else 'COMPRAS - SD PENDENTE'
    if wf.etapa_atual != etapa:
        wf.etapa_atual = etapa
        wf.save(update_fields=['etapa_atual', 'atualizado_em'])


def _registrar_sd_gerada(processo: Processo, formato: str):
    if not ProcessoMovimentacao.objects.filter(
        processo=processo,
        descricao=COMPRAS_SD_GERADA_DESCRICAO,
    ).exists():
        ProcessoMovimentacao.objects.create(
            processo=processo,
            modulo_origem=ModuloSistema.COMPRAS,
            modulo_destino=ModuloSistema.COMPRAS,
            descricao=COMPRAS_SD_GERADA_DESCRICAO,
            observacao=f'SD exportada em {str(formato or "").upper()}',
        )
    _atualizar_etapa_compras_sd(processo)


def _reverse_any(names, args=None, kwargs=None) -> str:
    for name in names or ():
        try:
            return reverse(name, args=args, kwargs=kwargs)
        except NoReverseMatch:
            continue
    return ''


def _licitacao_tipo_documento(codigo: str) -> str:
    return f'{LICITACAO_DOC_PREFIX}{codigo}'


def _licitacao_codigo_documento(tipo_documento: str) -> str:
    tipo = str(tipo_documento or '')
    if not tipo.startswith(LICITACAO_DOC_PREFIX):
        return ''
    return tipo.replace(LICITACAO_DOC_PREFIX, '', 1)


def _build_licitacao_checklist(processo: Processo):
    docs_qs = (
        DocumentoProcessoWorkflow.objects
        .filter(processo=processo, tipo_documento__startswith=LICITACAO_DOC_PREFIX)
        .order_by('-criado_em', '-id')
    )
    docs_por_codigo = {}
    for doc in docs_qs:
        codigo = _licitacao_codigo_documento(doc.tipo_documento)
        if codigo and codigo not in docs_por_codigo:
            docs_por_codigo[codigo] = doc

    checklist = []
    total = len(LICITACAO_CHECKLIST)
    total_ok = 0
    for item in LICITACAO_CHECKLIST:
        codigo = item['codigo']
        doc = docs_por_codigo.get(codigo)
        presente = bool(doc)
        if presente:
            total_ok += 1
        gerar_links = []
        if item.get('gerar_ci_html'):
            gerar_links.append(
                {
                    'label': 'Gerar HTML',
                    'url': reverse('workflow:licitacao_ci_documento', args=[processo.id, codigo]),
                }
            )
        if item.get('gerar_doc_names'):
            docx_url = _reverse_any(item['gerar_doc_names'], args=[processo.id])
            if docx_url:
                gerar_links.append(
                    {
                        'label': 'Gerar DOCX',
                        'url': docx_url,
                    }
                )
        checklist.append(
            {
                **item,
                'presente': presente,
                'documento': doc,
                'arquivo_url': doc.arquivo.url if doc and doc.arquivo else '',
                'arquivo_nome': doc.arquivo.name.split('/')[-1] if doc and doc.arquivo else '',
                'data_envio': doc.criado_em if doc else None,
                'gerar_links': gerar_links,
            }
        )

    percentual = int((total_ok / total) * 100) if total else 0
    pendencias = [row['label'] for row in checklist if not row['presente']]
    fases = {}
    for row in checklist:
        fases.setdefault(row['fase'], []).append(row)
    return {
        'itens': checklist,
        'fases': fases,
        'pendencias': pendencias,
        'total': total,
        'total_ok': total_ok,
        'percentual': percentual,
        'docs_por_codigo': docs_por_codigo,
    }


def _licitacao_documento_choices():
    return [
        (item['codigo'], f"{item['fase']} - {item['label']}")
        for item in LICITACAO_CHECKLIST
    ]


def _upsert_licitacao_documento(processo: Processo, codigo: str, arquivo):
    tipo = _licitacao_tipo_documento(codigo)
    ordem = (
        DocumentoProcessoWorkflow.objects.filter(processo=processo)
        .aggregate(maior=Max('ordem_cronologica'))
        .get('maior')
        or 0
    ) + 1
    existente = (
        DocumentoProcessoWorkflow.objects
        .filter(processo=processo, tipo_documento=tipo)
        .order_by('-id')
        .first()
    )
    if existente:
        existente.modulo = ModuloSistema.LICITACAO
        existente.arquivo = arquivo
        existente.ordem_cronologica = ordem
        existente.gerar_no_etcm = False
        existente.save(update_fields=['modulo', 'arquivo', 'ordem_cronologica', 'gerar_no_etcm'])
        return existente
    return DocumentoProcessoWorkflow.objects.create(
        processo=processo,
        modulo=ModuloSistema.LICITACAO,
        tipo_documento=tipo,
        arquivo=arquivo,
        ordem_cronologica=ordem,
        gerar_no_etcm=False,
    )


MESES_PT_BR = [
    'janeiro',
    'fevereiro',
    'março',
    'abril',
    'maio',
    'junho',
    'julho',
    'agosto',
    'setembro',
    'outubro',
    'novembro',
    'dezembro',
]


def _data_extenso_pt_br(data_ref: date) -> str:
    if isinstance(data_ref, datetime):
        data_ref = data_ref.date()
    return f'{data_ref.day} de {MESES_PT_BR[data_ref.month - 1]} de {data_ref.year}'


def _localidade_orgao_padrao() -> str:
    orgao = _orgao_ativo()
    if not orgao:
        return 'Município'
    cidade = str(getattr(orgao, 'cidade', '') or '').strip()
    uf = str(getattr(orgao, 'uf', '') or '').strip()
    if cidade and uf:
        return f'{cidade}-{uf}'
    return cidade or 'Município'


def _secretaria_composta(processo: Processo) -> str:
    secretaria = getattr(processo, 'secretaria', None)
    if not secretaria:
        return 'SECRETARIA RESPONSÁVEL'
    sigla = str(getattr(secretaria, 'sigla', '') or '').strip()
    nome = str(getattr(secretaria, 'nome', '') or '').strip()
    if sigla and nome:
        return f'{sigla} - {nome}'
    return sigla or nome or 'SECRETARIA RESPONSÁVEL'


def _slug_ascii(texto: str) -> str:
    raw = unicodedata.normalize('NFKD', str(texto or ''))
    raw = raw.encode('ascii', 'ignore').decode('ascii')
    out = ''.join(ch if ch.isalnum() else '_' for ch in raw)
    while '__' in out:
        out = out.replace('__', '_')
    out = out.strip('_').lower()
    return out or 'documento'


def _pdf_safe(valor) -> str:
    return escape(str(valor or '-')).replace('\n', '<br/>')


def _format_datetime_br(valor) -> str:
    if not valor:
        return '-'
    dt = valor
    if timezone.is_aware(dt):
        dt = timezone.localtime(dt)
    return dt.strftime('%d/%m/%Y %H:%M')


def _label_tipo_documento_workflow(tipo_documento: str) -> str:
    tipo = str(tipo_documento or '').strip()
    if not tipo:
        return 'Documento'
    if tipo.startswith(LICITACAO_DOC_PREFIX):
        codigo = _licitacao_codigo_documento(tipo)
        row = LICITACAO_CHECKLIST_MAP.get(codigo)
        return str((row or {}).get('label') or f'Documento de Licitação ({codigo or "N/D"})')
    if tipo.startswith(DOC_ASSINADO_PREFIX):
        key = tipo.replace(DOC_ASSINADO_PREFIX, '', 1)
        conf = DOC_ASSINADO_CONFIG.get(key, {})
        return str(conf.get('label') or f'Documento assinado ({key or "N/D"})')
    if tipo == COMPRAS_COMPROVANTE_TIPO:
        return 'Comprovante de pesquisa de preços'
    return tipo


def _display_user_name(usuario) -> str:
    if not usuario:
        return ''
    if hasattr(usuario, 'get_full_name'):
        nome = (usuario.get_full_name() or '').strip()
        if nome:
            return nome
    if hasattr(usuario, 'get_username'):
        username = str(usuario.get_username() or '').strip()
        if username:
            return username
    return str(usuario or '').strip()


def _documento_adicionado_por(obj) -> str:
    for attr in ('criado_por', 'usuario', 'user', 'uploaded_by', 'owner', 'autor'):
        nome = _display_user_name(getattr(obj, attr, None))
        if nome:
            return nome
    for attr in ('usuario_nome', 'responsavel_envio'):
        valor = str(getattr(obj, attr, '') or '').strip()
        if valor:
            return valor
    return 'Usuario nao registrado'


def _coletar_documentos_processo_consolidado(processo: Processo) -> list[dict]:
    itens = []
    qs_workflow = (
        DocumentoProcessoWorkflow.objects
        .filter(processo=processo)
        .order_by('criado_em', 'id')
    )
    for doc in qs_workflow:
        nome = Path(str(getattr(doc.arquivo, 'name', '') or '')).name or 'documento'
        anexado_em = getattr(doc, 'criado_em', None)
        anexado_ts = float(anexado_em.timestamp()) if anexado_em else 0.0
        itens.append(
            {
                'origem': 'workflow',
                'id': int(doc.id),
                'titulo': _label_tipo_documento_workflow(doc.tipo_documento),
                'descricao': str(doc.tipo_documento or ''),
                'arquivo': doc.arquivo,
                'arquivo_nome': nome,
                'anexado_em': anexado_em,
                'anexado_em_exibicao': _format_datetime_br(anexado_em),
                'adicionado_por': _documento_adicionado_por(doc),
                'ordem_cronologica': int(doc.ordem_cronologica or 0),
                'ordem_sort': float(doc.ordem_cronologica or 0),
                'anexado_ts': anexado_ts,
            }
        )

    processo_anexo_model = None
    try:
        processo_anexo_model = apps.get_model('docs', 'ProcessoAnexo')
    except LookupError:
        processo_anexo_model = None

    if processo_anexo_model is not None:
        qs_docs = processo_anexo_model.objects.filter(processo=processo).order_by('uploaded_at', 'id')
        for doc in qs_docs:
            nome = Path(str(getattr(doc.arquivo, 'name', '') or '')).name or 'documento'
            titulo = str(getattr(doc, 'descricao', '') or '').strip() or str(getattr(doc, 'get_tipo_display', lambda: 'Anexo')())
            anexado_em = getattr(doc, 'uploaded_at', None)
            anexado_ts = float(anexado_em.timestamp()) if anexado_em else 0.0
            itens.append(
                {
                    'origem': 'docs',
                    'id': int(doc.id),
                    'titulo': titulo or 'Anexo',
                    'descricao': str(getattr(doc, 'tipo', '') or ''),
                    'arquivo': doc.arquivo,
                    'arquivo_nome': nome,
                    'anexado_em': anexado_em,
                    'anexado_em_exibicao': _format_datetime_br(anexado_em),
                    'adicionado_por': _documento_adicionado_por(doc),
                    'ordem_cronologica': 0,
                    'ordem_sort': 10_000_000.0,
                    'anexado_ts': anexado_ts,
                }
            )

    itens.sort(key=lambda x: (float(x.get('anexado_ts') or 0.0), float(x.get('ordem_sort') or 0.0), int(x.get('id') or 0)))
    return itens


def _ler_bytes_arquivo_field(arquivo) -> bytes:
    if not arquivo:
        raise FileNotFoundError('Arquivo do anexo não encontrado.')
    with arquivo.open('rb') as fp:
        payload = fp.read()
    return payload or b''


def _decode_bytes_texto(payload: bytes) -> str:
    if not payload:
        return ''
    for enc in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            return payload.decode(enc)
        except UnicodeDecodeError:
            continue
    return payload.decode('utf-8', errors='ignore')


def _canvas_draw_wrapped(
    c: canvas.Canvas,
    texto: str,
    *,
    x: float,
    y: float,
    width: float,
    font: str = 'Helvetica',
    size: int = 11,
    leading: int = 15,
) -> float:
    c.setFont(font, size)
    for bloco in str(texto or '').replace('\r', '').split('\n'):
        palavras = bloco.split()
        if not palavras:
            y -= leading
            continue
        linha = []
        for palavra in palavras:
            tentativa = (' '.join(linha + [palavra])).strip()
            if c.stringWidth(tentativa, font, size) <= width:
                linha.append(palavra)
                continue
            if linha:
                c.drawString(x, y, ' '.join(linha))
                y -= leading
            linha = [palavra]
        if linha:
            c.drawString(x, y, ' '.join(linha))
            y -= leading
    return y


def _draw_dossie_banner(
    c: canvas.Canvas,
    *,
    width: float,
    height: float,
    title: str,
    subtitle: str,
    accent: str,
    tag: str = '',
) -> float:
    orgao = _orgao_ativo()
    orgao_nome = (
        str(getattr(orgao, 'nome_fantasia', '') or '').strip()
        or str(getattr(orgao, 'razao_social', '') or '').strip()
        or 'Órgão/Entidade'
    )
    cnpj = str(getattr(orgao, 'cnpj', '') or '').strip() or '-'
    endereco_parts = [
        str(getattr(orgao, 'endereco', '') or '').strip(),
        str(getattr(orgao, 'numero', '') or '').strip(),
        str(getattr(orgao, 'bairro', '') or '').strip(),
        str(getattr(orgao, 'cidade', '') or '').strip(),
        str(getattr(orgao, 'uf', '') or '').strip(),
    ]
    endereco = ', '.join([x for x in endereco_parts if x]) or '-'

    c.setFillColor(colors.HexColor('#0e2238'))
    c.roundRect(24, height - 170, width - 48, 132, 18, stroke=0, fill=1)
    c.setFillColor(colors.HexColor(accent))
    c.roundRect(24, height - 170, width - 48, 18, 18, stroke=0, fill=1)

    text_x = 44
    if orgao and getattr(orgao, 'logo', None):
        try:
            c.drawImage(orgao.logo.path, 42, height - 120, width=44, height=44, preserveAspectRatio=True, mask='auto')
            text_x = 96
        except Exception:
            text_x = 44

    if tag:
        c.setFillColor(colors.HexColor('#f4f8fc'))
        c.roundRect(width - 146, height - 164, 106, 14, 7, stroke=0, fill=1)
        c.setFillColor(colors.HexColor(accent))
        c.setFont('Helvetica-Bold', 7.6)
        c.drawCentredString(width - 93, height - 159, str(tag or '')[:18].upper())

    c.setFillColor(colors.white)
    c.setFont('Helvetica-Bold', 15)
    c.drawString(text_x, height - 72, orgao_nome[:82])
    c.setFont('Helvetica', 9.2)
    c.drawString(text_x, height - 88, f'CNPJ {cnpj}')
    c.drawString(text_x, height - 102, endereco[:96])
    c.setFont('Helvetica-Bold', 20)
    c.drawString(44, height - 134, str(title or '')[:92])
    c.setFont('Helvetica', 11)
    c.drawString(44, height - 150, str(subtitle or '')[:108])
    return height - 196


def _draw_info_card(
    c: canvas.Canvas,
    *,
    x: float,
    y: float,
    w: float,
    h: float,
    label: str,
    value: str,
    accent: str = '#1f5fa8',
) -> None:
    c.setFillColor(colors.HexColor('#f5f8fc'))
    c.roundRect(x, y, w, h, 12, stroke=0, fill=1)
    c.setFillColor(colors.HexColor(accent))
    c.roundRect(x, y + h - 8, w, 8, 12, stroke=0, fill=1)
    c.setFillColor(colors.HexColor('#60758f'))
    c.setFont('Helvetica-Bold', 8)
    c.drawString(x + 12, y + h - 22, str(label or '').upper()[:28])
    c.setFillColor(colors.HexColor('#10253c'))
    c.setFont('Helvetica-Bold', 12)
    c.drawString(x + 12, y + 16, str(value or '-')[:34])


def _build_capa_processo_pdf(processo: Processo, total_documentos: int) -> bytes:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    y = _draw_dossie_banner(
        c,
        width=width,
        height=height,
        title='Processo Completo',
        subtitle='Arquivo padronizado para envio institucional e conferência documental',
        accent='#1f5fa8',
        tag='e-TCM Bahia',
    )

    _draw_info_card(
        c,
        x=40,
        y=y - 88,
        w=120,
        h=64,
        label='Processo',
        value=str(processo.numero_processo_principal or '-'),
        accent='#1f5fa8',
    )
    _draw_info_card(
        c,
        x=174,
        y=y - 88,
        w=140,
        h=64,
        label='Modalidade',
        value=str(getattr(getattr(processo, 'modalidade', None), 'nome', '') or '-'),
        accent='#245a63',
    )
    _draw_info_card(
        c,
        x=328,
        y=y - 88,
        w=106,
        h=64,
        label='Anexos',
        value=str(int(total_documentos)),
        accent='#9b6b2b',
    )
    _draw_info_card(
        c,
        x=448,
        y=y - 88,
        w=108,
        h=64,
        label='Gerado em',
        value=timezone.localtime().strftime('%d/%m/%Y'),
        accent='#5c4aa8',
    )

    c.setFillColor(colors.HexColor('#11263f'))
    c.setFont('Helvetica-Bold', 18)
    y = y - 118
    y = _canvas_draw_wrapped(
        c,
        f'Processo {processo.numero_processo_principal}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica-Bold',
        size=18,
        leading=24,
    )

    c.setFont('Helvetica', 11)
    y -= 6
    y = _canvas_draw_wrapped(
        c,
        f'Objeto: {str(processo.objeto or "-").strip()}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y = _canvas_draw_wrapped(
        c,
        f'Secretaria: {_secretaria_composta(processo)}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y = _canvas_draw_wrapped(
        c,
        f'Modalidade: {str(getattr(getattr(processo, "modalidade", None), "nome", "") or "-")}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y = _canvas_draw_wrapped(
        c,
        f'Quantidade de anexos incorporados: {int(total_documentos)}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica-Bold',
        size=12,
        leading=16,
    )
    y = _canvas_draw_wrapped(
        c,
        f'Gerado em: {timezone.localtime().strftime("%d/%m/%Y %H:%M:%S")}',
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y -= 10
    c.setFillColor(colors.HexColor('#eef4fa'))
    c.roundRect(40, y - 74, width - 80, 64, 14, stroke=0, fill=1)
    c.setFillColor(colors.HexColor('#1b3a5a'))
    c.setFont('Helvetica-Bold', 10)
    c.drawString(54, y - 30, 'Padrões aplicados nesta exportação')
    c.setFont('Helvetica', 9.4)
    c.drawString(54, y - 46, 'A4 padronizado | paginação institucional | compactação documental | PDF pesquisável')
    c.drawString(54, y - 60, 'Estrutura preparada para conferência e remessa ao e-TCM/BA.')

    c.setStrokeColor(colors.HexColor('#d5deea'))
    c.line(40, 112, width - 40, 112)
    c.setFillColor(colors.HexColor('#4a607b'))
    c.setFont('Helvetica', 9)
    c.drawString(40, 94, 'Documento gerado automaticamente pelo módulo Documentos do SIREL.')

    c.showPage()
    c.save()
    return bio.getvalue()


def _build_capa_documento_pdf(processo: Processo, documento: dict, idx: int, total: int) -> bytes:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4

    y = _draw_dossie_banner(
        c,
        width=width,
        height=height,
        title='Capa do Documento',
        subtitle=f'Anexo {idx} de {total} no processo consolidado',
        accent='#245a63',
        tag='Anexo',
    )

    _draw_info_card(
        c,
        x=40,
        y=y - 88,
        w=112,
        h=64,
        label='Posição',
        value=f'{idx}/{total}',
        accent='#245a63',
    )
    _draw_info_card(
        c,
        x=166,
        y=y - 88,
        w=148,
        h=64,
        label='Processo',
        value=str(processo.numero_processo_principal or '-'),
        accent='#1f5fa8',
    )
    _draw_info_card(
        c,
        x=328,
        y=y - 88,
        w=106,
        h=64,
        label='Origem',
        value=str(documento.get('origem') or '-'),
        accent='#9b6b2b',
    )
    _draw_info_card(
        c,
        x=448,
        y=y - 88,
        w=108,
        h=64,
        label='Adicionado por',
        value=str(documento.get('adicionado_por') or '-')[:34],
        accent='#5c4aa8',
    )

    c.setFillColor(colors.HexColor('#11263f'))
    y = y - 118
    y = _canvas_draw_wrapped(
        c,
        str(documento.get('titulo') or 'Documento'),
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica-Bold',
        size=15,
        leading=20,
    )
    y -= 4
    y = _canvas_draw_wrapped(
        c,
        f"Arquivo: {str(documento.get('arquivo_nome') or '-')}",
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y = _canvas_draw_wrapped(
        c,
        f"Origem: {str(documento.get('origem') or '-')}",
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    y = _canvas_draw_wrapped(
        c,
        f"Tipo interno: {str(documento.get('descricao') or '-')}",
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    _canvas_draw_wrapped(
        c,
        f"Adicionado por: {str(documento.get('adicionado_por') or '-')}",
        x=40,
        y=y,
        width=width - 80,
        font='Helvetica',
        size=11,
        leading=15,
    )
    c.setFillColor(colors.HexColor('#eef4fa'))
    c.roundRect(40, 142, width - 80, 54, 14, stroke=0, fill=1)
    c.setFillColor(colors.HexColor('#245a63'))
    c.setFont('Helvetica-Bold', 10)
    c.drawString(54, 174, 'Finalidade desta capa')
    c.setFont('Helvetica', 9.3)
    c.drawString(54, 158, 'Identificar claramente o anexo seguinte, preservar a ordem cronológica e facilitar auditoria.')

    c.setStrokeColor(colors.HexColor('#d5deea'))
    c.line(40, 112, width - 40, 112)
    c.setFillColor(colors.HexColor('#4a607b'))
    c.setFont('Helvetica', 9)
    c.drawString(40, 94, 'A próxima(s) página(s) contém o conteúdo do anexo.')

    c.showPage()
    c.save()
    return bio.getvalue()


def _build_texto_pdf(titulo: str, texto: str) -> bytes:
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'DocTextoTitle',
        parent=styles['Heading1'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#153153'),
        alignment=0,
    )
    body_style = ParagraphStyle(
        'DocTextoBody',
        parent=styles['BodyText'],
        fontSize=10.4,
        leading=14,
        textColor=colors.HexColor('#1d2b40'),
        alignment=TA_JUSTIFY,
    )
    elementos = [Paragraph(escape(str(titulo or 'Documento')), title_style), Spacer(1, 8)]
    linhas = str(texto or '').replace('\r', '').split('\n')
    if not linhas:
        linhas = ['']
    for linha in linhas:
        if str(linha).strip():
            elementos.append(Paragraph(escape(str(linha)), body_style))
        else:
            elementos.append(Paragraph('\u00a0', body_style))
        elementos.append(Spacer(1, 2))
    pdf.build(elementos)
    return bio.getvalue()


def _docx_bytes_to_text(payload: bytes) -> str:
    doc = Document(BytesIO(payload))
    linhas = []
    for p in doc.paragraphs:
        linhas.append(str(p.text or '').rstrip())
    for table in doc.tables:
        for row in table.rows:
            celulas = [str(cell.text or '').strip() for cell in row.cells]
            celulas = [c for c in celulas if c]
            if celulas:
                linhas.append(' | '.join(celulas))
    out = '\n'.join(linhas).strip()
    return out or 'Documento sem conteúdo textual extraível.'


def _image_bytes_to_pdf(payload: bytes) -> bytes:
    src = PILImage.open(BytesIO(payload))
    frames = []
    for frame in ImageSequence.Iterator(src):
        frames.append(frame.convert('RGB').copy())
    if not frames:
        raise ValueError('Imagem sem frames válidos.')
    bio = BytesIO()
    frames[0].save(bio, format='PDF', save_all=True, append_images=frames[1:])
    for frame in frames:
        frame.close()
    src.close()
    return bio.getvalue()


def _arquivo_para_pdf_bytes(*, arquivo_nome: str, payload: bytes) -> bytes:
    nome = str(arquivo_nome or '').strip().lower()
    ext = Path(nome).suffix.lower()
    mime = str(mimetypes.guess_type(nome)[0] or '').lower()

    if ext == '.pdf' or mime == 'application/pdf':
        return payload

    if ext in {'.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff', '.gif', '.webp'} or mime.startswith('image/'):
        return _image_bytes_to_pdf(payload)

    if ext == '.docx':
        texto = _docx_bytes_to_text(payload)
        return _build_texto_pdf(f'Conteúdo textual de {arquivo_nome}', texto)

    if ext in {'.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm'} or mime.startswith('text/'):
        texto = _decode_bytes_texto(payload)
        if ext in {'.html', '.htm'}:
            texto = _rich_html_to_plain_text(texto)
        return _build_texto_pdf(f'Conteúdo textual de {arquivo_nome}', texto)

    raise ValueError(f'Formato não suportado para incorporação automática: {ext or "desconhecido"}')


def _merge_pdf_parts(pdf_parts: list[bytes]) -> bytes:
    try:
        from pypdf import PdfReader, PdfWriter
    except Exception as exc:
        raise RuntimeError(
            'Para consolidar anexos em PDF é necessário instalar a dependência "pypdf". '
            'Execute: pip install pypdf'
        ) from exc

    writer = PdfWriter()
    for idx, blob in enumerate(pdf_parts, start=1):
        if not blob:
            continue
        try:
            reader = PdfReader(BytesIO(blob))
        except Exception as exc:
            aviso = (
                f'Não foi possível ler uma das partes PDF da consolidação (parte {idx}).\n'
                f'Motivo: {exc}'
            )
            reader = PdfReader(BytesIO(_build_texto_pdf('Falha de leitura de PDF', aviso)))
        if getattr(reader, 'is_encrypted', False):
            try:
                reader.decrypt('')
            except Exception:
                pass
        for page in reader.pages:
            writer.add_page(page)

    out = BytesIO()
    writer.write(out)
    return out.getvalue()


@lru_cache(maxsize=1)
def _ocr_engine_easyocr():
    try:
        import easyocr
    except Exception:
        return None
    try:
        return easyocr.Reader(['pt', 'en'], gpu=False, verbose=False)
    except Exception:
        return None


def _extrair_texto_ocr(
    img: PILImage.Image,
    *,
    max_width: int = 220,
    canvas_size: int = 512,
    paragraph: bool = False,
) -> str:
    reader = _ocr_engine_easyocr()
    if reader is None:
        return ''

    base = img.convert('RGB')
    if base.width > max_width:
        ratio = max_width / float(base.width)
        base = base.resize((max_width, max(1, int(base.height * ratio))), PILImage.Resampling.LANCZOS)

    try:
        result = reader.readtext(
            np.array(base),
            detail=0,
            paragraph=paragraph,
            batch_size=1,
            workers=0,
            decoder='greedy',
            beamWidth=1,
            min_size=8,
            mag_ratio=1.0,
            canvas_size=canvas_size,
        )
    except Exception:
        return ''

    if not result:
        return ''
    linhas = []
    for row in result:
        texto = str(row or '').strip()
        if texto:
            linhas.append(texto)
    texto = '\n'.join(linhas)
    return re.sub(r'\n{3,}', '\n\n', texto).strip()


def _extrair_textos_ocr_batch(
    imagens: list[PILImage.Image],
    *,
    max_width: int = 220,
    canvas_size: int = 512,
    paragraph: bool = False,
) -> list[str]:
    reader = _ocr_engine_easyocr()
    if reader is None:
        return ['' for _ in imagens]
    if not imagens:
        return []

    prepared = []
    for img in imagens:
        base = img.convert('RGB')
        if base.width > max_width:
            ratio = max_width / float(base.width)
            base = base.resize((max_width, max(1, int(base.height * ratio))), PILImage.Resampling.LANCZOS)
        prepared.append(np.array(base))

    try:
        resultados = reader.readtext_batched(
            prepared,
            detail=0,
            paragraph=paragraph,
            batch_size=min(12, max(1, len(prepared))),
            workers=0,
            decoder='greedy',
            beamWidth=1,
            min_size=8,
            mag_ratio=1.0,
            canvas_size=canvas_size,
        )
    except Exception:
        return [
            _extrair_texto_ocr(
                img,
                max_width=max_width,
                canvas_size=canvas_size,
                paragraph=paragraph,
            )
            for img in imagens
        ]

    textos = []
    for row in resultados:
        if isinstance(row, (list, tuple)):
            linhas = [str(x or '').strip() for x in row if str(x or '').strip()]
            textos.append(re.sub(r'\n{3,}', '\n\n', '\n'.join(linhas)).strip())
        else:
            textos.append(str(row or '').strip())
    while len(textos) < len(imagens):
        textos.append('')
    return textos[:len(imagens)]


def _extrair_fragmentos_texto_paginas_pdf(pdf_data: bytes) -> list[dict]:
    try:
        import fitz
    except Exception:
        return []

    pages = []
    doc = fitz.open(stream=pdf_data, filetype='pdf')
    try:
        for page in doc:
            payload = {
                'page_width': float(page.rect.width),
                'page_height': float(page.rect.height),
                'text': '',
                'fragments': [],
            }
            textos = []
            try:
                data = page.get_text('dict')
            except Exception:
                data = {}
            for block in data.get('blocks', []):
                if int(block.get('type') or 0) != 0:
                    continue
                for line in block.get('lines', []):
                    partes = []
                    for span in line.get('spans', []):
                        texto = re.sub(r'\s+', ' ', str(span.get('text') or '')).strip()
                        bbox = span.get('bbox') or ()
                        if not texto or len(bbox) != 4:
                            continue
                        payload['fragments'].append(
                            {
                                'text': texto,
                                'bbox': tuple(float(v) for v in bbox),
                            }
                        )
                        partes.append(texto)
                    if partes:
                        textos.append(' '.join(partes))
            payload['text'] = '\n'.join(textos).strip()
            pages.append(payload)
    finally:
        doc.close()
    return pages


def _mapear_fragmentos_pdf_para_pagina_px(
    page_payload: dict,
    *,
    source_size_px: tuple[int, int],
    a4_px: tuple[int, int],
    content_scale: float,
) -> list[dict]:
    fragments = list(page_payload.get('fragments') or [])
    page_width = float(page_payload.get('page_width') or 0.0)
    page_height = float(page_payload.get('page_height') or 0.0)
    source_w_px, source_h_px = source_size_px
    if not fragments or page_width <= 0 or page_height <= 0 or source_w_px <= 0 or source_h_px <= 0:
        return []

    footer_h = 48
    content_w_px = int(a4_px[0])
    content_h_px = int(a4_px[1] - footer_h)
    scale_limit = min(max(float(content_scale or 1.0), 0.45), 1.0)
    target_w = content_w_px * scale_limit
    target_h = content_h_px * scale_limit
    ratio = min(target_w / float(source_w_px), target_h / float(source_h_px))
    scaled_w = float(source_w_px) * ratio
    scaled_h = float(source_h_px) * ratio
    x_pad = (content_w_px - scaled_w) / 2.0
    y_pad = (content_h_px - scaled_h) / 2.0

    fx = float(source_w_px) / page_width
    fy = float(source_h_px) / page_height
    out = []
    for fragment in fragments:
        texto = re.sub(r'\s+', ' ', str(fragment.get('text') or '')).strip()
        bbox = fragment.get('bbox') or ()
        if not texto or len(bbox) != 4:
            continue
        x0, y0, x1, y1 = [float(v) for v in bbox]
        px0 = x_pad + (x0 * fx * ratio)
        py0 = y_pad + (y0 * fy * ratio)
        px1 = x_pad + (x1 * fx * ratio)
        py1 = y_pad + (y1 * fy * ratio)
        if px1 <= px0 or py1 <= py0:
            continue
        out.append({'text': texto, 'bbox_px': (px0, py0, px1, py1)})
    return out


def _extrair_fragmentos_ocr_pagina(
    img: PILImage.Image,
    *,
    max_width: int = 380,
    canvas_size: int = 960,
) -> list[dict]:
    reader = _ocr_engine_easyocr()
    if reader is None:
        return []

    base = img.convert('RGB')
    scale = 1.0
    if base.width > max_width:
        scale = max_width / float(base.width)
        base = base.resize((max_width, max(1, int(base.height * scale))), PILImage.Resampling.LANCZOS)

    try:
        resultados = reader.readtext(
            np.array(base),
            detail=1,
            paragraph=False,
            batch_size=1,
            workers=0,
            decoder='greedy',
            beamWidth=1,
            min_size=8,
            mag_ratio=1.0,
            canvas_size=canvas_size,
        )
    except Exception:
        return []

    fragments = []
    inv_scale = 1.0 / scale if scale else 1.0
    for row in resultados or []:
        if not isinstance(row, (list, tuple)) or len(row) < 2:
            continue
        bbox = row[0] or ()
        texto = re.sub(r'\s+', ' ', str(row[1] or '')).strip()
        if not texto or not bbox:
            continue
        xs = [float(pt[0]) for pt in bbox if isinstance(pt, (list, tuple)) and len(pt) >= 2]
        ys = [float(pt[1]) for pt in bbox if isinstance(pt, (list, tuple)) and len(pt) >= 2]
        if not xs or not ys:
            continue
        x0, x1 = min(xs) * inv_scale, max(xs) * inv_scale
        y0, y1 = min(ys) * inv_scale, max(ys) * inv_scale
        if x1 <= x0 or y1 <= y0:
            continue
        fragments.append({'text': texto, 'bbox_px': (x0, y0, x1, y1)})
    return fragments


def _draw_hidden_text_fragments(
    c: canvas.Canvas,
    fragments: list[dict],
    *,
    page_width: float,
    page_height: float,
    page_px: tuple[int, int],
) -> None:
    if not fragments:
        return

    width_px, height_px = page_px
    if width_px <= 0 or height_px <= 0:
        return

    for fragment in fragments:
        texto = re.sub(r'\s+', ' ', str(fragment.get('text') or '')).strip()
        bbox_px = fragment.get('bbox_px') or ()
        if not texto or len(bbox_px) != 4:
            continue
        x0_px, y0_px, x1_px, y1_px = [float(v) for v in bbox_px]
        box_w_px = x1_px - x0_px
        box_h_px = y1_px - y0_px
        if box_w_px <= 1 or box_h_px <= 1:
            continue

        x = x0_px * page_width / float(width_px)
        y = page_height - (y1_px * page_height / float(height_px))
        box_w = box_w_px * page_width / float(width_px)
        box_h = box_h_px * page_height / float(height_px)
        font_size = max(4.0, min(18.0, box_h * 0.82))
        base_width = pdfmetrics.stringWidth(texto, 'Helvetica', font_size) or 1.0
        horiz_scale = max(35.0, min(400.0, (box_w / base_width) * 100.0))

        txt_obj = c.beginText()
        txt_obj.setTextOrigin(x, y)
        txt_obj.setFont('Helvetica', font_size)
        if hasattr(txt_obj, 'setTextRenderMode'):
            try:
                txt_obj.setTextRenderMode(3)
            except Exception:
                pass
        if hasattr(txt_obj, 'setHorizScale'):
            try:
                txt_obj.setHorizScale(horiz_scale)
            except Exception:
                pass
        txt_obj.textLine(texto[:500])
        c.drawText(txt_obj)


def _pdf_para_imagens(pdf_data: bytes, *, dpi: int = 145) -> list[PILImage.Image]:
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError(
            'Dependência ausente para padronização de páginas (PyMuPDF). '
            'Execute: pip install pymupdf'
        ) from exc

    out = []
    doc = fitz.open(stream=pdf_data, filetype='pdf')
    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    for page in doc:
        pix = page.get_pixmap(matrix=mat, alpha=False)
        mode = 'RGB'
        img = PILImage.frombytes(mode, [pix.width, pix.height], pix.samples)
        out.append(img)
    doc.close()
    return out


def _extrair_texto_paginas_pdf(pdf_data: bytes) -> list[str]:
    try:
        from pypdf import PdfReader
    except Exception:
        return []
    out = []
    try:
        reader = PdfReader(BytesIO(pdf_data))
    except Exception:
        return []
    for page in reader.pages:
        try:
            txt = str(page.extract_text() or '').strip()
        except Exception:
            txt = ''
        out.append(txt)
    return out


def _fit_image_a4_canvas(img: PILImage.Image, *, a4_px: tuple[int, int], content_scale: float = 1.0) -> PILImage.Image:
    base = img.convert('RGB')
    cw, ch = a4_px
    scale = min(max(content_scale, 0.45), 1.0)
    tw = int(cw * scale)
    th = int(ch * scale)

    ratio = min(tw / max(base.width, 1), th / max(base.height, 1))
    nw = max(1, int(base.width * ratio))
    nh = max(1, int(base.height * ratio))
    resized = base.resize((nw, nh), PILImage.Resampling.LANCZOS)

    canvas_img = PILImage.new('RGB', (cw, ch), 'white')
    x = (cw - nw) // 2
    y = (ch - nh) // 2
    canvas_img.paste(resized, (x, y))
    return canvas_img


def _render_a4_page_image(
    img: PILImage.Image,
    *,
    page_num: int,
    total_pages: int,
    a4_px: tuple[int, int],
    footer_left: str = 'SIREL Modular',
    footer_center: str = '',
) -> PILImage.Image:
    footer_h = 48
    width_px, height_px = a4_px
    base = img.convert('RGB')
    content_h = max(1, height_px - footer_h)
    content_canvas = PILImage.new('RGB', (width_px, content_h), 'white')

    ratio = min(width_px / max(base.width, 1), content_h / max(base.height, 1))
    nw = max(1, int(base.width * ratio))
    nh = max(1, int(base.height * ratio))
    resized = base.resize((nw, nh), PILImage.Resampling.LANCZOS)
    x = (width_px - nw) // 2
    y = (content_h - nh) // 2
    content_canvas.paste(resized, (x, y))

    page = PILImage.new('RGB', (width_px, height_px), 'white')
    page.paste(content_canvas, (0, 0))
    return page


def _compress_page_image(img: PILImage.Image, *, jpeg_quality: int) -> bytes:
    out = BytesIO()
    img.save(
        out,
        format='JPEG',
        quality=max(30, min(95, int(jpeg_quality))),
        optimize=True,
        progressive=True,
        dpi=(150, 150),
    )
    return out.getvalue()


def _build_a4_page_pdf(
    jpeg_data: bytes,
    *,
    page_num: int,
    total_pages: int,
    ocr_text: str = '',
    text_fragments: list[dict] | None = None,
    footer_left: str = 'SIREL Modular',
    footer_center: str = '',
) -> bytes:
    jpeg_buf = BytesIO(jpeg_data)
    out = BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    width, height = A4
    c.drawImage(ImageReader(jpeg_buf), 0, 0, width=width, height=height, preserveAspectRatio=False, mask='auto')
    c.setStrokeColor(colors.HexColor('#cfd7e3'))
    c.line(16, 24, width - 16, 24)
    c.setFillColor(colors.HexColor('#3f5570'))
    c.setFont('Helvetica-Bold', 9)
    c.drawString(16, 8, str(footer_left or 'SIREL Modular')[:64])
    if footer_center:
        c.setFont('Helvetica', 8)
        c.setFillColor(colors.HexColor('#60758f'))
        c.drawCentredString(width / 2, 8, str(footer_center or '')[:96])
        c.setFillColor(colors.HexColor('#3f5570'))
        c.setFont('Helvetica-Bold', 9)
    c.drawRightString(width - 16, 8, f'Página {int(page_num)} de {int(total_pages)}')

    if text_fragments:
        footer_fragments = [
            {'text': str(footer_left or 'SIREL Modular')[:64], 'bbox_px': (16, 1728, 320, 1748)},
            {'text': f'Página {int(page_num)} de {int(total_pages)}', 'bbox_px': (980, 1728, 1228, 1748)},
        ]
        if footer_center:
            footer_fragments.append({'text': str(footer_center or '')[:96], 'bbox_px': (430, 1728, 830, 1748)})
        _draw_hidden_text_fragments(
            c,
            list(text_fragments) + footer_fragments,
            page_width=width,
            page_height=height,
            page_px=(1240, 1754),
        )
    elif ocr_text:
        hidden_text = re.sub(r'\s+', ' ', str(ocr_text or '')).strip()
        footer_text = f'Página {int(page_num)} de {int(total_pages)}'
        if footer_text not in hidden_text:
            hidden_text = f'{hidden_text} {footer_text}'.strip()
        if footer_left and footer_left not in hidden_text:
            hidden_text = f'{hidden_text} {footer_left}'.strip()
        if footer_center and footer_center not in hidden_text:
            hidden_text = f'{hidden_text} {footer_center}'.strip()
        line_size = 4.0
        leading = 4.8
        max_lines = max(1, int((height - 20) / leading))
        linhas = [hidden_text[i:i + 180] for i in range(0, len(hidden_text), 180)][:max_lines]

        txt_obj = c.beginText(14, height - 16)
        txt_obj.setFont('Helvetica', line_size)
        txt_obj.setLeading(leading)
        if hasattr(txt_obj, 'setTextRenderMode'):
            try:
                txt_obj.setTextRenderMode(3)
            except Exception:
                pass
        for ln in linhas:
            line = ln.strip()
            if line:
                txt_obj.textLine(line[:220])
        c.drawText(txt_obj)

    c.showPage()
    c.save()
    return out.getvalue()


def _normalizar_pdf_para_padrao_documental(
    pdf_data: bytes,
    *,
    aplicar_ocr: bool = True,
    pagina_max_kb: int = 500,
    footer_left: str = 'SIREL Modular',
    footer_center: str = '',
) -> tuple[bytes, dict]:
    max_page_bytes = int(pagina_max_kb * 1024)
    a4_px = (1240, 1754)  # A4 em ~150 DPI
    imagens = _pdf_para_imagens(pdf_data, dpi=150)
    camadas_originais = _extrair_fragmentos_texto_paginas_pdf(pdf_data)
    textos_originais = [str(page.get('text') or '').strip() for page in camadas_originais]
    total = len(imagens) or 1
    parts = []
    ocr_reader = _ocr_engine_easyocr() if aplicar_ocr else None
    ocr_habilitado = bool(ocr_reader is not None)
    ocr_aplicado = 0
    jpeg_pages: list[bytes] = []
    ocr_texts: list[str] = ['' for _ in range(total)]
    page_fragments: list[list[dict]] = [[] for _ in range(total)]
    page_layouts: list[dict] = []

    for i, img in enumerate(imagens, start=1):
        best_jpeg = b''
        prepared = None
        paged = None
        chosen_scale = 1.0
        source_size_px = (int(img.width), int(img.height))
        try:
            for content_scale in (1.0, 0.95, 0.90, 0.85, 0.80, 0.75, 0.70):
                if prepared is not None:
                    try:
                        prepared.close()
                    except Exception:
                        pass
                    prepared = None
                if paged is not None:
                    try:
                        paged.close()
                    except Exception:
                        pass
                    paged = None

                prepared = _fit_image_a4_canvas(img, a4_px=(a4_px[0], a4_px[1] - 48), content_scale=content_scale)
                paged = _render_a4_page_image(
                    prepared,
                    page_num=i,
                    total_pages=total,
                    a4_px=a4_px,
                    footer_left=footer_left,
                    footer_center=footer_center,
                )
                for quality in (80, 72, 64, 56, 48, 40, 34):
                    jpeg_data = _compress_page_image(paged, jpeg_quality=quality)
                    best_jpeg = jpeg_data
                    if len(jpeg_data) <= max_page_bytes:
                        chosen_scale = content_scale
                        break
                if best_jpeg and len(best_jpeg) <= max_page_bytes:
                    break
        finally:
            try:
                img.close()
            except Exception:
                pass
            if prepared is not None:
                try:
                    prepared.close()
                except Exception:
                    pass
            if paged is not None:
                try:
                    paged.close()
                except Exception:
                    pass

        jpeg_pages.append(best_jpeg)
        page_layouts.append(
            {
                'source_size_px': source_size_px,
                'content_scale': float(chosen_scale),
            }
        )

    paginas_sem_texto = [
        idx for idx in range(total)
        if not (textos_originais[idx] if idx < len(textos_originais) else '').strip()
    ]
    if ocr_habilitado and paginas_sem_texto:
        repro_chunk_size = 4
        for offset in range(0, len(paginas_sem_texto), repro_chunk_size):
            idx_chunk = paginas_sem_texto[offset:offset + repro_chunk_size]
            chunk_images = []
            try:
                for idx in idx_chunk:
                    chunk_images.append(PILImage.open(BytesIO(jpeg_pages[idx])).convert('RGB'))
                chunk_texts = _extrair_textos_ocr_batch(
                    chunk_images,
                    max_width=380,
                    canvas_size=960,
                    paragraph=False,
                )
            finally:
                for chunk_img in chunk_images:
                    try:
                        chunk_img.close()
                    except Exception:
                        pass

            for idx, chunk_text in zip(idx_chunk, chunk_texts):
                ocr_texts[idx] = str(chunk_text or '').strip()
        ocr_aplicado = len(paginas_sem_texto)

    for idx in range(total):
        original_txt = textos_originais[idx] if idx < len(textos_originais) else ''
        if original_txt and idx < len(camadas_originais) and idx < len(page_layouts):
            page_fragments[idx] = _mapear_fragmentos_pdf_para_pagina_px(
                camadas_originais[idx],
                source_size_px=tuple(page_layouts[idx].get('source_size_px') or (0, 0)),
                a4_px=a4_px,
                content_scale=float(page_layouts[idx].get('content_scale') or 1.0),
            )
        merged_text = str(ocr_texts[idx] or '').strip()
        if original_txt:
            if merged_text:
                if original_txt not in merged_text:
                    merged_text = f'{merged_text}\n{original_txt}'.strip()
            else:
                merged_text = original_txt
        elif ocr_habilitado and idx in paginas_sem_texto and idx < len(jpeg_pages):
            try:
                with PILImage.open(BytesIO(jpeg_pages[idx])) as ocr_img:
                    ocr_rgb = ocr_img.convert('RGB')
                    page_fragments[idx] = _extrair_fragmentos_ocr_pagina(
                        ocr_rgb,
                        max_width=380,
                        canvas_size=960,
                    )
                    ocr_rgb.close()
            except Exception:
                page_fragments[idx] = []
        ocr_texts[idx] = merged_text

    for i, jpeg_data in enumerate(jpeg_pages, start=1):
        parts.append(
            _build_a4_page_pdf(
                jpeg_data,
                page_num=i,
                total_pages=total,
                ocr_text=ocr_texts[i - 1] if i - 1 < len(ocr_texts) else '',
                text_fragments=page_fragments[i - 1] if i - 1 < len(page_fragments) else None,
                footer_left=footer_left,
                footer_center=footer_center,
            )
        )

    normalizado = _merge_pdf_parts(parts)
    return normalizado, {
        'total_paginas': len(parts),
        'paginas_ocr': int(ocr_aplicado),
        'paginas_texto_nativo': max(0, len(parts) - int(ocr_aplicado)),
        'ocr_habilitado': ocr_habilitado,
        'ocr_engine': 'easyocr' if ocr_habilitado else '',
    }


def _split_pdf_by_size(pdf_data: bytes, *, max_bytes: int) -> list[bytes]:
    try:
        from pypdf import PdfReader, PdfWriter
    except Exception as exc:
        raise RuntimeError(
            'Para divisão do PDF por tamanho é necessário instalar "pypdf". '
            'Execute: pip install pypdf'
        ) from exc

    reader = PdfReader(BytesIO(pdf_data))

    def build_blob(pages) -> bytes:
        writer = PdfWriter()
        for pg in pages:
            writer.add_page(pg)
        out = BytesIO()
        writer.write(out)
        return out.getvalue()

    parts = []
    current_pages = []
    for page in reader.pages:
        trial = current_pages + [page]
        trial_blob = build_blob(trial)
        if current_pages and len(trial_blob) > int(max_bytes):
            parts.append(build_blob(current_pages))
            current_pages = [page]
        else:
            current_pages = trial
    if current_pages:
        parts.append(build_blob(current_pages))
    return parts or [pdf_data]


def _gerar_processo_etcm_zip_ou_pdf(processo: Processo) -> tuple[bytes, str, str, dict]:
    consolidado_raw, documentos = _gerar_processo_consolidado_pdf(processo)
    padrao_pdf, meta = _normalizar_pdf_para_padrao_documental(
        consolidado_raw,
        aplicar_ocr=True,
        pagina_max_kb=500,
        footer_left=f'Processo {processo.numero_processo_principal}',
        footer_center='Processo padronizado',
    )

    partes = _split_pdf_by_size(padrao_pdf, max_bytes=4 * 1024 * 1024)
    base = f"processo_etcm_{_slug_ascii(processo.numero_processo_principal)}_{processo.id}"
    meta['partes'] = len(partes)
    meta['documentos'] = len(documentos)

    if len(partes) <= 1:
        return partes[0], f'{base}.pdf', 'application/pdf', meta

    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
        for idx, part in enumerate(partes, start=1):
            zf.writestr(f'{base}_parte_{idx:02d}.pdf', part)
    return zip_buf.getvalue(), f'{base}.zip', 'application/zip', meta


def _gerar_processo_consolidado_pdf(processo: Processo) -> tuple[bytes, list[dict]]:
    documentos = _coletar_documentos_processo_consolidado(processo)
    pdf_parts = [_build_capa_processo_pdf(processo, len(documentos))]

    if not documentos:
        aviso = (
            'Este processo ainda não possui anexos em DocumentoProcessoWorkflow ou docs.ProcessoAnexo.\n'
            'A geração integral permanece disponível mesmo para processos parciais/incompletos.'
        )
        pdf_parts.append(_build_texto_pdf('Sem anexos incorporáveis no momento', aviso))
        return _merge_pdf_parts(pdf_parts), documentos

    for idx, documento in enumerate(documentos, start=1):
        pdf_parts.append(_build_capa_documento_pdf(processo, documento, idx, len(documentos)))
        try:
            payload = _ler_bytes_arquivo_field(documento.get('arquivo'))
            pdf_parts.append(_arquivo_para_pdf_bytes(arquivo_nome=documento.get('arquivo_nome', ''), payload=payload))
        except Exception as exc:
            aviso = (
                f'Não foi possível incorporar automaticamente o arquivo "{documento.get("arquivo_nome", "-")}".\n'
                f'Motivo técnico: {exc}\n'
                'O registro do anexo foi preservado nesta montagem para manter a sequência do processo.'
            )
            pdf_parts.append(_build_texto_pdf(f'Falha de incorporação - Documento {idx}', aviso))

    return _merge_pdf_parts(pdf_parts), documentos


def _balance_reportlab_markup_tokens(texto: str) -> str:
    if not texto:
        return ''
    token_re = re.compile(r'(\[\[(?:/)?[BIU]\]\])')
    open_tokens = {'[[B]]': 'B', '[[I]]': 'I', '[[U]]': 'U'}
    close_tokens = {'[[/B]]': 'B', '[[/I]]': 'I', '[[/U]]': 'U'}
    close_for = {'B': '[[/B]]', 'I': '[[/I]]', 'U': '[[/U]]'}

    out = []
    stack = []
    for part in token_re.split(texto):
        if not part:
            continue
        if part in open_tokens:
            out.append(part)
            stack.append(open_tokens[part])
            continue
        if part in close_tokens:
            tag = close_tokens[part]
            if tag not in stack:
                continue
            while stack and stack[-1] != tag:
                out.append(close_for[stack.pop()])
            if stack and stack[-1] == tag:
                out.append(part)
                stack.pop()
            continue
        out.append(part)

    while stack:
        out.append(close_for[stack.pop()])
    return ''.join(out)


def _rich_html_to_plain_text(valor: str) -> str:
    texto = str(valor or '')
    if not texto:
        return ''
    texto = re.sub(r'(?is)<\s*br\s*/?\s*>', '\n', texto)
    texto = re.sub(r'(?is)</\s*p\s*>', '\n\n', texto)
    texto = re.sub(r'(?is)</\s*div\s*>', '\n', texto)
    texto = re.sub(r'(?is)<\s*li[^>]*>', '\n- ', texto)
    texto = re.sub(r'(?is)</\s*li\s*>', '', texto)
    texto = re.sub(r'(?is)<[^>]+>', '', texto)
    texto = html_unescape(texto)
    texto = texto.replace('\r\n', '\n').replace('\r', '\n')
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto.strip()


def _rich_html_to_pdf_markup(valor: str) -> str:
    texto = str(valor or '')
    if not texto:
        return '-'

    substituicoes = [
        (r'(?is)<\s*strong[^>]*>', '[[B]]'),
        (r'(?is)</\s*strong\s*>', '[[/B]]'),
        (r'(?is)<\s*b(?:\s[^>]*)?>', '[[B]]'),
        (r'(?is)</\s*b\s*>', '[[/B]]'),
        (r'(?is)<\s*em[^>]*>', '[[I]]'),
        (r'(?is)</\s*em\s*>', '[[/I]]'),
        (r'(?is)<\s*i(?:\s[^>]*)?>', '[[I]]'),
        (r'(?is)</\s*i\s*>', '[[/I]]'),
        (r'(?is)<\s*u(?:\s[^>]*)?>', '[[U]]'),
        (r'(?is)</\s*u\s*>', '[[/U]]'),
        (r'(?is)<\s*br\s*/?\s*>', '\n'),
        (r'(?is)</\s*p\s*>', '\n\n'),
        (r'(?is)</\s*div\s*>', '\n'),
        (r'(?is)<\s*li[^>]*>', '\n- '),
        (r'(?is)</\s*li\s*>', ''),
    ]
    for pattern, repl in substituicoes:
        texto = re.sub(pattern, repl, texto)

    texto = re.sub(r'(?is)<[^>]+>', '', texto)
    texto = html_unescape(texto)
    texto = escape(texto)
    texto = _balance_reportlab_markup_tokens(texto)
    texto = texto.replace('[[B]]', '<b>').replace('[[/B]]', '</b>')
    texto = texto.replace('[[I]]', '<i>').replace('[[/I]]', '</i>')
    texto = texto.replace('[[U]]', '<u>').replace('[[/U]]', '</u>')
    texto = texto.replace('\r\n', '\n').replace('\r', '\n')
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = texto.strip()
    if not texto:
        return '-'
    return texto.replace('\n', '<br/>')


def _licitacao_ci_context(processo: Processo, codigo: str, *, form_data=None, usuario=None):
    item = LICITACAO_CHECKLIST_MAP.get(codigo)
    if not item or not item.get('gerar_ci_html'):
        return None

    hoje = timezone.localdate()
    ano_ref = int(hoje.year)
    objeto_proc = str(processo.objeto or '-').strip()
    secretaria_ref = _secretaria_composta(processo)
    numero_proc_adm = f'{processo.numero_processo_adm or processo.numero_processo_principal}/{processo.ano_referencia}'
    valor_estimado_brl = fmt_brl(getattr(processo, 'valor_estimado', 0))
    valor_estimado_extenso = valor_por_extenso(getattr(processo, 'valor_estimado', 0), moeda=True) or ''
    numero_ci_preview = f'{_proximo_numero_ci(ModuloSistema.LICITACAO, ano_ref):03d}/{ano_ref}'
    responsavel_default = (
        (usuario.get_full_name() or '').strip()
        if getattr(usuario, 'is_authenticated', False)
        else ''
    ) or (
        usuario.get_username()
        if getattr(usuario, 'is_authenticated', False)
        else ''
    ) or (
        str(getattr(getattr(processo, 'condutor_processo', None), 'nome', '') or '').strip()
    ) or 'Responsável pelo documento'
    cargo_default = (
        str(getattr(getattr(processo, 'condutor_processo', None), 'cargo', '') or '').strip()
        or 'Setor de Licitação'
    )

    modelo_map = {
        'ci_orcamento_reserva': {
            'tipo': 'CI',
            'titulo': 'Comunicação Interna',
            'numero_prefixo': 'CI n.',
            'numero_default': numero_ci_preview,
            'de_default': 'Assessoria de Licitação',
            'para_default': 'Setor de Contabilidade',
            'assunto_default': 'Solicitação de reserva orçamentária, conforme informações abaixo.',
            'corpo_default': (
                f'Prezado Senhor,\n\n'
                f'Em observância ao art. 150 da Lei 14.133/2021, solicitamos do setor contábil a '
                f'indicação dos recursos orçamentários para {objeto_proc}.\n\n'
                f'Caso exista previsão, favor indicar a fonte do recurso correspondente.\n\n'
                f'Atenciosamente,'
            ),
        },
        'ci_pgm_parecer': {
            'tipo': 'CI',
            'titulo': 'Comunicação Interna',
            'numero_prefixo': 'CI n.',
            'numero_default': numero_ci_preview,
            'de_default': 'Assessoria de Licitação',
            'para_default': 'Procuradoria Geral do Município',
            'assunto_default': f'Parecer jurídico sobre {objeto_proc}.',
            'corpo_default': (
                'Prezado Procurador,\n\n'
                'Solicito que seja previamente examinada a solicitação para contratação por dispensa '
                'de licitação, com emissão de parecer jurídico para que o processo transcorra dentro '
                'dos trâmites legais e com lisura administrativa.\n\n'
                'Segue anexo o processo administrativo contendo solicitação de despesa, termo de '
                'referência, cotações, mapa comparativo de preços e indicação de recursos orçamentários.\n\n'
                'Caso opine favoravelmente pela contratação, favor encaminhar parecer jurídico para '
                'ratificação da autoridade superior e publicidade do ato, nos termos da Lei 14.133/2021.'
            ),
        },
        'declaracao_nao_fracionamento': {
            'tipo': 'DECLARACAO',
            'titulo': 'Declaração de Não Fracionamento de Despesa',
            'numero_prefixo': 'Processo Administrativo n.º',
            'numero_default': numero_proc_adm,
            'de_default': secretaria_ref,
            'para_default': '',
            'assunto_default': f'Declaração sobre a contratação de {objeto_proc}.',
            'corpo_default': (
                f'Declaro, na qualidade de Ordenador de Despesas da {secretaria_ref}, que após pesquisa '
                'no banco de dados, constatamos que a presente contratação não ultrapassará o limite '
                'estabelecido no inciso II do artigo 75 da Lei Federal 14.133/2021 e que não haverá '
                f'fracionamento da despesa para contratações da mesma natureza no exercício financeiro de {ano_ref}.'
            ),
        },
        'ato_autorizacao': {
            'tipo': 'ATO',
            'titulo': 'Ato de Autorização',
            'numero_prefixo': 'Ato de Autorização n.º',
            'numero_default': f'____/{ano_ref}',
            'de_default': secretaria_ref,
            'para_default': '',
            'assunto_default': f'Autorização para contratação de {objeto_proc}.',
            'corpo_default': (
                f'Autorizo a contratação de {objeto_proc}, para atender as demandas da {secretaria_ref}.\n\n'
                f'O valor estimado para contratação será de aproximadamente {valor_estimado_brl} '
                f'({valor_estimado_extenso}).\n\n'
                f'Informamos que as despesas decorrentes correrão por conta das dotações orçamentárias '
                f'vinculadas ao Orçamento Anual/{ano_ref}, em compatibilidade com o PPA e a LDO.'
            ),
        },
        'termo_autuacao_agente': {
            'tipo': 'TERMO',
            'titulo': 'Termo de Autuação',
            'numero_prefixo': 'Autuação n.º',
            'numero_default': numero_proc_adm,
            'de_default': secretaria_ref,
            'para_default': 'Assessoria de Licitação',
            'assunto_default': objeto_proc,
            'corpo_default': (
                f'PROCESSO ADMINISTRATIVO: {numero_proc_adm}\n'
                f'SECRETARIA DE ORIGEM: {secretaria_ref}\n'
                f'SETOR DE DESTINO: Assessoria de Licitação\n'
                f'DATA DE ENTRADA: {_data_extenso_pt_br(hoje)}\n'
                f'ASSUNTO: {objeto_proc}\n\n'
                'Em data supra, eu, Agente de Contratação, nomeado por decreto vigente e lotado na '
                'Assessoria de Licitação, recebi o presente processo de dispensa para autuação, contendo '
                'solicitação de despesa e comunicação interna da unidade requisitante.\n\n'
                'O presente Termo de Autuação deverá ser juntado aos autos do referido processo.'
            ),
        },
    }
    modelo = modelo_map.get(codigo)
    if not modelo:
        return None

    pessoa_opts = list(Pessoa.objects.select_related('secretaria').order_by('nome')[:800])
    signatarios_catalogo = [
        {
            'id': p.id,
            'nome': p.nome,
            'cargo': str(getattr(p, 'cargo', '') or '').strip(),
            'secretaria': str(getattr(getattr(p, 'secretaria', None), 'sigla', '') or '').strip(),
        }
        for p in pessoa_opts
    ]
    signatarios_ids_default = []
    for pessoa in [getattr(processo, 'condutor_processo', None), getattr(processo, 'autoridade_competente', None)]:
        if pessoa and pessoa.id and str(pessoa.id) not in signatarios_ids_default:
            signatarios_ids_default.append(str(pessoa.id))

    data_documento = hoje
    if form_data:
        data_raw = str(form_data.get('data_documento') or '').strip()
        if data_raw:
            try:
                data_documento = datetime.fromisoformat(data_raw).date()
            except Exception:
                data_documento = hoje

    signatarios_ids = signatarios_ids_default
    if form_data:
        signatarios_ids = [str(x).strip() for x in form_data.getlist('signatarios') if str(x).strip()]

    pessoas_sel = []
    if signatarios_ids:
        pessoas_sel = list(
            Pessoa.objects
            .filter(id__in=[int(x) for x in signatarios_ids if x.isdigit()])
            .select_related('secretaria')
            .order_by('nome')
        )
    by_id = {str(p.id): p for p in pessoas_sel}
    signatarios_selecionados = []
    for sid in signatarios_ids:
        pessoa = by_id.get(str(sid))
        if not pessoa:
            continue
        signatarios_selecionados.append(
            {
                'id': pessoa.id,
                'nome': pessoa.nome,
                'cargo': str(getattr(pessoa, 'cargo', '') or '').strip(),
                'secretaria': str(getattr(getattr(pessoa, 'secretaria', None), 'sigla', '') or '').strip(),
            }
        )

    numero_documento = modelo['numero_default']
    localidade = _localidade_orgao_padrao()
    de_origem = modelo['de_default']
    para_destino = modelo['para_default']
    assunto = modelo['assunto_default']
    referencia = f'{processo.numero_processo_principal} - {objeto_proc}'
    corpo = modelo['corpo_default']
    responsavel_nome = responsavel_default
    responsavel_cargo = cargo_default

    if form_data:
        numero_documento = str(form_data.get('numero_documento') or numero_documento).strip() or numero_documento
        localidade = str(form_data.get('localidade') or localidade).strip() or localidade
        de_origem = str(form_data.get('de_origem') or de_origem).strip() or de_origem
        para_destino = str(form_data.get('para_destino') or para_destino).strip()
        assunto = str(form_data.get('assunto') or assunto).strip() or assunto
        referencia = str(form_data.get('referencia') or referencia).strip() or referencia
        corpo = str(form_data.get('corpo') or corpo).strip() or corpo
        responsavel_nome = str(form_data.get('responsavel_nome') or responsavel_nome).strip() or responsavel_nome
        responsavel_cargo = str(form_data.get('responsavel_cargo') or responsavel_cargo).strip() or responsavel_cargo

    return {
        'processo': processo,
        'orgao': _orgao_ativo(),
        'ci_documento_label': item['label'],
        'ci_fase': item['fase'],
        'titulo_pagina': f'Gerador de Documento - {item["label"]}',
        'voltar_label': 'Voltar para Licitação',
        'modelo_tipo': modelo['tipo'],
        'modelo_titulo': modelo['titulo'],
        'numero_prefixo': modelo['numero_prefixo'],
        'numero_documento': numero_documento,
        'localidade': localidade,
        'data_documento': data_documento,
        'data_documento_iso': data_documento.isoformat(),
        'data_documento_extenso': _data_extenso_pt_br(data_documento),
        'de_origem': de_origem,
        'para_destino': para_destino,
        'assunto': assunto,
        'referencia': referencia,
        'corpo': corpo,
        'responsavel_nome': responsavel_nome,
        'responsavel_cargo': responsavel_cargo,
        'signatarios_ids': signatarios_ids,
        'signatarios_catalogo': signatarios_catalogo,
        'signatarios_selecionados': signatarios_selecionados,
        'rodape_documento': 'Documento gerado pelo SIREL Modular.',
    }


def _export_licitacao_documento_pdf(*, processo: Processo, doc_ctx: dict):
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'LicitacaoDocTitle',
        parent=styles['Heading1'],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#153153'),
        alignment=1,
    )
    subtitle_style = ParagraphStyle(
        'LicitacaoDocSubtitle',
        parent=styles['BodyText'],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor('#5d6f8a'),
    )
    value_style = ParagraphStyle(
        'LicitacaoDocValue',
        parent=styles['BodyText'],
        fontSize=10.2,
        leading=14,
        textColor=colors.HexColor('#1d2b40'),
    )
    section_style = ParagraphStyle(
        'LicitacaoDocSection',
        parent=styles['Heading3'],
        fontSize=9.2,
        leading=11.8,
        textColor=colors.HexColor('#4f6483'),
    )

    orgao = doc_ctx.get('orgao')
    logo_path = ''
    if orgao and getattr(orgao, 'logo', None):
        try:
            logo_path = orgao.logo.path
        except Exception:
            logo_path = ''

    elements = []
    if logo_path:
        try:
            elements.append(Image(logo_path, width=56, height=56))
        except Exception:
            pass
    nome_orgao = ''
    if orgao:
        nome_orgao = str(getattr(orgao, 'nome_fantasia', '') or getattr(orgao, 'razao_social', '') or '').strip()
    if nome_orgao:
        elements.append(Paragraph(_pdf_safe(nome_orgao), ParagraphStyle('LicitacaoDocOrgao', parent=styles['Heading2'], fontSize=13, textColor=colors.HexColor('#153153'))))
    if orgao and getattr(orgao, 'cnpj', None):
        elements.append(Paragraph(_pdf_safe(f'CNPJ {orgao.cnpj}'), subtitle_style))
    if orgao and getattr(orgao, 'endereco_completo', None):
        elements.append(Paragraph(_pdf_safe(orgao.endereco_completo), subtitle_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph(_pdf_safe(doc_ctx.get('modelo_titulo', 'Documento')), title_style))
    numero_full = f"{doc_ctx.get('numero_prefixo', '')} {doc_ctx.get('numero_documento', '')}".strip()
    if numero_full:
        elements.append(Paragraph(_pdf_safe(numero_full), section_style))
    elements.append(Spacer(1, 8))

    meta_rows = [
        [Paragraph('<b>Processo:</b>', value_style), Paragraph(_pdf_safe(processo.numero_processo_principal), value_style)],
        [Paragraph('<b>Data:</b>', value_style), Paragraph(_pdf_safe(f"{doc_ctx.get('localidade', '-')}, {doc_ctx.get('data_documento_extenso', '-')}"), value_style)],
        [Paragraph('<b>Referência:</b>', value_style), Paragraph(_pdf_safe(doc_ctx.get('referencia', '-')), value_style)],
    ]
    if doc_ctx.get('de_origem'):
        meta_rows.append([Paragraph('<b>De:</b>', value_style), Paragraph(_pdf_safe(doc_ctx.get('de_origem')), value_style)])
    if doc_ctx.get('para_destino'):
        meta_rows.append([Paragraph('<b>Para:</b>', value_style), Paragraph(_pdf_safe(doc_ctx.get('para_destino')), value_style)])
    meta_table = Table(meta_rows, colWidths=[122, 398])
    meta_table.setStyle(
        TableStyle(
            [
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#edf3fb')),
                ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#d4dce8')),
                ('INNERGRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#d4dce8')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph('ASSUNTO', section_style))
    assunto_box = Table([[Paragraph(_pdf_safe(doc_ctx.get('assunto', '-')), value_style)]], colWidths=[520])
    assunto_box.setStyle(
        TableStyle(
            [
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f3f7fd')),
                ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#d4dce8')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 7),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ]
        )
    )
    elements.append(assunto_box)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph(_rich_html_to_pdf_markup(doc_ctx.get('corpo', '-')), value_style))
    elements.append(Spacer(1, 16))
    elements.append(Paragraph('ASSINATURA(S)', section_style))

    signatarios = list(doc_ctx.get('signatarios_selecionados') or [])
    if signatarios:
        sign_rows = []
        for sig in signatarios:
            nome = str(sig.get('nome') or '-')
            cargo = str(sig.get('cargo') or '').strip()
            sign_rows.append([Paragraph('________________________________________', value_style)])
            sign_rows.append([Paragraph(_pdf_safe(nome), value_style)])
            if cargo:
                sign_rows.append([Paragraph(_pdf_safe(cargo), subtitle_style)])
            sign_rows.append([Spacer(1, 6)])
        sign_table = Table(sign_rows, colWidths=[520])
        sign_table.setStyle(TableStyle([('ALIGN', (0, 0), (-1, -1), 'CENTER')]))
        elements.append(sign_table)
    else:
        elements.append(Paragraph('________________________________________', value_style))
        elements.append(Paragraph(_pdf_safe(doc_ctx.get('responsavel_nome', '-')), value_style))
        if doc_ctx.get('responsavel_cargo'):
            elements.append(Paragraph(_pdf_safe(doc_ctx.get('responsavel_cargo')), subtitle_style))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph(_pdf_safe(doc_ctx.get('rodape_documento', 'Documento gerado pelo SIREL Modular.')), subtitle_style))
    pdf.build(elements)

    filename = f"{_slug_ascii(doc_ctx.get('modelo_titulo', 'documento'))}_{processo.id}.pdf"
    response = HttpResponse(bio.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _export_licitacao_documento_docx(*, processo: Processo, doc_ctx: dict):
    doc = Document()
    orgao = doc_ctx.get('orgao')
    if orgao:
        nome_orgao = str(getattr(orgao, 'nome_fantasia', '') or getattr(orgao, 'razao_social', '') or '').strip()
        if nome_orgao:
            doc.add_paragraph(nome_orgao)
        if getattr(orgao, 'cnpj', None):
            doc.add_paragraph(f'CNPJ {orgao.cnpj}')
        if getattr(orgao, 'endereco_completo', None):
            doc.add_paragraph(str(orgao.endereco_completo))
        doc.add_paragraph('')

    doc.add_heading(str(doc_ctx.get('modelo_titulo') or 'Documento'), level=1)
    numero_full = f"{doc_ctx.get('numero_prefixo', '')} {doc_ctx.get('numero_documento', '')}".strip()
    if numero_full:
        doc.add_paragraph(numero_full)
    doc.add_paragraph(f"{doc_ctx.get('localidade', '-')}, {doc_ctx.get('data_documento_extenso', '-')}")
    doc.add_paragraph(f"Processo: {processo.numero_processo_principal}")
    if doc_ctx.get('referencia'):
        doc.add_paragraph(f"Referência: {doc_ctx.get('referencia')}")
    if doc_ctx.get('de_origem'):
        doc.add_paragraph(f"De: {doc_ctx.get('de_origem')}")
    if doc_ctx.get('para_destino'):
        doc.add_paragraph(f"Para: {doc_ctx.get('para_destino')}")
    doc.add_paragraph('')

    p_assunto = doc.add_paragraph('ASSUNTO: ')
    p_assunto.add_run(str(doc_ctx.get('assunto') or '-'))
    doc.add_paragraph('')
    doc.add_paragraph(_rich_html_to_plain_text(doc_ctx.get('corpo')) or '-')
    doc.add_paragraph('')

    signatarios = list(doc_ctx.get('signatarios_selecionados') or [])
    if signatarios:
        for sig in signatarios:
            nome = str(sig.get('nome') or '-')
            cargo = str(sig.get('cargo') or '').strip()
            doc.add_paragraph('________________________________________')
            doc.add_paragraph(nome)
            if cargo:
                doc.add_paragraph(cargo)
            doc.add_paragraph('')
    else:
        doc.add_paragraph('________________________________________')
        doc.add_paragraph(str(doc_ctx.get('responsavel_nome') or '-'))
        if doc_ctx.get('responsavel_cargo'):
            doc.add_paragraph(str(doc_ctx.get('responsavel_cargo')))

    bio = BytesIO()
    doc.save(bio)
    filename = f"{_slug_ascii(doc_ctx.get('modelo_titulo', 'documento'))}_{processo.id}.docx"
    response = HttpResponse(
        bio.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _post_bool(form_data, key: str, default: bool = False) -> bool:
    if form_data is None:
        return bool(default)
    raw = form_data.get(key)
    if raw is None:
        return False
    return str(raw).strip().lower() in {'1', 'true', 'on', 'sim', 'yes'}


def _parse_brl_decimal(raw_value, default: Decimal | None = None) -> Decimal:
    default_value = default if default is not None else Decimal('0')
    raw = str(raw_value or '').strip()
    if not raw:
        return default_value
    texto = raw.replace('R$', '').replace(' ', '').replace('.', '').replace(',', '.')
    texto = re.sub(r'[^0-9.\-]', '', texto)
    if not texto:
        return default_value
    try:
        return Decimal(texto)
    except (InvalidOperation, ValueError):
        return default_value


def _edital_modalidades_catalogo() -> list[str]:
    catalogo = []
    vistos = set()

    def add(nome: str):
        n = str(nome or '').strip()
        if not n:
            return
        key = n.lower()
        if key in vistos:
            return
        vistos.add(key)
        catalogo.append(n)

    try:
        for nome in Modalidade.objects.order_by('nome').values_list('nome', flat=True):
            add(nome)
    except Exception:
        pass

    for row in PNCP_MODALIDADES_PADRAO:
        add(str(row.get('nome') or '').split(' - ')[0])
    add('Pregão Eletrônico')
    add('Concorrência')
    add('Dispensa de Licitação')
    add('Inexigibilidade')
    add('Credenciamento')
    return catalogo


def _edital_modalidade_token(modalidade_nome: str) -> str:
    nome = str(modalidade_nome or '').lower()
    if 'dispensa' in nome and 'simplificada' in nome:
        return 'dispensa_simplificada'
    if 'dispensa' in nome:
        return 'dispensa'
    if 'inexig' in nome:
        return 'inexigibilidade'
    if 'credenc' in nome:
        return 'credenciamento'
    if 'concorr' in nome:
        return 'concorrencia'
    if 'preg' in nome:
        return 'pregao'
    if 'leil' in nome:
        return 'leilao'
    if 'concurso' in nome:
        return 'concurso'
    return 'geral'


@lru_cache(maxsize=32)
def _load_edital_template_text(template_name: str) -> str:
    safe_name = Path(str(template_name or '')).name
    if not safe_name:
        return ''
    fpath = EDITAL_TEMPLATES_DIR / safe_name
    if not fpath.exists():
        return ''
    for encoding in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
        try:
            return fpath.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError:
            return ''
    return ''


def _edital_section_text(texto: str, *, preserve_blank_lines: bool = False) -> list[str]:
    linhas = str(texto or '').replace('\r', '').split('\n')
    if preserve_blank_lines:
        if len(linhas) == 1 and linhas[0] == '':
            return ['-']
        return linhas or ['-']
    linhas = [ln.strip() for ln in linhas if ln.strip()]
    return linhas or ['-']


def _edital_sections_html(secoes: list[dict]) -> str:
    html_parts = []
    for sec in secoes:
        show_title = bool(sec.get('show_title', True))
        if show_title:
            html_parts.append(f"<h4>{escape(str(sec.get('titulo') or ''))}</h4>")
        for linha in _edital_section_text(
            sec.get('texto', ''),
            preserve_blank_lines=bool(sec.get('preserve_blank_lines')),
        ):
            if str(linha) == '':
                html_parts.append('<p>&nbsp;</p>')
            else:
                html_parts.append(f"<p>{escape(str(linha))}</p>")
    return ''.join(html_parts)


def _edital_sections_flatten(secoes: list[dict]) -> list[str]:
    out = []
    for sec in secoes:
        if bool(sec.get('show_title', True)):
            out.append(str(sec.get('titulo') or '-'))
        out.extend(
            _edital_section_text(
                sec.get('texto', ''),
                preserve_blank_lines=bool(sec.get('preserve_blank_lines')),
            )
        )
    return out


def _edital_default_sections(ctx: dict) -> list[dict]:
    modalidade = str(ctx.get('modalidade_nome') or 'Licitação')
    processo_adm = str(ctx.get('processo_adm') or '-')
    objeto = str(ctx.get('objeto') or '-')
    orgao = str(ctx.get('orgao_nome') or 'Órgão/Entidade')
    cnpj = str(ctx.get('cnpj') or '-')
    setor = str(ctx.get('setor_responsavel') or 'Setor de Licitação')
    endereco = str(ctx.get('endereco_completo_exibicao') or '-')
    criterio = str(ctx.get('criterio_julgamento') or '-')
    modo = str(ctx.get('modo_disputa') or '-')
    valor = str(ctx.get('valor_estimado_exibicao') or '-')
    data_sessao = str(ctx.get('data_sessao_extenso') or '-')
    hora_sessao = str(ctx.get('hora_sessao') or '-')
    email_dispensa = 'dispensas.pmtf@gmail.com'
    is_srp = bool(ctx.get('srp'))
    is_pref = bool(ctx.get('preferencia_me_epp'))

    token = _edital_modalidade_token(modalidade)

    if token in {'dispensa_simplificada', 'dispensa'}:
        minuta_integral = _load_edital_template_text('minuta_dispensa_simplificada.txt')
        if minuta_integral:
            return [
                {
                    'key': 'minuta_integral',
                    'titulo': 'MINUTA DE DISPENSA SIMPLIFICADA',
                    'texto': minuta_integral,
                    'preserve_blank_lines': True,
                    'show_title': False,
                },
            ]

        srp_line = (
            'Não se aplica Sistema de Registro de Preços para esta contratação direta.'
            if not is_srp else
            'Será aplicado Sistema de Registro de Preços, conforme regras do aviso.'
        )
        pref_line = (
            'Procedimento exclusivo para microempresa e empresa de pequeno porte.'
            if is_pref else
            'Sem reserva exclusiva para microempresa e empresa de pequeno porte.'
        )
        return [
            {
                'key': 'abertura',
                'titulo': 'ABERTURA DO AVISO',
                'texto': (
                    f'AVISO DE DISPENSA DE LICITAÇÃO SIMPLIFICADA - {modalidade}\n'
                    f'Processo Administrativo: {processo_adm}\n'
                    f'A Administração {orgao}, CNPJ {cnpj}, por meio de {setor}, com fundamento no art. 75 da Lei 14.133/2021, '
                    'torna público o presente procedimento de dispensa simplificada.'
                ),
            },
            {
                'key': 'data_hora',
                'titulo': 'DATA, LOCAL E HORÁRIO DA SESSÃO',
                'texto': (
                    'Referência de horário: Brasília-DF.\n'
                    f'Data da sessão: {data_sessao}\n'
                    f'Horário da sessão: {hora_sessao}\n'
                    f'Endereço para envio de proposta e habilitação: {email_dispensa}'
                ),
            },
            {
                'key': 'objeto',
                'titulo': '1. OBJETO DA CONTRATAÇÃO DIRETA',
                'texto': (
                    f'1.1. Objeto: {objeto}.\n'
                    f'1.2. Critério de julgamento: {criterio}.\n'
                    '1.3. As condições técnicas, quantitativos e exigências constam no Termo de Referência e anexos.'
                ),
            },
            {
                'key': 'participacao',
                'titulo': '2. PARTICIPAÇÃO NA DISPENSA SIMPLIFICADA',
                'texto': (
                    f'2.1. A participação ocorrerá mediante envio de proposta para {email_dispensa}.\n'
                    f'2.2. {pref_line}\n'
                    '2.3. Não poderão participar interessados com impedimentos previstos na Lei 14.133/2021.'
                ),
            },
            {
                'key': 'proposta',
                'titulo': '3. CADASTRAMENTO E ENVIO DA PROPOSTA',
                'texto': (
                    '3.1. A proposta deve conter identificação da empresa, descrição do objeto, marca quando aplicável e preço.\n'
                    '3.2. A validade mínima da proposta será de 90 (noventa) dias.\n'
                    '3.3. Os valores devem contemplar todos os custos diretos e indiretos da execução.'
                ),
            },
            {
                'key': 'julgamento',
                'titulo': '4. JULGAMENTO E ACEITAÇÃO DAS PROPOSTAS',
                'texto': (
                    '4.1. Encerrado o prazo, será verificada a conformidade das propostas e a compatibilidade de preços.\n'
                    f'4.2. Em negociação, o valor final deverá observar o limite de referência da contratação ({valor}).\n'
                    f'4.3. Modo de disputa de referência para condução do rito: {modo}.'
                ),
            },
            {
                'key': 'habilitacao',
                'titulo': '5. HABILITAÇÃO',
                'texto': (
                    '5.1. O fornecedor mais bem classificado será convocado para envio da documentação de habilitação.\n'
                    '5.2. Serão observados os documentos exigidos no Termo de Referência e na legislação aplicável.\n'
                    '5.3. A ausência de documentação obrigatória implicará inabilitação.'
                ),
            },
            {
                'key': 'contratacao',
                'titulo': '6. CONTRATAÇÃO E EXECUÇÃO',
                'texto': (
                    '6.1. Após adjudicação e homologação, será emitida ordem de fornecimento ou instrumento contratual equivalente.\n'
                    f'6.2. Endereço de execução/entrega: {endereco}.\n'
                    f'6.3. {srp_line}'
                ),
            },
            {
                'key': 'recursos_sancoes',
                'titulo': '7. RECURSOS, SANÇÕES E PENALIDADES',
                'texto': (
                    '7.1. Os recursos administrativos obedecerão aos prazos e hipóteses da Lei 14.133/2021.\n'
                    '7.2. O descumprimento contratual sujeitará o contratado às penalidades previstas no edital e na legislação.'
                ),
            },
            {
                'key': 'impugnacao',
                'titulo': '8. IMPUGNAÇÃO E ESCLARECIMENTOS',
                'texto': (
                    f'8.1. Pedidos de esclarecimentos e impugnações devem ser encaminhados ao setor {setor}.\n'
                    '8.2. As respostas serão juntadas aos autos e comunicadas aos interessados.'
                ),
            },
            {
                'key': 'disposicoes',
                'titulo': '9. DISPOSIÇÕES GERAIS E ANEXOS',
                'texto': (
                    '9.1. Integram este aviso o Termo de Referência e os modelos de declarações.\n'
                    '9.2. Casos omissos serão resolvidos pela autoridade competente, conforme Lei 14.133/2021.\n'
                    '9.3. Fica eleito o foro da comarca de Teixeira de Freitas-BA.'
                ),
            },
        ]

    return [
        {
            'key': 'abertura',
            'titulo': 'ABERTURA DA MINUTA',
            'texto': (
                f'Torna-se público que {orgao}, CNPJ {cnpj}, por meio do {setor}, realizará {modalidade}, '
                'na forma eletrônica, nos termos da Lei 14.133/2021 e normas complementares.'
            ),
        },
        {
            'key': 'data_hora',
            'titulo': 'DATA E HORA',
            'texto': (
                f'Data da sessão pública: {data_sessao}\n'
                f'Hora inicial: {hora_sessao}\n'
                'Referência de horário: Brasília-DF.'
            ),
        },
        {
            'key': 'objeto',
            'titulo': '1. DO OBJETO',
            'texto': (
                f'1.1. Objeto da licitação: {objeto}.\n'
                '1.2. As condições técnicas e quantitativos constam no Termo de Referência e anexos.'
            ),
        },
        {
            'key': 'participacao',
            'titulo': '2. DA PARTICIPAÇÃO NA LICITAÇÃO',
            'texto': (
                f'2.1. Poderão participar interessados com ramo compatível ao objeto e previamente credenciados.\n'
                f'2.2. {"Haverá" if is_pref else "Não haverá"} tratamento favorecido para microempresas e empresas de pequeno porte.\n'
                '2.3. Não poderão participar os impedidos de contratar com a Administração.'
            ),
        },
        {
            'key': 'propostas_habilitacao',
            'titulo': '3. DA APRESENTAÇÃO DA PROPOSTA E DOS DOCUMENTOS DE HABILITAÇÃO',
            'texto': (
                '3.1. A proposta deverá ser apresentada conforme exigências do edital e anexos.\n'
                '3.2. A documentação de habilitação será exigida da licitante melhor classificada.'
            ),
        },
        {
            'key': 'abertura_lances',
            'titulo': '4. DA ABERTURA DA SESSÃO, CLASSIFICAÇÃO DAS PROPOSTAS E FORMULAÇÃO DE LANCES',
            'texto': (
                f'4.1. Critério de julgamento: {criterio}.\n'
                f'4.2. Modo de disputa: {modo}.\n'
                f'4.3. Valor de referência da contratação: {valor}.'
            ),
        },
        {
            'key': 'julgamento',
            'titulo': '5. DA FASE DE JULGAMENTO',
            'texto': (
                '5.1. Encerrada a etapa competitiva, será analisada a proposta melhor classificada.\n'
                '5.2. Poderá haver negociação para obtenção da proposta mais vantajosa.'
            ),
        },
        {
            'key': 'habilitacao',
            'titulo': '6. DA FASE DE HABILITAÇÃO',
            'texto': (
                '6.1. A habilitação observará requisitos jurídicos, fiscais, trabalhistas, técnicos e econômico-financeiros.\n'
                '6.2. A não apresentação da documentação obrigatória resultará na inabilitação.'
            ),
        },
        {
            'key': 'contrato',
            'titulo': '7. DO TERMO DE CONTRATO',
            'texto': (
                f'7.1. Após homologação, será celebrado contrato ou instrumento equivalente para execução em {endereco}.\n'
                f'7.2. {"Será" if is_srp else "Não será"} adotado Sistema de Registro de Preços (SRP).'
            ),
        },
        {
            'key': 'recursos',
            'titulo': '8. DOS RECURSOS',
            'texto': '8.1. Os recursos administrativos obedecerão aos prazos e procedimentos previstos na Lei 14.133/2021.',
        },
        {
            'key': 'infracoes',
            'titulo': '9. DAS INFRAÇÕES ADMINISTRATIVAS E SANÇÕES',
            'texto': (
                '9.1. A prática de infrações sujeitará a contratada às sanções de advertência, multa, impedimento e declaração de inidoneidade,\n'
                'conforme previsão legal e editalícia.'
            ),
        },
        {
            'key': 'impugnacao',
            'titulo': '10. DA IMPUGNAÇÃO AO EDITAL E DOS PEDIDOS DE ESCLARECIMENTO',
            'texto': (
                f'10.1. Impugnações e esclarecimentos serão recebidos pelo {setor}.\n'
                '10.2. As respostas serão publicadas no processo e terão efeito vinculante.'
            ),
        },
        {
            'key': 'disposicoes',
            'titulo': '11. DAS DISPOSIÇÕES GERAIS',
            'texto': (
                f'11.1. Processo administrativo: {processo_adm}.\n'
                '11.2. Integram a presente minuta os anexos obrigatórios definidos pela Administração.\n'
                '11.3. Casos omissos serão resolvidos pela autoridade competente.'
            ),
        },
    ]


def _licitacao_edital_context(processo: Processo, *, form_data=None, usuario=None):
    orgao = _orgao_ativo()
    hoje = timezone.localdate()

    orgao_nome_default = (
        str(getattr(orgao, 'nome_fantasia', '') or '').strip()
        or str(getattr(orgao, 'razao_social', '') or '').strip()
        or 'Órgão/Entidade'
    )
    cnpj_default = str(getattr(orgao, 'cnpj', '') or '').strip()
    endereco_default = str(getattr(orgao, 'endereco', '') or '').strip()
    numero_default = str(getattr(orgao, 'numero', '') or '').strip()
    complemento_default = str(getattr(orgao, 'complemento', '') or '').strip()
    bairro_default = str(getattr(orgao, 'bairro', '') or '').strip()
    cidade_default = str(getattr(orgao, 'cidade', '') or '').strip()
    uf_default = str(getattr(orgao, 'uf', '') or '').strip()
    cep_default = str(getattr(orgao, 'cep', '') or '').strip()
    localidade_default = _localidade_orgao_padrao()

    modalidade_nome_default = (
        str(getattr(getattr(processo, 'modalidade', None), 'nome', '') or '').strip()
        or 'Pregão Eletrônico'
    )
    numero_edital_default = str(processo.numero_edital or '').strip() or f'00001/{hoje.year}'
    processo_adm_base = str(processo.numero_processo_adm or processo.numero_processo_principal or '').strip() or '-'
    if re.search(r'/\d{4}$', processo_adm_base):
        processo_adm_default = processo_adm_base
    else:
        processo_adm_default = f'{processo_adm_base}/{processo.ano_referencia}'

    criterio_default = str(processo.get_criterio_julgamento_display() or '').strip() or 'Menor preço'
    modo_default = str(processo.get_modo_disputa_display() or '').strip() or 'Aberto'
    setor_default = (
        str(getattr(getattr(processo, 'secretaria', None), 'nome', '') or '').strip()
        or 'Setor de Licitação'
    )
    objeto_default = str(processo.objeto or '').strip() or '-'
    valor_estimado_default = Decimal(str(getattr(processo, 'valor_estimado', 0) or 0))
    codigo_contratante_default = str(processo.identificador_bll or '').strip()

    data_sessao_default = hoje
    hora_sessao_default = '09:00'
    if processo.data_hora_abertura:
        dt_open = timezone.localtime(processo.data_hora_abertura)
        data_sessao_default = dt_open.date()
        hora_sessao_default = dt_open.strftime('%H:%M')

    responsavel_nome_default = (
        str(getattr(getattr(processo, 'condutor_processo', None), 'nome', '') or '').strip()
        or (
            (usuario.get_full_name() or '').strip()
            if getattr(usuario, 'is_authenticated', False)
            else ''
        )
        or (
            usuario.get_username()
            if getattr(usuario, 'is_authenticated', False)
            else ''
        )
        or 'Responsável'
    )
    responsavel_cargo_default = (
        str(getattr(getattr(processo, 'condutor_processo', None), 'cargo', '') or '').strip()
        or 'Agente de Contratação'
    )

    modalidades_catalogo = _edital_modalidades_catalogo()
    criterios_catalogo = [label for _, label in Processo.CriterioJulgamento.choices]
    modos_catalogo = [label for _, label in Processo.ModoDisputa.choices]

    ctx = {
        'processo': processo,
        'orgao': orgao,
        'titulo_pagina': 'Gerador de Minuta de Edital',
        'voltar_url': reverse('workflow:licitacao_detail', args=[processo.id]),
        'modalidades_catalogo': modalidades_catalogo,
        'criterios_catalogo': criterios_catalogo,
        'modos_catalogo': modos_catalogo,
        'modalidade_nome': modalidade_nome_default,
        'numero_edital_modelo': numero_edital_default,
        'processo_adm': processo_adm_default,
        'codigo_contratante': codigo_contratante_default,
        'cnpj': cnpj_default,
        'orgao_nome': orgao_nome_default,
        'setor_responsavel': setor_default,
        'objeto': objeto_default,
        'logradouro': endereco_default,
        'numero_endereco': numero_default,
        'complemento': complemento_default,
        'bairro': bairro_default,
        'cidade': cidade_default,
        'uf': uf_default,
        'cep': cep_default,
        'localidade': localidade_default,
        'criterio_julgamento': criterio_default,
        'modo_disputa': modo_default,
        'contratacao_tic': False,
        'srp': processo.tipo_contratacao == Processo.TipoContratacao.REGISTRO_PRECO,
        'preferencia_me_epp': True,
        'margem_preferencia': False,
        'orcamento_sigiloso': False,
        'inversao_fases': False,
        'veda_consorcio': False,
        'dedicacao_exclusiva': False,
        'valor_estimado_input': fmt_brl(valor_estimado_default),
        'data_sessao': data_sessao_default.isoformat(),
        'hora_sessao': hora_sessao_default,
        'clausulas_complementares': '',
        'signatario_nome': responsavel_nome_default,
        'signatario_cargo': responsavel_cargo_default,
    }

    if form_data:
        ctx['modalidade_nome'] = str(form_data.get('modalidade_nome') or ctx['modalidade_nome']).strip() or ctx['modalidade_nome']
        ctx['numero_edital_modelo'] = str(form_data.get('numero_edital_modelo') or ctx['numero_edital_modelo']).strip() or ctx['numero_edital_modelo']
        ctx['processo_adm'] = str(form_data.get('processo_adm') or ctx['processo_adm']).strip() or ctx['processo_adm']
        ctx['codigo_contratante'] = str(form_data.get('codigo_contratante') or '').strip()
        ctx['cnpj'] = str(form_data.get('cnpj') or ctx['cnpj']).strip()
        ctx['orgao_nome'] = str(form_data.get('orgao_nome') or ctx['orgao_nome']).strip()
        ctx['setor_responsavel'] = str(form_data.get('setor_responsavel') or ctx['setor_responsavel']).strip()
        ctx['objeto'] = str(form_data.get('objeto') or ctx['objeto']).strip()
        ctx['logradouro'] = str(form_data.get('logradouro') or ctx['logradouro']).strip()
        ctx['numero_endereco'] = str(form_data.get('numero_endereco') or ctx['numero_endereco']).strip()
        ctx['complemento'] = str(form_data.get('complemento') or '').strip()
        ctx['bairro'] = str(form_data.get('bairro') or ctx['bairro']).strip()
        ctx['cidade'] = str(form_data.get('cidade') or ctx['cidade']).strip()
        ctx['uf'] = str(form_data.get('uf') or ctx['uf']).strip()
        ctx['cep'] = str(form_data.get('cep') or ctx['cep']).strip()
        ctx['localidade'] = str(form_data.get('localidade') or ctx['localidade']).strip()
        ctx['criterio_julgamento'] = str(form_data.get('criterio_julgamento') or ctx['criterio_julgamento']).strip()
        ctx['modo_disputa'] = str(form_data.get('modo_disputa') or ctx['modo_disputa']).strip()
        ctx['valor_estimado_input'] = str(form_data.get('valor_estimado_input') or ctx['valor_estimado_input']).strip()
        ctx['data_sessao'] = str(form_data.get('data_sessao') or ctx['data_sessao']).strip()
        ctx['hora_sessao'] = str(form_data.get('hora_sessao') or ctx['hora_sessao']).strip()
        ctx['clausulas_complementares'] = str(form_data.get('clausulas_complementares') or '').strip()
        ctx['signatario_nome'] = str(form_data.get('signatario_nome') or ctx['signatario_nome']).strip()
        ctx['signatario_cargo'] = str(form_data.get('signatario_cargo') or ctx['signatario_cargo']).strip()

        ctx['contratacao_tic'] = _post_bool(form_data, 'contratacao_tic', default=ctx['contratacao_tic'])
        ctx['srp'] = _post_bool(form_data, 'srp', default=ctx['srp'])
        ctx['preferencia_me_epp'] = _post_bool(form_data, 'preferencia_me_epp', default=ctx['preferencia_me_epp'])
        ctx['margem_preferencia'] = _post_bool(form_data, 'margem_preferencia', default=ctx['margem_preferencia'])
        ctx['orcamento_sigiloso'] = _post_bool(form_data, 'orcamento_sigiloso', default=ctx['orcamento_sigiloso'])
        ctx['inversao_fases'] = _post_bool(form_data, 'inversao_fases', default=ctx['inversao_fases'])
        ctx['veda_consorcio'] = _post_bool(form_data, 'veda_consorcio', default=ctx['veda_consorcio'])
        ctx['dedicacao_exclusiva'] = _post_bool(form_data, 'dedicacao_exclusiva', default=ctx['dedicacao_exclusiva'])

    valor_estimado = _parse_brl_decimal(ctx.get('valor_estimado_input'), default=valor_estimado_default)
    if valor_estimado < 0:
        valor_estimado = Decimal('0')
    ctx['valor_estimado'] = valor_estimado
    ctx['valor_estimado_exibicao'] = 'Sigiloso' if ctx['orcamento_sigiloso'] else fmt_brl(valor_estimado)

    try:
        data_sessao_obj = datetime.fromisoformat(str(ctx.get('data_sessao') or '')).date()
    except Exception:
        data_sessao_obj = data_sessao_default
        ctx['data_sessao'] = data_sessao_obj.isoformat()
    ctx['data_sessao_extenso'] = _data_extenso_pt_br(data_sessao_obj)

    endereco_parts = [ctx.get('logradouro'), ctx.get('numero_endereco'), ctx.get('complemento'), ctx.get('bairro')]
    endereco_txt = ', '.join([str(p).strip() for p in endereco_parts if str(p or '').strip()])
    cidade_uf = ' - '.join([x for x in [str(ctx.get('cidade') or '').strip(), str(ctx.get('uf') or '').strip()] if x])
    if cidade_uf:
        endereco_txt = f'{endereco_txt} | {cidade_uf}' if endereco_txt else cidade_uf
    if str(ctx.get('cep') or '').strip():
        endereco_txt = f'{endereco_txt} | CEP {ctx["cep"]}' if endereco_txt else f'CEP {ctx["cep"]}'
    ctx['endereco_completo_exibicao'] = endereco_txt or '-'

    action = ''
    if form_data is not None:
        action = str(form_data.get('action') or '').strip().lower()

    secoes = _edital_default_sections(ctx)
    if form_data is not None and action != 'reload_template':
        for sec in secoes:
            key = str(sec.get('key') or '').strip()
            if not key:
                continue
            posted = str(form_data.get(f'sec_{key}') or '').strip()
            if posted:
                sec['texto'] = posted

    for sec in secoes:
        sec['preserve_blank_lines'] = bool(sec.get('preserve_blank_lines'))
        sec['show_title'] = bool(sec.get('show_title', True))
        sec['paragrafos'] = _edital_section_text(
            sec.get('texto', ''),
            preserve_blank_lines=sec['preserve_blank_lines'],
        )

    ctx['edital_secoes'] = secoes
    ctx['edital_corpo_html'] = _edital_sections_html(secoes)
    ctx['edital_paragrafos'] = _edital_sections_flatten(secoes)
    return ctx


def _export_licitacao_edital_pdf(*, processo: Processo, edital_ctx: dict):
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'EditalTitle',
        parent=styles['Heading1'],
        fontSize=15.8,
        leading=19,
        textColor=colors.HexColor('#153153'),
        alignment=1,
    )
    value_style = ParagraphStyle(
        'EditalValue',
        parent=styles['BodyText'],
        fontSize=10.2,
        leading=14,
        textColor=colors.HexColor('#1d2b40'),
        alignment=TA_JUSTIFY,
    )
    subtitle_style = ParagraphStyle(
        'EditalSubtitle',
        parent=styles['BodyText'],
        fontSize=9.4,
        leading=12,
        textColor=colors.HexColor('#5d6f8a'),
    )

    elements = []
    orgao = edital_ctx.get('orgao')
    if orgao and getattr(orgao, 'logo', None):
        try:
            img = Image(orgao.logo.path, width=95, height=95)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(Spacer(1, 8))
        except Exception:
            pass

    elements.append(Paragraph(_pdf_safe(edital_ctx.get('orgao_nome', 'Órgão/Entidade')), title_style))
    elements.append(Spacer(1, 2))
    elements.append(Paragraph(_pdf_safe(f"CNPJ {edital_ctx.get('cnpj', '-')}"), subtitle_style))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(_pdf_safe(f"MINUTA DE EDITAL - {edital_ctx.get('modalidade_nome', 'Licitação')}"), title_style))
    elements.append(Spacer(1, 6))

    meta_rows = [
        [Paragraph('<b>Edital:</b>', value_style), Paragraph(_pdf_safe(edital_ctx.get('numero_edital_modelo', '-')), value_style)],
        [Paragraph('<b>Processo:</b>', value_style), Paragraph(_pdf_safe(edital_ctx.get('processo_adm', '-')), value_style)],
        [Paragraph('<b>Setor:</b>', value_style), Paragraph(_pdf_safe(edital_ctx.get('setor_responsavel', '-')), value_style)],
        [Paragraph('<b>Endereço:</b>', value_style), Paragraph(_pdf_safe(edital_ctx.get('endereco_completo_exibicao', '-')), value_style)],
        [Paragraph('<b>Sessão:</b>', value_style), Paragraph(_pdf_safe(f"{edital_ctx.get('data_sessao_extenso', '-')}, às {edital_ctx.get('hora_sessao', '-')}"), value_style)],
    ]
    meta_table = Table(meta_rows, colWidths=[120, 400])
    meta_table.setStyle(
        TableStyle(
            [
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#edf3fb')),
                ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#d4dce8')),
                ('INNERGRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#d4dce8')),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 10))

    for par in edital_ctx.get('edital_paragrafos') or []:
        txt = str(par)
        elements.append(Paragraph(_pdf_safe(txt if txt.strip() else '\u00a0'), value_style))
        elements.append(Spacer(1, 4))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph(_pdf_safe('________________________________________'), value_style))
    elements.append(Paragraph(_pdf_safe(edital_ctx.get('signatario_nome', '-')), value_style))
    if edital_ctx.get('signatario_cargo'):
        elements.append(Paragraph(_pdf_safe(edital_ctx.get('signatario_cargo')), subtitle_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(_pdf_safe('Documento gerado pelo SIREL Modular.'), subtitle_style))

    pdf.build(elements)
    filename = f"minuta_edital_{_slug_ascii(edital_ctx.get('modalidade_nome', 'licitacao'))}_{processo.id}.pdf"
    response = HttpResponse(bio.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _export_licitacao_edital_docx(*, processo: Processo, edital_ctx: dict):
    doc = Document()
    doc.add_heading(str(edital_ctx.get('orgao_nome') or 'Órgão/Entidade'), level=1)
    doc.add_paragraph(f"CNPJ {edital_ctx.get('cnpj') or '-'}")
    doc.add_paragraph(f"MINUTA DE EDITAL - {edital_ctx.get('modalidade_nome') or 'Licitação'}")
    doc.add_paragraph(f"Edital: {edital_ctx.get('numero_edital_modelo') or '-'}")
    doc.add_paragraph(f"Processo: {edital_ctx.get('processo_adm') or '-'}")
    doc.add_paragraph(f"Setor responsável: {edital_ctx.get('setor_responsavel') or '-'}")
    doc.add_paragraph(f"Endereço: {edital_ctx.get('endereco_completo_exibicao') or '-'}")
    doc.add_paragraph(f"Sessão pública: {edital_ctx.get('data_sessao_extenso') or '-'}, às {edital_ctx.get('hora_sessao') or '-'}")
    doc.add_paragraph('')

    for par in edital_ctx.get('edital_paragrafos') or []:
        txt = str(par)
        p = doc.add_paragraph(txt if txt.strip() else ' ')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph('')
    doc.add_paragraph('________________________________________')
    doc.add_paragraph(str(edital_ctx.get('signatario_nome') or '-'))
    if edital_ctx.get('signatario_cargo'):
        doc.add_paragraph(str(edital_ctx.get('signatario_cargo')))
    doc.add_paragraph('Documento gerado pelo SIREL Modular.')

    bio = BytesIO()
    doc.save(bio)
    filename = f"minuta_edital_{_slug_ascii(edital_ctx.get('modalidade_nome', 'licitacao'))}_{processo.id}.docx"
    response = HttpResponse(
        bio.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _export_licitacao_edital_html(*, processo: Processo, edital_ctx: dict):
    title = escape(f"Minuta de Edital - {edital_ctx.get('modalidade_nome', 'Licitação')}")
    corpo_html = edital_ctx.get('edital_corpo_html') or ''
    html_data = (
        '<!doctype html>'
        '<html lang="pt-BR"><head><meta charset="utf-8">'
        f'<title>{title}</title>'
        '<style>'
        'body{font-family:Arial,Helvetica,sans-serif;color:#1b2b40;margin:30px;}'
        '.head{margin-bottom:16px;text-align:center;}'
        '.meta{border-collapse:collapse;width:100%;margin-bottom:16px;}'
        '.meta td{border:1px solid #d6e0ef;padding:6px;font-size:13px;vertical-align:top;}'
        '.meta td:first-child{width:170px;background:#eef3fb;font-weight:700;}'
        '.doc h4{font-size:15px;margin:16px 0 6px;color:#153153;}'
        '.doc p{line-height:1.45;font-size:14px;margin:0 0 10px;text-align:justify;}'
        '.sign{margin-top:28px;text-align:center;}'
        '</style></head><body>'
        f'<div class="head"><h1>{escape(str(edital_ctx.get("orgao_nome") or "Órgão/Entidade"))}</h1>'
        f'<div>CNPJ {escape(str(edital_ctx.get("cnpj") or "-"))}</div>'
        f'<h2>MINUTA DE EDITAL - {escape(str(edital_ctx.get("modalidade_nome") or "Licitação"))}</h2></div>'
        '<table class="meta">'
        f'<tr><td>Edital</td><td>{escape(str(edital_ctx.get("numero_edital_modelo") or "-"))}</td></tr>'
        f'<tr><td>Processo</td><td>{escape(str(edital_ctx.get("processo_adm") or "-"))}</td></tr>'
        f'<tr><td>Setor</td><td>{escape(str(edital_ctx.get("setor_responsavel") or "-"))}</td></tr>'
        f'<tr><td>Endereço</td><td>{escape(str(edital_ctx.get("endereco_completo_exibicao") or "-"))}</td></tr>'
        f'<tr><td>Sessão</td><td>{escape(str(edital_ctx.get("data_sessao_extenso") or "-"))}, às {escape(str(edital_ctx.get("hora_sessao") or "-"))}</td></tr>'
        '</table>'
        f'<div class="doc">{corpo_html}</div>'
        '<div class="sign">________________________________________<br>'
        f'{escape(str(edital_ctx.get("signatario_nome") or "-"))}<br>'
        f'{escape(str(edital_ctx.get("signatario_cargo") or ""))}'
        '</div>'
        '</body></html>'
    )

    filename = f"minuta_edital_{_slug_ascii(edital_ctx.get('modalidade_nome', 'licitacao'))}_{processo.id}.html"
    response = HttpResponse(html_data, content_type='text/html; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def licitacao_edital_generator(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method == 'POST':
        edital_ctx = _licitacao_edital_context(processo, form_data=request.POST, usuario=request.user)
        action = str(request.POST.get('action') or '').strip().lower()
        if action == 'preview_pdf':
            return _export_licitacao_edital_pdf(processo=processo, edital_ctx=edital_ctx)
        if action == 'download_docx':
            return _export_licitacao_edital_docx(processo=processo, edital_ctx=edital_ctx)
        if action == 'download_html':
            return _export_licitacao_edital_html(processo=processo, edital_ctx=edital_ctx)
    else:
        edital_ctx = _licitacao_edital_context(processo, usuario=request.user)
    return render(request, 'workflow/licitacao_edital_generator.html', edital_ctx)


def _get_status(nome='EM PLANEJAMENTO'):
    return StatusProcesso.objects.get_or_create(nome=nome)[0]


def _get_modalidade(nome='Pregão'):
    return Modalidade.objects.get_or_create(nome=nome)[0]


def _pncp_safe_dict(value):
    return value if isinstance(value, dict) else {}


def _pncp_safe_text(value):
    if value is None:
        return ''
    if isinstance(value, dict):
        for key in ('nome', 'descricao', 'sigla', 'codigo', 'valor'):
            raw = value.get(key)
            if raw not in (None, ''):
                return str(raw).strip()
        return ''
    return str(value).strip()


def _normalizar_entrada_pncp_numero_controle(valor: str) -> str:
    raw = (valor or '').strip()
    if not raw:
        return raw
    if 'pncp.gov.br/app/editais/' in raw:
        try:
            after = raw.split('/app/editais/', 1)[1]
            partes = [p for p in after.split('/') if p]
            if len(partes) >= 3:
                cnpj = ''.join(ch for ch in partes[0] if ch.isdigit())
                ano = ''.join(ch for ch in partes[1] if ch.isdigit())
                sequencial = ''.join(ch for ch in partes[2] if ch.isdigit())
                if cnpj and ano and sequencial:
                    return f'{cnpj}-1-{int(sequencial)}/{ano}'
        except Exception:
            return raw
    return raw


def _parse_pncp_date(value):
    if not value:
        return None
    if isinstance(value, date) and not isinstance(value, datetime):
        return value
    if isinstance(value, datetime):
        return value.date()
    text = str(value).strip()
    if not text:
        return None
    text = text.replace('Z', '+00:00')
    try:
        return datetime.fromisoformat(text).date()
    except Exception:
        try:
            return datetime.strptime(text[:10], '%Y-%m-%d').date()
        except Exception:
            return None


def _parse_pncp_datetime(value):
    if not value:
        return None
    if isinstance(value, datetime):
        if timezone.is_naive(value):
            return timezone.make_aware(value, timezone.get_current_timezone())
        return value
    text = str(value).strip()
    if not text:
        return None
    text = text.replace('Z', '+00:00')
    try:
        parsed = datetime.fromisoformat(text)
        if timezone.is_naive(parsed):
            return timezone.make_aware(parsed, timezone.get_current_timezone())
        return parsed
    except Exception:
        for fmt in ('%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M', '%d/%m/%Y %H:%M:%S', '%d/%m/%Y %H:%M'):
            try:
                parsed = datetime.strptime(text, fmt)
                return timezone.make_aware(parsed, timezone.get_current_timezone())
            except Exception:
                continue
        try:
            parsed = datetime.strptime(text[:19], '%Y-%m-%dT%H:%M:%S')
            return timezone.make_aware(parsed, timezone.get_current_timezone())
        except Exception:
            pass
    return None


def _pncp_to_int(value):
    if value is None or value == '':
        return None
    try:
        return int(str(value).strip())
    except Exception:
        return None


def _pncp_to_decimal(value):
    if value is None or value == '':
        return Decimal('0')
    try:
        text = str(value).strip()
        if not text:
            return Decimal('0')
        text = text.replace('R$', '').replace('r$', '').replace(' ', '')
        text = ''.join(ch for ch in text if ch.isdigit() or ch in {'.', ',', '-'})
        if not text:
            return Decimal('0')
        # Ex.: 1.500.000 -> 1500000
        if ',' not in text and text.count('.') >= 1:
            partes = text.split('.')
            if all(p.isdigit() for p in partes) and all(len(p) == 3 for p in partes[1:]):
                text = ''.join(partes)
        if ',' in text and '.' in text:
            text = text.replace('.', '').replace(',', '.')
        elif ',' in text:
            text = text.replace(',', '.')
        return Decimal(text)
    except Exception:
        return Decimal('0')


def _pncp_to_bool(value):
    if isinstance(value, bool):
        return value
    text = str(value or '').strip().lower()
    return text in {'1', 'true', 't', 'sim', 's', 'yes', 'y'}


def _pncp_first_payload(data):
    if isinstance(data, dict):
        if isinstance(data.get('data'), list):
            return data['data'][0] if data['data'] else {}
        if isinstance(data.get('data'), dict):
            return data['data']
        if isinstance(data.get('resultado'), list):
            return data['resultado'][0] if data['resultado'] else {}
        if isinstance(data.get('resultado'), dict):
            return data['resultado']
        return data
    if isinstance(data, list):
        return data[0] if data else {}
    return {}


def _pncp_value(payload: dict, key: str, default=''):
    value = (payload or {}).get(key, default)
    if value is None:
        return default
    return value


def _pncp_detectar_processo_existente(numero_controle: str, payload: dict):
    snap = PNCPContratacaoSnapshot.objects.filter(numero_controle_pncp=numero_controle).select_related('processo').first()
    if snap and snap.processo_id:
        return snap.processo
    wf = ProcessoWorkflow.objects.filter(pncp_numero_controle=numero_controle).select_related('processo').first()
    if wf and wf.processo_id:
        return wf.processo

    numero_proc = _pncp_safe_text(_pncp_value(payload, 'processo')) or _pncp_safe_text(_pncp_value(payload, 'numeroCompra'))
    ano_ref = _pncp_to_int(_pncp_value(payload, 'anoCompra'))
    if numero_proc and ano_ref:
        numero_proc_norm = ''.join(ch for ch in numero_proc.upper() if ch.isalnum())
        match = Processo.objects.filter(
            Q(numero_processo_sirel__iexact=numero_proc, ano_referencia=ano_ref)
            | Q(numero_processo_adm__iexact=numero_proc, ano_referencia=ano_ref)
            | Q(numero_edital__iexact=numero_proc, ano_referencia=ano_ref)
        ).order_by('-id').first()
        if match:
            return match
        candidatos = (
            Processo.objects
            .filter(ano_referencia=ano_ref)
            .only('id', 'numero_processo_sirel', 'numero_processo_adm', 'numero_edital')
            .order_by('-id')[:300]
        )
        for candidato in candidatos:
            for numero_candidato in (
                candidato.numero_processo_sirel,
                candidato.numero_processo_adm,
                candidato.numero_edital,
            ):
                numero_candidato_norm = ''.join(ch for ch in str(numero_candidato or '').upper() if ch.isalnum())
                if numero_candidato_norm and numero_candidato_norm == numero_proc_norm:
                    return candidato
    return None


def _normalizar_modalidade_pncp(nome):
    nome_up = (nome or '').upper()
    if 'PREGAO' in nome_up:
        return 'Pregão'
    if 'DISPENSA' in nome_up:
        return 'Dispensa Eletrônica' if 'ELETRONICA' in nome_up else 'Dispensa'
    if 'INEXIG' in nome_up:
        return 'Inexigibilidade'
    if 'CONCORR' in nome_up:
        return 'Concorrência'
    if 'CREDEN' in nome_up:
        return 'Credenciamento'
    if 'LEILAO' in nome_up:
        return 'Leilão'
    if 'CONCURSO' in nome_up:
        return 'Concurso'
    if 'DIALOGO' in nome_up:
        return 'Diálogo competitivo'
    return (nome or 'Pregão').strip() or 'Pregão'


def _mapear_modo_disputa_pncp(nome):
    texto = (nome or '').upper()
    if 'ABERTO' in texto and 'FECHADO' in texto:
        if texto.index('ABERTO') < texto.index('FECHADO'):
            return Processo.ModoDisputa.ABERTO_FECHADO
        return Processo.ModoDisputa.FECHADO_ABERTO
    if 'ABERTO' in texto:
        return Processo.ModoDisputa.ABERTO
    if 'FECHADO' in texto:
        return Processo.ModoDisputa.FECHADO
    return ''


def _mapear_criterio_pncp(nome):
    texto = (nome or '').upper()
    if 'LOTE' in texto:
        return Processo.CriterioJulgamento.MENOR_PRECO_POR_LOTE
    if 'GLOBAL' in texto:
        return Processo.CriterioJulgamento.MENOR_PRECO_GLOBAL
    if 'DESCONTO' in texto:
        return Processo.CriterioJulgamento.MAIOR_DESCONTO
    if 'TECNICA E PRECO' in texto or 'TÉCNICA E PREÇO' in texto:
        return Processo.CriterioJulgamento.TECNICA_E_PRECO
    if 'MELHOR TECNICA' in texto or 'MELHOR TÉCNICA' in texto:
        return Processo.CriterioJulgamento.MELHOR_TECNICA
    if 'ITEM' in texto or 'MENOR PRECO' in texto:
        return Processo.CriterioJulgamento.MENOR_PRECO_POR_ITEM
    return ''


def _upsert_pncp_snapshot(processo: Processo, numero_controle: str, payload: dict):
    payload = payload if isinstance(payload, dict) else {}
    numero_controle = _pncp_safe_text(numero_controle or payload.get('numeroControlePNCP'))
    if not numero_controle:
        raise ValueError('Numero de controle PNCP nao informado no retorno da API.')

    orgao = _pncp_safe_dict(payload.get('orgaoEntidade'))
    unidade = _pncp_safe_dict(payload.get('unidadeOrgao'))
    municipio = unidade.get('municipioNome')
    uf = unidade.get('ufNome')
    orgao_sub = _pncp_safe_dict(payload.get('orgaoSubRogado'))
    unidade_sub = _pncp_safe_dict(payload.get('unidadeSubRogada'))
    municipio_sub = unidade_sub.get('municipioNome')
    uf_sub = unidade_sub.get('ufNome')
    amparo = payload.get('amparoLegal')
    amparo_dict = _pncp_safe_dict(amparo)

    municipio_nome = _pncp_safe_text(municipio.get('nome') if isinstance(municipio, dict) else municipio)
    uf_sigla = _pncp_safe_text(uf.get('sigla') if isinstance(uf, dict) else '')
    uf_nome = _pncp_safe_text(uf.get('nome') if isinstance(uf, dict) else uf)
    municipio_sub_nome = _pncp_safe_text(municipio_sub.get('nome') if isinstance(municipio_sub, dict) else municipio_sub)
    uf_sub_sigla = _pncp_safe_text(uf_sub.get('sigla') if isinstance(uf_sub, dict) else '')
    uf_sub_nome = _pncp_safe_text(uf_sub.get('nome') if isinstance(uf_sub, dict) else uf_sub)

    amparo_codigo = _pncp_to_int(amparo_dict.get('codigo'))
    amparo_nome = _pncp_safe_text(amparo_dict.get('nome'))
    amparo_descricao = _pncp_safe_text(amparo_dict.get('descricao'))
    if not amparo_nome and not amparo_descricao and not amparo_codigo and amparo and not isinstance(amparo, dict):
        amparo_descricao = _pncp_safe_text(amparo)

    snap, _ = PNCPContratacaoSnapshot.objects.update_or_create(
        processo=processo,
        defaults={
            'numero_controle_pncp': numero_controle,
            'numero_compra': _pncp_safe_text(payload.get('numeroCompra')),
            'ano_compra': _pncp_to_int(payload.get('anoCompra')),
            'processo_origem': _pncp_safe_text(payload.get('processo')),
            'tipo_instrumento_convocatorio_id': _pncp_to_int(
                payload.get('tipoInstrumentoConvocatorioCodigo')
                or payload.get('tipoInstrumentoConvocatorioId')
            ),
            'tipo_instrumento_convocatorio_nome': _pncp_safe_text(payload.get('tipoInstrumentoConvocatorioNome')),
            'modalidade_id': _pncp_to_int(payload.get('modalidadeId') or payload.get('modalidadeCodigo')),
            'modalidade_nome': _pncp_safe_text(payload.get('modalidadeNome')),
            'modo_disputa_id': _pncp_to_int(payload.get('modoDisputaId') or payload.get('modoDisputaCodigo')),
            'modo_disputa_nome': _pncp_safe_text(payload.get('modoDisputaNome')),
            'criterio_julgamento_id': _pncp_to_int(payload.get('criterioJulgamentoId') or payload.get('criterioJulgamentoCodigo')),
            'criterio_julgamento_nome': _pncp_safe_text(payload.get('criterioJulgamentoNome')),
            'situacao_compra_id': _pncp_to_int(payload.get('situacaoCompraId') or payload.get('situacaoCompraCodigo')),
            'situacao_compra_nome': _pncp_safe_text(payload.get('situacaoCompraNome')),
            'objeto_compra': _pncp_safe_text(payload.get('objetoCompra')),
            'informacao_complementar': _pncp_safe_text(payload.get('informacaoComplementar')),
            'srp': _pncp_to_bool(payload.get('srp')),
            'amparo_legal_codigo': amparo_codigo,
            'amparo_legal_nome': amparo_nome,
            'amparo_legal_descricao': amparo_descricao,
            'valor_total_estimado': _pncp_to_decimal(payload.get('valorTotalEstimado')),
            'valor_total_homologado': _pncp_to_decimal(payload.get('valorTotalHomologado')),
            'data_abertura_proposta': _parse_pncp_datetime(payload.get('dataAberturaProposta')),
            'data_encerramento_proposta': _parse_pncp_datetime(payload.get('dataEncerramentoProposta')),
            'data_publicacao_pncp': _parse_pncp_date(payload.get('dataPublicacaoPncp')),
            'data_inclusao': _parse_pncp_date(payload.get('dataInclusao')),
            'data_atualizacao': _parse_pncp_date(payload.get('dataAtualizacao')),
            'sequencial_compra': _pncp_to_int(payload.get('sequencialCompra')),
            'orgao_cnpj': _pncp_safe_text(orgao.get('cnpj')),
            'orgao_razao_social': _pncp_safe_text(orgao.get('razaoSocial')),
            'orgao_poder_id': _pncp_safe_text(orgao.get('poderId')),
            'orgao_esfera_id': _pncp_safe_text(orgao.get('esferaId')),
            'unidade_codigo': _pncp_safe_text(unidade.get('codigoUnidade')),
            'unidade_nome': _pncp_safe_text(unidade.get('nomeUnidade')),
            'unidade_codigo_ibge': _pncp_safe_text(unidade.get('codigoIbge')),
            'unidade_municipio_nome': municipio_nome,
            'unidade_uf_sigla': uf_sigla,
            'unidade_uf_nome': uf_nome,
            'orgao_subrogado_cnpj': _pncp_safe_text(orgao_sub.get('cnpj')),
            'orgao_subrogado_razao_social': _pncp_safe_text(orgao_sub.get('razaoSocial')),
            'orgao_subrogado_poder_id': _pncp_safe_text(orgao_sub.get('poderId')),
            'orgao_subrogado_esfera_id': _pncp_safe_text(orgao_sub.get('esferaId')),
            'unidade_subrogada_codigo': _pncp_safe_text(unidade_sub.get('codigoUnidade')),
            'unidade_subrogada_nome': _pncp_safe_text(unidade_sub.get('nomeUnidade')),
            'unidade_subrogada_codigo_ibge': _pncp_safe_text(unidade_sub.get('codigoIbge')),
            'unidade_subrogada_municipio_nome': municipio_sub_nome,
            'unidade_subrogada_uf_sigla': uf_sub_sigla,
            'unidade_subrogada_uf_nome': uf_sub_nome,
            'usuario_nome': _pncp_safe_text(payload.get('usuarioNome')),
            'link_sistema_origem': _pncp_safe_text(payload.get('linkSistemaOrigem')),
            'justificativa_presencial': _pncp_safe_text(payload.get('justificativaPresencial')),
            'payload_completo': payload,
        },
    )
    return snap


def _sincronizar_processo_com_snapshot_pncp(processo: Processo, snapshot: PNCPContratacaoSnapshot):
    modalidade_nome = _normalizar_modalidade_pncp(snapshot.modalidade_nome)
    modalidade_obj = _get_modalidade(modalidade_nome)

    numero_proc = snapshot.processo_origem or snapshot.numero_compra or processo.numero_processo_adm
    ano_ref = snapshot.ano_compra or processo.ano_referencia or timezone.localdate().year
    objeto = snapshot.objeto_compra or processo.objeto

    processo.numero_processo_adm = numero_proc
    processo.ano_referencia = int(ano_ref)
    processo.objeto = objeto
    processo.modalidade = modalidade_obj
    processo.modo_disputa = _mapear_modo_disputa_pncp(snapshot.modo_disputa_nome)
    processo.criterio_julgamento = _mapear_criterio_pncp(snapshot.criterio_julgamento_nome)
    processo.valor_estimado = (snapshot.valor_total_estimado or Decimal('0')).quantize(Decimal('0.01'))
    processo.valor_homologado = (snapshot.valor_total_homologado or Decimal('0')).quantize(Decimal('0.01'))
    if snapshot.data_publicacao_pncp:
        processo.data_publicacao = snapshot.data_publicacao_pncp
    if snapshot.data_abertura_proposta:
        processo.data_hora_abertura = snapshot.data_abertura_proposta
    if snapshot.srp:
        processo.tipo_contratacao = Processo.TipoContratacao.REGISTRO_PRECO
    elif not processo.tipo_contratacao:
        processo.tipo_contratacao = Processo.TipoContratacao.AQUISICAO

    situacao_raw = snapshot.situacao_compra_nome or ''
    situacao = ''.join(
        ch for ch in unicodedata.normalize('NFKD', str(situacao_raw))
        if not unicodedata.combining(ch)
    ).upper()

    status_nome = 'EM PLANEJAMENTO'
    wf_situacao = SituacaoWorkflow.EM_ANDAMENTO
    etapa_pncp = 'IMPORTADO DO PNCP'
    wf_publicado = bool(snapshot.data_publicacao_pncp)
    wf_homologado = False
    wf_finalizado_licitacao = False

    if 'HOMOLOG' in situacao:
        status_nome = 'HOMOLOGADO'
        wf_situacao = SituacaoWorkflow.CONCLUIDO
        etapa_pncp = 'HOMOLOGADO NO PNCP'
        wf_homologado = True
        wf_finalizado_licitacao = True
        wf_publicado = True
    elif 'CONCL' in situacao:
        status_nome = 'CONCLUIDO'
        wf_situacao = SituacaoWorkflow.CONCLUIDO
        etapa_pncp = 'CONCLUIDO NO PNCP'
        wf_finalizado_licitacao = True
    elif 'DIVULG' in situacao or 'PUBLIC' in situacao:
        status_nome = 'PUBLICADO'
        wf_situacao = SituacaoWorkflow.EM_ANDAMENTO
        etapa_pncp = 'PUBLICADO NO PNCP'
        wf_publicado = True
    elif 'SUSPENS' in situacao or 'ANUL' in situacao or 'REVOG' in situacao:
        status_nome = 'SUSPENSO'
        wf_situacao = SituacaoWorkflow.SUSPENSO
        etapa_pncp = 'SUSPENSO NO PNCP'
    elif 'FRACASS' in situacao or 'DESERT' in situacao:
        status_nome = 'FRACASSADO'
        wf_situacao = SituacaoWorkflow.CONCLUIDO
        etapa_pncp = 'RESULTADO NO PNCP'
        wf_finalizado_licitacao = True
    elif 'AGUARD' in situacao:
        status_nome = 'AGUARDANDO'
        wf_situacao = SituacaoWorkflow.AGUARDANDO
        etapa_pncp = 'AGUARDANDO NO PNCP'

    # Forca status com base na camada canonica de itens quando houver divergencia.
    status_itens = list(
        ProcessoItem.objects.filter(processo=processo).values_list('status_consolidado', flat=True)
    )
    if status_itens:
        if any(s == ProcessoItem.StatusConsolidado.HOMOLOGADO for s in status_itens):
            status_nome = 'HOMOLOGADO'
            wf_situacao = SituacaoWorkflow.CONCLUIDO
            etapa_pncp = 'HOMOLOGADO NO PNCP'
            wf_homologado = True
            wf_finalizado_licitacao = True
            wf_publicado = True
        elif all(
            s in {ProcessoItem.StatusConsolidado.FRACASSADO, ProcessoItem.StatusConsolidado.CANCELADO}
            for s in status_itens
        ):
            status_nome = 'FRACASSADO'
            wf_situacao = SituacaoWorkflow.CONCLUIDO
            etapa_pncp = 'RESULTADO NO PNCP'
            wf_finalizado_licitacao = True

    processo.status = _get_status(status_nome)
    processo.save()

    wf = _ensure_workflow(processo)
    wf.pncp_numero_controle = snapshot.numero_controle_pncp
    wf.pncp_ultima_importacao = timezone.now()
    if wf.modulo_atual == ModuloSistema.PLANEJAMENTO and (snapshot.situacao_compra_nome or snapshot.data_publicacao_pncp):
        wf.modulo_atual = ModuloSistema.LICITACAO
    wf.etapa_atual = etapa_pncp
    wf.situacao = wf_situacao
    wf.publicado = wf_publicado
    wf.homologado = wf_homologado
    wf.finalizado_licitacao = wf_finalizado_licitacao
    wf.save(
        update_fields=[
            'pncp_numero_controle',
            'pncp_ultima_importacao',
            'modulo_atual',
            'etapa_atual',
            'situacao',
            'publicado',
            'homologado',
            'finalizado_licitacao',
            'atualizado_em',
        ]
    )


def _pncp_mapping_report(snapshot: PNCPContratacaoSnapshot, processo: Processo):
    payload = snapshot.payload_completo if isinstance(snapshot.payload_completo, dict) else {}
    payload_keys = set(payload.keys())
    mapped_payload_keys = {
        'numeroControlePNCP', 'numeroCompra', 'anoCompra', 'processo',
        'tipoInstrumentoConvocatorioCodigo', 'tipoInstrumentoConvocatorioId', 'tipoInstrumentoConvocatorioNome',
        'modalidadeId', 'modalidadeCodigo', 'modalidadeNome', 'modoDisputaId', 'modoDisputaCodigo', 'modoDisputaNome',
        'criterioJulgamentoId', 'criterioJulgamentoCodigo', 'criterioJulgamentoNome',
        'situacaoCompraId', 'situacaoCompraCodigo', 'situacaoCompraNome',
        'objetoCompra', 'informacaoComplementar', 'srp',
        'amparoLegal', 'valorTotalEstimado', 'valorTotalHomologado',
        'dataAberturaProposta', 'dataEncerramentoProposta',
        'dataPublicacaoPncp', 'dataInclusao', 'dataAtualizacao',
        'sequencialCompra', 'orgaoEntidade', 'unidadeOrgao',
        'orgaoSubRogado', 'unidadeSubRogada', 'usuarioNome',
        'linkSistemaOrigem', 'justificativaPresencial',
    }
    extras = sorted(payload_keys - mapped_payload_keys)
    return {
        'payload_total_campos': len(payload_keys),
        'payload_campos_mapeados': len(payload_keys.intersection(mapped_payload_keys)),
        'payload_campos_extras': extras,
        'constantes': {
            'modalidade_pncp': snapshot.modalidade_nome or '-',
            'modalidade_sistema': getattr(processo.modalidade, 'nome', '-') or '-',
            'modo_disputa_pncp': snapshot.modo_disputa_nome or '-',
            'modo_disputa_sistema': processo.get_modo_disputa_display() if processo.modo_disputa else '-',
            'criterio_pncp': snapshot.criterio_julgamento_nome or '-',
            'criterio_sistema': processo.get_criterio_julgamento_display() if processo.criterio_julgamento else '-',
            'situacao_compra_pncp': snapshot.situacao_compra_nome or '-',
            'status_sistema': getattr(processo.status, 'nome', '-') or '-',
            'tipo_instrumento_pncp': snapshot.tipo_instrumento_convocatorio_nome or '-',
            'tipo_contratacao_sistema': processo.get_tipo_contratacao_display() if processo.tipo_contratacao else '-',
        },
    }


def _pncp_pick(payload: dict, *keys):
    if not isinstance(payload, dict):
        return None
    for key in keys:
        value = payload.get(key)
        if value not in (None, '', [], {}):
            return value
    return None


def _pncp_digits(value):
    return ''.join(ch for ch in str(value or '') if ch.isdigit())


def _pncp_normalizar_upper(value):
    raw = _pncp_safe_text(value)
    if not raw:
        return ''
    return ''.join(
        ch for ch in unicodedata.normalize('NFKD', raw)
        if not unicodedata.combining(ch)
    ).upper()


def _pncp_extract_embedded_list(payload, keys: list[str]):
    keys_norm = {str(k or '').strip().lower() for k in (keys or []) if str(k or '').strip()}
    if not keys_norm:
        return []

    def _as_dict_list(value):
        if isinstance(value, list):
            return [x for x in value if isinstance(x, dict)]
        if isinstance(value, dict):
            for inner in ('data', 'resultado', 'resultados', 'itens', 'items', 'content'):
                val = value.get(inner)
                if isinstance(val, list):
                    return [x for x in val if isinstance(x, dict)]
        return []

    fila = [payload]
    visitados = 0
    while fila and visitados < 300:
        atual = fila.pop(0)
        visitados += 1
        if isinstance(atual, dict):
            for chave, valor in atual.items():
                chave_norm = str(chave or '').strip().lower()
                if chave_norm in keys_norm:
                    lista = _as_dict_list(valor)
                    if lista:
                        return lista
                if isinstance(valor, (dict, list)):
                    fila.append(valor)
        elif isinstance(atual, list):
            for valor in atual[:120]:
                if isinstance(valor, (dict, list)):
                    fila.append(valor)
    return []


def _mapear_status_item_pncp(*, situacao_item: str = '', situacao_resultado: str = '', situacao_compra: str = ''):
    texto = ' | '.join(
        [
            _pncp_normalizar_upper(situacao_item),
            _pncp_normalizar_upper(situacao_resultado),
            _pncp_normalizar_upper(situacao_compra),
        ]
    )
    if not texto.strip():
        return FornecimentoItem.StatusItem.PLANEJADO
    if any(token in texto for token in ['FRACASS', 'DESERT', 'SEM VENCEDOR', 'ITEM MAL SUCEDIDO']):
        return FornecimentoItem.StatusItem.FRACASSADO
    if any(token in texto for token in ['CANCEL', 'ANUL', 'REVOG']):
        return FornecimentoItem.StatusItem.CANCELADO
    if any(token in texto for token in ['HOMOLOG', 'ADJUDIC', 'VENCEDOR']):
        return FornecimentoItem.StatusItem.HOMOLOGADO
    return FornecimentoItem.StatusItem.PLANEJADO


def _legacy_status_para_consolidado(status_item: str) -> str:
    if status_item == FornecimentoItem.StatusItem.HOMOLOGADO:
        return ProcessoItem.StatusConsolidado.HOMOLOGADO
    if status_item == FornecimentoItem.StatusItem.FRACASSADO:
        return ProcessoItem.StatusConsolidado.FRACASSADO
    if status_item == FornecimentoItem.StatusItem.CANCELADO:
        return ProcessoItem.StatusConsolidado.CANCELADO
    return ProcessoItem.StatusConsolidado.EM_COTACAO


def _legacy_status_para_resultado(status_item: str) -> str:
    if status_item == FornecimentoItem.StatusItem.HOMOLOGADO:
        return ProcessoItemResultado.StatusResultado.HOMOLOGADO
    if status_item == FornecimentoItem.StatusItem.FRACASSADO:
        return ProcessoItemResultado.StatusResultado.FRACASSADO
    if status_item == FornecimentoItem.StatusItem.CANCELADO:
        return ProcessoItemResultado.StatusResultado.CANCELADO
    return ProcessoItemResultado.StatusResultado.CLASSIFICADO


def _pncp_unidade_text(value):
    if isinstance(value, dict):
        return _pncp_safe_text(
            value.get('sigla')
            or value.get('nome')
            or value.get('descricao')
            or value.get('codigo')
        )
    return _pncp_safe_text(value)


def _buscar_fornecedor_por_documento(documento_digits: str, cache_por_doc: dict | None = None):
    if not documento_digits:
        return None
    if cache_por_doc is not None and documento_digits in cache_por_doc:
        return cache_por_doc[documento_digits]
    for fornecedor in Fornecedor.objects.only('id', 'cnpj'):
        if _pncp_digits(fornecedor.cnpj) == documento_digits:
            if cache_por_doc is not None:
                cache_por_doc[documento_digits] = fornecedor
            return fornecedor
    return None


def _upsert_fornecedor_from_pncp(
    documento: str,
    nome: str,
    cidade: str = '',
    estado: str = '',
    cache_por_doc: dict | None = None,
    cache_por_razao: dict | None = None,
):
    doc_digits = _pncp_digits(documento)
    fornecedor = _buscar_fornecedor_por_documento(doc_digits, cache_por_doc=cache_por_doc) if doc_digits else None
    razao = (nome or '').strip()
    if not fornecedor and cache_por_razao is not None and razao and razao in cache_por_razao:
        fornecedor = cache_por_razao[razao]

    if not razao:
        razao = f'Fornecedor PNCP {doc_digits or "N/I"}'

    if not fornecedor:
        cnpj_salvar = doc_digits if doc_digits else _fornecedor_cnpj_fallback()
        fornecedor = Fornecedor.objects.create(
            razao_social=razao[:255],
            cnpj=cnpj_salvar[:18],
            cidade=(cidade or '').strip()[:100],
            estado=(estado or '').strip()[:2],
        )
        ensure_fornecedor_documento_externo(
            fornecedor,
            origem="PNCP",
            documento=doc_digits,
            identificador=doc_digits,
            payload={"nome": razao or "", "cidade": cidade or "", "estado": estado or ""},
        )
        if doc_digits and cache_por_doc is not None:
            cache_por_doc[doc_digits] = fornecedor
        if cache_por_razao is not None and razao:
            cache_por_razao[razao] = fornecedor
        return fornecedor

    changed = False
    if razao and fornecedor.razao_social != razao[:255]:
        fornecedor.razao_social = razao[:255]
        changed = True
    if cidade and (fornecedor.cidade or '').strip() != cidade.strip()[:100]:
        fornecedor.cidade = cidade.strip()[:100]
        changed = True
    if estado and (fornecedor.estado or '').strip() != estado.strip()[:2]:
        fornecedor.estado = estado.strip()[:2]
        changed = True
    if changed:
        fornecedor.save(update_fields=['razao_social', 'cidade', 'estado'])
    ensure_fornecedor_documento_externo(
        fornecedor,
        origem="PNCP",
        documento=doc_digits,
        identificador=doc_digits,
        payload={"nome": razao or "", "cidade": cidade or "", "estado": estado or ""},
    )
    if doc_digits and cache_por_doc is not None:
        cache_por_doc[doc_digits] = fornecedor
    if cache_por_razao is not None and razao:
        cache_por_razao[razao] = fornecedor
    return fornecedor


def _sincronizar_itens_fornecedores_pncp(
    processo: Processo,
    itens_pncp: list[dict],
    resultados_pncp: list[dict],
    situacao_compra_nome: str = '',
    numero_controle: str = '',
):
    itens_importados = 0
    fornecedores_atualizados = 0
    status_forcados = 0
    itens_por_numero = {}
    numero_controle = _pncp_safe_text(numero_controle)
    itens_com_homologacao = 0
    itens_com_fracasso = 0
    itens_com_cancelamento = 0
    now = timezone.now()
    situacao_upper = _pncp_normalizar_upper(situacao_compra_nome)
    cache_fornecedor_por_doc = {}
    cache_fornecedor_por_razao = {}

    for idx, item in enumerate(itens_pncp, start=1):
        if not isinstance(item, dict):
            continue
        numero_item = _pncp_to_int(_pncp_pick(item, 'numeroItem', 'numeroItemCompra', 'numeroSequencialItem', 'itemNumero')) or idx
        codigo_item_externo = _pncp_safe_text(_pncp_pick(item, 'codigoItem', 'codigo', 'idItem', 'itemId'))
        descricao = _pncp_safe_text(_pncp_pick(item, 'descricaoItem', 'descricao', 'objeto', 'objetoCompra')) or f'Item {numero_item} importado do PNCP'
        unidade = _pncp_unidade_text(_pncp_pick(item, 'unidadeMedida', 'unidade', 'unidadeFornecimento'))
        quantidade = _pncp_to_decimal(_pncp_pick(item, 'quantidade', 'quantidadeTotal', 'qtd', 'quantidadeItem'))
        valor_unitario_estimado = _pncp_to_decimal(_pncp_pick(item, 'valorUnitarioEstimado', 'valorUnitario', 'valorEstimado', 'valor'))
        valor_total_estimado = _pncp_to_decimal(_pncp_pick(item, 'valorTotalEstimado', 'valorTotal', 'valorGlobal', 'valorItem'))
        if valor_total_estimado <= 0 and quantidade > 0 and valor_unitario_estimado > 0:
            valor_total_estimado = (quantidade * valor_unitario_estimado).quantize(Decimal('0.01'))
        situacao_item_pncp = _pncp_safe_text(
            _pncp_pick(
                item,
                'situacaoCompraItemNome',
                'situacaoCompraItem',
                'situacaoItemNome',
                'situacaoItem',
                'situacaoNome',
                'situacao',
            )
        )
        criterio_item_pncp = _pncp_safe_text(_pncp_pick(item, 'criterioJulgamentoNome', 'criterioJulgamento', 'criterio'))
        tipo_item_pncp = _pncp_safe_text(
            _pncp_pick(
                item,
                'materialOuServicoNome',
                'materialOuServico',
                'tipoItem',
                'tipo',
                'tipoObjeto',
            )
        )
        categoria_item_pncp = _pncp_safe_text(
            _pncp_pick(
                item,
                'itemCategoriaNome',
                'itemCategoriaId',
                'categoriaItem',
                'categoriaItemLeilao',
                'categoria',
            )
        )
        status_item_forcado = _mapear_status_item_pncp(
            situacao_item=situacao_item_pncp,
            situacao_resultado='',
            situacao_compra=situacao_compra_nome,
        )
        status_consolidado = _legacy_status_para_consolidado(status_item_forcado)
        status_resultado = _legacy_status_para_resultado(status_item_forcado)

        item_core = ProcessoItem.objects.filter(processo=processo, numero_item=numero_item).first()
        if not item_core:
            item_core = ProcessoItem.objects.create(
                processo=processo,
                numero_item=numero_item,
                descricao_snapshot=descricao,
                unidade_snapshot=unidade,
                quantidade=quantidade,
                status_consolidado=status_consolidado,
                valor_referencia_unitario=_decimal_2(valor_unitario_estimado),
                valor_referencia_total=_decimal_2(valor_total_estimado),
                pncp_ultima_atualizacao=now,
            )
            if status_item_forcado != FornecimentoItem.StatusItem.PLANEJADO:
                status_forcados += 1
        else:
            fields = []
            if descricao and item_core.descricao_snapshot != descricao:
                item_core.descricao_snapshot = descricao
                fields.append('descricao_snapshot')
            if unidade and item_core.unidade_snapshot != unidade:
                item_core.unidade_snapshot = unidade
                fields.append('unidade_snapshot')
            if quantidade > 0 and item_core.quantidade != quantidade:
                item_core.quantidade = quantidade
                fields.append('quantidade')
            if valor_unitario_estimado > 0 and item_core.valor_referencia_unitario != _decimal_2(valor_unitario_estimado):
                item_core.valor_referencia_unitario = _decimal_2(valor_unitario_estimado)
                fields.append('valor_referencia_unitario')
            if valor_total_estimado > 0 and item_core.valor_referencia_total != _decimal_2(valor_total_estimado):
                item_core.valor_referencia_total = _decimal_2(valor_total_estimado)
                fields.append('valor_referencia_total')
            if item_core.status_consolidado != status_consolidado:
                item_core.status_consolidado = status_consolidado
                fields.append('status_consolidado')
                status_forcados += 1
            item_core.pncp_ultima_atualizacao = now
            fields.append('pncp_ultima_atualizacao')
            if fields:
                item_core.save(update_fields=list(dict.fromkeys(fields)))

        ProcessoItemResultado.objects.update_or_create(
            processo=processo,
            processo_item=item_core,
            origem=ProcessoItemResultado.Origem.PNCP,
            chave_origem=f'pncp-item:{numero_item}',
            defaults={
                'status_resultado': status_resultado,
                'valor_unitario': _decimal_2(valor_unitario_estimado),
                'valor_total': _decimal_2(valor_total_estimado),
                'situacao_texto': (situacao_item_pncp or situacao_compra_nome or '')[:140],
                'payload_resumo': {
                    'numero_controle_pncp': numero_controle,
                    'codigo_item_externo': codigo_item_externo,
                    'criterio_julgamento_item_pncp': criterio_item_pncp,
                    'tipo_item_pncp': tipo_item_pncp,
                    'categoria_item_pncp': categoria_item_pncp,
                    'payload_item_pncp': item,
                },
                'ativo': True,
            },
        )

        lote_numero = _pncp_to_int(_pncp_pick(item, 'numeroLote', 'lote', 'numeroGrupo'))
        if lote_numero and lote_numero > 0:
            lote_obj, _ = Lote.objects.get_or_create(
                processo=processo,
                numero=lote_numero,
                defaults={'titulo': f'Lote {lote_numero}', 'status': 'PNCP'},
            )
            if lote_obj.status != 'PNCP':
                lote_obj.status = 'PNCP'
                lote_obj.save(update_fields=['status'])
            ProcessoLoteItem.objects.update_or_create(
                processo=processo,
                lote=lote_obj,
                item=item_core,
                defaults={'ativo': True},
            )

        itens_por_numero[numero_item] = item_core
        itens_importados += 1

    for resultado in resultados_pncp:
        if not isinstance(resultado, dict):
            continue
        numero_item = _pncp_to_int(_pncp_pick(resultado, 'numeroItem', 'numeroItemCompra', 'itemNumero', 'item'))
        item_core = itens_por_numero.get(numero_item) if numero_item else None
        if not item_core and numero_item:
            item_core = ProcessoItem.objects.filter(processo=processo, numero_item=numero_item).first()
            if item_core:
                itens_por_numero[numero_item] = item_core
        if not item_core:
            continue

        doc = _pncp_pick(
            resultado,
            'cnpjCpfFornecedor',
            'niFornecedor',
            'cnpjFornecedor',
            'cpfFornecedor',
            'numeroDocumentoFornecedor',
        )
        nome = _pncp_safe_text(_pncp_pick(
            resultado,
            'nomeRazaoSocialFornecedor',
            'razaoSocialFornecedor',
            'fornecedorNome',
            'nomeFornecedor',
            'fornecedor',
        ))
        cidade = _pncp_safe_text(_pncp_pick(resultado, 'municipioFornecedor', 'cidadeFornecedor'))
        estado = _pncp_safe_text(_pncp_pick(resultado, 'ufFornecedor', 'estadoFornecedor'))
        fornecedor = _upsert_fornecedor_from_pncp(
            doc,
            nome,
            cidade,
            estado,
            cache_por_doc=cache_fornecedor_por_doc,
            cache_por_razao=cache_fornecedor_por_razao,
        )

        situacao_resultado_pncp = _pncp_safe_text(
            _pncp_pick(
                resultado,
                'situacaoCompraItemResultadoNome',
                'situacaoCompraItemResultadoId',
                'situacaoResultadoNome',
                'situacaoResultado',
                'situacaoNome',
                'situacao',
                'status',
            )
        )
        ordem_classificacao = _pncp_to_int(
            _pncp_pick(
                resultado,
                'ordemClassificacaoSrp',
                'ordemClassificacao',
                'ordemDeClassificacao',
                'ordemClassificacaoFornecedor',
            )
        )
        data_homologacao = _parse_pncp_date(_pncp_pick(resultado, 'dataResultadoHomologacao', 'dataHomologacao', 'dataResultado'))

        valor_unitario_homologado = _pncp_to_decimal(_pncp_pick(
            resultado,
            'valorUnitarioHomologado',
            'valorHomologado',
            'valorUnitario',
            'valorFinal',
            'valor',
        ))
        valor_total_homologado = _pncp_to_decimal(_pncp_pick(resultado, 'valorTotalHomologado', 'valorTotalHomologadoItem', 'valorTotal'))
        if valor_total_homologado <= 0 and item_core.quantidade and valor_unitario_homologado > 0:
            valor_total_homologado = (item_core.quantidade * valor_unitario_homologado).quantize(Decimal('0.01'))

        status_resultado_forcado = _mapear_status_item_pncp(
            situacao_item='',
            situacao_resultado=situacao_resultado_pncp,
            situacao_compra=situacao_compra_nome,
        )
        chave_doc = _pncp_digits(doc) or _pncp_safe_text(doc)
        chave_resultado = f'pncp-res:{numero_item}:{ordem_classificacao or 0}:{chave_doc or "na"}'
        ProcessoItemResultado.objects.update_or_create(
            processo=processo,
            processo_item=item_core,
            origem=ProcessoItemResultado.Origem.PNCP,
            chave_origem=chave_resultado,
            defaults={
                'fornecedor': fornecedor,
                'status_resultado': _legacy_status_para_resultado(status_resultado_forcado),
                'classificacao': ordem_classificacao,
                'valor_unitario': _decimal_2(valor_unitario_homologado),
                'valor_total': _decimal_2(valor_total_homologado),
                'data_resultado': data_homologacao,
                'situacao_texto': (situacao_resultado_pncp or '')[:140],
                'payload_resumo': {
                    'documento_fornecedor': doc,
                    'nome_fornecedor': nome,
                    'cidade_fornecedor': cidade,
                    'estado_fornecedor': estado,
                    'payload_resultado_pncp': resultado,
                },
                'ativo': True,
            },
        )
        if fornecedor:
            fornecedores_atualizados += 1

        fields = []
        if fornecedor and item_core.fornecedor_homologado_id != fornecedor.id:
            item_core.fornecedor_homologado = fornecedor
            fields.append('fornecedor_homologado')
        if valor_unitario_homologado > 0 and item_core.valor_homologado_unitario != _decimal_2(valor_unitario_homologado):
            item_core.valor_homologado_unitario = _decimal_2(valor_unitario_homologado)
            fields.append('valor_homologado_unitario')
        if valor_total_homologado > 0 and item_core.valor_homologado_total != _decimal_2(valor_total_homologado):
            item_core.valor_homologado_total = _decimal_2(valor_total_homologado)
            fields.append('valor_homologado_total')
        novo_status_consolidado = _legacy_status_para_consolidado(status_resultado_forcado)
        if item_core.status_consolidado != novo_status_consolidado:
            item_core.status_consolidado = novo_status_consolidado
            fields.append('status_consolidado')
            status_forcados += 1
        item_core.pncp_ultima_atualizacao = now
        fields.append('pncp_ultima_atualizacao')
        if fields:
            item_core.save(update_fields=list(dict.fromkeys(fields)))

    if not resultados_pncp and 'HOMOLOG' in situacao_upper:
        alvo = list(itens_por_numero.values()) or list(ProcessoItem.objects.filter(processo=processo))
        for item_core in alvo:
            if item_core.status_consolidado != ProcessoItem.StatusConsolidado.HOMOLOGADO:
                item_core.status_consolidado = ProcessoItem.StatusConsolidado.HOMOLOGADO
                item_core.pncp_ultima_atualizacao = now
                item_core.save(update_fields=['status_consolidado', 'pncp_ultima_atualizacao', 'atualizado_em'])
                ProcessoItemResultado.objects.update_or_create(
                    processo=processo,
                    processo_item=item_core,
                    origem=ProcessoItemResultado.Origem.PNCP,
                    chave_origem=f'pncp-item:{item_core.numero_item}',
                    defaults={
                        'status_resultado': ProcessoItemResultado.StatusResultado.HOMOLOGADO,
                        'situacao_texto': 'Homologado (inferido)',
                        'ativo': True,
                    },
                )
                status_forcados += 1

    atualizados_fallback = 0
    itens_processo = list(ProcessoItem.objects.filter(processo=processo))
    total_estimado = Decimal('0.00')
    total_homologado = Decimal('0.00')
    for item_core in itens_processo:
        est = item_core.valor_referencia_total
        if (not est or est <= 0) and item_core.quantidade and item_core.valor_referencia_unitario:
            est = (item_core.quantidade * item_core.valor_referencia_unitario).quantize(Decimal('0.01'))
        total_estimado += (est or Decimal('0.00'))

        hom = item_core.valor_homologado_total
        if (not hom or hom <= 0) and item_core.status_consolidado == ProcessoItem.StatusConsolidado.HOMOLOGADO:
            hom = (Decimal(str(item_core.quantidade or 0)) * Decimal(str(item_core.valor_homologado_unitario or 0))).quantize(Decimal('0.01'))
        total_homologado += (hom or Decimal('0.00'))

        if item_core.status_consolidado == ProcessoItem.StatusConsolidado.HOMOLOGADO:
            itens_com_homologacao += 1
        elif item_core.status_consolidado == ProcessoItem.StatusConsolidado.FRACASSADO:
            itens_com_fracasso += 1
        elif item_core.status_consolidado == ProcessoItem.StatusConsolidado.CANCELADO:
            itens_com_cancelamento += 1

    update_processo_fields = []
    total_estimado = (total_estimado or Decimal('0.00')).quantize(Decimal('0.01'))
    total_homologado = (total_homologado or Decimal('0.00')).quantize(Decimal('0.01'))
    if total_estimado > 0 and processo.valor_estimado != total_estimado:
        processo.valor_estimado = total_estimado
        update_processo_fields.append('valor_estimado')
    if total_homologado > 0 and processo.valor_homologado != total_homologado:
        processo.valor_homologado = total_homologado
        update_processo_fields.append('valor_homologado')
    if update_processo_fields:
        processo.save(update_fields=update_processo_fields + ['atualizado_em'])

    for lote in Lote.objects.filter(processo=processo):
        qtd = ProcessoLoteItem.objects.filter(processo=processo, lote=lote, ativo=True).count()
        if lote.qtd_itens != qtd:
            lote.qtd_itens = qtd
            lote.save(update_fields=['qtd_itens'])
    sync_canonico = sync_canonical_items_for_processo(processo)
    for item in ProcessoItem.objects.filter(processo=processo):
        _espelhar_item_canonico_legacy(item)
    _aplicar_lote_espelho_legacy(processo)

    return {
        'itens_importados': itens_importados,
        'itens_atualizados_fallback': atualizados_fallback,
        'fornecedores_atualizados': fornecedores_atualizados,
        'status_forcados': status_forcados,
        'itens_homologados': itens_com_homologacao,
        'itens_fracassados': itens_com_fracasso,
        'itens_cancelados': itens_com_cancelamento,
        'total_estimado_itens': str(total_estimado),
        'total_homologado_itens': str(total_homologado),
        'resultados_recebidos': len(resultados_pncp or []),
        'sincronizacao_canonica': sync_canonico,
    }


def _modulos_vizinhos(modulo_atual):
    try:
        idx = FLUXO_MODULOS_PROCESSO.index(modulo_atual)
    except ValueError:
        return None, None
    modulo_anterior = FLUXO_MODULOS_PROCESSO[idx - 1] if idx > 0 else None
    modulo_proximo = FLUXO_MODULOS_PROCESSO[idx + 1] if idx < (len(FLUXO_MODULOS_PROCESSO) - 1) else None
    return modulo_anterior, modulo_proximo


def _proximo_numero_ci(modulo_origem: str, ano_exercicio: int) -> int:
    agg = (
        ComunicacaoInterna.objects.filter(
            modulo_origem=modulo_origem,
            ano_exercicio=ano_exercicio,
        ).aggregate(ultimo=Max('numero_sequencial'))
    )
    ultimo = agg.get('ultimo') or 0
    return int(ultimo) + 1


def _secretaria_adm_principal():
    adm = Secretaria.objects.filter(sigla__iexact='ADM').first()
    if adm:
        return adm
    adm = Secretaria.objects.filter(nome__icontains='ADMINISTRA').first()
    if adm:
        return adm
    return Secretaria.objects.order_by('sigla').first()


def _orgao_ativo():
    try:
        return OrgaoEntidade.objects.order_by('-atualizado_em', '-id').first()
    except OperationalError:
        return None


def _decimal_2(valor):
    return Decimal(str(valor or 0)).quantize(Decimal('0.01'))


def _decimal_3(valor):
    return Decimal(str(valor or 0)).quantize(Decimal('0.001'))


def _fornecedor_cnpj_fallback():
    base = int(timezone.now().timestamp() * 1000000)
    for idx in range(0, 1000):
        cnpj = f'{(base + idx) % 10**14:014d}'
        if not Fornecedor.objects.filter(cnpj=cnpj).exists():
            return cnpj
    return f'{timezone.now().strftime("%H%M%S%f")[:14]:0>14}'


def _sincronizar_fornecedor_fonte(fonte: ETPCotacaoFonte):
    if fonte.tipo_fonte != ETPCotacaoFonte.TipoFonte.FORNECEDOR:
        return
    fornecedor = fonte.fornecedor
    if not fornecedor and fonte.nome_fonte:
        fornecedor = Fornecedor.objects.filter(razao_social__iexact=fonte.nome_fonte.strip()).first()
    if not fornecedor:
        nome = (fonte.nome_fonte or '').strip()
        if not nome:
            return
        fornecedor = Fornecedor.objects.create(razao_social=nome, cnpj=_fornecedor_cnpj_fallback())
    if fonte.fornecedor_id != fornecedor.id:
        fonte.fornecedor = fornecedor
        if not fonte.nome_fonte:
            fonte.nome_fonte = fornecedor.razao_social
        fonte.save(update_fields=['fornecedor', 'nome_fonte'])


def _espelhar_item_canonico_legacy(item: ProcessoItem):
    # Compatibilidade temporaria para telas/exportacoes ainda acopladas ao modelo legado.
    legacy = (
        FornecimentoItem.objects.filter(processo=item.processo, numero_item=item.numero_item)
        .order_by('-id')
        .first()
    )
    if not legacy:
        legacy = FornecimentoItem(processo=item.processo, numero_item=item.numero_item)
    legacy.descricao = item.descricao_snapshot
    legacy.unidade = item.unidade_snapshot or ""
    legacy.quantidade = _decimal_3(item.quantidade)
    legacy.fornecedor = item.fornecedor_homologado
    legacy.valor_unitario = _decimal_2(item.valor_homologado_unitario or item.valor_referencia_unitario or 0)
    legacy.valor_total = _decimal_2(item.valor_homologado_total or item.valor_referencia_total or 0)
    legacy.valor_unitario_estimado = _decimal_2(item.valor_referencia_unitario or 0)
    legacy.valor_total_estimado = _decimal_2(item.valor_referencia_total or 0)
    legacy.valor_unitario_homologado = _decimal_2(item.valor_homologado_unitario or 0)
    legacy.valor_total_homologado = _decimal_2(item.valor_homologado_total or 0)
    if item.status_consolidado == ProcessoItem.StatusConsolidado.HOMOLOGADO:
        legacy.status_item = FornecimentoItem.StatusItem.HOMOLOGADO
    elif item.status_consolidado == ProcessoItem.StatusConsolidado.FRACASSADO:
        legacy.status_item = FornecimentoItem.StatusItem.FRACASSADO
    elif item.status_consolidado == ProcessoItem.StatusConsolidado.CANCELADO:
        legacy.status_item = FornecimentoItem.StatusItem.CANCELADO
    else:
        legacy.status_item = FornecimentoItem.StatusItem.PLANEJADO
    legacy.pncp_ultima_atualizacao = item.pncp_ultima_atualizacao
    legacy.save()
    return legacy


def _aplicar_lote_espelho_legacy(processo: Processo):
    lote_por_numero_item = {
        row['item__numero_item']: row['lote_id']
        for row in ProcessoLoteItem.objects.filter(processo=processo, ativo=True).values('item__numero_item', 'lote_id')
    }
    for legacy in FornecimentoItem.objects.filter(processo=processo):
        novo_lote_id = lote_por_numero_item.get(legacy.numero_item)
        if legacy.lote_id != novo_lote_id:
            legacy.lote_id = novo_lote_id
            legacy.save(update_fields=['lote'])


def _referencia_cotacao_por_metodologia(etp: ETPPlanejamento, item: DFDItem):
    valores = sorted(
        [Decimal(str(v)) for v in etp.cotacoes.filter(item=item, considerar_no_calculo=True).values_list('valor_unitario', flat=True)]
    )
    if not valores:
        return None
    if etp.metodologia_cotacao == ETPPlanejamento.MetodologiaCotacao.MENOR:
        return min(valores)
    if etp.metodologia_cotacao == ETPPlanejamento.MetodologiaCotacao.MEDIANA:
        n = len(valores)
        meio = n // 2
        if n % 2:
            return valores[meio]
        return (valores[meio - 1] + valores[meio]) / Decimal('2')
    return sum(valores) / Decimal(len(valores))


def _item_core_base(processo: Processo, numero_item: int):
    return ProcessoItem.objects.filter(processo=processo, numero_item=numero_item).first()


def _sincronizar_item_core(processo: Processo, item_dfd: DFDItem):
    item_core = _item_core_base(processo, item_dfd.codigo)
    if not item_core:
        item_core = ProcessoItem(processo=processo, numero_item=item_dfd.codigo)

    item_core.descricao_snapshot = item_dfd.descricao
    item_core.unidade_snapshot = item_dfd.unidade or ''
    item_core.quantidade = _decimal_3(item_dfd.quantidade)
    if item_dfd.catalogo_id:
        item_core.item_catalogo_id = item_dfd.catalogo_id
    if not item_core.status_consolidado:
        item_core.status_consolidado = ProcessoItem.StatusConsolidado.PLANEJADO
    if (item_core.valor_referencia_unitario or Decimal('0')) > 0:
        item_core.valor_referencia_total = _decimal_2(
            Decimal(str(item_core.valor_referencia_unitario or 0)) * _decimal_3(item_dfd.quantidade)
        )
    if (item_core.valor_homologado_unitario or Decimal('0')) > 0:
        item_core.valor_homologado_total = _decimal_2(
            Decimal(str(item_core.valor_homologado_unitario or 0)) * _decimal_3(item_dfd.quantidade)
        )
    item_core.save()
    _espelhar_item_canonico_legacy(item_core)
    return item_core


def _remover_item_core(processo: Processo, numero_item: int):
    item_core = ProcessoItem.objects.filter(processo=processo, numero_item=numero_item).first()
    if not item_core:
        return
    if item_core.resultados.exclude(origem=ProcessoItemResultado.Origem.MANUAL).exists():
        item_core.status_consolidado = ProcessoItem.StatusConsolidado.CANCELADO
        item_core.save(update_fields=['status_consolidado', 'atualizado_em'])
        ProcessoItemResultado.objects.update_or_create(
            processo=processo,
            processo_item=item_core,
            origem=ProcessoItemResultado.Origem.MANUAL,
            chave_origem='manual:cancelamento',
            defaults={
                'status_resultado': ProcessoItemResultado.StatusResultado.CANCELADO,
                'situacao_texto': 'Item removido do DFD',
                'ativo': True,
            },
        )
        _espelhar_item_canonico_legacy(item_core)
    else:
        ProcessoLoteItem.objects.filter(processo=processo, item=item_core).delete()
        ProcessoItemResultado.objects.filter(processo=processo, processo_item=item_core).delete()
        item_core.delete()
        FornecimentoItem.objects.filter(processo=processo, numero_item=numero_item).delete()


def _sincronizar_itens_core(processo: Processo):
    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    if not dfd:
        return
    codigos = set()
    for item_dfd in dfd.itens.all():
        codigos.add(item_dfd.codigo)
        _sincronizar_item_core(processo, item_dfd)
    for item_core in ProcessoItem.objects.filter(processo=processo):
        if item_core.numero_item not in codigos:
            _remover_item_core(processo, item_core.numero_item)
    sync_canonical_items_for_processo(processo)
    for item in ProcessoItem.objects.filter(processo=processo):
        _espelhar_item_canonico_legacy(item)
    _aplicar_lote_espelho_legacy(processo)


def _sincronizar_lotes_core(processo: Processo):
    tr = TRPlanejamento.objects.filter(processo=processo).first()
    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    if not dfd:
        return

    desejados = {}
    if tr and tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE:
        for lote_tr in tr.lotes.prefetch_related('itens').all():
            desejados[lote_tr.numero] = {
                'titulo': lote_tr.titulo or f'Lote {lote_tr.numero}',
                'itens': {i.codigo for i in lote_tr.itens.all()},
            }
    elif tr and tr.criterio_julgamento in [
        TRPlanejamento.CriterioJulgamento.MENOR_PRECO_GLOBAL,
        TRPlanejamento.CriterioJulgamento.MAIOR_PERCENTUAL_DESCONTO,
        TRPlanejamento.CriterioJulgamento.MENOR_TAXA_ADMINISTRATIVA,
    ]:
        desejados[1] = {
            'titulo': 'Lote unico',
            'itens': {i.codigo for i in dfd.itens.all()},
        }

    lotes_core = {l.numero: l for l in Lote.objects.filter(processo=processo).order_by('numero')}
    lotes_ativos = {}
    for numero, payload in desejados.items():
        lote = lotes_core.get(numero)
        if not lote:
            lote = Lote.objects.create(processo=processo, numero=numero)
        lote.titulo = payload['titulo']
        lote.status = lote.status or 'PLANEJAMENTO'
        lote.qtd_itens = len(payload['itens'])
        lote.save(update_fields=['titulo', 'status', 'qtd_itens'])
        lotes_ativos[numero] = lote

    for numero, lote in lotes_core.items():
        if numero not in lotes_ativos and (not lote.status or lote.status == 'PLANEJAMENTO'):
            lote.delete()

    mapa_item_lote = {}
    for numero, payload in desejados.items():
        for codigo_item in payload['itens']:
            mapa_item_lote[codigo_item] = lotes_ativos[numero]

    ProcessoLoteItem.objects.filter(processo=processo).delete()
    for item_core in ProcessoItem.objects.filter(processo=processo):
        novo_lote = mapa_item_lote.get(item_core.numero_item)
        if novo_lote:
            ProcessoLoteItem.objects.create(
                processo=processo,
                lote=novo_lote,
                item=item_core,
                ativo=True,
            )

    for lote in lotes_ativos.values():
        qtd_itens = ProcessoLoteItem.objects.filter(processo=processo, lote=lote, ativo=True).count()
        if lote.qtd_itens != qtd_itens:
            lote.qtd_itens = qtd_itens
            lote.save(update_fields=['qtd_itens'])
    sync_canonical_items_for_processo(processo)
    _aplicar_lote_espelho_legacy(processo)


def _atualizar_estimativa_item_core(etp: ETPPlanejamento, item: DFDItem):
    referencia = _referencia_cotacao_por_metodologia(etp, item)
    if referencia is None:
        return
    valor_unitario = _decimal_2(referencia)
    item_core = ProcessoItem.objects.filter(processo=etp.processo, numero_item=item.codigo).first()
    if not item_core:
        return
    item_core.valor_referencia_unitario = valor_unitario
    item_core.valor_referencia_total = _decimal_2(valor_unitario * _decimal_3(item_core.quantidade))
    if item_core.status_consolidado == ProcessoItem.StatusConsolidado.PLANEJADO:
        item_core.status_consolidado = ProcessoItem.StatusConsolidado.EM_COTACAO
    item_core.save(update_fields=['valor_referencia_unitario', 'valor_referencia_total', 'status_consolidado', 'atualizado_em'])
    ProcessoItemResultado.objects.update_or_create(
        processo=etp.processo,
        processo_item=item_core,
        origem=ProcessoItemResultado.Origem.MANUAL,
        chave_origem='manual:estimativa',
        defaults={
            'status_resultado': ProcessoItemResultado.StatusResultado.CLASSIFICADO,
            'valor_unitario': valor_unitario,
            'valor_total': _decimal_2(valor_unitario * _decimal_3(item_core.quantidade)),
            'situacao_texto': f'Estimativa ETP ({etp.get_metodologia_cotacao_display()})',
            'ativo': True,
        },
    )
    _espelhar_item_canonico_legacy(item_core)
    sync_canonical_items_for_processo(etp.processo)


def _recalcular_alertas_cotacoes(etp: ETPPlanejamento, item: DFDItem):
    qs = etp.cotacoes.filter(item=item)
    base_calculo = qs.filter(considerar_no_calculo=True)
    valores = [Decimal(str(x.valor_unitario)) for x in base_calculo]
    if not valores:
        qs.update(inexequivel=False, sobrepreco=False)
        _atualizar_estimativa_item_core(etp, item)
        return
    media = sum(valores) / len(valores)
    piso = media * Decimal('0.5')
    teto = media * Decimal('1.5')
    for cot in qs:
        v = Decimal(str(cot.valor_unitario))
        cot.inexequivel = v < piso
        cot.sobrepreco = v > teto
        cot.save(update_fields=['inexequivel', 'sobrepreco'])
    _atualizar_estimativa_item_core(etp, item)


def _atendente_context(user):
    if not getattr(user, 'is_authenticated', False):
        return {'atendente_nome': '-', 'atendente_identificacao': '-'}
    nome = (user.get_full_name() or '').strip() or user.get_username()
    return {'atendente_nome': nome, 'atendente_identificacao': user.get_username()}


def _format_date_label(value):
    if not value:
        return '-'
    if hasattr(value, 'strftime'):
        return value.strftime('%d/%m/%Y')
    try:
        return datetime.fromisoformat(str(value)).strftime('%d/%m/%Y')
    except Exception:
        return str(value)


def _mapa_preview_context(processo: Processo):
    etp = ETPPlanejamento.objects.filter(processo=processo).first()
    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    if not etp:
        return {
            'etp': None,
            'cotacoes': [],
            'cotacoes_media': [],
            'itens_estimados': [],
            'total_estimado_processo': Decimal('0.00'),
        }
    cotacoes = list(
        etp.cotacoes.select_related('item', 'fonte').order_by('item__codigo', 'fonte__nome_fonte')
    )
    cotacoes_media = list(
        etp.cotacoes.filter(considerar_no_calculo=True)
        .values('item_id', 'item__codigo', 'item__descricao')
        .annotate(media=Avg('valor_unitario'))
        .order_by('item__codigo')
    )
    medias_por_item = {row['item_id']: Decimal(str(row['media'])) for row in cotacoes_media}
    itens_estimados = []
    total_estimado_processo = Decimal('0.00')
    for item in (dfd.itens.all().order_by('codigo') if dfd else []):
        quantidade = Decimal(str(item.quantidade or 0))
        media_item = medias_por_item.get(item.id)
        total_item = None
        if media_item is not None:
            total_item = (media_item * quantidade).quantize(Decimal('0.01'))
            total_estimado_processo += total_item
        itens_estimados.append({
            'item': item,
            'media': media_item,
            'total_estimado': total_item,
        })
    total_estimado_processo = total_estimado_processo.quantize(Decimal('0.01'))
    return {
        'etp': etp,
        'cotacoes': cotacoes,
        'cotacoes_media': cotacoes_media,
        'itens_estimados': itens_estimados,
        'total_estimado_processo': total_estimado_processo,
    }


def _status_lotes_tr(processo: Processo, tr: TRPlanejamento):
    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    itens_dfd = list(dfd.itens.all().order_by('codigo')) if dfd else []
    lotes = list(tr.lotes.prefetch_related('itens').all()) if tr else []
    lotes_por_item = {}
    for lote in lotes:
        for item in lote.itens.all():
            lotes_por_item.setdefault(item.id, []).append(lote.numero)

    duplicados = []
    for item in itens_dfd:
        lotes_item = lotes_por_item.get(item.id, [])
        if len(lotes_item) > 1:
            duplicados.append({'item': item, 'lotes': sorted(lotes_item)})

    nao_alocados = [item for item in itens_dfd if item.id not in lotes_por_item]
    return {
        'itens_total': len(itens_dfd),
        'nao_alocados': nao_alocados,
        'duplicados': duplicados,
        'todos_alocados': bool(itens_dfd) and not nao_alocados,
        'sem_duplicidade': not duplicados,
    }


def _parse_lote_item_ids(request):
    raw_ids = []
    payload = (request.POST.get('itens_payload') or '').strip()
    if payload:
        try:
            parsed = json.loads(payload)
            if isinstance(parsed, list):
                raw_ids = parsed
        except Exception:
            raw_ids = []
    if not raw_ids:
        raw_ids = request.POST.getlist('itens')

    item_ids = []
    for value in raw_ids:
        try:
            item_id = int(value)
        except (TypeError, ValueError):
            continue
        if item_id not in item_ids:
            item_ids.append(item_id)
    return item_ids


def _mapa_distribuicao_context(processo: Processo, tr: TRPlanejamento):
    if not tr or tr.nao_aplica_divisao_secretaria:
        return {'linhas': [], 'resumo_itens': []}
    linhas = list(
        tr.distribuicoes.select_related('secretaria', 'item')
        .order_by('secretaria__sigla', 'item__codigo')
    )
    totais_por_item = {}
    for linha in linhas:
        item_id = linha.item_id
        totais_por_item[item_id] = totais_por_item.get(item_id, Decimal('0')) + Decimal(str(linha.quantidade or 0))

    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    itens_dfd = dfd.itens.all().order_by('codigo') if dfd else []
    resumo_itens = []
    for item in itens_dfd:
        qtd_dfd = Decimal(str(item.quantidade or 0))
        qtd_distribuida = totais_por_item.get(item.id, Decimal('0'))
        resumo_itens.append({
            'item': item,
            'quantidade_dfd': qtd_dfd,
            'quantidade_distribuida': qtd_distribuida,
            'quantidade_pendente': (qtd_dfd - qtd_distribuida),
        })
    return {'linhas': linhas, 'resumo_itens': resumo_itens}


def _distribuicao_preview_context(processo: Processo):
    tr = TRPlanejamento.objects.filter(processo=processo).first()
    distribuicao_ctx = _mapa_distribuicao_context(processo, tr) if tr else {'linhas': [], 'resumo_itens': []}
    total_qtd_dfd = Decimal('0.000')
    total_qtd_distribuida = Decimal('0.000')
    for row in distribuicao_ctx['resumo_itens']:
        total_qtd_dfd += Decimal(str(row.get('quantidade_dfd') or 0))
        total_qtd_distribuida += Decimal(str(row.get('quantidade_distribuida') or 0))
    return {
        'tr': tr,
        'distribuicao_habilitada': bool(tr and not tr.nao_aplica_divisao_secretaria),
        'distribuicao_linhas': distribuicao_ctx['linhas'],
        'distribuicao_resumo_itens': distribuicao_ctx['resumo_itens'],
        'total_qtd_dfd': total_qtd_dfd,
        'total_qtd_distribuida': total_qtd_distribuida,
    }


def home(request):
    total_processos = Processo.objects.count()
    recentes = Processo.objects.order_by('-criado_em')[:10]
    try:
        total_workflows = ProcessoWorkflow.objects.count()
        por_modulo = ProcessoWorkflow.objects.values('modulo_atual').annotate(total=Count('id')).order_by('modulo_atual')
        aviso_migracao = ''
    except OperationalError:
        total_workflows = 0
        por_modulo = []
        aviso_migracao = 'As tabelas do módulo workflow ainda não existem neste banco. Execute: python manage.py migrate'
    context = {
        'total_processos': total_processos,
        'total_workflows': total_workflows,
        'por_modulo': por_modulo,
        'modulos': _modulos_ui_data(),
        'recentes': recentes,
        'aviso_migracao': aviso_migracao,
    }
    return render(request, 'workflow/home.html', context)


def modulos(request):
    return render(request, 'workflow/modulos.html', {'modulos': _modulos_ui_data()})


def _contexto_modulo_documentos(request, card: dict) -> dict:
    busca = (request.GET.get('q') or '').strip()
    processos_qs = Processo.objects.select_related('secretaria', 'modalidade', 'status').order_by('-atualizado_em', '-id')

    if busca:
        filtros = (
            Q(numero_processo_sirel__icontains=busca)
            | Q(numero_processo_adm__icontains=busca)
            | Q(numero_edital__icontains=busca)
            | Q(objeto__icontains=busca)
        )
        processos_qs = processos_qs.filter(filtros)

    processos = list(processos_qs[:220])
    processo_ids = [p.id for p in processos]

    workflow_docs_map = {
        row['processo_id']: row['total']
        for row in DocumentoProcessoWorkflow.objects.filter(processo_id__in=processo_ids)
        .values('processo_id')
        .annotate(total=Count('id'))
    }

    anexos_docs_map = {}
    try:
        processo_anexo_model = apps.get_model('docs', 'ProcessoAnexo')
    except LookupError:
        processo_anexo_model = None
    if processo_anexo_model is not None:
        anexos_docs_map = {
            row['processo_id']: row['total']
            for row in processo_anexo_model.objects.filter(processo_id__in=processo_ids)
            .values('processo_id')
            .annotate(total=Count('id'))
        }

    workflows_map = {
        wf.processo_id: wf
        for wf in ProcessoWorkflow.objects.filter(processo_id__in=processo_ids).select_related('processo')
    }

    rows = []
    for processo in processos:
        qtd_workflow = int(workflow_docs_map.get(processo.id, 0))
        qtd_docs = int(anexos_docs_map.get(processo.id, 0))
        rows.append(
            {
                'processo': processo,
                'workflow': workflows_map.get(processo.id),
                'qtd_docs_workflow': qtd_workflow,
                'qtd_docs_app_docs': qtd_docs,
                'qtd_docs_total': qtd_workflow + qtd_docs,
                'gerar_pdf_url': reverse('workflow:documentos_gerar_processo_pdf', args=[processo.id]),
                'gerar_etcm_url': reverse('workflow:documentos_gerar_processo_etcm', args=[processo.id]),
                'resumo_url': reverse('workflow:processo_resumo', args=[processo.id]),
            }
        )

    total_docs = sum(r['qtd_docs_total'] for r in rows)
    return {
        'card': card,
        'filtro_q': busca,
        'rows': rows,
        'total_processos': len(rows),
        'total_docs': total_docs,
    }


def documentos_gerar_processo_pdf(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    try:
        pdf_raw, documentos = _gerar_processo_consolidado_pdf(processo)
        pdf_data, meta = _normalizar_pdf_para_padrao_documental(
            pdf_raw,
            aplicar_ocr=True,
            pagina_max_kb=500,
            footer_left=f'Processo {processo.numero_processo_principal}',
            footer_center='Processo padronizado',
        )
    except Exception as exc:
        messages.error(request, f'Falha ao gerar processo integral em PDF: {exc}')
        back = f"{reverse('workflow:modulo_detail', args=['DOCUMENTOS'])}?q={processo.numero_processo_principal}"
        return redirect(back)

    filename = f"processo_integral_{_slug_ascii(processo.numero_processo_principal)}_{processo.id}.pdf"
    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response['X-Total-Documentos'] = str(len(documentos))
    response['X-Total-Paginas'] = str(meta.get('total_paginas', 0))
    response['X-OCR-Paginas'] = str(meta.get('paginas_ocr', 0))
    response['X-OCR-Habilitado'] = '1' if meta.get('ocr_habilitado') else '0'
    return response


def documentos_gerar_processo_etcm(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    try:
        payload, filename, content_type, meta = _gerar_processo_etcm_zip_ou_pdf(processo)
    except Exception as exc:
        messages.error(request, f'Falha ao gerar pacote e-TCM: {exc}')
        back = f"{reverse('workflow:modulo_detail', args=['DOCUMENTOS'])}?q={processo.numero_processo_principal}"
        return redirect(back)

    response = HttpResponse(payload, content_type=content_type)
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    response['X-Total-Documentos'] = str(meta.get('documentos', 0))
    response['X-Total-Paginas'] = str(meta.get('total_paginas', 0))
    response['X-OCR-Paginas'] = str(meta.get('paginas_ocr', 0))
    response['X-Partes'] = str(meta.get('partes', 1))
    return response


def modulo_detail(request, modulo: str):
    modulo = modulo.upper()
    card = next((m for m in MODULOS if m['slug'] == modulo), None)
    if not card:
        raise Http404('Modulo invalido.')
    if modulo == 'DASHBOARDS':
        return redirect('workflow:dashboards_geral')
    if modulo == 'FREQUENCIA':
        return redirect('workflow:frequencia')
    if modulo == 'CADASTROS':
        return redirect('workflow:cadastros_dashboard')
    if modulo == 'DOCUMENTOS':
        ctx = _contexto_modulo_documentos(request, card)
        return render(request, 'workflow/modulo_documentos.html', ctx)

    workflows_qs = (
        ProcessoWorkflow.objects
        .filter(modulo_atual=modulo)
        .select_related('processo', 'processo__secretaria', 'processo__status')
    )

    busca = (request.GET.get('q') or '').strip()
    campo = (request.GET.get('campo') or 'todos').strip().lower()
    situacao = (request.GET.get('situacao') or '').strip().upper()
    situacoes_validas = {valor for valor, _ in SituacaoWorkflow.choices}
    if situacao and situacao not in situacoes_validas:
        situacao = ''

    if busca:
        filtros = Q()
        if campo == 'numero':
            numero = busca
            ano = ''
            if '/' in busca:
                numero, ano = (busca.split('/', 1) + [''])[:2]
            numero = numero.strip()
            ano = ano.strip()
            if numero:
                filtros |= (
                    Q(processo__numero_processo_sirel__icontains=numero)
                    | Q(processo__numero_processo_adm__icontains=numero)
                    | Q(processo__numero_edital__icontains=numero)
                )
            if ano.isdigit():
                filtros |= Q(processo__ano_referencia=int(ano))
        elif campo == 'objeto':
            filtros = Q(processo__objeto__icontains=busca)
        elif campo == 'nome':
            filtros = (
                Q(processo__secretaria__nome__icontains=busca)
                | Q(processo__secretaria__sigla__icontains=busca)
                | Q(processo__status__nome__icontains=busca)
            )
        else:
            filtros = (
                Q(processo__numero_processo_sirel__icontains=busca)
                | Q(processo__numero_processo_adm__icontains=busca)
                | Q(processo__numero_edital__icontains=busca)
                | Q(processo__objeto__icontains=busca)
                | Q(processo__secretaria__nome__icontains=busca)
                | Q(processo__secretaria__sigla__icontains=busca)
                | Q(processo__status__nome__icontains=busca)
            )
        workflows_qs = workflows_qs.filter(filtros)

    if situacao:
        workflows_qs = workflows_qs.filter(situacao=situacao)

    workflows = list(workflows_qs.order_by('-atualizado_em')[:300])
    processo_ids = [wf.processo_id for wf in workflows]
    itens_map = {
        row['processo_id']: row['total']
        for row in ProcessoItem.objects.filter(processo_id__in=processo_ids).values('processo_id').annotate(total=Count('id'))
    }
    lotes_map = {
        row['processo_id']: row['total']
        for row in Lote.objects.filter(processo_id__in=processo_ids).values('processo_id').annotate(total=Count('id'))
    }
    fornecedores_map = {
        row['processo_id']: row['total']
        for row in ProcessoItem.objects.filter(processo_id__in=processo_ids, fornecedor_homologado__isnull=False)
        .values('processo_id')
        .annotate(total=Count('fornecedor_homologado_id', distinct=True))
    }
    fornecedores_result_map = {
        row['processo_id']: row['total']
        for row in ProcessoItemResultado.objects.filter(
            processo_id__in=processo_ids,
            ativo=True,
            fornecedor__isnull=False,
        )
        .values('processo_id')
        .annotate(total=Count('fornecedor_id', distinct=True))
    }

    por_secretaria = (
        Processo.objects
        .filter(id__in=processo_ids)
        .values('secretaria__sigla')
        .annotate(total=Count('id'))
        .order_by('-total', 'secretaria__sigla')[:8]
    )
    por_etapa = (
        workflows_qs
        .values('etapa_atual')
        .annotate(total=Count('id'))
        .order_by('-total', 'etapa_atual')[:8]
    )
    por_situacao = (
        workflows_qs
        .values('situacao')
        .annotate(total=Count('id'))
        .order_by('-total', 'situacao')
    )

    total_itens = sum(itens_map.values()) if itens_map else 0
    total_lotes = sum(lotes_map.values()) if lotes_map else 0
    total_fornecedores = sum(
        max(fornecedores_map.get(pid, 0), fornecedores_result_map.get(pid, 0))
        for pid in processo_ids
    ) if processo_ids else 0

    workflows_data = []
    for wf in workflows:
        if modulo == ModuloSistema.PLANEJAMENTO:
            acao_url = reverse('workflow:planejamento_detail', args=[wf.processo_id])
            acao_label = 'Abrir planejamento'
        elif modulo == ModuloSistema.COMPRAS:
            acao_url = reverse('workflow:compras_detail', args=[wf.processo_id])
            acao_label = 'Abrir compras'
        elif modulo == ModuloSistema.LICITACAO:
            acao_url = reverse('workflow:licitacao_detail', args=[wf.processo_id])
            acao_label = 'Abrir licitação'
        else:
            acao_url = reverse('workflow:processo_resumo', args=[wf.processo_id])
            acao_label = 'Abrir resumo'
        workflows_data.append({
            'wf': wf,
            'total_itens': itens_map.get(wf.processo_id, 0),
            'total_lotes': lotes_map.get(wf.processo_id, 0),
            'total_fornecedores': max(
                fornecedores_map.get(wf.processo_id, 0),
                fornecedores_result_map.get(wf.processo_id, 0),
            ),
            'acao_url': acao_url,
            'acao_label': acao_label,
        })
    return render(
        request,
        'workflow/modulo_detail.html',
        {
            'card': card,
            'workflows_data': workflows_data,
            'filtro_q': busca,
            'filtro_campo': campo,
            'filtro_situacao': situacao,
            'situacoes_workflow': SituacaoWorkflow.choices,
            'por_secretaria': por_secretaria,
            'por_etapa': por_etapa,
            'por_situacao': por_situacao,
            'total_processos': len(workflows),
            'total_itens': total_itens,
            'total_lotes': total_lotes,
            'total_fornecedores': total_fornecedores,
        },
    )


def licitacao_novo_externo(request):
    if not Modalidade.objects.exists():
        _get_modalidade('Dispensa de licitação - Lei 14.133/2021')

    ano_preview = timezone.localdate().year
    pendencias_detectadas = []
    if request.method == 'POST':
        ano_raw = (request.POST.get('ano_referencia') or '').strip()
        if ano_raw.isdigit():
            ano_preview = int(ano_raw)
        form = LicitacaoProcessoExternoForm(request.POST)
        if form.is_valid():
            uploads = {}
            for item in LICITACAO_CHECKLIST:
                arquivo = request.FILES.get(f"arquivo_{item['codigo']}")
                if arquivo:
                    uploads[item['codigo']] = arquivo
            pendencias_detectadas = [
                item['label']
                for item in LICITACAO_CHECKLIST
                if item['codigo'] not in uploads
            ]
            confirmou = bool(form.cleaned_data.get('confirmar_pendencias'))
            if pendencias_detectadas and not confirmou:
                messages.warning(
                    request,
                    'Foram detectadas pendências documentais. Confirme o registro com pendências para continuar.',
                )
            else:
                modalidade = form.cleaned_data['modalidade']
                status = form.cleaned_data.get('status') or _get_status('EM LICITAÇÃO')
                with transaction.atomic():
                    processo = Processo.objects.create(
                        numero_processo_adm=(form.cleaned_data['numero_processo_externo'] or '').strip(),
                        numero_edital=(form.cleaned_data.get('numero_edital') or '').strip(),
                        ano_referencia=form.cleaned_data['ano_referencia'],
                        secretaria=form.cleaned_data.get('secretaria'),
                        modalidade=modalidade,
                        status=status,
                        objeto=form.cleaned_data['objeto'],
                        data_publicacao=form.cleaned_data.get('data_publicacao'),
                        valor_estimado=form.cleaned_data.get('valor_estimado') or Decimal('0'),
                    )
                    ProcessoWorkflow.objects.create(
                        processo=processo,
                        modulo_atual=ModuloSistema.LICITACAO,
                        etapa_atual=LICITACAO_ETAPA_FASE_INTERNA,
                        situacao=SituacaoWorkflow.EM_ANDAMENTO,
                    )
                    observacao = (
                        f'Registro externo criado diretamente no módulo Licitação. '
                        f'Documentos anexados no cadastro: {len(uploads)}.'
                    )
                    if pendencias_detectadas:
                        observacao += (
                            '\nPendências assumidas no registro inicial: '
                            + '; '.join(pendencias_detectadas)
                            + '.'
                        )
                    ProcessoMovimentacao.objects.create(
                        processo=processo,
                        modulo_origem='',
                        modulo_destino=ModuloSistema.LICITACAO,
                        descricao='Processo externo registrado diretamente no módulo de Licitação.',
                        observacao=observacao,
                    )
                    for codigo, arquivo in uploads.items():
                        _upsert_licitacao_documento(processo, codigo, arquivo)
                if pendencias_detectadas:
                    messages.warning(
                        request,
                        (
                            f'Processo {processo.numero_processo_principal} registrado com pendências '
                            f'({len(pendencias_detectadas)} documento(s) ausente(s)).'
                        ),
                    )
                else:
                    messages.success(
                        request,
                        f'Processo {processo.numero_processo_principal} registrado na Licitação com checklist completo.',
                    )
                return redirect('workflow:licitacao_detail', processo.id)
    else:
        form = LicitacaoProcessoExternoForm()

    checklist_fases = {}
    for item in LICITACAO_CHECKLIST:
        checklist_fases.setdefault(item['fase'], []).append(item)
    proximo_numero_sirel = Processo.gerar_numero_processo_sirel(ano_preview)
    return render(
        request,
        'workflow/licitacao_novo_externo.html',
        {
            'form': form,
            'checklist_fases': checklist_fases,
            'pendencias_detectadas': pendencias_detectadas,
            'ano_preview': ano_preview,
            'proximo_numero_sirel': proximo_numero_sirel,
        },
    )


def licitacao_detail(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    checklist_ctx = _build_licitacao_checklist(processo)
    if workflow.modulo_atual == ModuloSistema.LICITACAO and not (workflow.etapa_atual or '').strip():
        workflow.etapa_atual = LICITACAO_ETAPA_FASE_INTERNA
        workflow.save(update_fields=['etapa_atual', 'atualizado_em'])
    pendencias = checklist_ctx['pendencias']
    form_upload = LicitacaoDocumentoUploadForm(
        doc_choices=_licitacao_documento_choices(),
        initial={'doc_codigo': LICITACAO_CHECKLIST[0]['codigo']},
    )
    docs_extras = (
        DocumentoProcessoWorkflow.objects
        .filter(processo=processo, modulo=ModuloSistema.LICITACAO)
        .exclude(tipo_documento__startswith=LICITACAO_DOC_PREFIX)
        .order_by('-criado_em', '-id')[:20]
    )
    return render(
        request,
        'workflow/licitacao_detail.html',
        {
            'processo': processo,
            'workflow': workflow,
            'checklist': checklist_ctx['itens'],
            'checklist_fases': checklist_ctx['fases'],
            'pendencias': pendencias,
            'total_docs': checklist_ctx['total'],
            'total_docs_ok': checklist_ctx['total_ok'],
            'percentual_completude': checklist_ctx['percentual'],
            'form_upload': form_upload,
            'docs_extras': docs_extras,
            'pode_encaminhar_procuradoria': len(pendencias) == 0,
        },
    )


def licitacao_upload_documento(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:licitacao_detail', processo.id)
    form = LicitacaoDocumentoUploadForm(
        request.POST,
        request.FILES,
        doc_choices=_licitacao_documento_choices(),
    )
    if form.is_valid():
        codigo = form.cleaned_data['doc_codigo']
        if codigo not in LICITACAO_CHECKLIST_MAP:
            return _ajax_or_redirect(
                request,
                ok=False,
                message='Documento de checklist inválido.',
                redirect_name='workflow:licitacao_detail',
                redirect_args=[processo.id],
            )
        _upsert_licitacao_documento(processo, codigo, form.cleaned_data['arquivo'])
        workflow = _ensure_workflow(processo)
        if workflow.modulo_atual == ModuloSistema.LICITACAO:
            workflow.etapa_atual = LICITACAO_ETAPA_FASE_INTERNA
            if workflow.situacao == SituacaoWorkflow.RASCUNHO:
                workflow.situacao = SituacaoWorkflow.EM_ANDAMENTO
                workflow.save(update_fields=['etapa_atual', 'situacao', 'atualizado_em'])
            else:
                workflow.save(update_fields=['etapa_atual', 'atualizado_em'])
        msg = f"Documento '{LICITACAO_CHECKLIST_MAP[codigo]['label']}' anexado/atualizado com sucesso."
        return _ajax_or_redirect(
            request,
            ok=True,
            message=msg,
            redirect_name='workflow:licitacao_detail',
            redirect_args=[processo.id],
        )
    return _ajax_or_redirect(
        request,
        ok=False,
        message=f'Não foi possível anexar o documento. {_form_errors_text(form)}',
        redirect_name='workflow:licitacao_detail',
        redirect_args=[processo.id],
    )


def licitacao_ci_documento(request, processo_id: int, codigo: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    codigo = (codigo or '').strip().lower()
    ci_ctx = _licitacao_ci_context(
        processo,
        codigo,
        form_data=request.POST if request.method == 'POST' else None,
        usuario=request.user,
    )
    if not ci_ctx:
        raise Http404('Documento não disponível para este código.')

    if request.method == 'POST':
        action = (request.POST.get('action') or '').strip().lower()
        if action == 'preview_pdf':
            return _export_licitacao_documento_pdf(processo=processo, doc_ctx=ci_ctx)
        if action == 'download_docx':
            return _export_licitacao_documento_docx(processo=processo, doc_ctx=ci_ctx)

    ci_ctx['voltar_url'] = reverse('workflow:licitacao_detail', args=[processo.id])
    return render(request, 'workflow/licitacao_ci_documento.html', ci_ctx)


def licitacao_relatorio_pendencias(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    checklist_ctx = _build_licitacao_checklist(processo)
    rows_atuais = []
    for row in checklist_ctx['itens']:
        rows_atuais.append(
            {
                'fase': row['fase'],
                'documento': row['label'],
                'codigo': row['codigo'],
                'status': 'CONCLUÍDO' if row['presente'] else 'PENDENTE',
                'presente': row['presente'],
                'data': row['data_envio'],
                'tipo': 'ATUAL',
            }
        )

    rows_futuras = []
    for fase in LICITACAO_FASES_FUTURAS:
        for idx, doc_nome in enumerate(fase['documentos'], start=1):
            rows_futuras.append(
                {
                    'fase': fase['fase'],
                    'documento': doc_nome,
                    'codigo': f"futuro_{idx:02d}",
                    'status': 'PLANEJADO',
                    'presente': False,
                    'data': None,
                    'tipo': 'FUTURO',
                }
            )

    todos = rows_atuais + rows_futuras
    agrupado = {}
    for row in todos:
        agrupado.setdefault(row['fase'], []).append(row)

    total_concluidos = sum(1 for row in rows_atuais if row['presente'])
    total_pendentes = sum(1 for row in rows_atuais if not row['presente'])
    total_planejados = len(rows_futuras)

    docs_licitacao = (
        DocumentoProcessoWorkflow.objects
        .filter(processo=processo, modulo=ModuloSistema.LICITACAO)
        .order_by('ordem_cronologica', 'id')
    )
    return render(
        request,
        'workflow/licitacao_relatorio_pendencias.html',
        {
            'processo': processo,
            'workflow': workflow,
            'agrupado': agrupado,
            'gerado_em': timezone.localtime(),
            'total_concluidos': total_concluidos,
            'total_pendentes': total_pendentes,
            'total_planejados': total_planejados,
            'checklist_ctx': checklist_ctx,
            'docs_licitacao': docs_licitacao,
        },
    )


def itens_rastreamento(request):
    filtros = {
        'q': (request.GET.get('q') or '').strip(),
        'status_item': (request.GET.get('status_item') or '').strip(),
        'situacao_workflow': (request.GET.get('situacao_workflow') or '').strip().upper(),
        'homologado': (request.GET.get('homologado') or 'todos').strip().lower(),
    }
    status_validos = {valor for valor, _ in ProcessoItem.StatusConsolidado.choices}
    if filtros['status_item'] and filtros['status_item'] not in status_validos:
        filtros['status_item'] = ''
    situacoes_validas = {valor for valor, _ in SituacaoWorkflow.choices}
    if filtros['situacao_workflow'] and filtros['situacao_workflow'] not in situacoes_validas:
        filtros['situacao_workflow'] = ''

    latest_result_pncp = (
        ProcessoItemResultado.objects
        .filter(
            processo_item_id=OuterRef('pk'),
            origem=ProcessoItemResultado.Origem.PNCP,
            ativo=True,
        )
        .order_by('-data_resultado', '-atualizado_em', '-id')
    )
    qs = (
        ProcessoItem.objects
        .select_related('processo', 'processo__workflow', 'fornecedor_homologado', 'item_catalogo')
        .annotate(
            situacao_item_pncp=Subquery(latest_result_pncp.values('situacao_texto')[:1]),
            situacao_resultado_pncp=Subquery(latest_result_pncp.values('status_resultado')[:1]),
            data_resultado_homologacao=Subquery(latest_result_pncp.values('data_resultado')[:1]),
        )
        .all()
    )

    if filtros['q']:
        termo = filtros['q']
        filtros_q = (
            Q(descricao_snapshot__icontains=termo)
            | Q(unidade_snapshot__icontains=termo)
            | Q(processo__numero_processo_sirel__icontains=termo)
            | Q(processo__numero_processo_adm__icontains=termo)
            | Q(processo__numero_edital__icontains=termo)
            | Q(processo__objeto__icontains=termo)
            | Q(fornecedor_homologado__razao_social__icontains=termo)
            | Q(fornecedor_homologado__cnpj__icontains=termo)
            | Q(item_catalogo__descricao_padrao__icontains=termo)
        )
        if termo.isdigit():
            filtros_q |= Q(numero_item=int(termo))
        qs = qs.filter(filtros_q)

    if filtros['status_item']:
        qs = qs.filter(status_consolidado=filtros['status_item'])
    if filtros['situacao_workflow']:
        qs = qs.filter(processo__workflow__situacao=filtros['situacao_workflow'])
    if filtros['homologado'] == 'sim':
        qs = qs.filter(status_consolidado=ProcessoItem.StatusConsolidado.HOMOLOGADO)
    elif filtros['homologado'] == 'nao':
        qs = qs.exclude(status_consolidado=ProcessoItem.StatusConsolidado.HOMOLOGADO)

    itens = list(
        qs.order_by('-pncp_ultima_atualizacao', '-processo__atualizado_em', '-processo_id', 'numero_item')[:1200]
    )

    total_estimado = Decimal('0.00')
    total_homologado = Decimal('0.00')
    homologados = 0
    fracassados = 0
    cancelados = 0
    conflitos_lote = 0
    processos_ids = set()
    for item in itens:
        processos_ids.add(item.processo_id)
        est = Decimal(str(item.valor_referencia_total or 0))
        if est <= 0:
            est = (Decimal(str(item.valor_referencia_unitario or 0)) * Decimal(str(item.quantidade or 0))).quantize(Decimal('0.01'))
        hom = Decimal(str(item.valor_homologado_total or 0))
        if hom <= 0 and item.status_consolidado == ProcessoItem.StatusConsolidado.HOMOLOGADO:
            hom = (Decimal(str(item.valor_homologado_unitario or 0)) * Decimal(str(item.quantidade or 0))).quantize(Decimal('0.01'))
        total_estimado += est
        total_homologado += hom
        if item.status_consolidado == ProcessoItem.StatusConsolidado.HOMOLOGADO:
            homologados += 1
        elif item.status_consolidado == ProcessoItem.StatusConsolidado.FRACASSADO:
            fracassados += 1
        elif item.status_consolidado == ProcessoItem.StatusConsolidado.CANCELADO:
            cancelados += 1
        if item.conflito_lote:
            conflitos_lote += 1

    ranking_itens = (
        qs.values('descricao_snapshot')
        .annotate(total=Count('id'), processos=Count('processo_id', distinct=True))
        .order_by('-processos', '-total', 'descricao_snapshot')[:15]
    )

    return render(
        request,
        'workflow/itens_rastreamento.html',
        {
            'itens': itens,
            'ranking_itens': ranking_itens,
            'filtros': filtros,
            'status_choices': ProcessoItem.StatusConsolidado.choices,
            'situacoes_workflow': SituacaoWorkflow.choices,
            'total_itens': len(itens),
            'total_processos': len(processos_ids),
            'homologados': homologados,
            'fracassados': fracassados,
            'cancelados': cancelados,
            'conflitos_lote': conflitos_lote,
            'total_estimado': total_estimado.quantize(Decimal('0.01')),
            'total_homologado': total_homologado.quantize(Decimal('0.01')),
        },
    )


def _format_horas_minutos(minutos):
    total = int(minutos or 0)
    sinal = '-' if total < 0 else ''
    total = abs(total)
    horas = total // 60
    mins = total % 60
    return f'{sinal}{horas:02d}:{mins:02d}'


def _parse_mes_referencia(raw):
    texto = (raw or '').strip()
    if len(texto) >= 7 and texto[4] == '-':
        try:
            ano = int(texto[:4])
            mes = int(texto[5:7])
            if 1 <= mes <= 12:
                return ano, mes
        except Exception:
            pass
    hoje = timezone.localdate()
    return hoje.year, hoje.month
_DASHBOARD_EVENT_SPECS = (
    ('data_publicacao', 'Publicacao', 'date', 'pub'),
    ('inicio_recolhimento_propostas', 'Inicio de propostas', 'datetime', 'start'),
    ('fim_impugnacao_esclarecimentos', 'Impugnacao/esclarecimentos', 'datetime', 'warn'),
    ('fim_recolhimento_propostas', 'Fim de propostas', 'datetime', 'deadline'),
    ('data_hora_abertura', 'Abertura / sessao', 'datetime', 'open'),
)


def _format_date_br(valor) -> str:
    if not valor:
        return '-'
    if isinstance(valor, datetime):
        valor = timezone.localtime(valor).date() if timezone.is_aware(valor) else valor.date()
    if isinstance(valor, date):
        return valor.strftime('%d/%m/%Y')
    try:
        return datetime.fromisoformat(str(valor)).strftime('%d/%m/%Y')
    except Exception:
        return str(valor)


def _dashboard_eventos_processo(processo: Processo, *, hoje: date) -> list[dict]:
    eventos = []
    for ordem, (field, label, tipo, tone) in enumerate(_DASHBOARD_EVENT_SPECS, start=1):
        valor = getattr(processo, field, None)
        if not valor:
            continue
        if tipo == 'datetime':
            if isinstance(valor, datetime):
                valor_dt = timezone.localtime(valor) if timezone.is_aware(valor) else valor
            else:
                try:
                    valor_dt = datetime.fromisoformat(str(valor))
                except Exception:
                    continue
            data_ref = valor_dt.date()
            hora_ref = valor_dt.time().replace(second=0, microsecond=0)
            exibicao = _format_datetime_br(valor_dt)
        else:
            data_ref = valor if isinstance(valor, date) else None
            if data_ref is None:
                try:
                    data_ref = datetime.fromisoformat(str(valor)).date()
                except Exception:
                    continue
            hora_ref = None
            exibicao = _format_date_br(data_ref)

        if data_ref < hoje:
            situacao = 'atrasado'
        elif data_ref == hoje:
            situacao = 'hoje'
        else:
            situacao = 'proximo'

        eventos.append(
            {
                'field': field,
                'label': label,
                'tipo': tipo,
                'tone': tone,
                'data': data_ref,
                'hora': hora_ref,
                'exibicao': exibicao,
                'situacao': situacao,
                'ordem': ordem,
            }
        )
    eventos.sort(key=lambda e: (e['data'], e['hora'] or time(0, 0), e['ordem']))
    return eventos


def _dashboard_calendar_weeks(ano: int, mes: int, eventos_por_dia: dict[date, list[dict]], hoje: date) -> list[list[dict]]:
    semanas = []
    cal = calendar.Calendar(firstweekday=0)
    for semana in cal.monthdatescalendar(ano, mes):
        linha = []
        for dia_ref in semana:
            eventos = list(eventos_por_dia.get(dia_ref, []))
            linha.append(
                {
                    'date': dia_ref,
                    'day': dia_ref.day,
                    'in_month': dia_ref.month == mes,
                    'is_today': dia_ref == hoje,
                    'events': eventos[:3],
                    'count': len(eventos),
                }
            )
        semanas.append(linha)
    return semanas


def dashboards_geral(request):
    filtros = {
        'q': (request.GET.get('q') or '').strip(),
        'processo_id': (request.GET.get('processo_id') or '').strip(),
        'status_id': (request.GET.get('status_id') or '').strip(),
        'modalidade_id': (request.GET.get('modalidade_id') or '').strip(),
        'secretaria_id': (request.GET.get('secretaria_id') or '').strip(),
        'condutor_id': (request.GET.get('condutor_id') or '').strip(),
        'modulo': (request.GET.get('modulo') or '').strip().upper(),
        'situacao': (request.GET.get('situacao') or '').strip().upper(),
        'homologado': (request.GET.get('homologado') or 'todos').strip().lower(),
        'parados_dias': (request.GET.get('parados_dias') or '15').strip(),
        'somente_parados': (request.GET.get('somente_parados') or '').strip(),
        'ordem': (request.GET.get('ordem') or 'atualizado_desc').strip(),
        'agenda_mes': (request.GET.get('agenda_mes') or '').strip(),
        'agenda_ano': (request.GET.get('agenda_ano') or '').strip(),
    }

    try:
        parados_dias = max(1, int(filtros['parados_dias']))
    except Exception:
        parados_dias = 15
    somente_parados = filtros['somente_parados'] in {'1', 'true', 'True', 'on'}
    hoje = timezone.localdate()
    try:
        agenda_mes = min(12, max(1, int(filtros['agenda_mes'] or hoje.month)))
    except Exception:
        agenda_mes = hoje.month
    try:
        agenda_ano = min(2100, max(2000, int(filtros['agenda_ano'] or hoje.year)))
    except Exception:
        agenda_ano = hoje.year

    processos_qs = Processo.objects.select_related(
        'secretaria', 'status', 'modalidade', 'condutor_processo', 'autoridade_competente'
    ).all()

    if filtros['q']:
        termo = filtros['q']
        processos_qs = processos_qs.filter(
            Q(numero_processo_sirel__icontains=termo)
            | Q(numero_processo_adm__icontains=termo)
            | Q(numero_edital__icontains=termo)
            | Q(objeto__icontains=termo)
            | Q(protocolo__icontains=termo)
        )
    if filtros['processo_id'].isdigit():
        processos_qs = processos_qs.filter(pk=int(filtros['processo_id']))
    if filtros['status_id'].isdigit():
        processos_qs = processos_qs.filter(status_id=int(filtros['status_id']))
    if filtros['modalidade_id'].isdigit():
        processos_qs = processos_qs.filter(modalidade_id=int(filtros['modalidade_id']))
    if filtros['secretaria_id'].isdigit():
        processos_qs = processos_qs.filter(secretaria_id=int(filtros['secretaria_id']))
    if filtros['condutor_id'].isdigit():
        processos_qs = processos_qs.filter(condutor_processo_id=int(filtros['condutor_id']))
    if filtros['modulo']:
        processos_qs = processos_qs.filter(workflow__modulo_atual=filtros['modulo'])
    if filtros['situacao']:
        processos_qs = processos_qs.filter(workflow__situacao=filtros['situacao'])
    if filtros['homologado'] == 'sim':
        processos_qs = processos_qs.filter(
            Q(workflow__homologado=True) | Q(status__nome__icontains='HOMOLOG')
        )
    elif filtros['homologado'] == 'nao':
        processos_qs = processos_qs.exclude(
            Q(workflow__homologado=True) | Q(status__nome__icontains='HOMOLOG')
        )

    processos = list(processos_qs.order_by('-atualizado_em')[:1200])
    processo_ids = [p.id for p in processos]
    workflows_map = {
        wf.processo_id: wf
        for wf in ProcessoWorkflow.objects.filter(processo_id__in=processo_ids).select_related('processo')
    }
    lotes_map = {
        row['processo_id']: row['total']
        for row in Lote.objects.filter(processo_id__in=processo_ids).values('processo_id').annotate(total=Count('id'))
    }
    fornecedores_map = {
        row['processo_id']: row['total']
        for row in (
            ProcessoItem.objects.filter(processo_id__in=processo_ids, fornecedor_homologado__isnull=False)
            .values('processo_id')
            .annotate(total=Count('fornecedor_homologado_id', distinct=True))
        )
    }
    itens_core_map = {
        row['processo_id']: row
        for row in (
            ProcessoItem.objects.filter(processo_id__in=processo_ids)
            .values('processo_id')
            .annotate(
                total_itens=Count('id'),
                quantidade_total=Coalesce(
                    Sum('quantidade'),
                    Value(Decimal('0.000')),
                    output_field=DecimalField(max_digits=18, decimal_places=3),
                ),
                valor_total=Coalesce(
                    Sum('valor_referencia_total'),
                    Value(Decimal('0.00')),
                    output_field=DecimalField(max_digits=18, decimal_places=2),
                ),
                homologados=Count('id', filter=Q(status_consolidado=ProcessoItem.StatusConsolidado.HOMOLOGADO)),
            )
        )
    }
    itens_dfd_map = {
        row['dfd__processo_id']: row
        for row in (
            DFDItem.objects.filter(dfd__processo_id__in=processo_ids)
            .values('dfd__processo_id')
            .annotate(
                total_itens=Count('id'),
                quantidade_total=Coalesce(
                    Sum('quantidade'),
                    Value(Decimal('0.000')),
                    output_field=DecimalField(max_digits=18, decimal_places=3),
                ),
            )
        )
    }
    movimentos_map = {
        row['processo_id']: row['ultima']
        for row in ProcessoMovimentacao.objects.filter(processo_id__in=processo_ids).values('processo_id').annotate(ultima=Max('criado_em'))
    }
    rows = []
    for p in processos:
        wf = workflows_map.get(p.id)
        eventos_agenda = _dashboard_eventos_processo(p, hoje=hoje)
        proximos = [ev for ev in eventos_agenda if ev['data'] >= hoje]
        proximo_evento = proximos[0] if proximos else (eventos_agenda[-1] if eventos_agenda else None)
        row_core = itens_core_map.get(p.id, {})
        row_dfd = itens_dfd_map.get(p.id, {})
        total_itens = int(row_core.get('total_itens') or 0) or int(row_dfd.get('total_itens') or 0)
        quantidade_total = row_core.get('quantidade_total') if row_core.get('total_itens') else row_dfd.get('quantidade_total', Decimal('0.000'))
        quantidade_total = Decimal(str(quantidade_total or 0))
        valor_total_itens = Decimal(str(row_core.get('valor_total') or 0))
        homologados_itens = int(row_core.get('homologados') or 0)
        status_nome = (getattr(p.status, 'nome', '') or '').strip()
        is_homologado = bool(
            (wf and wf.homologado)
            or ('HOMOLOG' in status_nome.upper())
            or homologados_itens > 0
        )
        referencia_atualizacao = p.atualizado_em
        if wf and wf.atualizado_em and wf.atualizado_em > referencia_atualizacao:
            referencia_atualizacao = wf.atualizado_em
        ultima_mov = movimentos_map.get(p.id)
        if ultima_mov and ultima_mov > referencia_atualizacao:
            referencia_atualizacao = ultima_mov
        dias_sem_atualizacao = (hoje - referencia_atualizacao.date()).days if referencia_atualizacao else 0
        is_concluido = bool(
            (wf and wf.situacao == SituacaoWorkflow.CONCLUIDO)
            or ('CONCL' in status_nome.upper())
            or is_homologado
        )
        is_parado = bool(dias_sem_atualizacao >= parados_dias and not is_concluido)
        row = {
            'processo': p,
            'workflow': wf,
            'condutor': getattr(p, 'condutor_processo', None),
            'modulo_atual': wf.modulo_atual if wf else '',
            'situacao': wf.situacao if wf else '',
            'total_itens': total_itens,
            'total_lotes': lotes_map.get(p.id, 0),
            'total_fornecedores': fornecedores_map.get(p.id, 0),
            'quantidade_total': quantidade_total,
            'valor_total_itens': valor_total_itens,
            'homologados_itens': homologados_itens,
            'is_homologado': is_homologado,
            'is_parado': is_parado,
            'dias_sem_atualizacao': dias_sem_atualizacao,
            'ultima_atualizacao': referencia_atualizacao,
            'agenda_eventos': eventos_agenda,
            'proximo_evento': proximo_evento,
            'sem_condutor': getattr(p, 'condutor_processo_id', None) is None,
            'sem_data_critica': len(eventos_agenda) == 0,
        }
        rows.append(row)

    if somente_parados:
        rows = [r for r in rows if r['is_parado']]

    ordem = filtros['ordem']
    if ordem == 'valor_desc':
        rows.sort(key=lambda r: Decimal(str(r['processo'].valor_estimado or 0)), reverse=True)
    elif ordem == 'valor_homologado_desc':
        rows.sort(key=lambda r: Decimal(str(r['processo'].valor_homologado or 0)), reverse=True)
    elif ordem == 'itens_desc':
        rows.sort(key=lambda r: r['total_itens'], reverse=True)
    elif ordem == 'parado_desc':
        rows.sort(key=lambda r: r['dias_sem_atualizacao'], reverse=True)
    else:
        rows.sort(key=lambda r: r['ultima_atualizacao'] or timezone.now(), reverse=True)

    total_processos = len(rows)
    total_homologados = sum(1 for r in rows if r['is_homologado'])
    total_parados = sum(1 for r in rows if r['is_parado'])
    total_itens = sum(int(r['total_itens'] or 0) for r in rows)
    total_qtd = sum(Decimal(str(r['quantidade_total'] or 0)) for r in rows)
    total_valor_estimado = sum(Decimal(str(r['processo'].valor_estimado or 0)) for r in rows)
    total_valor_homologado = sum(Decimal(str(r['processo'].valor_homologado or 0)) for r in rows)
    total_sem_condutor = sum(1 for r in rows if r['sem_condutor'] and not r['is_homologado'])
    total_sem_agenda = sum(1 for r in rows if r['sem_data_critica'] and not r['is_homologado'])

    agenda_eventos_hoje = []
    agenda_proximos_eventos = []
    agenda_atrasados = []
    eventos_por_dia = {}
    ranking_condutores_map = {}
    for r in rows:
        processo = r['processo']
        condutor = r['condutor']
        if condutor:
            chave = condutor.nome
            item = ranking_condutores_map.setdefault(
                chave,
                {
                    'nome': condutor.nome,
                    'cargo': condutor.cargo or '',
                    'processos': 0,
                    'aberturas': 0,
                    'eventos_hoje': 0,
                }
            )
            item['processos'] += 1
            if processo.data_hora_abertura:
                item['aberturas'] += 1
        for evento in r['agenda_eventos']:
            payload = {
                'processo': processo,
                'row': r,
                'label': evento['label'],
                'tone': evento['tone'],
                'data': evento['data'],
                'hora': evento['hora'],
                'exibicao': evento['exibicao'],
                'situacao': evento['situacao'],
                'condutor': condutor,
            }
            eventos_por_dia.setdefault(evento['data'], []).append(payload)
            if evento['situacao'] == 'hoje':
                agenda_eventos_hoje.append(payload)
                if condutor:
                    ranking_condutores_map[condutor.nome]['eventos_hoje'] += 1
            elif evento['situacao'] == 'proximo':
                agenda_proximos_eventos.append(payload)
            elif evento['situacao'] == 'atrasado' and evento['field'] != 'data_publicacao' and not r['is_homologado']:
                agenda_atrasados.append(payload)

    for lista in (agenda_eventos_hoje, agenda_proximos_eventos, agenda_atrasados):
        lista.sort(key=lambda e: (e['data'], e['hora'] or time(0, 0), e['processo'].numero_processo_principal))
    for data_ref, lista in eventos_por_dia.items():
        eventos_por_dia[data_ref] = sorted(
            lista,
            key=lambda e: (e['hora'] or time(0, 0), e['label'], e['processo'].numero_processo_principal),
        )

    ranking_condutores = sorted(
        ranking_condutores_map.values(),
        key=lambda x: (x['processos'], x['eventos_hoje'], x['aberturas'], x['nome']),
        reverse=True,
    )[:15]
    processos_sem_condutor = [r for r in rows if r['sem_condutor'] and not r['is_homologado']][:20]
    processos_sem_agenda = [r for r in rows if r['sem_data_critica'] and not r['is_homologado']][:20]
    agenda_calendar_weeks = _dashboard_calendar_weeks(agenda_ano, agenda_mes, eventos_por_dia, hoje)

    ranking_secretarias_map = {}
    ranking_status_map = {}
    ranking_modulo_map = {}
    for r in rows:
        sec = r['processo'].secretaria
        sec_key = sec.sigla if sec else 'N/A'
        sec_obj = ranking_secretarias_map.setdefault(sec_key, {
            'secretaria': sec,
            'sigla': sec_key,
            'processos': 0,
            'valor_estimado': Decimal('0.00'),
            'valor_homologado': Decimal('0.00'),
            'itens': 0,
            'quantidade': Decimal('0.000'),
        })
        sec_obj['processos'] += 1
        sec_obj['valor_estimado'] += Decimal(str(r['processo'].valor_estimado or 0))
        sec_obj['valor_homologado'] += Decimal(str(r['processo'].valor_homologado or 0))
        sec_obj['itens'] += int(r['total_itens'] or 0)
        sec_obj['quantidade'] += Decimal(str(r['quantidade_total'] or 0))

        status_key = (getattr(r['processo'].status, 'nome', '') or 'Sem status').strip()
        ranking_status_map[status_key] = ranking_status_map.get(status_key, 0) + 1

        modulo_key = r['workflow'].get_modulo_atual_display() if r['workflow'] else 'Sem workflow'
        ranking_modulo_map[modulo_key] = ranking_modulo_map.get(modulo_key, 0) + 1

    ranking_secretarias = sorted(
        ranking_secretarias_map.values(),
        key=lambda x: (x['processos'], x['valor_estimado']),
        reverse=True,
    )[:15]
    ranking_status = sorted(
        [{'status': k, 'total': v} for k, v in ranking_status_map.items()],
        key=lambda x: x['total'],
        reverse=True,
    )
    ranking_modulos = sorted(
        [{'modulo': k, 'total': v} for k, v in ranking_modulo_map.items()],
        key=lambda x: x['total'],
        reverse=True,
    )
    top_valores = sorted(rows, key=lambda r: Decimal(str(r['processo'].valor_estimado or 0)), reverse=True)[:15]
    processos_parados = [r for r in rows if r['is_parado']][:30]
    processos_homologados = [r for r in rows if r['is_homologado']][:30]

    ids_filtrados = [r['processo'].id for r in rows]
    top_itens_core = (
        ProcessoItem.objects.filter(processo_id__in=ids_filtrados)
        .values('descricao_snapshot', 'unidade_snapshot')
        .annotate(
            processos=Count('processo_id', distinct=True),
            quantidade=Coalesce(
                Sum('quantidade'),
                Value(Decimal('0.000')),
                output_field=DecimalField(max_digits=18, decimal_places=3),
            ),
            valor_total=Coalesce(
                Sum('valor_referencia_total'),
                Value(Decimal('0.00')),
                output_field=DecimalField(max_digits=18, decimal_places=2),
            ),
        )
        .order_by('-processos', '-quantidade', 'descricao_snapshot')[:20]
    )
    top_itens_dfd = (
        DFDItem.objects.filter(dfd__processo_id__in=ids_filtrados)
        .values('descricao', 'unidade')
        .annotate(
            processos=Count('dfd__processo_id', distinct=True),
            quantidade=Coalesce(
                Sum('quantidade'),
                Value(Decimal('0.000')),
                output_field=DecimalField(max_digits=18, decimal_places=3),
            ),
        )
        .order_by('-processos', '-quantidade', 'descricao')[:20]
    )

    context = {
        'rows': rows[:500],
        'rows_total': total_processos,
        'filtros': filtros,
        'parados_dias': parados_dias,
        'somente_parados': somente_parados,
        'total_homologados': total_homologados,
        'total_parados': total_parados,
        'total_itens': total_itens,
        'total_quantidade': total_qtd,
        'total_valor_estimado': total_valor_estimado,
        'total_valor_homologado': total_valor_homologado,
        'total_sem_condutor': total_sem_condutor,
        'total_sem_agenda': total_sem_agenda,
        'ranking_secretarias': ranking_secretarias,
        'ranking_condutores': ranking_condutores,
        'ranking_status': ranking_status,
        'ranking_modulos': ranking_modulos,
        'top_valores': top_valores,
        'processos_parados': processos_parados,
        'processos_homologados': processos_homologados,
        'agenda_eventos_hoje': agenda_eventos_hoje[:25],
        'agenda_proximos_eventos': agenda_proximos_eventos[:30],
        'agenda_atrasados': agenda_atrasados[:20],
        'processos_sem_condutor': processos_sem_condutor,
        'processos_sem_agenda': processos_sem_agenda,
        'agenda_calendar_weeks': agenda_calendar_weeks,
        'agenda_mes': agenda_mes,
        'agenda_ano': agenda_ano,
        'agenda_mes_nome': str(MESES_PT_BR[agenda_mes - 1]).title(),
        'agenda_mes_opts': [
            {'value': idx, 'label': str(MESES_PT_BR[idx - 1]).title()}
            for idx in range(1, 13)
        ],
        'top_itens_core': top_itens_core,
        'top_itens_dfd': top_itens_dfd,
        'processos_opts': Processo.objects.order_by('-ano_referencia', '-numero_processo_sirel', '-numero_processo_adm')[:600],
        'status_opts': StatusProcesso.objects.order_by('nome'),
        'modalidade_opts': Modalidade.objects.order_by('nome'),
        'secretaria_opts': Secretaria.objects.order_by('sigla'),
        'condutor_opts': Pessoa.objects.filter(proc_condutor__isnull=False).distinct().order_by('nome')[:300],
        'modulo_opts': ModuloSistema.choices,
        'situacao_opts': SituacaoWorkflow.choices,
    }
    return render(request, 'workflow/dashboards_geral.html', context)


def _frequencia_bool(value) -> bool:
    return str(value or '').strip().lower() in {'1', 'true', 'on', 'sim', 's', 'yes'}


def _frequencia_parse_hora(raw: str):
    texto = (raw or '').strip()
    if not texto:
        return None
    for fmt in ('%H:%M', '%H:%M:%S'):
        try:
            return datetime.strptime(texto, fmt).time()
        except Exception:
            continue
    raise ValueError(f'Horário inválido: {texto}')


def _frequencia_resolver_usuario(request):
    usuario_alvo = request.user
    usuario_id = (request.GET.get('usuario_id') or request.POST.get('usuario_id') or '').strip()
    usuarios_opts = None
    if request.user.is_staff:
        usuarios_opts = User.objects.order_by('first_name', 'username')[:500]
        if usuario_id.isdigit():
            usuario_alvo = get_object_or_404(User, pk=int(usuario_id))
            usuario_id = str(usuario_alvo.id)
        else:
            usuario_id = str(request.user.id)
            usuario_alvo = request.user
    else:
        usuario_id = str(request.user.id)
    return usuario_alvo, usuario_id, usuarios_opts


def _frequencia_payload_mensal(usuario_alvo, ano_ref: int, mes_ref: int):
    dias_semana = ['Seg', 'Ter', 'Qua', 'Qui', 'Sex', 'Sab', 'Dom']
    meses = [
        'Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho',
        'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro',
    ]
    _, ultimo_dia = calendar.monthrange(ano_ref, mes_ref)
    data_inicio = date(ano_ref, mes_ref, 1)
    data_fim = date(ano_ref, mes_ref, ultimo_dia)
    registros_mes = list(
        FrequenciaRegistro.objects.filter(usuario=usuario_alvo, data__range=(data_inicio, data_fim))
        .order_by('data')
    )
    registros_map = {r.data: r for r in registros_mes}

    linhas = []
    pendencias = 0
    dias_uteis = 0
    dias_trabalhados = 0
    dias_justificados = 0
    total_minutos = 0
    total_extras = 0
    extras_justificativas = []

    for dia in range(1, ultimo_dia + 1):
        data_atual = date(ano_ref, mes_ref, dia)
        is_util = data_atual.weekday() < 5
        if is_util:
            dias_uteis += 1
        reg = registros_map.get(data_atual)
        sem_registro = reg is None
        horarios_completos = bool(
            reg and all([reg.entrada, reg.inicio_intervalo, reg.fim_intervalo, reg.saida])
        )
        possui_algum_horario = bool(
            reg and any([reg.entrada, reg.inicio_intervalo, reg.fim_intervalo, reg.saida])
        )
        nao_trabalhado = bool(reg.nao_trabalhado_util) if reg else False
        em_preenchimento = bool(reg and (not nao_trabalhado) and possui_algum_horario and (not horarios_completos))
        pendente = bool(is_util and (sem_registro or em_preenchimento))
        if pendente:
            pendencias += 1

        entrada_val = reg.entrada.strftime('%H:%M') if reg and reg.entrada else ''
        inicio_val = reg.inicio_intervalo.strftime('%H:%M') if reg and reg.inicio_intervalo else ''
        fim_val = reg.fim_intervalo.strftime('%H:%M') if reg and reg.fim_intervalo else ''
        saida_val = reg.saida.strftime('%H:%M') if reg and reg.saida else ''
        horas_total = _format_horas_minutos(reg.horas_trabalhadas_minutos if reg else 0)
        horas_extras = _format_horas_minutos(reg.horas_extras_minutos if reg else 0)
        just_extras = (reg.justificativa_horas_extras if reg else '') or ''
        just_nao = (reg.justificativa_nao_trabalhado if reg else '') or ''
        observacao = (reg.observacao if reg else '') or ''

        if reg:
            total_minutos += int(reg.horas_trabalhadas_minutos or 0)
            total_extras += int(reg.horas_extras_minutos or 0)
            if reg.nao_trabalhado_util:
                dias_justificados += 1
            elif horarios_completos:
                dias_trabalhados += 1
            if int(reg.horas_extras_minutos or 0) > 0:
                extras_justificativas.append({
                    'data': data_atual,
                    'dia_semana': dias_semana[data_atual.weekday()],
                    'horas_total': horas_total,
                    'horas_extras': horas_extras,
                    'justificativa': just_extras,
                })

        status_label = 'Pendente' if pendente else (
            'Sem registro' if sem_registro else (
                'Não trabalhado (justificado)' if nao_trabalhado else (
                    'Em preenchimento' if em_preenchimento else 'Registrado'
                )
            )
        )
        partes_just = []
        if just_extras:
            partes_just.append(f'Extras: {just_extras}')
        if just_nao:
            partes_just.append(f'Não trabalhado: {just_nao}')
        if observacao:
            partes_just.append(f'Obs.: {observacao}')
        justificativa_resumo = ' | '.join(partes_just)

        linhas.append({
            'data': data_atual,
            'date_key': data_atual.strftime('%Y%m%d'),
            'dia_semana': dias_semana[data_atual.weekday()],
            'is_util': is_util,
            'registro': reg,
            'sem_registro': sem_registro,
            'pendente': pendente,
            'status_label': status_label,
            'entrada': entrada_val or '-',
            'inicio_intervalo': inicio_val or '-',
            'fim_intervalo': fim_val or '-',
            'saida': saida_val or '-',
            'entrada_val': entrada_val,
            'inicio_intervalo_val': inicio_val,
            'fim_intervalo_val': fim_val,
            'saida_val': saida_val,
            'nao_trabalhado_val': '1' if nao_trabalhado else '0',
            'horas_total': horas_total,
            'horas_extras': horas_extras,
            'justificativa_horas_extras': just_extras,
            'justificativa_nao_trabalhado': just_nao,
            'observacao': observacao,
            'justificativa_resumo': justificativa_resumo,
            'justificativa_resumo_curta': (
                (justificativa_resumo[:76] + '...') if len(justificativa_resumo) > 79 else (justificativa_resumo or '-')
            ),
        })

    return {
        'linhas': linhas,
        'dias_uteis': dias_uteis,
        'dias_trabalhados': dias_trabalhados,
        'dias_justificados': dias_justificados,
        'pendencias': pendencias,
        'total_horas': _format_horas_minutos(total_minutos),
        'total_extras': _format_horas_minutos(total_extras),
        'extras_justificativas': extras_justificativas,
        'data_inicio': data_inicio,
        'data_fim': data_fim,
        'mes_label': f'{meses[mes_ref - 1]} de {ano_ref}',
    }


def _frequencia_salvar_mes(request, usuario_alvo, ano_ref: int, mes_ref: int):
    _, ultimo_dia = calendar.monthrange(ano_ref, mes_ref)
    data_inicio = date(ano_ref, mes_ref, 1)
    data_fim = date(ano_ref, mes_ref, ultimo_dia)
    registros_map = {
        r.data: r
        for r in FrequenciaRegistro.objects.filter(usuario=usuario_alvo, data__range=(data_inicio, data_fim))
    }

    erros = []
    operacoes = []
    excluir_ids = []
    row_only_key = (request.POST.get('row_only_key') or '').strip()

    for dia in range(1, ultimo_dia + 1):
        data_atual = date(ano_ref, mes_ref, dia)
        key = data_atual.strftime('%Y%m%d')
        if row_only_key and key != row_only_key:
            continue
        reg_existente = registros_map.get(data_atual)

        entrada_raw = (request.POST.get(f'entrada_{key}') or '').strip()
        inicio_raw = (request.POST.get(f'inicio_intervalo_{key}') or '').strip()
        fim_raw = (request.POST.get(f'fim_intervalo_{key}') or '').strip()
        saida_raw = (request.POST.get(f'saida_{key}') or '').strip()
        nao_trabalhado = _frequencia_bool(request.POST.get(f'nao_{key}'))
        just_extras = (request.POST.get(f'just_extras_{key}') or '').strip()
        just_nao = (request.POST.get(f'just_nao_{key}') or '').strip()
        observacao = (request.POST.get(f'observacao_{key}') or '').strip()

        possui_horarios = any([entrada_raw, inicio_raw, fim_raw, saida_raw])
        possui_justificativas = any([just_extras, just_nao, observacao])
        possui_algum_dado = nao_trabalhado or possui_horarios or possui_justificativas

        if not possui_algum_dado:
            if reg_existente:
                excluir_ids.append(reg_existente.id)
            continue

        valores = {
            'nao_trabalhado_util': nao_trabalhado,
            'entrada': None,
            'inicio_intervalo': None,
            'fim_intervalo': None,
            'saida': None,
            'justificativa_horas_extras': '',
            'justificativa_nao_trabalhado': '',
            'observacao': observacao,
        }

        etiqueta_data = data_atual.strftime('%d/%m/%Y')

        if nao_trabalhado:
            if not just_nao:
                erros.append(f'{etiqueta_data}: informe justificativa para dia útil não trabalhado.')
            valores['justificativa_nao_trabalhado'] = just_nao
        else:
            parsed_times = {}
            for nome, raw in {
                'entrada': entrada_raw,
                'inicio_intervalo': inicio_raw,
                'fim_intervalo': fim_raw,
                'saida': saida_raw,
            }.items():
                if not raw:
                    parsed_times[nome] = None
                    continue
                try:
                    parsed_times[nome] = _frequencia_parse_hora(raw)
                except Exception:
                    erros.append(f'{etiqueta_data}: horário inválido em "{nome.replace("_", " ")}".')

            if all([entrada_raw, inicio_raw, fim_raw, saida_raw]):
                entrada = parsed_times['entrada']
                inicio_int = parsed_times['inicio_intervalo']
                fim_int = parsed_times['fim_intervalo']
                saida = parsed_times['saida']
                if all([entrada, inicio_int, fim_int, saida]):
                    if not (entrada < inicio_int < fim_int < saida):
                        erros.append(f'{etiqueta_data}: ordem dos horários inválida (entrada < início < volta < saída).')
                    else:
                        dummy = FrequenciaRegistro(
                            usuario=usuario_alvo,
                            data=data_atual,
                            entrada=entrada,
                            inicio_intervalo=inicio_int,
                            fim_intervalo=fim_int,
                            saida=saida,
                            nao_trabalhado_util=False,
                        )
                        _, extras_min = dummy.calcular_totais()
                        if extras_min > 0 and not just_extras:
                            erros.append(f'{etiqueta_data}: informe justificativa para horas extras.')

            valores.update({
                'entrada': parsed_times.get('entrada'),
                'inicio_intervalo': parsed_times.get('inicio_intervalo'),
                'fim_intervalo': parsed_times.get('fim_intervalo'),
                'saida': parsed_times.get('saida'),
                'justificativa_horas_extras': just_extras,
            })

        if nao_trabalhado and possui_horarios:
            erros.append(f'{etiqueta_data}: limpe os horários ao marcar "Não trabalhado".')

        operacoes.append((reg_existente, data_atual, valores))

    if erros:
        return False, 'Erros ao salvar a folha mensal: ' + ' | '.join(erros[:10]), 0, 0, 0

    criados = 0
    atualizados = 0
    excluidos = 0
    with transaction.atomic():
        if excluir_ids:
            excluidos = FrequenciaRegistro.objects.filter(id__in=excluir_ids).delete()[0]

        for reg_existente, data_atual, valores in operacoes:
            reg = reg_existente or FrequenciaRegistro(usuario=usuario_alvo, data=data_atual)
            for campo, valor in valores.items():
                setattr(reg, campo, valor)
            reg.usuario = usuario_alvo
            reg.data = data_atual
            reg.save()
            if reg_existente:
                atualizados += 1
            else:
                criados += 1

    return True, 'Folha mensal salva com sucesso.', criados, atualizados, excluidos


def frequencia(request):
    hoje = timezone.localdate()
    mes_referencia = (request.GET.get('mes') or request.POST.get('mes') or hoje.strftime('%Y-%m')).strip()
    ano_ref, mes_ref = _parse_mes_referencia(mes_referencia)
    mes_referencia = f'{ano_ref:04d}-{mes_ref:02d}'
    usuario_alvo, usuario_id, usuarios_opts = _frequencia_resolver_usuario(request)

    if request.method == 'POST' and (request.POST.get('action') or '').strip() == 'save_month':
        ok, msg, criados, atualizados, excluidos = _frequencia_salvar_mes(request, usuario_alvo, ano_ref, mes_ref)
        msg_full = (
            f'{msg} Registros criados: {criados}, atualizados: {atualizados}, removidos: {excluidos}.'
            if ok else msg
        )
        if _is_ajax(request):
            return JsonResponse({'ok': bool(ok), 'message': msg_full})
        if ok:
            messages.success(request, msg_full)
        else:
            messages.error(request, msg)
        params = {'mes': mes_referencia}
        if request.user.is_staff:
            params['usuario_id'] = str(usuario_alvo.id)
        return redirect(f'{reverse("workflow:frequencia")}?{"&".join([f"{k}={v}" for k, v in params.items()])}')

    payload = _frequencia_payload_mensal(usuario_alvo, ano_ref, mes_ref)
    params = {'mes': mes_referencia}
    if request.user.is_staff:
        params['usuario_id'] = str(usuario_alvo.id)
    query = '&'.join([f'{k}={v}' for k, v in params.items()])

    context = {
        'usuario_alvo': usuario_alvo,
        'usuarios_opts': usuarios_opts,
        'usuario_id': usuario_id,
        'mes_referencia': mes_referencia,
        'preview_url': f'{reverse("workflow:frequencia_preview_embed")}?{query}',
        'export_pdf_url': f'{reverse("workflow:frequencia_exportar", args=["pdf"])}?{query}',
        'export_xlsx_url': f'{reverse("workflow:frequencia_exportar", args=["xlsx"])}?{query}',
        **payload,
    }
    return render(request, 'workflow/frequencia.html', context)


@xframe_options_exempt
def frequencia_preview_embed(request):
    hoje = timezone.localdate()
    mes_referencia = (request.GET.get('mes') or hoje.strftime('%Y-%m')).strip()
    ano_ref, mes_ref = _parse_mes_referencia(mes_referencia)
    mes_referencia = f'{ano_ref:04d}-{mes_ref:02d}'
    usuario_alvo, usuario_id, _usuarios_opts = _frequencia_resolver_usuario(request)
    payload = _frequencia_payload_mensal(usuario_alvo, ano_ref, mes_ref)
    context = {
        'mes_referencia': mes_referencia,
        'usuario_id': usuario_id,
        'usuario_alvo': usuario_alvo,
        'orgao': _orgao_ativo(),
        **payload,
    }
    return render(request, 'workflow/frequencia_preview_embed.html', context)


def frequencia_exportar(request, formato: str):
    fmt = (formato or '').strip().lower()
    if fmt not in {'pdf', 'xlsx'}:
        raise Http404('Formato de exportação inválido.')
    hoje = timezone.localdate()
    mes_referencia = (request.GET.get('mes') or hoje.strftime('%Y-%m')).strip()
    ano_ref, mes_ref = _parse_mes_referencia(mes_referencia)
    mes_referencia = f'{ano_ref:04d}-{mes_ref:02d}'
    usuario_alvo, _usuario_id, _usuarios_opts = _frequencia_resolver_usuario(request)
    payload = _frequencia_payload_mensal(usuario_alvo, ano_ref, mes_ref)
    payload['usuario_alvo'] = usuario_alvo
    payload['mes_referencia'] = mes_referencia
    if fmt == 'pdf':
        return export_frequencia_pdf(payload)
    return export_frequencia_xlsx(payload)


def compras_detail(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    _atualizar_etapa_compras_sd(processo)
    workflow.refresh_from_db()
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    dfd = PlanejamentoDFD.objects.filter(processo=processo).first()
    tr = TRPlanejamento.objects.filter(processo=processo).first()

    form_pesquisa = ComprasPesquisaForm(instance=etp)
    form_fonte = ETPCotacaoFonteForm()
    form_cotacao = ETPCotacaoItemForm()
    if dfd:
        form_cotacao.fields['item'].queryset = dfd.itens.all().order_by('codigo')
        form_cotacao.fields['item'].empty_label = 'Selecione um item'
    else:
        form_cotacao.fields['item'].queryset = DFDItem.objects.none()
    form_cotacao.fields['fonte'].queryset = etp.fontes.all().order_by('nome_fonte')
    form_cotacao.fields['fonte'].empty_label = 'Selecione uma fonte'

    cotacoes = etp.cotacoes.select_related('item', 'fonte').order_by('item__codigo', 'fonte__nome_fonte')
    cotacoes_media = (
        etp.cotacoes.filter(considerar_no_calculo=True)
        .values('item__codigo', 'item__descricao')
        .annotate(media=Avg('valor_unitario'))
        .order_by('item__codigo')
    )
    mapa_ctx = _mapa_preview_context(processo)
    totais_canonicos = ProcessoItem.objects.filter(processo=processo).aggregate(
        total_ref=Coalesce(Sum('valor_referencia_total'), Decimal('0.00')),
        total_hom=Coalesce(Sum('valor_homologado_total'), Decimal('0.00')),
        itens=Count('id'),
    )
    total_estimado_canonico = Decimal(str(totais_canonicos.get('total_ref') or 0))
    total_estimado_mapa = Decimal(str(mapa_ctx.get('total_estimado_processo') or 0))
    total_estimado = total_estimado_canonico if total_estimado_canonico > 0 else total_estimado_mapa
    comprovantes = DocumentoProcessoWorkflow.objects.filter(
        processo=processo,
        tipo_documento=COMPRAS_COMPROVANTE_TIPO,
    ).order_by('-criado_em', '-id')

    sd_ctx = build_sd_payload(processo)
    pncp_snapshot = PNCPContratacaoSnapshot.objects.filter(processo=processo).first()
    pncp_constantes = _pncp_mapping_report(pncp_snapshot, processo)['constantes'] if pncp_snapshot else None

    docs_assinados = [
        d for d in _listar_documentos_assinados(processo)
        if d['key'] in {'sd', 'mapa_compras', 'declaracao_desconsideracao'}
    ]
    sd_status = _compras_sd_status(processo)
    etapa_status = {
        'cotacoes': etp.cotacoes.exists(),
        'fontes': etp.fontes.exists() or comprovantes.exists(),
        'sd': sd_status['concluida'],
    }

    context = {
        'processo': processo,
        'workflow': workflow,
        'etp': etp,
        'dfd': dfd,
        'tr': tr,
        'form_pesquisa': form_pesquisa,
        'form_fonte': form_fonte,
        'form_cotacao': form_cotacao,
        'cotacoes': cotacoes,
        'cotacoes_media': cotacoes_media,
        'comprovantes': comprovantes,
        'form_comprovante': ComprasComprovanteUploadForm(),
        'form_doc_assinado': DocumentoAssinadoUploadForm(initial={'doc_key': 'sd'}),
        'documentos_assinados': docs_assinados,
        'dotacoes': sd_ctx['dotacoes'],
        'secretaria_assinatura': sd_ctx['secretaria_assinatura'],
        'signatario_sd': sd_ctx['signatario'],
        'multiplas_secretarias': sd_ctx['multiplas_secretarias'],
        'sd_ctx': sd_ctx,
        'sd_preview_popup_url': reverse('workflow:compras_preview_embed', args=[processo.id, 'sd']),
        'sd_preview_page_url': reverse('workflow:compras_exportar', args=[processo.id, 'sd', 'html']),
        'sd_docx_url': reverse('workflow:compras_exportar', args=[processo.id, 'sd', 'docx']),
        'sd_pdf_url': reverse('workflow:compras_exportar', args=[processo.id, 'sd', 'pdf']),
        'sd_xlsx_url': reverse('workflow:compras_exportar', args=[processo.id, 'sd', 'xlsx']),
        'mapa_preview_url': reverse('workflow:planejamento_exportar', args=[processo.id, 'mapa', 'html']),
        'mapa_docx_url': reverse('workflow:planejamento_exportar', args=[processo.id, 'mapa', 'docx']),
        'mapa_pdf_url': reverse('workflow:planejamento_exportar', args=[processo.id, 'mapa', 'pdf']),
        'mapa_xlsx_url': reverse('workflow:planejamento_exportar', args=[processo.id, 'mapa', 'xlsx']),
        'total_estimado': total_estimado,
        'total_estimado_canonico': total_estimado_canonico,
        'total_homologado_canonico': totais_canonicos.get('total_hom') or Decimal('0'),
        'itens_canonicos_count': totais_canonicos.get('itens') or 0,
        'sd_status': sd_status,
        'etapa_status': etapa_status,
        'pode_encaminhar_licitacao': sd_status['concluida'],
        'pncp_snapshot': pncp_snapshot,
        'pncp_constantes': pncp_constantes,
    }
    return render(request, 'workflow/compras_detail.html', context)


@xframe_options_exempt
def compras_preview_embed(request, processo_id: int, doc: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    doc = (doc or '').lower()
    if doc == 'sd':
        return render(
            request,
            'workflow/sd_preview_embed.html',
            {'processo': processo, **build_sd_payload(processo)},
        )
    return HttpResponseNotFound('Pre-visualizacao nao disponivel para este documento.')


def compras_exportar(request, processo_id: int, doc: str, formato: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    doc = (doc or '').lower()
    formato = (formato or '').lower()

    if doc == 'sd':
        if formato == 'html':
            return render(
                request,
                'workflow/sd_preview.html',
                {'processo': processo, **build_sd_payload(processo)},
            )
        if formato == 'docx':
            _registrar_sd_gerada(processo, formato)
            return export_sd_docx(processo)
        if formato == 'pdf':
            _registrar_sd_gerada(processo, formato)
            return export_sd_pdf(processo)
        if formato == 'xlsx':
            _registrar_sd_gerada(processo, formato)
            return export_sd_xlsx(processo)

    messages.error(request, 'Exportacao nao disponivel para esta combinacao.')
    return redirect('workflow:compras_detail', processo.id)


def compras_salvar_pesquisa(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ComprasPesquisaForm(request.POST, instance=etp)
    if form.is_valid():
        etp = form.save()
        for item in processo.planejamento_dfd.itens.all():
            _recalcular_alertas_cotacoes(etp, item)
        wf = _ensure_workflow(processo)
        wf.situacao = SituacaoWorkflow.EM_ANDAMENTO
        wf.save(update_fields=['situacao', 'atualizado_em'])
        _atualizar_etapa_compras_sd(processo)
        ok = True
        msg = 'Dados de pesquisa de precos salvos.'
    else:
        ok = False
        msg = f'Revise os campos da pesquisa. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_adicionar_fonte(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ETPCotacaoFonteForm(request.POST)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.etp = etp
        obj.save()
        _sincronizar_fornecedor_fonte(obj)
        ok = True
        msg = 'Fonte adicionada com sucesso.'
    else:
        ok = False
        msg = f'Nao foi possivel adicionar a fonte. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_excluir_fonte(request, processo_id: int, fonte_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    fonte = get_object_or_404(ETPCotacaoFonte, pk=fonte_id, etp__processo=processo)
    fonte.delete()
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Fonte removida com sucesso.',
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_anexar_comprovante(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    form = ComprasComprovanteUploadForm(request.POST, request.FILES)
    if form.is_valid():
        ordem = (
            DocumentoProcessoWorkflow.objects.filter(processo=processo)
            .aggregate(maior=Max('ordem_cronologica'))
            .get('maior')
            or 0
        ) + 1
        DocumentoProcessoWorkflow.objects.create(
            processo=processo,
            modulo=ModuloSistema.COMPRAS,
            tipo_documento=COMPRAS_COMPROVANTE_TIPO,
            arquivo=form.cleaned_data['arquivo'],
            ordem_cronologica=ordem,
            gerar_no_etcm=False,
        )
        ok = True
        msg = 'Comprovante anexado.'
    else:
        ok = False
        msg = f'Nao foi possivel anexar o comprovante. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_excluir_comprovante(request, processo_id: int, doc_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    doc = get_object_or_404(
        DocumentoProcessoWorkflow,
        pk=doc_id,
        processo=processo,
        tipo_documento=COMPRAS_COMPROVANTE_TIPO,
    )
    doc.delete()
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Comprovante removido.',
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_adicionar_cotacao(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ETPCotacaoItemForm(request.POST)
    form.fields['item'].queryset = processo.planejamento_dfd.itens.all().order_by('codigo')
    form.fields['item'].empty_label = 'Selecione um item'
    form.fields['fonte'].queryset = etp.fontes.all()
    form.fields['fonte'].empty_label = 'Selecione uma fonte'
    if form.is_valid():
        obj = form.save(commit=False)
        obj.etp = etp
        obj.save()
        _recalcular_alertas_cotacoes(etp, obj.item)
        ok = True
        msg = 'Cotacao registrada e validada.'
    else:
        ok = False
        msg = f'Nao foi possivel registrar a cotacao. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_editar_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    form = ETPCotacaoEdicaoForm(request.POST, instance=cotacao)
    if form.is_valid():
        cotacao = form.save()
        _recalcular_alertas_cotacoes(cotacao.etp, cotacao.item)
        ok = True
        msg = 'Cotacao atualizada.'
    else:
        ok = False
        msg = f'Nao foi possivel atualizar a cotacao. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_alterar_consideracao_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    considerar = request.POST.get('considerar_no_calculo') in {'on', '1', 'true', 'True', 'sim', 'Sim'}
    cotacao.considerar_no_calculo = considerar
    cotacao.save(update_fields=['considerar_no_calculo'])
    _recalcular_alertas_cotacoes(cotacao.etp, cotacao.item)
    msg = 'Cotacao marcada para considerar no calculo.' if considerar else 'Cotacao desconsiderada no calculo.'
    return _ajax_or_redirect(
        request,
        ok=True,
        message=msg,
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def compras_excluir_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:compras_detail', processo.id)
    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    etp = cotacao.etp
    item = cotacao.item
    cotacao.delete()
    _recalcular_alertas_cotacoes(etp, item)
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Cotacao excluida com sucesso.',
        redirect_name='workflow:compras_detail',
        redirect_args=[processo.id],
    )


def perfil_usuario(request):
    if request.method == 'POST':
        form = PerfilUsuarioForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            return _ajax_or_redirect(
                request,
                ok=True,
                message='Dados do usuario atualizados com sucesso.',
                redirect_name='workflow:perfil',
                redirect_args=[],
            )
        return _ajax_or_redirect(
            request,
            ok=False,
            message=f'Nao foi possivel salvar seu perfil. {_form_errors_text(form)}',
            redirect_name='workflow:perfil',
            redirect_args=[],
        )
    form = PerfilUsuarioForm(instance=request.user)
    return render(request, 'workflow/perfil.html', {'form': form})


def _cadastro_configs():
    return {
        'orgao': {
            'titulo': 'Órgão',
            'descricao': 'Identidade institucional para documentos e cabeçalhos.',
            'model': OrgaoEntidade,
            'form': CadastroOrgaoForm,
            'order_by': '-atualizado_em',
            'columns': [
                ('Nome fantasia', lambda o: o.nome_fantasia or '-'),
                ('Razão social', lambda o: o.razao_social),
                ('CNPJ', lambda o: o.cnpj or '-'),
                ('Cidade/UF', lambda o: f'{o.cidade}/{o.uf}' if o.cidade or o.uf else '-'),
                ('Logo', lambda o: 'Sim' if o.logo else 'Não'),
            ],
        },
        'usuarios': {
            'titulo': 'Usuários',
            'descricao': 'Cadastro de contas de acesso ao sistema.',
            'model': User,
            'form': CadastroUsuarioForm,
            'order_by': 'username',
            'columns': [
                ('Usuário', lambda o: o.username),
                ('Nome', lambda o: f'{o.first_name} {o.last_name}'.strip() or '-'),
                ('E-mail', lambda o: o.email or '-'),
                ('Ativo', lambda o: 'Sim' if o.is_active else 'Não'),
                ('Equipe/Admin', lambda o: 'Sim' if o.is_staff else 'Não'),
            ],
        },
        'pessoas': {
            'titulo': 'Pessoas',
            'descricao': 'Responsáveis, fiscais e demais pessoas vinculadas ao processo.',
            'model': Pessoa,
            'form': CadastroPessoaForm,
            'order_by': 'nome',
            'columns': [
                ('Nome', lambda o: o.nome),
                ('CPF', lambda o: o.cpf or '-'),
                ('Cargo', lambda o: o.cargo or '-'),
                ('Secretaria', lambda o: o.secretaria.sigla if o.secretaria else '-'),
            ],
        },
        'secretarias': {
            'titulo': 'Secretarias',
            'descricao': 'Cadastro base de secretarias.',
            'model': Secretaria,
            'form': CadastroSecretariaForm,
            'order_by': 'sigla',
            'columns': [
                ('Sigla', lambda o: o.sigla),
                ('Nome', lambda o: o.nome),
            ],
        },
        'unidades': {
            'titulo': 'Unidades orçamentárias',
            'descricao': 'Cadastro de unidades orçamentárias.',
            'model': UnidadeOrcamentaria,
            'form': CadastroUnidadeOrcamentariaForm,
            'order_by': 'sigla',
            'columns': [
                ('Sigla', lambda o: o.sigla),
                ('Nome', lambda o: o.nome),
            ],
        },
        'projetos': {
            'titulo': 'Projetos/Atividades',
            'descricao': 'Cadastro de projetos e atividades orçamentárias.',
            'model': ProjetoAtividade,
            'form': CadastroProjetoAtividadeForm,
            'order_by': 'codigo',
            'columns': [
                ('Código', lambda o: o.codigo),
                ('Descrição', lambda o: o.descricao),
            ],
        },
        'elementos': {
            'titulo': 'Elementos de despesa',
            'descricao': 'Cadastro de elementos de despesa.',
            'model': ElementoDespesa,
            'form': CadastroElementoDespesaForm,
            'order_by': 'codigo',
            'columns': [
                ('Código', lambda o: o.codigo),
                ('Descrição', lambda o: o.descricao),
            ],
        },
        'fontes': {
            'titulo': 'Fontes de recurso',
            'descricao': 'Cadastro de fontes de recurso.',
            'model': FonteRecurso,
            'form': CadastroFonteRecursoForm,
            'order_by': 'codigo',
            'columns': [
                ('Código', lambda o: o.codigo),
                ('Descrição', lambda o: o.descricao),
            ],
        },
        'fornecedores': {
            'titulo': 'Fornecedores',
            'descricao': 'Cadastro central de fornecedores do processo.',
            'model': Fornecedor,
            'form': CadastroFornecedorForm,
            'order_by': 'razao_social',
            'columns': [
                ('Razão social', lambda o: o.razao_social),
                ('CNPJ', lambda o: o.cnpj),
                ('Cidade', lambda o: o.cidade or '-'),
                ('UF', lambda o: o.estado or '-'),
                ('Contato', lambda o: o.telefone or o.email or '-'),
            ],
        },
        'itens-catalogo': {
            'titulo': 'Itens de catálogo',
            'descricao': 'Cadastro base de itens reutilizáveis na DFD.',
            'model': DFDItemCatalogo,
            'form': CadastroCatalogoItemForm,
            'order_by': 'codigo',
            'columns': [
                ('Código', lambda o: o.codigo),
                ('Descrição', lambda o: o.descricao),
                ('Unidade', lambda o: o.unidade or '-'),
            ],
        },
    }


def _get_cadastro_config(tipo: str):
    configs = _cadastro_configs()
    config = configs.get(tipo)
    if not config:
        raise Http404('Tipo de cadastro inválido.')
    return config


def cadastros_dashboard(request):
    configs = _cadastro_configs()
    cards = []
    for slug, conf in configs.items():
        cards.append({
            'slug': slug,
            'pill': conf['titulo'],
            'titulo': conf['titulo'],
            'descricao': conf['descricao'],
            'total': conf['model'].objects.count(),
        })
    cards.sort(key=lambda c: c['titulo'])
    return render(request, 'workflow/cadastros_dashboard.html', {'cards': cards})


def cadastros_tipo(request, tipo: str):
    config = _get_cadastro_config(tipo)
    model = config['model']
    form_class = config['form']
    edit_id = request.GET.get('edit')
    edit_obj = model.objects.filter(pk=edit_id).first() if edit_id else None

    if request.method == 'POST':
        action = (request.POST.get('action') or 'save').strip().lower()
        obj_id = request.POST.get('id')
        obj = model.objects.filter(pk=obj_id).first() if obj_id else None

        if action == 'delete':
            if not obj:
                ok, msg = False, 'Registro não encontrado para exclusão.'
            else:
                try:
                    obj.delete()
                    ok, msg = True, 'Registro excluído com sucesso.'
                except Exception as exc:
                    ok, msg = False, f'Não foi possível excluir o registro. {exc}'
            return _ajax_or_redirect(
                request,
                ok=ok,
                message=msg,
                redirect_name='workflow:cadastros_tipo',
                redirect_args=[tipo],
            )

        form = form_class(request.POST, request.FILES, instance=obj)
        if form.is_valid():
            is_new = obj is None
            form.save()
            ok = True
            msg = 'Cadastro criado com sucesso.' if is_new else 'Cadastro atualizado com sucesso.'
            return _ajax_or_redirect(
                request,
                ok=ok,
                message=msg,
                redirect_name='workflow:cadastros_tipo',
                redirect_args=[tipo],
            )
        edit_obj = obj
    else:
        form = form_class(instance=edit_obj)

    if request.method != 'POST':
        form = form_class(instance=edit_obj)

    rows = model.objects.order_by(config['order_by'])[:500]
    rows_data = []
    for row in rows:
        values = [getter(row) for _, getter in config['columns']]
        rows_data.append({'id': row.pk, 'values': values})

    context = {
        'tipo': tipo,
        'titulo': config['titulo'],
        'descricao': config['descricao'],
        'columns': [label for label, _ in config['columns']],
        'rows_data': rows_data,
        'form': form,
        'form_is_multipart': form.is_multipart(),
        'edit_obj': edit_obj,
        'tipos': [{'slug': s, 'titulo': c['titulo']} for s, c in _cadastro_configs().items()],
    }
    return render(request, 'workflow/cadastros_tipo.html', context)


def integracoes(request):
    logs = IntegracaoProcesso.objects.select_related('processo')[:50]
    processos = Processo.objects.order_by('-atualizado_em', '-id')[:1500]
    pncp_snapshots = PNCPContratacaoSnapshot.objects.select_related('processo')[:30]
    fila_detalhamento = PNCPDetalhamentoFila.objects.select_related('processo').order_by('prioridade', '-agendado_em')[:40]
    fila_status_total = {
        'pendente': PNCPDetalhamentoFila.objects.filter(status=PNCPDetalhamentoFila.Status.PENDENTE).count(),
        'processando': PNCPDetalhamentoFila.objects.filter(status=PNCPDetalhamentoFila.Status.PROCESSANDO).count(),
        'concluido': PNCPDetalhamentoFila.objects.filter(status=PNCPDetalhamentoFila.Status.CONCLUIDO).count(),
        'parcial': PNCPDetalhamentoFila.objects.filter(status=PNCPDetalhamentoFila.Status.PARCIAL).count(),
        'erro': PNCPDetalhamentoFila.objects.filter(status=PNCPDetalhamentoFila.Status.ERRO).count(),
    }
    pncp_erro_popup = request.session.pop('pncp_erro_popup', None)
    pncp_sucesso_popup = request.session.pop('pncp_sucesso_popup', None)
    orgao = _orgao_ativo()
    cnpj_orgao_default = ''.join(ch for ch in str(getattr(orgao, 'cnpj', '') or '') if ch.isdigit())
    uf_default = (getattr(orgao, 'uf', '') or '').strip().upper()

    modalidades = {(m['codigo'], m['nome']) for m in PNCP_MODALIDADES_PADRAO}
    for row in (
        PNCPContratacaoSnapshot.objects.exclude(modalidade_id__isnull=True)
        .exclude(modalidade_nome='')
        .values('modalidade_id', 'modalidade_nome')
        .order_by('modalidade_id')
        .distinct()
    ):
        modalidades.add((int(row['modalidade_id']), row['modalidade_nome']))
    pncp_modalidades = [
        {'codigo': codigo, 'nome': nome}
        for codigo, nome in sorted(modalidades, key=lambda x: x[0])
    ]

    return render(
        request,
        'workflow/integracoes.html',
        {
            'logs': logs,
            'processos': processos,
            'pncp_snapshots': pncp_snapshots,
            'today': timezone.localdate(),
            'pncp_modalidades': pncp_modalidades,
            'cnpj_orgao_default': cnpj_orgao_default,
            'uf_default': uf_default,
            'orgao_ativo': orgao,
            'pncp_erro_popup': pncp_erro_popup,
            'pncp_sucesso_popup': pncp_sucesso_popup,
            'pncp_envio_habilitado': bool(getattr(settings, 'PNCP_ENVIO_HABILITADO', False)),
            'pncp_envio_dry_run': bool(getattr(settings, 'PNCP_ENVIO_DRY_RUN', True)),
            'fila_detalhamento': fila_detalhamento,
            'fila_status_total': fila_status_total,
        },
    )


def processar_fila_pncp_manual(request):
    if request.method != 'POST':
        return redirect('workflow:integracoes')
    limite = _pncp_to_int(request.POST.get('limite')) or 10
    try:
        resumo = processar_fila_pncp_service(limit=max(1, int(limite)))
        messages.success(
            request,
            (
                'Fila PNCP processada. '
                f"Capturados: {resumo.get('capturados', 0)}, "
                f"Concluidos: {resumo.get('concluidos', 0)}, "
                f"Parciais: {resumo.get('parciais', 0)}, "
                f"Erros: {resumo.get('erros', 0)}."
            ),
        )
    except Exception as exc:
        messages.error(request, f'Falha ao processar fila PNCP: {exc}')
    return redirect('workflow:integracoes')


def importar_pncp_publicacoes_status(request):
    op_id = (request.GET.get('op_id') or '').strip()
    if not op_id:
        return JsonResponse({'ok': False, 'error': 'op_id obrigatorio'}, status=400)
    if not request.user.is_authenticated:
        return JsonResponse({'ok': False, 'error': 'nao autenticado'}, status=401)
    payload = cache.get(_pncp_progress_key(request.user.id, op_id))
    if not payload:
        return JsonResponse({'ok': False, 'error': 'status nao encontrado'}, status=404)
    return JsonResponse({'ok': True, 'progress': payload})


def importar_pncp_publicacoes_auto(request):
    if request.method != 'POST':
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    data_inicial = (request.POST.get('data_inicial') or '').strip()
    data_final = (request.POST.get('data_final') or '').strip()
    modo_importacao = (request.POST.get('modo_importacao') or '').strip().lower()
    importacao_rapida = modo_importacao in {'rapida', 'rápida', 'rapido', 'rápido', '1', 'true', 'on'}
    op_id = (request.POST.get('op_id') or '').strip() or uuid4().hex

    _set_pncp_import_progress(
        request,
        op_id,
        status='running',
        message=(
            'Iniciando importacao automatica PNCP (modo rapido)...'
            if importacao_rapida else
            'Iniciando importacao automatica PNCP...'
        ),
        modalidade_atual=0,
        modalidades_total=0,
        modalidade_nome='',
        page_current=0,
        paginas_lidas=0,
        max_paginas=0,
        criados=0,
        atualizados=0,
        ignorados=0,
    )

    dt_ini = _parse_pncp_date(data_inicial)
    dt_fim = _parse_pncp_date(data_final)
    if not data_inicial or not data_final:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - importacao PNCP em lote',
            detalhes='Informe data inicial e data final.',
            contexto={
                'data_inicial': data_inicial or '-',
                'data_final': data_final or '-',
            },
        )
        _set_pncp_import_progress(request, op_id, status='error', message='Falha: informe data inicial e final.')
        messages.error(request, 'Informe data inicial e data final para importar.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)
    if dt_ini and dt_fim and dt_fim < dt_ini:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - periodo invalido',
            detalhes='Data final menor que data inicial.',
            contexto={'data_inicial': data_inicial, 'data_final': data_final},
        )
        _set_pncp_import_progress(
            request,
            op_id,
            status='error',
            message='Falha: periodo invalido (data final menor que inicial).',
        )
        messages.error(request, 'Data final deve ser maior ou igual a data inicial.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    orgao = _orgao_ativo()
    cnpj = ''.join(ch for ch in str(getattr(orgao, 'cnpj', '') or '') if ch.isdigit())
    uf = (getattr(orgao, 'uf', '') or '').strip().upper()
    if len(cnpj) != 14:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - CNPJ do orgao',
            detalhes='O CNPJ do orgao cadastrado no sistema e obrigatorio e deve conter 14 digitos.',
            contexto={
                'cnpj_orgao': cnpj or '-',
                'orgao': getattr(orgao, 'razao_social', '') or getattr(orgao, 'nome_fantasia', '') or '-',
            },
        )
        _set_pncp_import_progress(
            request,
            op_id,
            status='error',
            message='Falha: CNPJ do orgao nao cadastrado ou invalido.',
        )
        messages.error(request, 'Cadastre um CNPJ valido (14 digitos) no cadastro do orgao para importar em lote.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    tamanho_pagina_auto, max_paginas_auto = _pncp_auto_paginacao(dt_ini, dt_fim)
    modalidades = sorted(
        {(int(m['codigo']), str(m['nome'])) for m in PNCP_MODALIDADES_PADRAO},
        key=lambda x: x[0],
    )
    modalidades_total = len(modalidades)
    if not modalidades_total:
        _set_pncp_import_progress(request, op_id, status='error', message='Nenhuma modalidade PNCP configurada.')
        messages.error(request, 'Nenhuma modalidade PNCP configurada para importacao.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    _set_pncp_import_progress(
        request,
        op_id,
        status='running',
        message='Configuracao automatica aplicada.',
        modalidade_atual=0,
        modalidades_total=modalidades_total,
        modalidade_nome='Todas',
        max_paginas=max_paginas_auto,
    )

    processo_log = (
        Processo.objects.order_by('-id').first()
        or Processo.objects.create(
            numero_processo_adm='PNCP-LOTE',
            ano_referencia=timezone.localdate().year,
            objeto='Importacao em lote PNCP',
            modalidade=_get_modalidade('Pregao'),
            status=_get_status('EM PLANEJAMENTO'),
            secretaria=_secretaria_adm_principal(),
        )
    )
    log = IntegracaoProcesso.objects.create(
        processo=processo_log,
        tipo=IntegracaoProcesso.Tipo.PNCP,
        identificador_externo=f'LOTE:{data_inicial}:{data_final}:ALL:{"RAPIDA" if importacao_rapida else "COMPLETA"}',
    )

    client = PNCPClient(timeout=60, max_retries=2, rate_limit_retries=3)
    criados = 0
    atualizados = 0
    ignorados = 0
    paginas_lidas = 0
    retries_timeout = 0
    itens_detalhados_total = 0
    fornecedores_detalhados_total = 0
    erros_detalhes_total = 0
    detalhamento_enfileirado = 0
    fila_auto_disparada = False
    processos_importados_map = {}
    erros = []
    ok_result = True

    try:
        for modalidade_idx, (modalidade_code, modalidade_nome) in enumerate(modalidades, start=1):
            limite_modalidade_atingido = True
            for pagina in range(1, max_paginas_auto + 1):
                _set_pncp_import_progress(
                    request,
                    op_id,
                    status='running',
                    message=(
                        f'Consultando modalidade {modalidade_nome} ({modalidade_idx}/{modalidades_total}) '
                        f'- pagina {pagina}...'
                    ),
                    modalidade_atual=modalidade_idx,
                    modalidades_total=modalidades_total,
                    modalidade_nome=modalidade_nome,
                    page_current=pagina,
                    paginas_lidas=paginas_lidas,
                    max_paginas=max_paginas_auto,
                    criados=criados,
                    atualizados=atualizados,
                    ignorados=ignorados,
                )
                tamanho_pagina_efetivo = tamanho_pagina_auto
                try:
                    resposta = client.listar_publicacoes(
                        data_inicial=data_inicial,
                        data_final=data_final,
                        codigo_modalidade_contratacao=modalidade_code,
                        uf=uf or None,
                        codigo_municipio_ibge=None,
                        cnpj=cnpj,
                        pagina=pagina,
                        tamanho_pagina=tamanho_pagina_efetivo,
                    )
                except Exception as exc_lote:
                    erro_lote_texto = str(exc_lote).lower()
                    timeout_detectado = ('timeout' in erro_lote_texto) or ('timed out' in erro_lote_texto)
                    if timeout_detectado and tamanho_pagina_auto > 10:
                        tamanho_pagina_efetivo = max(10, int(tamanho_pagina_auto / 2))
                        retries_timeout += 1
                        try:
                            resposta = client.listar_publicacoes(
                                data_inicial=data_inicial,
                                data_final=data_final,
                                codigo_modalidade_contratacao=modalidade_code,
                                uf=uf or None,
                                codigo_municipio_ibge=None,
                                cnpj=cnpj,
                                pagina=pagina,
                                tamanho_pagina=tamanho_pagina_efetivo,
                            )
                        except Exception as exc_retry:
                            erros.append(
                                f'[{modalidade_nome}] pagina {pagina}: {exc_retry}'
                            )
                            break
                    else:
                        erros.append(
                            f'[{modalidade_nome}] pagina {pagina}: {exc_lote}'
                        )
                        break

                paginas_lidas += 1
                registros = resposta.get('data') if isinstance(resposta, dict) else []
                if not registros:
                    limite_modalidade_atingido = False
                    break

                if pagina < max_paginas_auto and not importacao_rapida:
                    try:
                        import time as _time
                        _time.sleep(0.12)
                    except Exception:
                        pass

                for payload in registros:
                    try:
                        numero = _pncp_safe_text(payload.get('numeroControlePNCP'))
                        if not numero:
                            ignorados += 1
                            continue
                        processo = _pncp_detectar_processo_existente(numero, payload)
                        processo_existente = processo is not None

                        if not processo:
                            numero_proc = _pncp_safe_text(payload.get('processo')) or _pncp_safe_text(payload.get('numeroCompra')) or numero
                            ano_ref = _pncp_to_int(payload.get('anoCompra')) or timezone.localdate().year
                            objeto = _pncp_safe_text(payload.get('objetoCompra')) or 'Objeto importado do PNCP'
                            processo = Processo.objects.create(
                                numero_processo_adm=numero_proc,
                                ano_referencia=ano_ref,
                                objeto=objeto,
                                modalidade=_get_modalidade(_normalizar_modalidade_pncp(_pncp_safe_text(payload.get('modalidadeNome')) or 'Pregao')),
                                status=_get_status('EM PLANEJAMENTO'),
                                secretaria=_secretaria_adm_principal(),
                            )

                        payload_enriquecido = dict(payload)
                        if importacao_rapida:
                            snap = _upsert_pncp_snapshot(processo, numero, payload_enriquecido)
                            _sincronizar_processo_com_snapshot_pncp(processo, snap)
                            _marcar_pendente_detalhamento_pncp(processo, numero_controle=numero)
                            detalhamento_enfileirado += 1
                        else:
                            itens_detalhados = []
                            resultados_detalhados = []
                            try:
                                itens_detalhados = client.consultar_itens(numero, strict=True)
                            except Exception as exc_itens:
                                erros_detalhes_total += 1
                                erros.append(f'[{numero}] itens: {exc_itens}')
                            try:
                                resultados_detalhados = client.consultar_resultados(numero, strict=True)
                            except Exception as exc_res:
                                erros_detalhes_total += 1
                                erros.append(f'[{numero}] resultados: {exc_res}')

                            if not itens_detalhados:
                                itens_detalhados = _pncp_extract_embedded_list(
                                    payload,
                                    ['sirel_pncp_itens', 'itens', 'items', 'itensCompra'],
                                )
                            if not resultados_detalhados:
                                resultados_detalhados = _pncp_extract_embedded_list(
                                    payload,
                                    ['sirel_pncp_resultados', 'resultados', 'resultado', 'resultadoItens', 'resultado-itens'],
                                )

                            if itens_detalhados:
                                payload_enriquecido['sirel_pncp_itens'] = itens_detalhados
                            if resultados_detalhados:
                                payload_enriquecido['sirel_pncp_resultados'] = resultados_detalhados

                            sync_detalhada = _sincronizar_itens_fornecedores_pncp(
                                processo,
                                itens_detalhados,
                                resultados_detalhados,
                                situacao_compra_nome=_pncp_safe_text(payload.get('situacaoCompraNome')),
                                numero_controle=numero,
                            )
                            itens_detalhados_total += int(sync_detalhada.get('itens_importados') or 0)
                            fornecedores_detalhados_total += int(sync_detalhada.get('fornecedores_atualizados') or 0)

                            snap = _upsert_pncp_snapshot(processo, numero, payload_enriquecido)
                            _sincronizar_processo_com_snapshot_pncp(processo, snap)

                        processo_key = str(processo.id)
                        processo_acao = 'Atualizado' if processo_existente else 'Criado'
                        atual = processos_importados_map.get(processo_key)
                        if not atual:
                            processos_importados_map[processo_key] = {
                                'numero': f'{processo.numero_processo_adm}/{processo.ano_referencia}',
                                'objeto': processo.objeto or '',
                                'acao': processo_acao,
                            }
                            if processo_acao == 'Criado':
                                criados += 1
                            else:
                                atualizados += 1
                        elif atual.get('acao') != 'Criado' and processo_acao == 'Criado':
                            processos_importados_map[processo_key]['acao'] = 'Criado'
                    except Exception as exc_item:
                        erros.append(str(exc_item))
                        continue

            if limite_modalidade_atingido:
                erros.append(
                    f'[{modalidade_nome}] limite automatico de {max_paginas_auto} paginas atingido; '
                    'revise o periodo se houver necessidade de maior abrangencia.'
                )

        if importacao_rapida and detalhamento_enfileirado > 0:
            fila_auto_disparada = _disparar_fila_pncp_background(
                limit=getattr(settings, 'PNCP_DETALHAMENTO_AUTOSTART_LIMIT', 20),
            )

        log.status = 'SUCESSO' if not erros else 'PARCIAL'
        log.payload_resumo = {
            'criados': criados,
            'atualizados': atualizados,
            'ignorados': ignorados,
            'paginas_lidas': paginas_lidas,
            'retries_timeout': retries_timeout,
            'itens_detalhados_importados': itens_detalhados_total,
            'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
            'erros_detalhes_total': erros_detalhes_total,
            'modo_importacao': 'RAPIDA' if importacao_rapida else 'COMPLETA',
            'detalhamento_enfileirado': detalhamento_enfileirado,
            'fila_auto_disparada': bool(fila_auto_disparada),
            'tamanho_pagina_auto': tamanho_pagina_auto,
            'max_paginas_auto': max_paginas_auto,
            'modalidades_consultadas': [codigo for codigo, _ in modalidades],
            'erros': erros[:20],
        }
        log.mensagem = (
            f'Importacao automatica PNCP em lote concluida. Criados: {criados}, '
            f'Atualizados: {atualizados}, Ignorados: {ignorados}. '
            f'Paginas lidas: {paginas_lidas}.'
        )
        log.save(update_fields=['status', 'payload_resumo', 'mensagem'])

        if not criados and not atualizados and not erros:
            messages.warning(
                request,
                'Nenhum registro retornado pelo PNCP para o periodo informado.',
            )
            _set_pncp_import_progress(
                request,
                op_id,
                status='done',
                message='Importacao finalizada sem registros para o periodo informado.',
                modalidade_atual=modalidades_total,
                modalidades_total=modalidades_total,
                page_current=0,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas_auto,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
                detalhamento_enfileirado=detalhamento_enfileirado,
                fila_auto_disparada=bool(fila_auto_disparada),
            )
        elif erros:
            _set_pncp_error_popup(
                request,
                titulo='Importacao automatica PNCP em lote com falhas parciais',
                detalhes=erros[:20],
                contexto={
                    'data_inicial': data_inicial,
                    'data_final': data_final,
                    'cnpj_orgao': cnpj,
                    'uf_orgao': uf or '-',
                    'modalidades_consultadas': ', '.join([str(cod) for cod, _ in modalidades]),
                    'tamanho_pagina_auto': tamanho_pagina_auto,
                    'max_paginas_auto': max_paginas_auto,
                    'paginas_lidas': paginas_lidas,
                    'retries_timeout': retries_timeout,
                    'itens_detalhados_importados': itens_detalhados_total,
                    'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
                    'modo_importacao': 'RAPIDA' if importacao_rapida else 'COMPLETA',
                    'detalhamento_enfileirado': detalhamento_enfileirado,
                    'fila_auto_disparada': bool(fila_auto_disparada),
                    'criados': criados,
                    'atualizados': atualizados,
                    'ignorados': ignorados,
                },
            )
            messages.warning(request, log.mensagem + ' Houve falhas em parte das consultas.')
            _set_pncp_import_progress(
                request,
                op_id,
                status='partial',
                message='Importacao finalizada com falhas parciais.',
                modalidade_atual=modalidades_total,
                modalidades_total=modalidades_total,
                page_current=0,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas_auto,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
                detalhamento_enfileirado=detalhamento_enfileirado,
                fila_auto_disparada=bool(fila_auto_disparada),
            )
        else:
            _set_pncp_success_popup(
                request,
                titulo='Importacao automatica PNCP em lote concluida com sucesso',
                processos=list(processos_importados_map.values()),
                contexto={
                    'processos_importados': len(processos_importados_map),
                    'criados': criados,
                    'atualizados': atualizados,
                    'ignorados': ignorados,
                    'modalidades_consultadas': modalidades_total,
                    'paginas_lidas': paginas_lidas,
                    'tamanho_pagina_auto': tamanho_pagina_auto,
                    'max_paginas_auto': max_paginas_auto,
                    'itens_detalhados_importados': itens_detalhados_total,
                    'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
                    'modo_importacao': 'RAPIDA' if importacao_rapida else 'COMPLETA',
                    'detalhamento_enfileirado': detalhamento_enfileirado,
                    'fila_auto_disparada': bool(fila_auto_disparada),
                },
            )
            messages.success(request, log.mensagem)
            _set_pncp_import_progress(
                request,
                op_id,
                status='done',
                message='Importacao finalizada com sucesso.',
                modalidade_atual=modalidades_total,
                modalidades_total=modalidades_total,
                page_current=0,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas_auto,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
                detalhamento_enfileirado=detalhamento_enfileirado,
                fila_auto_disparada=bool(fila_auto_disparada),
            )
        if importacao_rapida and detalhamento_enfileirado > 0:
            if fila_auto_disparada:
                messages.info(
                    request,
                    f'Detalhamento PNCP enfileirado para {detalhamento_enfileirado} processo(s) e processamento em segundo plano iniciado.',
                )
            else:
                messages.info(
                    request,
                    f'Detalhamento PNCP enfileirado para {detalhamento_enfileirado} processo(s). Execute "Processar fila PNCP" para concluir agora.',
                )
    except Exception as exc:
        ok_result = False
        log.status = 'ERRO'
        log.mensagem = str(exc)
        log.save(update_fields=['status', 'mensagem'])
        _set_pncp_error_popup(
            request,
            titulo='Falha na importacao automatica PNCP em lote',
            detalhes=str(exc),
            contexto={
                'data_inicial': data_inicial,
                'data_final': data_final,
                'cnpj_orgao': cnpj,
                'uf_orgao': uf or '-',
                'paginas_lidas': paginas_lidas,
            },
        )
        messages.error(request, f'Falha na importacao em lote do PNCP: {exc}')
        _set_pncp_import_progress(
            request,
            op_id,
            status='error',
            message=f'Falha na importacao: {exc}',
            modalidade_atual=0,
            modalidades_total=modalidades_total,
            page_current=0,
            paginas_lidas=paginas_lidas,
            max_paginas=max_paginas_auto,
            criados=criados,
            atualizados=atualizados,
            ignorados=ignorados,
        )

    return _ajax_redirect_payload(request, 'workflow:integracoes', ok=ok_result)


def importar_pncp(request):
    if request.method == 'POST':
        processo_id = (request.POST.get('processo_id') or '').strip()
        numero = _normalizar_entrada_pncp_numero_controle(request.POST.get('numero_controle'))
        criar_novo = (request.POST.get('criar_novo_se_ausente') or '') in {'on', '1', 'true', 'True'}
        modo_importacao = (request.POST.get('modo_importacao') or '').strip().lower()
        importacao_rapida = modo_importacao in {'rapida', 'rápida', 'rapido', 'rápido', '1', 'true', 'on'}
        fila_auto_disparada = False
        if not numero:
            messages.error(request, 'Informe o numero de controle do PNCP.')
            return redirect('workflow:integracoes')

        processo = None
        if processo_id:
            processo = get_object_or_404(Processo, pk=processo_id)

        processo_log = processo
        if not processo_log:
            processo_log = (
                Processo.objects.order_by('-id').first()
                or Processo.objects.create(
                    numero_processo_adm='PNCP-TEMP',
                    ano_referencia=timezone.localdate().year,
                    objeto='Importacao PNCP em andamento',
                    modalidade=_get_modalidade('Pregao'),
                    status=_get_status('EM PLANEJAMENTO'),
                    secretaria=_secretaria_adm_principal(),
                )
            )

        log = IntegracaoProcesso.objects.create(
            processo=processo_log,
            tipo=IntegracaoProcesso.Tipo.PNCP,
            identificador_externo=numero,
        )
        try:
            client = PNCPClient(timeout=60, max_retries=2)
            payload_raw = client.consultar(numero)
            payload = _pncp_first_payload(payload_raw)
            if not payload:
                raise ValueError('PNCP nao retornou dados para o numero de controle informado.')
            if not processo:
                processo = _pncp_detectar_processo_existente(numero, payload)
            if not processo and not criar_novo:
                raise ValueError(
                    'Processo nao localizado no sistema para este numero PNCP. '
                    'Marque a opcao para criar processo novo.'
                )

            if not processo:
                ano_ref = _pncp_to_int(payload.get('anoCompra')) or timezone.localdate().year
                numero_proc = (
                    _pncp_safe_text(payload.get('processo'))
                    or _pncp_safe_text(payload.get('numeroCompra'))
                    or numero
                )
                objeto = _pncp_safe_text(payload.get('objetoCompra')) or 'Objeto importado do PNCP'
                processo = Processo.objects.create(
                    numero_processo_adm=numero_proc,
                    ano_referencia=ano_ref,
                    objeto=objeto,
                    modalidade=_get_modalidade(_normalizar_modalidade_pncp(_pncp_safe_text(payload.get('modalidadeNome')) or 'Pregao')),
                    status=_get_status('EM PLANEJAMENTO'),
                    secretaria=_secretaria_adm_principal(),
                )
                criacao = True
                log.processo = processo
            else:
                criacao = False

            payload_enriquecido = dict(payload)
            erros_detalhes = []
            if importacao_rapida:
                sync_detalhada = {
                    'itens_importados': 0,
                    'fornecedores_atualizados': 0,
                    'status_forcados': 0,
                    'resultados_recebidos': 0,
                    'detalhamento_pendente': True,
                }
            else:
                itens_detalhados = []
                resultados_detalhados = []
                try:
                    itens_detalhados = client.consultar_itens(numero, strict=True)
                except Exception as exc_itens:
                    erros_detalhes.append(f'itens: {exc_itens}')
                try:
                    resultados_detalhados = client.consultar_resultados(numero, strict=True)
                except Exception as exc_res:
                    erros_detalhes.append(f'resultados: {exc_res}')

                if not itens_detalhados:
                    itens_detalhados = _pncp_extract_embedded_list(
                        payload_raw,
                        ['sirel_pncp_itens', 'itens', 'items', 'itensCompra'],
                    )
                if not itens_detalhados:
                    itens_detalhados = _pncp_extract_embedded_list(
                        payload,
                        ['sirel_pncp_itens', 'itens', 'items', 'itensCompra'],
                    )
                if not resultados_detalhados:
                    resultados_detalhados = _pncp_extract_embedded_list(
                        payload_raw,
                        ['sirel_pncp_resultados', 'resultados', 'resultado', 'resultadoItens', 'resultado-itens'],
                    )
                if not resultados_detalhados:
                    resultados_detalhados = _pncp_extract_embedded_list(
                        payload,
                        ['sirel_pncp_resultados', 'resultados', 'resultado', 'resultadoItens', 'resultado-itens'],
                    )

                if itens_detalhados:
                    payload_enriquecido['sirel_pncp_itens'] = itens_detalhados
                if resultados_detalhados:
                    payload_enriquecido['sirel_pncp_resultados'] = resultados_detalhados

                sync_detalhada = _sincronizar_itens_fornecedores_pncp(
                    processo,
                    itens_detalhados,
                    resultados_detalhados,
                    situacao_compra_nome=_pncp_safe_text(payload.get('situacaoCompraNome')),
                    numero_controle=numero,
                )

            snap = _upsert_pncp_snapshot(processo, numero, payload_enriquecido)
            _sincronizar_processo_com_snapshot_pncp(processo, snap)
            if importacao_rapida:
                _marcar_pendente_detalhamento_pncp(processo, numero_controle=numero)
                fila_auto_disparada = _disparar_fila_pncp_background(
                    limit=getattr(settings, 'PNCP_DETALHAMENTO_AUTOSTART_LIMIT', 20),
                )
            map_report = _pncp_mapping_report(snap, processo)

            log.status = 'SUCESSO'
            log.payload_resumo = {
                'numeroControlePNCP': snap.numero_controle_pncp,
                'processo_id': processo.id,
                'processo': f'{processo.numero_processo_adm}/{processo.ano_referencia}',
                'objetoCompra': snap.objeto_compra,
                'modalidadeNome': snap.modalidade_nome,
                'situacaoCompraNome': snap.situacao_compra_nome,
                'modo_importacao': 'RAPIDA' if importacao_rapida else 'COMPLETA',
                'fila_auto_disparada': bool(fila_auto_disparada),
                'valorTotalEstimado': str(snap.valor_total_estimado),
                'valorTotalHomologado': str(snap.valor_total_homologado),
                'sincronizacao_detalhada': sync_detalhada,
                'erros_dados_detalhados': erros_detalhes[:10],
                'mapeamento': {
                    'payload_total_campos': map_report['payload_total_campos'],
                    'payload_campos_mapeados': map_report['payload_campos_mapeados'],
                    'payload_campos_extras': map_report['payload_campos_extras'][:20],
                },
            }
            log.mensagem = (
                'Importacao PNCP concluida em modo rapido (detalhamento enfileirado).'
                if importacao_rapida else
                'Importacao PNCP concluida e dados mapeados para o processo.'
                if criacao else
                'Importacao PNCP concluida em modo rapido com atualizacao de processo existente.'
                if importacao_rapida else
                'Importacao PNCP concluida com atualizacao de processo existente.'
            )
            log.save(update_fields=['processo', 'status', 'payload_resumo', 'mensagem'])
            _set_pncp_success_popup(
                request,
                titulo='Importacao PNCP concluida',
                processos=[
                    {
                        'numero': f'{processo.numero_processo_adm}/{processo.ano_referencia}',
                        'objeto': processo.objeto,
                        'acao': 'Criado' if criacao else 'Atualizado',
                    }
                ],
                contexto={
                    'processos_importados': 1,
                    'criados': 1 if criacao else 0,
                    'atualizados': 0 if criacao else 1,
                    'modo_importacao': 'RAPIDA' if importacao_rapida else 'COMPLETA',
                    'fila_auto_disparada': bool(fila_auto_disparada),
                    'itens_detalhados_importados': sync_detalhada.get('itens_importados', 0),
                    'fornecedores_atualizados': sync_detalhada.get('fornecedores_atualizados', 0),
                },
            )
            if importacao_rapida:
                detalhe_msg = (
                    'Detalhamento foi enfileirado e o processamento em segundo plano foi iniciado.'
                    if fila_auto_disparada
                    else 'Detalhamento foi enfileirado. Use "Processar fila PNCP" para executar agora.'
                )
                messages.success(request, f'Importacao rapida PNCP concluida para {processo.numero_processo_adm}/{processo.ano_referencia}. {detalhe_msg}')
            else:
                messages.success(request, f'Importacao PNCP concluida para o processo {processo.numero_processo_adm}/{processo.ano_referencia}.')
            if erros_detalhes and not importacao_rapida:
                messages.warning(
                    request,
                    'Importacao principal concluida, mas houve falha parcial ao buscar itens/resultados detalhados do PNCP.',
                )
        except Exception as exc:
            log.status = 'ERRO'
            log.mensagem = str(exc)
            log.save(update_fields=['status', 'mensagem'])
            _set_pncp_error_popup(
                request,
                titulo='Falha na importacao PNCP (numero de controle)',
                detalhes=str(exc),
                contexto={
                    'numero_controle': numero,
                    'processo_id_informado': processo_id or '-',
                    'criar_novo_se_ausente': criar_novo,
                },
            )
            messages.error(request, f'Falha ao importar PNCP: {exc}')
    return redirect('workflow:integracoes')


def importar_pncp_publicacoes(request):
    # Compatibilidade: fluxo antigo agora delega para o modo automatico.
    return importar_pncp_publicacoes_auto(request)

    # Implementacao legada (mantida apenas para referencia).
    if request.method != 'POST':
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    data_inicial = (request.POST.get('data_inicial') or '').strip()
    data_final = (request.POST.get('data_final') or '').strip()
    codigo_modalidade = (request.POST.get('codigo_modalidade') or '').strip()
    tamanho_pagina = _pncp_to_int(request.POST.get('tamanho_pagina')) or 20
    max_paginas = _pncp_to_int(request.POST.get('max_paginas')) or 3
    uf = (request.POST.get('uf') or '').strip().upper()
    codigo_municipio_ibge = (request.POST.get('codigo_municipio_ibge') or '').strip()
    cnpj = ''.join(ch for ch in (request.POST.get('cnpj') or '').strip() if ch.isdigit())
    importar_somente_novos = (request.POST.get('somente_novos') or '') in {'on', '1', 'true', 'True'}
    op_id = (request.POST.get('op_id') or '').strip() or uuid4().hex

    _set_pncp_import_progress(
        request,
        op_id,
        status='running',
        message='Iniciando importacao PNCP em lote...',
        page_current=0,
        paginas_lidas=0,
        max_paginas=max_paginas,
        criados=0,
        atualizados=0,
        ignorados=0,
    )

    if not cnpj:
        orgao = _orgao_ativo()
        cnpj = ''.join(ch for ch in str(getattr(orgao, 'cnpj', '') or '') if ch.isdigit())

    if not data_inicial or not data_final or not codigo_modalidade:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - importacao PNCP em lote',
            detalhes='Informe data inicial, data final e modalidade.',
            contexto={
                'data_inicial': data_inicial,
                'data_final': data_final,
                'codigo_modalidade': codigo_modalidade or '-',
            },
        )
        messages.error(request, 'Informe data inicial, data final e codigo da modalidade.')
        _set_pncp_import_progress(request, op_id, status='error', message='Falha de validacao dos campos obrigatorios.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)
    dt_ini = _parse_pncp_date(data_inicial)
    dt_fim = _parse_pncp_date(data_final)
    if dt_ini and dt_fim and dt_fim < dt_ini:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - periodo invalido',
            detalhes='Data final menor que data inicial.',
            contexto={'data_inicial': data_inicial, 'data_final': data_final},
        )
        messages.error(request, 'Data final deve ser maior ou igual a data inicial.')
        _set_pncp_import_progress(request, op_id, status='error', message='Falha: periodo invalido (data final menor que inicial).')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    modalidade_code = _pncp_to_int(codigo_modalidade)
    if modalidade_code is None:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - modalidade invalida',
            detalhes=f'Codigo de modalidade nao reconhecido: {codigo_modalidade}',
            contexto={'codigo_modalidade': codigo_modalidade},
        )
        messages.error(request, 'Codigo da modalidade invalido.')
        _set_pncp_import_progress(request, op_id, status='error', message='Falha: modalidade PNCP invalida.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)
    tamanho_pagina = max(1, min(200, int(tamanho_pagina)))
    max_paginas = max(1, min(100, int(max_paginas)))
    if codigo_municipio_ibge and not codigo_municipio_ibge.isdigit():
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - municipio IBGE',
            detalhes='Codigo do municipio IBGE deve conter apenas numeros.',
            contexto={'codigo_municipio_ibge': codigo_municipio_ibge},
        )
        messages.error(request, 'Codigo do municipio IBGE deve conter apenas numeros.')
        _set_pncp_import_progress(request, op_id, status='error', message='Falha: codigo de municipio IBGE invalido.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)
    if cnpj and len(cnpj) != 14:
        _set_pncp_error_popup(
            request,
            titulo='Falha de validacao - CNPJ do orgao',
            detalhes='CNPJ do orgao deve conter 14 digitos numericos.',
            contexto={'cnpj': cnpj},
        )
        messages.error(request, 'CNPJ do orgao deve conter 14 digitos.')
        _set_pncp_import_progress(request, op_id, status='error', message='Falha: CNPJ do orgao invalido.')
        return _ajax_redirect_payload(request, 'workflow:integracoes', ok=False)

    processo_log = (
        Processo.objects.order_by('-id').first()
        or Processo.objects.create(
            numero_processo_adm='PNCP-LOTE',
            ano_referencia=timezone.localdate().year,
            objeto='Importacao em lote PNCP',
            modalidade=_get_modalidade('Pregao'),
            status=_get_status('EM PLANEJAMENTO'),
            secretaria=_secretaria_adm_principal(),
        )
    )
    log = IntegracaoProcesso.objects.create(
        processo=processo_log,
        tipo=IntegracaoProcesso.Tipo.PNCP,
        identificador_externo=f'LOTE:{data_inicial}:{data_final}:{modalidade_code}',
    )

    client = PNCPClient(timeout=60, max_retries=2, rate_limit_retries=3)
    criados = 0
    atualizados = 0
    ignorados = 0
    paginas_lidas = 0
    retries_timeout = 0
    itens_detalhados_total = 0
    fornecedores_detalhados_total = 0
    erros_detalhes_total = 0
    processos_importados_map = {}
    erros = []
    ok_result = True

    try:
        for pagina in range(1, max_paginas + 1):
            _set_pncp_import_progress(
                request,
                op_id,
                status='running',
                message=f'Consultando pagina {pagina} de {max_paginas} no PNCP...',
                page_current=pagina,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
            )
            tamanho_pagina_efetivo = tamanho_pagina
            try:
                resposta = client.listar_publicacoes(
                    data_inicial=data_inicial,
                    data_final=data_final,
                    codigo_modalidade_contratacao=modalidade_code,
                    uf=uf or None,
                    codigo_municipio_ibge=codigo_municipio_ibge or None,
                    cnpj=cnpj or None,
                    pagina=pagina,
                    tamanho_pagina=tamanho_pagina_efetivo,
                )
            except Exception as exc_lote:
                erro_lote_texto = str(exc_lote).lower()
                timeout_detectado = ('timeout' in erro_lote_texto) or ('timed out' in erro_lote_texto)
                if timeout_detectado and tamanho_pagina > 10:
                    tamanho_pagina_efetivo = max(10, int(tamanho_pagina / 2))
                    retries_timeout += 1
                    resposta = client.listar_publicacoes(
                        data_inicial=data_inicial,
                        data_final=data_final,
                        codigo_modalidade_contratacao=modalidade_code,
                        uf=uf or None,
                        codigo_municipio_ibge=codigo_municipio_ibge or None,
                        cnpj=cnpj or None,
                        pagina=pagina,
                        tamanho_pagina=tamanho_pagina_efetivo,
                    )
                else:
                    if '429' in erro_lote_texto and paginas_lidas > 0:
                        erros.append(
                            f'Rate limit (429) na pagina {pagina}. '
                            f'Importacao encerrada como parcial. Detalhe: {exc_lote}'
                        )
                        _set_pncp_import_progress(
                            request,
                            op_id,
                            status='partial',
                            message=f'Importacao parcial por limite de requisicoes na pagina {pagina}.',
                            page_current=pagina,
                            paginas_lidas=paginas_lidas,
                            max_paginas=max_paginas,
                            criados=criados,
                            atualizados=atualizados,
                            ignorados=ignorados,
                        )
                        break
                    raise
            paginas_lidas += 1
            registros = resposta.get('data') if isinstance(resposta, dict) else []
            if not registros:
                _set_pncp_import_progress(
                    request,
                    op_id,
                    status='running',
                    message=f'Nenhum registro na pagina {pagina}. Encerrando busca.',
                    page_current=pagina,
                    paginas_lidas=paginas_lidas,
                    max_paginas=max_paginas,
                    criados=criados,
                    atualizados=atualizados,
                    ignorados=ignorados,
                )
                break
            if pagina < max_paginas:
                # Pequeno intervalo entre paginas para reduzir chance de 429.
                try:
                    import time as _time
                    _time.sleep(0.35)
                except Exception:
                    pass

            for payload in registros:
                try:
                    numero = _pncp_safe_text(payload.get('numeroControlePNCP'))
                    if not numero:
                        ignorados += 1
                        continue
                    processo = _pncp_detectar_processo_existente(numero, payload)
                    processo_existente = processo is not None
                    if processo_existente and importar_somente_novos:
                        ignorados += 1
                        continue

                    if not processo:
                        numero_proc = _pncp_safe_text(payload.get('processo')) or _pncp_safe_text(payload.get('numeroCompra')) or numero
                        ano_ref = _pncp_to_int(payload.get('anoCompra')) or timezone.localdate().year
                        objeto = _pncp_safe_text(payload.get('objetoCompra')) or 'Objeto importado do PNCP'
                        processo = Processo.objects.create(
                            numero_processo_adm=numero_proc,
                            ano_referencia=ano_ref,
                            objeto=objeto,
                            modalidade=_get_modalidade(_normalizar_modalidade_pncp(_pncp_safe_text(payload.get('modalidadeNome')) or 'Pregao')),
                            status=_get_status('EM PLANEJAMENTO'),
                            secretaria=_secretaria_adm_principal(),
                        )
                        criados += 1

                    itens_detalhados = []
                    resultados_detalhados = []
                    try:
                        itens_detalhados = client.consultar_itens(numero, strict=True)
                    except Exception as exc_itens:
                        erros_detalhes_total += 1
                        erros.append(f'[{numero}] itens: {exc_itens}')
                    try:
                        resultados_detalhados = client.consultar_resultados(numero, strict=True)
                    except Exception as exc_res:
                        erros_detalhes_total += 1
                        erros.append(f'[{numero}] resultados: {exc_res}')

                    if not itens_detalhados:
                        itens_detalhados = _pncp_extract_embedded_list(
                            payload,
                            ['sirel_pncp_itens', 'itens', 'items', 'itensCompra'],
                        )
                    if not resultados_detalhados:
                        resultados_detalhados = _pncp_extract_embedded_list(
                            payload,
                            ['sirel_pncp_resultados', 'resultados', 'resultado', 'resultadoItens', 'resultado-itens'],
                        )

                    payload_enriquecido = dict(payload)
                    if itens_detalhados:
                        payload_enriquecido['sirel_pncp_itens'] = itens_detalhados
                    if resultados_detalhados:
                        payload_enriquecido['sirel_pncp_resultados'] = resultados_detalhados

                    sync_detalhada = _sincronizar_itens_fornecedores_pncp(
                        processo,
                        itens_detalhados,
                        resultados_detalhados,
                        situacao_compra_nome=_pncp_safe_text(payload.get('situacaoCompraNome')),
                        numero_controle=numero,
                    )
                    itens_detalhados_total += int(sync_detalhada.get('itens_importados') or 0)
                    fornecedores_detalhados_total += int(sync_detalhada.get('fornecedores_atualizados') or 0)

                    snap = _upsert_pncp_snapshot(processo, numero, payload_enriquecido)
                    _sincronizar_processo_com_snapshot_pncp(processo, snap)
                    if processo_existente:
                        atualizados += 1
                    processo_key = str(processo.id)
                    processo_acao = 'Atualizado' if processo_existente else 'Criado'
                    processo_reg = {
                        'numero': f'{processo.numero_processo_adm}/{processo.ano_referencia}',
                        'objeto': processo.objeto or '',
                        'acao': processo_acao,
                    }
                    atual = processos_importados_map.get(processo_key)
                    if not atual:
                        processos_importados_map[processo_key] = processo_reg
                    elif atual.get('acao') != 'Criado' and processo_acao == 'Criado':
                        processos_importados_map[processo_key] = processo_reg
                except Exception as exc_item:
                    erros.append(str(exc_item))
                    continue

            _set_pncp_import_progress(
                request,
                op_id,
                status='running',
                message=f'Pagina {pagina} processada.',
                page_current=pagina,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
            )

        log.status = 'SUCESSO' if not erros else 'PARCIAL'
        log.payload_resumo = {
            'criados': criados,
            'atualizados': atualizados,
            'ignorados': ignorados,
            'paginas_lidas': paginas_lidas,
            'retries_timeout': retries_timeout,
            'itens_detalhados_importados': itens_detalhados_total,
            'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
            'erros_detalhes_total': erros_detalhes_total,
            'erros': erros[:20],
        }
        log.mensagem = (
            f'Importacao PNCP em lote concluida. Criados: {criados}, '
            f'Atualizados: {atualizados}, Ignorados: {ignorados}. '
            f'Paginas lidas: {paginas_lidas}. Retentativas por timeout: {retries_timeout}. '
            f'Itens detalhados importados: {itens_detalhados_total}.'
        )
        log.save(update_fields=['status', 'payload_resumo', 'mensagem'])
        if not criados and not atualizados and not erros:
            messages.warning(
                request,
                'Nenhum registro retornado pelo PNCP para os filtros informados. '
                'Revise periodo/modalidade e tente novamente.'
            )
            _set_pncp_import_progress(
                request,
                op_id,
                status='done',
                message='Importacao finalizada sem registros para os filtros informados.',
                page_current=paginas_lidas,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
            )
        elif erros:
            _set_pncp_error_popup(
                request,
                titulo='Importacao PNCP em lote com falhas parciais',
                detalhes=erros[:20],
                contexto={
                    'data_inicial': data_inicial,
                    'data_final': data_final,
                    'codigo_modalidade': modalidade_code,
                    'uf': uf or '-',
                    'codigo_municipio_ibge': codigo_municipio_ibge or '-',
                    'cnpj': cnpj or '-',
                    'tamanho_pagina': tamanho_pagina,
                    'max_paginas': max_paginas,
                    'paginas_lidas': paginas_lidas,
                    'retries_timeout': retries_timeout,
                    'itens_detalhados_importados': itens_detalhados_total,
                    'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
                    'criados': criados,
                    'atualizados': atualizados,
                    'ignorados': ignorados,
                },
            )
            messages.warning(request, log.mensagem + ' Houve erros em alguns registros.')
            _set_pncp_import_progress(
                request,
                op_id,
                status='partial',
                message='Importacao finalizada com falhas parciais.',
                page_current=paginas_lidas,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
            )
        else:
            _set_pncp_success_popup(
                request,
                titulo='Importacao PNCP em lote concluida com sucesso',
                processos=list(processos_importados_map.values()),
                contexto={
                    'processos_importados': len(processos_importados_map),
                    'criados': criados,
                    'atualizados': atualizados,
                    'ignorados': ignorados,
                    'paginas_lidas': paginas_lidas,
                    'itens_detalhados_importados': itens_detalhados_total,
                    'fornecedores_detalhados_atualizados': fornecedores_detalhados_total,
                },
            )
            messages.success(request, log.mensagem)
            _set_pncp_import_progress(
                request,
                op_id,
                status='done',
                message='Importacao finalizada com sucesso.',
                page_current=paginas_lidas,
                paginas_lidas=paginas_lidas,
                max_paginas=max_paginas,
                criados=criados,
                atualizados=atualizados,
                ignorados=ignorados,
                processos_importados=len(processos_importados_map),
            )
    except Exception as exc:
        ok_result = False
        log.status = 'ERRO'
        log.mensagem = str(exc)
        log.save(update_fields=['status', 'mensagem'])
        _set_pncp_error_popup(
            request,
            titulo='Falha na importacao PNCP em lote',
            detalhes=str(exc),
            contexto={
                'data_inicial': data_inicial,
                'data_final': data_final,
                'codigo_modalidade': modalidade_code,
                'uf': uf or '-',
                'codigo_municipio_ibge': codigo_municipio_ibge or '-',
                'cnpj': cnpj or '-',
                'tamanho_pagina': tamanho_pagina,
                'max_paginas': max_paginas,
                'paginas_lidas': paginas_lidas,
                'retries_timeout': retries_timeout,
            },
        )
        messages.error(request, f'Falha na importacao em lote do PNCP: {exc}')
        _set_pncp_import_progress(
            request,
            op_id,
            status='error',
            message=f'Falha na importacao: {exc}',
            page_current=paginas_lidas,
            paginas_lidas=paginas_lidas,
            max_paginas=max_paginas,
            criados=criados,
            atualizados=atualizados,
            ignorados=ignorados,
        )

    return _ajax_redirect_payload(request, 'workflow:integracoes', ok=ok_result)


def importar_bll(request):
    if request.method == 'POST':
        processo_id = request.POST.get('processo_id')
        processo = get_object_or_404(Processo, pk=processo_id)
        arquivo = request.FILES.get('arquivo')
        log = IntegracaoProcesso.objects.create(processo=processo, tipo=IntegracaoProcesso.Tipo.BLL_IMPORTACAO)
        if not arquivo:
            log.status = 'ERRO'
            log.mensagem = 'Nenhum arquivo BLL foi enviado.'
            log.save(update_fields=['status', 'mensagem'])
            messages.error(request, 'Selecione o arquivo da BLL para importar.')
            return redirect('workflow:integracoes')
        try:
            from core.utils.bll_import import import_bll_file

            resultado_importacao = import_bll_file(processo, arquivo)
            sync_canonico = sync_canonical_items_for_processo(processo)
            wf = _ensure_workflow(processo)
            wf.bll_ultima_importacao = timezone.now()
            wf.save(update_fields=['bll_ultima_importacao', 'atualizado_em'])
            log.status = 'SUCESSO'
            log.payload_resumo = {
                'importacao_bll': resultado_importacao,
                'sincronizacao_canonica': sync_canonico,
            }
            log.mensagem = (
                'Arquivo BLL processado e aplicado ao processo. '
                'Itens/ofertas/resultados foram sincronizados com a camada canonica.'
            )
            log.save(update_fields=['status', 'payload_resumo', 'mensagem'])
            messages.success(request, 'Arquivo BLL processado e sincronizado com sucesso.')
        except Exception as exc:
            log.status = 'ERRO'
            log.mensagem = str(exc)
            log.save(update_fields=['status', 'mensagem'])
            messages.error(request, f'Falha ao importar BLL: {exc}')
    return redirect('workflow:integracoes')


def _arquivo_base_processo(processo: Processo) -> str:
    base = (processo.numero_edital or processo.numero_processo_adm or f'processo_{processo.id}').strip()
    safe = ''.join(ch if ch.isalnum() or ch in ('-', '_') else '_' for ch in base)
    safe = safe.strip('_') or f'processo_{processo.id}'
    return f'{safe}_{int(processo.ano_referencia or timezone.localdate().year)}'


def processo_exportar_bll_csv(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    messages.error(request, 'Exportação BLL em CSV está desabilitada. Utilize apenas XLSX.')
    return redirect('workflow:processo_resumo', processo.id)


def processo_exportar_bll_xlsx(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    log = IntegracaoProcesso.objects.create(processo=processo, tipo=IntegracaoProcesso.Tipo.BLL_EXPORTACAO)
    try:
        from io import BytesIO

        from core.utils.bll_export import export_bll_xlsx

        buf = BytesIO()
        export_bll_xlsx(processo, buf)
        payload = buf.getvalue()
        log.status = 'SUCESSO'
        log.mensagem = 'Exportação BLL XLSX concluída.'
        log.payload_resumo = {
            'formato': 'xlsx',
            'processo_id': processo.id,
            'arquivo': f'{_arquivo_base_processo(processo)}.xlsx',
            'bytes': len(payload),
        }
        log.save(update_fields=['status', 'mensagem', 'payload_resumo'])
        response = HttpResponse(
            payload,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = f'attachment; filename="{_arquivo_base_processo(processo)}.xlsx"'
        return response
    except Exception as exc:
        log.status = 'ERRO'
        log.mensagem = str(exc)
        log.save(update_fields=['status', 'mensagem'])
        messages.error(request, f'Falha ao exportar XLSX BLL: {exc}')
        return redirect('workflow:processo_resumo', processo.id)


def processo_enviar_pncp(request, processo_id: int):
    if request.method != 'POST':
        return redirect('workflow:processo_resumo', processo_id)

    processo = get_object_or_404(Processo, pk=processo_id)
    log = IntegracaoProcesso.objects.create(processo=processo, tipo=IntegracaoProcesso.Tipo.PNCP, status='PENDENTE')
    try:
        client = PNCPPublishClient()
        resultado = client.enviar(processo)
        status_execucao = resultado.get('status', 'DRY_RUN')
        resumo = {
            'status': status_execucao,
            'mensagem': resultado.get('mensagem', ''),
            'config': resultado.get('config', {}),
            'contexto': resultado.get('contexto', {}),
            'totais': resultado.get('totais', {}),
        }
        if resultado.get('respostas'):
            resumo['respostas'] = resultado.get('respostas')

        log.payload_resumo = resumo
        if status_execucao in ('DRY_RUN', 'SKIPPED_DISABLED'):
            log.status = 'SUCESSO'
            log.mensagem = resultado.get('mensagem', 'Envio PNCP não executado (modo opcional).')
            if status_execucao == 'SKIPPED_DISABLED':
                messages.warning(request, log.mensagem)
            else:
                messages.info(request, log.mensagem)
        else:
            log.status = 'SUCESSO'
            log.mensagem = 'Envio PNCP concluído com sucesso.'
            messages.success(request, log.mensagem)
        log.save(update_fields=['status', 'mensagem', 'payload_resumo'])
    except Exception as exc:
        log.status = 'ERRO'
        log.mensagem = str(exc)
        log.save(update_fields=['status', 'mensagem'])
        messages.error(request, f'Falha no envio opcional ao PNCP: {exc}')

    return redirect('workflow:processo_resumo', processo.id)


def processo_resumo(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    modulo_anterior, modulo_proximo = _modulos_vizinhos(workflow.modulo_atual)
    labels_modulo = dict(ModuloSistema.choices)
    movimentos = ProcessoMovimentacao.objects.filter(processo=processo)[:30]
    integracoes = IntegracaoProcesso.objects.filter(processo=processo)[:20]
    itens = list(
        ProcessoItem.objects
        .filter(processo=processo)
        .select_related('fornecedor_homologado')
        .order_by('numero_item')[:200]
    )
    lote_por_item = {
        row['item_id']: row['lote__numero']
        for row in ProcessoLoteItem.objects.filter(processo=processo, ativo=True).values('item_id', 'lote__numero')
    }
    for item in itens:
        item.lote_numero = lote_por_item.get(item.id)
    lotes = Lote.objects.filter(processo=processo).order_by('numero')
    fornecedores_ids = set(
        ProcessoItem.objects.filter(processo=processo, fornecedor_homologado__isnull=False).values_list(
            'fornecedor_homologado_id', flat=True
        )
    )
    fornecedores_ids.update(
        ProcessoItemResultado.objects.filter(
            processo=processo,
            ativo=True,
            fornecedor__isnull=False,
        ).values_list('fornecedor_id', flat=True)
    )
    fornecedores = Fornecedor.objects.filter(id__in=list(fornecedores_ids)).order_by('razao_social')[:80]
    pncp_snapshot = PNCPContratacaoSnapshot.objects.filter(processo=processo).first()
    pncp_constantes = None
    pncp_mapping = None
    if pncp_snapshot:
        pncp_mapping = _pncp_mapping_report(pncp_snapshot, processo)
        pncp_constantes = pncp_mapping['constantes']
    form_doc_assinado = DocumentoAssinadoUploadForm()
    documentos_assinados = _listar_documentos_assinados(processo)[:30]
    fila_detalhamento = PNCPDetalhamentoFila.objects.filter(processo=processo).first()
    return render(request, 'workflow/processo_resumo.html', {
        'processo': processo,
        'workflow': workflow,
        'movimentos': movimentos,
        'integracoes': integracoes,
        'itens': itens,
        'lotes': lotes,
        'fornecedores': fornecedores,
        'modulo_anterior': modulo_anterior,
        'modulo_proximo': modulo_proximo,
        'modulo_anterior_label': labels_modulo.get(modulo_anterior, ''),
        'modulo_proximo_label': labels_modulo.get(modulo_proximo, ''),
        'form_doc_assinado': form_doc_assinado,
        'documentos_assinados': documentos_assinados,
        'is_modulo_compras': workflow.modulo_atual == ModuloSistema.COMPRAS,
        'is_modulo_licitacao': workflow.modulo_atual == ModuloSistema.LICITACAO,
        'pncp_snapshot': pncp_snapshot,
        'pncp_constantes': pncp_constantes,
        'pncp_mapping': pncp_mapping,
        'pncp_envio_habilitado': bool(getattr(settings, 'PNCP_ENVIO_HABILITADO', False)),
        'pncp_envio_dry_run': bool(getattr(settings, 'PNCP_ENVIO_DRY_RUN', True)),
        'fila_detalhamento': fila_detalhamento,
    })


def processo_comunicacao_interna(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    modulo_origem = workflow.modulo_atual if workflow.modulo_atual in dict(ModuloSistema.choices) else ModuloSistema.PLANEJAMENTO
    labels_modulo = dict(ModuloSistema.choices)
    modulo_anterior, modulo_proximo = _modulos_vizinhos(modulo_origem)
    destino_query = (request.GET.get('destino') or request.POST.get('destinatario_modulo') or '').strip().upper()

    numero_proc = f'{processo.numero_processo_adm}/{processo.ano_referencia}'
    responsavel_padrao = (request.user.get_full_name() or '').strip() or request.user.get_username()
    data_padrao = timezone.localdate()

    form_kwargs = {
        'modulo_origem': modulo_origem,
        'processo_atual': processo,
        'destino_inicial': destino_query,
    }

    if request.method == 'POST':
        form = ComunicacaoInternaForm(request.POST, **form_kwargs)
        if form.is_valid():
            action = (request.POST.get('action') or '').strip().lower()
            destino = form.cleaned_data['destinatario_modulo']
            if modulo_origem == ModuloSistema.COMPRAS and destino == ModuloSistema.LICITACAO:
                sd_status = _compras_sd_status(processo)
                if not sd_status['concluida']:
                    messages.error(
                        request,
                        'Para encaminhar para Licitacao, gere a SD (DOCX/PDF/XLSX) ou anexe a SD assinada.',
                    )
                    return redirect('workflow:compras_detail', processo.id)
            data_comunicacao = form.cleaned_data['data_comunicacao']
            ano_exercicio = data_comunicacao.year
            ref_processo = form.cleaned_data.get('processo_referencia') or processo
            signatarios = form.cleaned_data.get('signatarios')
            signatarios_lista = list(signatarios) if signatarios is not None else []
            signatarios_payload = [
                {
                    'nome': p.nome,
                    'cargo': str(getattr(p, 'cargo', '') or '').strip(),
                }
                for p in signatarios_lista
            ]
            signatarios_nomes = [p['nome'] for p in signatarios_payload]
            if action == 'preview_pdf':
                numero_preview = _proximo_numero_ci(modulo_origem, ano_exercicio)
                return export_ci_pdf_preview(
                    processo=processo,
                    modulo_origem_label=labels_modulo.get(modulo_origem, modulo_origem),
                    modulo_destino_label=labels_modulo.get(destino, destino),
                    numero_formatado=f'{numero_preview:03d}/{ano_exercicio}',
                    data_comunicacao=data_comunicacao,
                    assunto=form.cleaned_data['assunto'].strip(),
                    mensagem=form.cleaned_data['mensagem'].strip(),
                    referencia=f'{ref_processo.numero_processo_adm}/{ref_processo.ano_referencia} - {ref_processo.objeto or "-"}',
                    responsavel_envio=form.cleaned_data['responsavel_envio'].strip(),
                    signatarios=signatarios_payload,
                    observacao=(form.cleaned_data.get('observacao') or '').strip(),
                )

            ci = None
            for _ in range(5):
                try:
                    with transaction.atomic():
                        numero_seq = _proximo_numero_ci(modulo_origem, ano_exercicio)
                        ci = ComunicacaoInterna.objects.create(
                            processo=processo,
                            processo_referencia=ref_processo,
                            modulo_origem=modulo_origem,
                            modulo_destino=destino,
                            ano_exercicio=ano_exercicio,
                            numero_sequencial=numero_seq,
                            data_comunicacao=data_comunicacao,
                            assunto=form.cleaned_data['assunto'].strip(),
                            mensagem=form.cleaned_data['mensagem'].strip(),
                            observacao=(form.cleaned_data.get('observacao') or '').strip(),
                            responsavel_envio=form.cleaned_data['responsavel_envio'].strip(),
                            criado_por=request.user if request.user.is_authenticated else None,
                        )
                    break
                except IntegrityError:
                    ci = None
                    continue

            if not ci:
                messages.error(request, 'Nao foi possivel gerar o numero da comunicacao interna. Tente novamente.')
                return redirect('workflow:processo_comunicacao_interna', processo.id)

            if signatarios is not None:
                ci.signatarios.set(signatarios)

            workflow.modulo_atual = destino
            if destino == ModuloSistema.PLANEJAMENTO:
                workflow.etapa_atual = 'RETORNADO PARA AJUSTES'
                workflow.situacao = SituacaoWorkflow.EM_ANDAMENTO
            else:
                workflow.etapa_atual = 'AGUARDANDO RECEBIMENTO'
                workflow.situacao = SituacaoWorkflow.AGUARDANDO
            workflow.save(update_fields=['modulo_atual', 'etapa_atual', 'situacao', 'atualizado_em'])

            ref = ci.processo_referencia or processo
            ref_texto = f'{ref.numero_processo_adm}/{ref.ano_referencia}'
            descricao_mov = f'CI n. {ci.numero_formatado} gerada para {labels_modulo.get(destino, destino)}.'
            observacao_linhas = [
                f'De: {labels_modulo.get(modulo_origem, modulo_origem)}',
                f'Para: {labels_modulo.get(destino, destino)}',
                f'Data: {ci.data_comunicacao.strftime("%d/%m/%Y")}',
                f'Ref.: {ref_texto}',
                f'Assunto: {ci.assunto}',
                '',
                _rich_html_to_plain_text(ci.mensagem) or '-',
            ]
            if ci.observacao:
                observacao_linhas.extend(['', f'Observacoes: {ci.observacao}'])
            if signatarios_nomes:
                observacao_linhas.extend(['', 'Signatários: ' + ', '.join(signatarios_nomes)])

            ProcessoMovimentacao.objects.create(
                processo=processo,
                modulo_origem=modulo_origem,
                modulo_destino=destino,
                descricao=descricao_mov[:255],
                observacao='\n'.join(observacao_linhas).strip(),
            )

            messages.success(
                request,
                f'CI n. {ci.numero_formatado} gerada e processo encaminhado para {labels_modulo.get(destino, destino)}.',
            )
            return redirect('workflow:processo_resumo', processo.id)
    else:
        destino_inicial = destino_query or (modulo_proximo or modulo_anterior or '')
        assunto_padrao = f'Encaminhamento do processo {numero_proc}'
        mensagem_padrao = (
            f'Prezados,\n\n'
            f'Encaminhamos para conhecimento e providencias o processo administrativo n. {numero_proc}, '
            f'referente ao objeto: {processo.objeto or "-"}.\n\n'
            f'Atenciosamente.'
        )
        form = ComunicacaoInternaForm(
            initial={
                'destinatario_modulo': destino_inicial,
                'processo_referencia': processo.id,
                'data_comunicacao': data_padrao,
                'assunto': assunto_padrao,
                'mensagem': mensagem_padrao,
                'responsavel_envio': responsavel_padrao,
            },
            **form_kwargs,
        )

    data_ref = None
    try:
        data_ref = form['data_comunicacao'].value()
        if isinstance(data_ref, str) and data_ref:
            data_ref = datetime.fromisoformat(data_ref).date()
    except Exception:
        data_ref = None
    data_ref = data_ref or data_padrao
    numero_preview = _proximo_numero_ci(modulo_origem, data_ref.year)
    form_doc_assinado_ci = DocumentoAssinadoUploadForm(initial={'doc_key': 'ci'})
    docs_ci_assinados = [
        d for d in _listar_documentos_assinados(processo)
        if d['key'] == 'ci'
    ]
    selected_sign_vals = form['signatarios'].value() if form else []
    if isinstance(selected_sign_vals, str):
        selected_sign_vals = [selected_sign_vals]
    selected_sign_ids = [str(x).strip() for x in (selected_sign_vals or []) if str(x).strip()]
    sign_catalog_qs = list(form.fields['signatarios'].queryset.select_related('secretaria')) if form else []
    signatarios_catalogo = [
        {
            'id': pessoa.id,
            'nome': pessoa.nome,
            'cargo': str(getattr(pessoa, 'cargo', '') or '').strip(),
            'secretaria': str(getattr(getattr(pessoa, 'secretaria', None), 'sigla', '') or '').strip(),
        }
        for pessoa in sign_catalog_qs
    ]
    catalogo_por_id = {str(row['id']): row for row in signatarios_catalogo}
    signatarios_selecionados = [catalogo_por_id[sid] for sid in selected_sign_ids if sid in catalogo_por_id]

    return render(
        request,
        'workflow/comunicacao_interna.html',
        {
            'processo': processo,
            'workflow': workflow,
            'form': form,
            'orgao': _orgao_ativo(),
            'modulo_origem': modulo_origem,
            'modulo_origem_label': labels_modulo.get(modulo_origem, modulo_origem),
            'modulo_anterior': modulo_anterior,
            'modulo_proximo': modulo_proximo,
            'numero_ci_preview': f'{numero_preview:03d}/{data_ref.year}',
            'form_doc_assinado_ci': form_doc_assinado_ci,
            'docs_ci_assinados': docs_ci_assinados,
            'signatarios_catalogo': signatarios_catalogo,
            'signatarios_selecionados': signatarios_selecionados,
        },
    )


def processo_movimentar(request, processo_id: int, acao: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    modulo_atual = workflow.modulo_atual
    modulo_anterior, modulo_proximo = _modulos_vizinhos(modulo_atual if modulo_atual in dict(ModuloSistema.choices) else ModuloSistema.PLANEJAMENTO)

    acao_norm = (acao or '').strip().lower()
    if acao_norm == 'enviar':
        if not modulo_proximo:
            messages.error(request, 'Nao ha modulo seguinte para envio deste processo.')
            return redirect('workflow:processo_resumo', processo.id)
        destino = modulo_proximo
    elif acao_norm == 'devolver':
        if not modulo_anterior:
            messages.error(request, 'Nao ha modulo anterior para devolucao deste processo.')
            return redirect('workflow:processo_resumo', processo.id)
        destino = modulo_anterior
    else:
        messages.error(request, 'Acao de movimentacao invalida.')
        return redirect('workflow:processo_resumo', processo.id)

    return redirect(f"{reverse('workflow:processo_comunicacao_interna', args=[processo.id])}?destino={destino}&acao={acao_norm}")


def processo_upload_documento_assinado(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    workflow = _ensure_workflow(processo)
    next_url = (
        request.POST.get('next_url')
        or request.META.get('HTTP_REFERER')
        or reverse('workflow:processo_resumo', args=[processo.id])
    )
    if request.method != 'POST':
        return redirect(next_url)

    form = DocumentoAssinadoUploadForm(request.POST, request.FILES)
    if not form.is_valid():
        msg = f'Nao foi possivel anexar o documento assinado. {_form_errors_text(form)}'
        if _is_ajax(request):
            return JsonResponse({'ok': False, 'message': msg})
        messages.error(request, msg)
        return redirect(next_url)

    key = form.cleaned_data['doc_key']
    conf = DOC_ASSINADO_CONFIG.get(key)
    if not conf:
        msg = 'Tipo de documento assinado invalido.'
        if _is_ajax(request):
            return JsonResponse({'ok': False, 'message': msg})
        messages.error(request, msg)
        return redirect(next_url)

    arquivo = form.cleaned_data['arquivo']
    ordem = (
        DocumentoProcessoWorkflow.objects.filter(processo=processo)
        .aggregate(maior=Max('ordem_cronologica'))
        .get('maior')
        or 0
    ) + 1
    modulo_doc = conf.get('modulo') or workflow.modulo_atual or ModuloSistema.PLANEJAMENTO
    DocumentoProcessoWorkflow.objects.create(
        processo=processo,
        modulo=modulo_doc,
        tipo_documento=_doc_assinado_tipo(key),
        arquivo=arquivo,
        ordem_cronologica=ordem,
        gerar_no_etcm=True,
    )
    if key == 'sd':
        _atualizar_etapa_compras_sd(processo)
    msg = f"{conf['label']} assinada anexada com sucesso."
    if _is_ajax(request):
        return JsonResponse({'ok': True, 'message': msg})
    messages.success(request, msg)
    return redirect(next_url)


def _reprocessar_pncp_processo(processo: Processo):
    wf = _ensure_workflow(processo)
    snap_existente = PNCPContratacaoSnapshot.objects.filter(processo=processo).first()
    numero_controle = (
        _pncp_safe_text(getattr(wf, 'pncp_numero_controle', ''))
        or _pncp_safe_text(getattr(snap_existente, 'numero_controle_pncp', ''))
    )
    if not numero_controle:
        raise ValueError('Processo sem numero de controle PNCP vinculado.')

    payload_base = (
        snap_existente.payload_completo
        if snap_existente and isinstance(snap_existente.payload_completo, dict)
        else {}
    )
    payload_raw = {}
    erros = []
    client = PNCPClient(timeout=60, max_retries=2, rate_limit_retries=3)

    try:
        payload_raw = client.consultar(numero_controle)
    except Exception as exc:
        erros.append(f'consulta principal: {exc}')

    payload = _pncp_first_payload(payload_raw) if payload_raw else {}
    if not payload:
        payload = _pncp_first_payload(payload_base) if payload_base else {}
    if not isinstance(payload, dict):
        payload = {}

    itens_detalhados = []
    resultados_detalhados = []
    try:
        itens_detalhados = client.consultar_itens(numero_controle, strict=True)
    except Exception as exc_itens:
        erros.append(f'itens: {exc_itens}')
    try:
        resultados_detalhados = client.consultar_resultados(numero_controle, strict=True)
    except Exception as exc_res:
        erros.append(f'resultados: {exc_res}')

    if not itens_detalhados:
        for origem in (payload_raw, payload, payload_base):
            itens_detalhados = _pncp_extract_embedded_list(
                origem,
                ['sirel_pncp_itens', 'itens', 'items', 'itensCompra'],
            )
            if itens_detalhados:
                break

    if not resultados_detalhados:
        for origem in (payload_raw, payload, payload_base):
            resultados_detalhados = _pncp_extract_embedded_list(
                origem,
                ['sirel_pncp_resultados', 'resultados', 'resultado', 'resultadoItens', 'resultado-itens'],
            )
            if resultados_detalhados:
                break

    payload_enriquecido = {}
    if isinstance(payload_base, dict):
        payload_enriquecido.update(payload_base)
    payload_enriquecido.update(payload)
    if itens_detalhados:
        payload_enriquecido['sirel_pncp_itens'] = itens_detalhados
    if resultados_detalhados:
        payload_enriquecido['sirel_pncp_resultados'] = resultados_detalhados

    situacao_compra_nome = _pncp_safe_text(
        payload_enriquecido.get('situacaoCompraNome')
        or payload_base.get('situacaoCompraNome')
        or getattr(snap_existente, 'situacao_compra_nome', '')
    )
    sync_detalhada = _sincronizar_itens_fornecedores_pncp(
        processo,
        itens_detalhados,
        resultados_detalhados,
        situacao_compra_nome=situacao_compra_nome,
        numero_controle=numero_controle,
    )

    if payload_enriquecido:
        snap = _upsert_pncp_snapshot(processo, numero_controle, payload_enriquecido)
        _sincronizar_processo_com_snapshot_pncp(processo, snap)
    elif snap_existente:
        _sincronizar_processo_com_snapshot_pncp(processo, snap_existente)

    return {
        'numero_controle': numero_controle,
        'sync': sync_detalhada,
        'itens_detalhados': len(itens_detalhados or []),
        'resultados_detalhados': len(resultados_detalhados or []),
        'erros': erros,
    }


def processo_reprocessar_pncp(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    next_url = (
        request.POST.get('next_url')
        or request.META.get('HTTP_REFERER')
        or reverse('workflow:processo_resumo', args=[processo.id])
    )
    if request.method != 'POST':
        return redirect(next_url)

    wf = _ensure_workflow(processo)
    numero_controle = _pncp_safe_text(getattr(wf, 'pncp_numero_controle', ''))
    log = IntegracaoProcesso.objects.create(
        processo=processo,
        tipo=IntegracaoProcesso.Tipo.PNCP,
        identificador_externo=f'REPROCESSAR:{numero_controle or "-"}',
        status='PENDENTE',
    )
    try:
        resultado = _reprocessar_pncp_processo(processo)
        erros = resultado.get('erros') or []
        sync = resultado.get('sync') or {}
        log.status = 'PARCIAL' if erros else 'SUCESSO'
        log.payload_resumo = {
            'numeroControlePNCP': resultado.get('numero_controle', ''),
            'itens_detalhados': resultado.get('itens_detalhados', 0),
            'resultados_detalhados': resultado.get('resultados_detalhados', 0),
            'sincronizacao_detalhada': sync,
            'erros': erros[:20],
        }
        log.mensagem = (
            'Reprocessamento PNCP concluido.'
            if not erros else
            'Reprocessamento PNCP concluido com falhas parciais.'
        )
        log.save(update_fields=['status', 'payload_resumo', 'mensagem'])
        msg = (
            'Reprocessamento PNCP concluido. '
            f"Itens importados: {sync.get('itens_importados', 0)}, "
            f"resultados recebidos: {sync.get('resultados_recebidos', 0)}, "
            f"status forcados: {sync.get('status_forcados', 0)}."
        )
        if erros:
            msg += f" Falhas parciais: {len(erros)}."
        if _is_ajax(request):
            return JsonResponse({'ok': True, 'message': msg})
        if erros:
            messages.warning(request, msg)
        else:
            messages.success(request, msg)
        return redirect(next_url)
    except Exception as exc:
        log.status = 'ERRO'
        log.mensagem = str(exc)
        log.save(update_fields=['status', 'mensagem'])
        msg = f'Falha no reprocessamento PNCP: {exc}'
        if _is_ajax(request):
            return JsonResponse({'ok': False, 'message': msg})
        messages.error(request, msg)
        return redirect(next_url)


def planejamento_dashboard(request):
    processos_qs = ProcessoWorkflow.objects.filter(modulo_atual=ModuloSistema.PLANEJAMENTO).select_related('processo')

    busca = (request.GET.get('q') or '').strip()
    campo = (request.GET.get('campo') or 'todos').strip().lower()
    situacao = (request.GET.get('situacao') or '').strip().upper()
    situacoes_validas = {valor for valor, _ in SituacaoWorkflow.choices}
    if situacao and situacao not in situacoes_validas:
        situacao = ''

    if busca:
        filtros = Q()
        if campo == 'numero':
            numero = busca
            ano = ''
            if '/' in busca:
                numero, ano = (busca.split('/', 1) + [''])[:2]
            numero = numero.strip()
            ano = ano.strip()
            if numero:
                filtros |= (
                    Q(processo__numero_processo_sirel__icontains=numero)
                    | Q(processo__numero_processo_adm__icontains=numero)
                    | Q(processo__numero_edital__icontains=numero)
                )
            if ano.isdigit():
                filtros |= Q(processo__ano_referencia=int(ano))
        elif campo == 'objeto':
            filtros = Q(processo__objeto__icontains=busca)
        elif campo == 'nome':
            filtros = Q(processo__planejamento_dfd__responsavel_demanda__icontains=busca)
        else:
            filtros = (
                Q(processo__numero_processo_sirel__icontains=busca)
                | Q(processo__numero_processo_adm__icontains=busca)
                | Q(processo__numero_edital__icontains=busca)
                | Q(processo__objeto__icontains=busca)
                | Q(processo__planejamento_dfd__responsavel_demanda__icontains=busca)
            )
        processos_qs = processos_qs.filter(filtros)

    if situacao:
        processos_qs = processos_qs.filter(situacao=situacao)

    processos = processos_qs.order_by('-atualizado_em')
    processo_ids = processos.values_list('processo_id', flat=True)
    por_secretaria = (
        DFDSecretaria.objects
        .filter(dfd__processo_id__in=processo_ids)
        .values('secretaria__sigla')
        .annotate(total=Count('dfd_id', distinct=True))
        .order_by('-total', 'secretaria__sigla')[:10]
    )
    por_modalidade = (
        PlanejamentoDFD.objects
        .filter(processo_id__in=processo_ids)
        .values('modalidade_pretendida')
        .annotate(total=Count('id'))
        .order_by('-total')
    )
    return render(request, 'workflow/planejamento_dashboard.html', {
        'processos': processos,
        'por_secretaria': por_secretaria,
        'por_modalidade': por_modalidade,
        'filtro_q': busca,
        'filtro_campo': campo,
        'filtro_situacao': situacao,
        'situacoes_workflow': SituacaoWorkflow.choices,
    })


def planejamento_novo(request):
    ano_preview = timezone.localdate().year
    if request.method == 'POST':
        form = ProcessoPlanejamentoForm(request.POST)
        ano_raw = (request.POST.get('ano_referencia') or '').strip()
        if ano_raw.isdigit():
            ano_preview = int(ano_raw)
        if form.is_valid():
            modalidade = _get_modalidade('Pregão')
            status = _get_status('EM PLANEJAMENTO')
            numero_externo = (form.cleaned_data.get('numero_processo_externo') or '').strip()
            processo = Processo.objects.create(
                numero_processo_adm=numero_externo,
                ano_referencia=form.cleaned_data['ano_referencia'],
                objeto=form.cleaned_data['objeto'],
                modalidade=modalidade,
                status=status,
                secretaria=_secretaria_adm_principal(),
            )
            ProcessoWorkflow.objects.create(processo=processo, modulo_atual=ModuloSistema.PLANEJAMENTO, etapa_atual='DFD', situacao='EM_ANDAMENTO')
            PlanejamentoDFD.objects.create(processo=processo, objeto_resumido=form.cleaned_data['objeto'])
            ProcessoMovimentacao.objects.create(processo=processo, modulo_destino=ModuloSistema.PLANEJAMENTO, descricao='Processo criado no módulo de Planejamento.')
            messages.success(
                request,
                f'Processo criado no Planejamento. Número SIREL: {processo.numero_processo_sirel}.',
            )
            return redirect('workflow:planejamento_detail', processo.id)
    else:
        form = ProcessoPlanejamentoForm()
    proximo_numero_sirel = Processo.gerar_numero_processo_sirel(ano_preview)
    return render(
        request,
        'workflow/planejamento_novo.html',
        {
            'form': form,
            'proximo_numero_sirel': proximo_numero_sirel,
            'ano_preview': ano_preview,
        },
    )


def planejamento_detail(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    wf = _ensure_workflow(processo)
    pncp_snapshot = PNCPContratacaoSnapshot.objects.filter(processo=processo).first()
    pncp_constantes = _pncp_mapping_report(pncp_snapshot, processo)['constantes'] if pncp_snapshot else None
    dfd, _ = PlanejamentoDFD.objects.get_or_create(processo=processo, defaults={'objeto_resumido': processo.objeto or ''})
    etp = ETPPlanejamento.objects.filter(processo=processo).first()
    tr = TRPlanejamento.objects.filter(processo=processo).first()
    cotacoes_media = ETPCotacaoItem.objects.filter(etp__processo=processo, considerar_no_calculo=True).values('item__codigo', 'item__descricao').annotate(media=Avg('valor_unitario')).order_by('item__codigo')
    form_cotacao = ETPCotacaoItemForm()
    form_cotacao.fields['item'].queryset = dfd.itens.all().order_by('codigo')
    form_cotacao.fields['item'].empty_label = 'Selecione um item'
    form_cotacao.fields['fonte'].queryset = etp.fontes.all() if etp else ETPCotacaoFonte.objects.none()
    form_cotacao.fields['fonte'].empty_label = 'Selecione uma fonte'
    form_distribuicao = TRDistribuicaoSecretariaForm()
    form_distribuicao.fields['secretaria'].queryset = Secretaria.objects.filter(id__in=dfd.secretarias_vinculadas.values_list('secretaria_id', flat=True))
    form_distribuicao.fields['item'].queryset = dfd.itens.all()
    form_dotacao = TRDotacaoForm()
    dotacao_refs_vazias = {
        'secretaria': not form_dotacao.fields['secretaria'].queryset.exists(),
        'unidade_orcamentaria': not form_dotacao.fields['unidade_orcamentaria'].queryset.exists(),
        'projeto_atividade': not form_dotacao.fields['projeto_atividade'].queryset.exists(),
        'elemento_despesa': not form_dotacao.fields['elemento_despesa'].queryset.exists(),
        'fonte_recurso': not form_dotacao.fields['fonte_recurso'].queryset.exists(),
    }
    tr_criterio = tr.criterio_julgamento if tr else TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_ITEM
    lotes_status = _status_lotes_tr(processo, tr) if tr else {
        'itens_total': dfd.itens.count(),
        'nao_alocados': list(dfd.itens.all().order_by('codigo')),
        'duplicados': [],
        'todos_alocados': False,
        'sem_duplicidade': True,
    }
    itens_lote_builder = [
        {
            'id': item.id,
            'codigo': item.codigo,
            'descricao': item.descricao,
            'unidade': item.unidade or '',
        }
        for item in dfd.itens.all().order_by('codigo')
    ]
    lotes_lote_builder = []
    itens_ocupados_lotes = set()
    if tr:
        for lote in tr.lotes.prefetch_related('itens').all().order_by('numero'):
            item_ids_lote = list(lote.itens.values_list('id', flat=True))
            lotes_lote_builder.append(
                {
                    'id': lote.id,
                    'numero': lote.numero,
                    'titulo': lote.titulo,
                    'item_ids': item_ids_lote,
                }
            )
            itens_ocupados_lotes.update(item_ids_lote)
    itens_disponiveis_lote_ids = [
        item['id'] for item in itens_lote_builder if item['id'] not in itens_ocupados_lotes
    ]
    distribuicao_ctx = _mapa_distribuicao_context(processo, tr) if tr else {'linhas': [], 'resumo_itens': []}
    return render(request, 'workflow/planejamento_detail.html', {
        'processo': processo, 'workflow': wf, 'dfd': dfd, 'etp': etp, 'tr': tr,
        'cotacoes_media': cotacoes_media,
        'form_dfd': DFDForm(instance=dfd, initial={'secretarias': [x.secretaria_id for x in dfd.secretarias_vinculadas.all()]}),
        'pessoas_autocomplete_url': reverse('workflow:planejamento_pessoas_autocomplete'),
        'catalogo_autocomplete_url': reverse('workflow:planejamento_catalogo_autocomplete'),
        'catalogo_unidades': list(
            DFDItemCatalogo.objects.exclude(unidade='').values_list('unidade', flat=True).distinct().order_by('unidade')
        ),
        'dfd_preview_popup_url': reverse('workflow:planejamento_preview_embed', args=[processo.id, 'dfd']),
        'mapa_preview_popup_url': reverse('workflow:planejamento_preview_embed', args=[processo.id, 'mapa']),
        'distribuicao_preview_popup_url': reverse('workflow:planejamento_preview_embed', args=[processo.id, 'distribuicao']),
        'form_catalogo': DFDItemCatalogoForm(),
        'form_etp': ETPForm(instance=etp),
        'form_fonte': ETPCotacaoFonteForm(),
        'form_cotacao': form_cotacao,
        'form_tr': TRForm(instance=tr),
        'form_lote': TRLoteForm(),
        'form_dotacao': form_dotacao,
        'dotacao_refs_vazias': dotacao_refs_vazias,
        'etp_havera_irp': bool(etp and etp.havera_irp),
        'tr_criterio': tr_criterio,
        'tr_lotes_status': lotes_status,
        'lote_builder_items': itens_lote_builder,
        'lote_builder_lotes': lotes_lote_builder,
        'lote_builder_available_ids': itens_disponiveis_lote_ids,
        'distribuicao_linhas': distribuicao_ctx['linhas'],
        'distribuicao_resumo_itens': distribuicao_ctx['resumo_itens'],
        'dotacao_ref_tipos': [
            ('SECRETARIA', 'Secretaria'),
            ('UNIDADE_ORCAMENTARIA', 'Unidade orçamentária'),
            ('PROJETO_ATIVIDADE', 'Projeto/Atividade'),
            ('ELEMENTO_DESPESA', 'Elemento de despesa'),
            ('FONTE_RECURSO', 'Fonte de recurso'),
        ],
        'form_distribuicao': form_distribuicao,
        'form_doc_assinado': DocumentoAssinadoUploadForm(),
        'documentos_assinados': _listar_documentos_assinados(processo)[:12],
        'pncp_snapshot': pncp_snapshot,
        'pncp_constantes': pncp_constantes,
    })


def planejamento_salvar_dfd(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    dfd, _ = PlanejamentoDFD.objects.get_or_create(processo=processo, defaults={'objeto_resumido': processo.objeto or ''})
    form = DFDForm(request.POST, instance=dfd)
    if form.is_valid():
        dfd = form.save()
        processo.objeto = dfd.objeto_resumido
        processo.modalidade = _get_modalidade(dfd.get_modalidade_pretendida_display())
        processo.save(update_fields=['objeto', 'modalidade'])
        DFDSecretaria.objects.filter(dfd=dfd).delete()
        secretarias = list(form.cleaned_data['secretarias'])
        principal = _secretaria_adm_principal() if len(secretarias) > 1 else (secretarias[0] if secretarias else _secretaria_adm_principal())
        if principal and principal not in secretarias:
            secretarias = [principal] + secretarias
        for s in secretarias:
            DFDSecretaria.objects.get_or_create(dfd=dfd, secretaria=s, defaults={'principal': s == principal})
        DFDSecretaria.objects.filter(dfd=dfd, secretaria=principal).update(principal=True)
        processo.secretaria = principal
        processo.save(update_fields=['secretaria'])
        wf = _ensure_workflow(processo)
        wf.etapa_atual = 'DFD'
        wf.situacao = 'EM_ANDAMENTO'
        wf.save(update_fields=['etapa_atual', 'situacao', 'atualizado_em'])
        _sincronizar_itens_core(processo)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'DFD salva com sucesso.'
    else:
        ok = False
        msg = f'Revise os campos da DFD. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_item(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    dfd = processo.planejamento_dfd
    form = DFDItemForm(request.POST)
    if form.is_valid():
        usar_catalogo = form.cleaned_data.get('usar_catalogo')
        codigo = (DFDItemCatalogo.objects.order_by('-codigo').first().codigo + 1) if DFDItemCatalogo.objects.exists() else 1
        descricao = form.cleaned_data['descricao']
        unidade = form.cleaned_data['unidade']
        catalogo = None
        if usar_catalogo:
            catalogo = usar_catalogo
            codigo = usar_catalogo.codigo
            descricao = usar_catalogo.descricao
            unidade = usar_catalogo.unidade
        else:
            catalogo = DFDItemCatalogo.objects.create(codigo=codigo, descricao=descricao, unidade=unidade)
        item_dfd = DFDItem.objects.create(
            dfd=dfd,
            catalogo=catalogo,
            codigo=codigo,
            descricao=descricao,
            unidade=unidade,
            quantidade=form.cleaned_data['quantidade'],
        )
        _sincronizar_item_core(processo, item_dfd)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'Item adicionado à DFD.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível adicionar o item. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_itens_catalogo(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    dfd = processo.planejamento_dfd
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    payload_raw = (request.POST.get('itens_payload') or '').strip()
    if not payload_raw:
        return _ajax_or_redirect(
            request,
            ok=False,
            message='Nenhum item foi enviado para adicao.',
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )

    try:
        payload = json.loads(payload_raw)
    except Exception:
        payload = None

    if not isinstance(payload, list):
        return _ajax_or_redirect(
            request,
            ok=False,
            message='Formato invalido dos itens selecionados.',
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )

    adicionados = 0
    atualizados = 0
    erros = []
    for entry in payload:
        try:
            catalogo_id = int(entry.get('catalogo_id'))
        except Exception:
            erros.append('Item sem identificador valido.')
            continue
        catalogo_item = DFDItemCatalogo.objects.filter(pk=catalogo_id).first()
        if not catalogo_item:
            erros.append(f'Item de catalogo {catalogo_id} nao encontrado.')
            continue

        unidade = (entry.get('unidade') or catalogo_item.unidade or '').strip()[:30]
        quantidade_raw = str(entry.get('quantidade') or '1').replace(',', '.').strip()
        try:
            quantidade = Decimal(quantidade_raw)
        except (InvalidOperation, ValueError):
            erros.append(f'Quantidade invalida para o item {catalogo_item.codigo}.')
            continue
        if quantidade <= 0:
            erros.append(f'Quantidade deve ser maior que zero para o item {catalogo_item.codigo}.')
            continue

        item_dfd = DFDItem.objects.filter(dfd=dfd, codigo=catalogo_item.codigo).first()
        if item_dfd:
            item_dfd.catalogo = catalogo_item
            item_dfd.descricao = catalogo_item.descricao
            item_dfd.unidade = unidade
            item_dfd.quantidade = quantidade
            item_dfd.save(update_fields=['catalogo', 'descricao', 'unidade', 'quantidade'])
            atualizados += 1
        else:
            item_dfd = DFDItem.objects.create(
                dfd=dfd,
                catalogo=catalogo_item,
                codigo=catalogo_item.codigo,
                descricao=catalogo_item.descricao,
                unidade=unidade,
                quantidade=quantidade,
            )
            adicionados += 1
        _sincronizar_item_core(processo, item_dfd)

    if adicionados or atualizados:
        _sincronizar_lotes_core(processo)
        sufixo = ''
        if erros:
            sufixo = f' ({len(erros)} pendencia(s): ' + '; '.join(erros[:3]) + ')'
        return _ajax_or_redirect(
            request,
            ok=True,
            message=f'Itens sincronizados: {adicionados} adicionados e {atualizados} atualizados{sufixo}.',
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )

    return _ajax_or_redirect(
        request,
        ok=False,
        message='Nenhum item foi processado. ' + ('; '.join(erros) if erros else ''),
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_editar_item(request, processo_id: int, item_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    item = get_object_or_404(DFDItem, pk=item_id, dfd__processo=processo)
    form = DFDItemEdicaoForm(request.POST, instance=item)
    if form.is_valid():
        item = form.save()
        _sincronizar_item_core(processo, item)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'Item da DFD atualizado com sucesso.'
    else:
        ok = False
        msg = f'Não foi possível atualizar o item. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_excluir_item(request, processo_id: int, item_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    item = get_object_or_404(DFDItem, pk=item_id, dfd__processo=processo)
    numero_item = item.codigo
    item.delete()
    _remover_item_core(processo, numero_item)
    _sincronizar_lotes_core(processo)
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Item da DFD excluído com sucesso.',
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_criar_catalogo(request, processo_id: int):
    form = DFDItemCatalogoForm(request.POST)
    if form.is_valid():
        codigo = (DFDItemCatalogo.objects.order_by('-codigo').first().codigo + 1) if DFDItemCatalogo.objects.exists() else 1
        obj = form.save(commit=False)
        obj.codigo = codigo
        obj.save()
        ok = True
        msg = f'Item de catálogo criado com código {codigo}.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível criar o item no catálogo. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo_id],
    )


def planejamento_salvar_etp(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ETPForm(request.POST, instance=etp)
    if form.is_valid():
        etp = form.save()
        for item in processo.planejamento_dfd.itens.all():
            _recalcular_alertas_cotacoes(etp, item)
        wf = _ensure_workflow(processo)
        wf.etapa_atual = 'ETP/COTAÇÕES'
        wf.irp_aplicavel = etp.havera_irp
        wf.save(update_fields=['etapa_atual', 'irp_aplicavel', 'atualizado_em'])
        ok = True
        msg = 'ETP salvo com sucesso.'
    else:
        ok = False
        msg = f'Revise os campos do ETP. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_fonte(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ETPCotacaoFonteForm(request.POST)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.etp = etp
        obj.save()
        _sincronizar_fornecedor_fonte(obj)
        ok = True
        msg = 'Fonte de cotação adicionada.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível adicionar a fonte. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_cotacao(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    etp, _ = ETPPlanejamento.objects.get_or_create(processo=processo)
    form = ETPCotacaoItemForm(request.POST)
    form.fields['item'].queryset = processo.planejamento_dfd.itens.all().order_by('codigo')
    form.fields['item'].empty_label = 'Selecione um item'
    form.fields['fonte'].queryset = etp.fontes.all()
    form.fields['fonte'].empty_label = 'Selecione uma fonte'
    if form.is_valid():
        obj = form.save(commit=False)
        obj.etp = etp
        obj.save()
        _recalcular_alertas_cotacoes(etp, obj.item)
        ok = True
        msg = 'Cotação registrada e faixas de alerta recalculadas.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível registrar a cotação. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_editar_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    form = ETPCotacaoEdicaoForm(request.POST, instance=cotacao)
    if form.is_valid():
        cotacao = form.save()
        _recalcular_alertas_cotacoes(cotacao.etp, cotacao.item)
        ok = True
        msg = 'Cotação atualizada com sucesso.'
    else:
        ok = False
        msg = f'Não foi possível atualizar a cotação. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_alterar_consideracao_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    considerar = request.POST.get('considerar_no_calculo') in {'on', '1', 'true', 'True', 'sim', 'Sim'}
    cotacao.considerar_no_calculo = considerar
    cotacao.save(update_fields=['considerar_no_calculo'])
    _recalcular_alertas_cotacoes(cotacao.etp, cotacao.item)

    msg = 'Cotação marcada para considerar no cálculo.' if considerar else 'Cotação marcada para não considerar no cálculo.'
    return _ajax_or_redirect(
        request,
        ok=True,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_excluir_cotacao(request, processo_id: int, cotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    cotacao = get_object_or_404(ETPCotacaoItem, pk=cotacao_id, etp__processo=processo)
    etp = cotacao.etp
    item = cotacao.item
    cotacao.delete()
    _recalcular_alertas_cotacoes(etp, item)
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Cotação excluída com sucesso.',
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_salvar_tr(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    form = TRForm(request.POST, request.FILES, instance=tr)
    if form.is_valid():
        tr = form.save()
        avisos = []
        etp = ETPPlanejamento.objects.filter(processo=processo).first()
        if not (etp and etp.havera_irp):
            if tr.arquivo_irp_pdf:
                try:
                    tr.arquivo_irp_pdf.delete(save=False)
                except Exception:
                    pass
            tr.arquivo_irp_pdf = None
            tr.save(update_fields=['arquivo_irp_pdf'])

        wf = _ensure_workflow(processo)
        wf.etapa_atual = 'TERMO DE REFERENCIA'
        wf.divisao_por_secretaria = not tr.nao_aplica_divisao_secretaria
        if tr.criterio_julgamento in [
            TRPlanejamento.CriterioJulgamento.MENOR_PRECO_GLOBAL,
            TRPlanejamento.CriterioJulgamento.MAIOR_PERCENTUAL_DESCONTO,
            TRPlanejamento.CriterioJulgamento.MENOR_TAXA_ADMINISTRATIVA,
        ]:
            tr.permite_cota_reservada = False
            tr.save(update_fields=['permite_cota_reservada'])

        if tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_ITEM:
            if tr.lotes.exists():
                tr.lotes.all().delete()
                avisos.append('Lotes desabilitados para menor preco por item.')
        elif tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_GLOBAL:
            tr.lotes.all().delete()
            lote_unico = TRLote.objects.create(tr=tr, numero=1, titulo='Lote unico (global)')
            lote_unico.itens.set(processo.planejamento_dfd.itens.all())
            avisos.append('Lote unico gerado automaticamente com todos os itens da DFD.')
        elif tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE:
            status_lotes = _status_lotes_tr(processo, tr)
            pendentes = len(status_lotes['nao_alocados'])
            duplicados = len(status_lotes['duplicados'])
            if pendentes or duplicados:
                avisos.append(
                    f'Criterio por lote ativo: faltam alocar {pendentes} item(ns) e ha {duplicados} item(ns) duplicado(s) em lotes.'
                )

        if tr.nao_aplica_divisao_secretaria and tr.distribuicoes.exists():
            tr.distribuicoes.all().delete()
            avisos.append('Distribuicao por secretaria desabilitada: lancamentos removidos.')

        processo.criterio_julgamento = {
            TRPlanejamento.CriterioJulgamento.MENOR_PRECO_GLOBAL: 'MENOR_PRECO_GLOBAL',
            TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_ITEM: 'MENOR_PRECO_POR_ITEM',
            TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE: 'MENOR_PRECO_POR_LOTE',
            TRPlanejamento.CriterioJulgamento.MAIOR_PERCENTUAL_DESCONTO: 'MAIOR_DESCONTO',
            TRPlanejamento.CriterioJulgamento.MENOR_TAXA_ADMINISTRATIVA: 'MENOR_PRECO_GLOBAL',
        }[tr.criterio_julgamento]
        processo.escopo_disputa = (
            'LOTE'
            if tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE
            else (
                'GLOBAL'
                if tr.criterio_julgamento in [
                    TRPlanejamento.CriterioJulgamento.MENOR_PRECO_GLOBAL,
                    TRPlanejamento.CriterioJulgamento.MAIOR_PERCENTUAL_DESCONTO,
                    TRPlanejamento.CriterioJulgamento.MENOR_TAXA_ADMINISTRATIVA,
                ]
                else 'ITEM'
            )
        )
        processo.save(update_fields=['criterio_julgamento', 'escopo_disputa'])
        wf.save(update_fields=['etapa_atual', 'divisao_por_secretaria', 'atualizado_em'])
        _sincronizar_itens_core(processo)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'Termo de Referencia salvo com sucesso.'
        if avisos:
            msg = msg + ' ' + ' '.join(avisos)
    else:
        ok = False
        msg = f'Revise os dados do Termo de Referencia. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_lote(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    if tr.criterio_julgamento != TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE:
        if tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_ITEM:
            msg = 'Criacao de lotes desabilitada para criterio menor preco por item.'
        else:
            msg = 'Para este criterio de julgamento, a alocacao de itens em lotes nao e manual.'
        return _ajax_or_redirect(
            request,
            ok=False,
            message=msg,
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )

    form = TRLoteForm(request.POST)
    if form.is_valid():
        item_ids = _parse_lote_item_ids(request)
        if not item_ids:
            return _ajax_or_redirect(
                request,
                ok=False,
                message='Selecione ao menos um item para criar o lote.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )
        itens_selecionados = list(processo.planejamento_dfd.itens.filter(id__in=item_ids).distinct())
        itens_ja_loteados = DFDItem.objects.filter(id__in=item_ids, lotes_tr__tr=tr).distinct()
        if itens_ja_loteados.exists():
            codigos = ', '.join(str(i.codigo) for i in itens_ja_loteados.order_by('codigo'))
            return _ajax_or_redirect(
                request,
                ok=False,
                message=f'Os itens {codigos} ja estao vinculados a outro lote.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )
        if len(itens_selecionados) != len(item_ids):
            return _ajax_or_redirect(
                request,
                ok=False,
                message='Um ou mais itens selecionados nao pertencem a DFD do processo.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )

        numero = (tr.lotes.order_by('-numero').first().numero + 1) if tr.lotes.exists() else 1
        lote = form.save(commit=False)
        lote.tr = tr
        lote.numero = numero
        lote.save()
        lote.itens.set(itens_selecionados)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'Lote criado com sucesso.'
        status_lotes = _status_lotes_tr(processo, tr)
        if status_lotes['nao_alocados']:
            msg += f" Ainda faltam alocar {len(status_lotes['nao_alocados'])} item(ns)."
    else:
        ok = False
        msg = f'Não foi possível criar o lote. {_form_errors_text(form)}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_editar_lote(request, processo_id: int, lote_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    lote = get_object_or_404(TRLote, pk=lote_id, tr=tr)

    if tr.criterio_julgamento != TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE:
        return _ajax_or_redirect(
            request,
            ok=False,
            message='Edicao manual de lotes disponivel apenas para criterio por lote.',
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )

    form = TRLoteForm(request.POST, instance=lote)
    if form.is_valid():
        item_ids = _parse_lote_item_ids(request)
        if not item_ids:
            return _ajax_or_redirect(
                request,
                ok=False,
                message='Selecione ao menos um item no lote.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )

        itens_selecionados = list(processo.planejamento_dfd.itens.filter(id__in=item_ids).distinct())
        if len(itens_selecionados) != len(item_ids):
            return _ajax_or_redirect(
                request,
                ok=False,
                message='Um ou mais itens selecionados nao pertencem a DFD do processo.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )

        itens_ja_loteados = (
            DFDItem.objects.filter(id__in=item_ids, lotes_tr__tr=tr)
            .exclude(lotes_tr=lote)
            .distinct()
        )
        if itens_ja_loteados.exists():
            codigos = ', '.join(str(i.codigo) for i in itens_ja_loteados.order_by('codigo'))
            return _ajax_or_redirect(
                request,
                ok=False,
                message=f'Os itens {codigos} ja estao vinculados a outro lote.',
                redirect_name='workflow:planejamento_detail',
                redirect_args=[processo.id],
            )

        lote = form.save()
        lote.itens.set(itens_selecionados)
        _sincronizar_lotes_core(processo)
        ok = True
        msg = 'Lote atualizado com sucesso.'
        status_lotes = _status_lotes_tr(processo, tr)
        if status_lotes['nao_alocados']:
            msg += f" Ainda faltam alocar {len(status_lotes['nao_alocados'])} item(ns)."
    else:
        ok = False
        msg = f'Nao foi possivel atualizar o lote. {_form_errors_text(form)}'

    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_excluir_lote(request, processo_id: int, lote_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    lote = get_object_or_404(TRLote, pk=lote_id, tr=tr)
    lote.delete()
    _sincronizar_lotes_core(processo)
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Lote excluido com sucesso. Os itens voltaram para a lista de disponiveis.',
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_dotacao(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    form = TRDotacaoForm(request.POST)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.tr = tr
        obj.save()
        ok = True
        msg = 'Dotação adicionada.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível adicionar a dotação. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_excluir_dotacao(request, processo_id: int, dotacao_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    dotacao = get_object_or_404(TRDotacao, pk=dotacao_id, tr__processo=processo)
    dotacao.delete()
    return _ajax_or_redirect(
        request,
        ok=True,
        message='Dotacao excluida com sucesso.',
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_criar_referencia_dotacao(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    if request.method != 'POST':
        return redirect('workflow:planejamento_detail', processo.id)

    tipo = (request.POST.get('tipo_referencia') or '').strip().upper()
    codigo = (request.POST.get('codigo') or '').strip()
    descricao = (request.POST.get('descricao') or '').strip()
    sigla = (request.POST.get('sigla') or '').strip().upper()
    nome = (request.POST.get('nome') or '').strip()

    try:
        if tipo == 'SECRETARIA':
            if not sigla or not nome:
                raise ValueError('Informe sigla e nome da secretaria.')
            obj, criado = Secretaria.objects.get_or_create(sigla=sigla, defaults={'nome': nome})
            if not criado and nome and obj.nome != nome:
                obj.nome = nome
                obj.save(update_fields=['nome'])
            ok = True
            msg = f'Secretaria {sigla} salva com sucesso.'
        elif tipo == 'UNIDADE_ORCAMENTARIA':
            if not sigla or not nome:
                raise ValueError('Informe sigla e nome da unidade orçamentária.')
            obj, criado = UnidadeOrcamentaria.objects.get_or_create(sigla=sigla, defaults={'nome': nome})
            if not criado and nome and obj.nome != nome:
                obj.nome = nome
                obj.save(update_fields=['nome'])
            ok = True
            msg = f'Unidade orçamentária {sigla} salva com sucesso.'
        elif tipo == 'PROJETO_ATIVIDADE':
            if not codigo or not descricao:
                raise ValueError('Informe código e descrição do projeto/atividade.')
            obj, criado = ProjetoAtividade.objects.get_or_create(codigo=codigo, defaults={'descricao': descricao})
            if not criado and descricao and obj.descricao != descricao:
                obj.descricao = descricao
                obj.save(update_fields=['descricao'])
            ok = True
            msg = f'Projeto/Atividade {codigo} salvo com sucesso.'
        elif tipo == 'ELEMENTO_DESPESA':
            if not codigo or not descricao:
                raise ValueError('Informe código e descrição do elemento de despesa.')
            obj, criado = ElementoDespesa.objects.get_or_create(codigo=codigo, defaults={'descricao': descricao})
            if not criado and descricao and obj.descricao != descricao:
                obj.descricao = descricao
                obj.save(update_fields=['descricao'])
            ok = True
            msg = f'Elemento de despesa {codigo} salvo com sucesso.'
        elif tipo == 'FONTE_RECURSO':
            if not codigo or not descricao:
                raise ValueError('Informe código e descrição da fonte de recurso.')
            obj, criado = FonteRecurso.objects.get_or_create(codigo=codigo, defaults={'descricao': descricao})
            if not criado and descricao and obj.descricao != descricao:
                obj.descricao = descricao
                obj.save(update_fields=['descricao'])
            ok = True
            msg = f'Fonte de recurso {codigo} salva com sucesso.'
        else:
            raise ValueError('Tipo de referência inválido.')
    except Exception as exc:
        ok = False
        msg = f'Não foi possível salvar a referência de dotação. {exc}'

    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_adicionar_distribuicao(request, processo_id: int):
    processo = get_object_or_404(Processo, pk=processo_id)
    tr, _ = TRPlanejamento.objects.get_or_create(processo=processo)
    if tr.nao_aplica_divisao_secretaria:
        return _ajax_or_redirect(
            request,
            ok=False,
            message='Distribuicao por secretaria esta desabilitada no TR.',
            redirect_name='workflow:planejamento_detail',
            redirect_args=[processo.id],
        )
    form = TRDistribuicaoSecretariaForm(request.POST)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.tr = tr
        obj.save()
        ok = True
        msg = 'Distribuição por secretaria adicionada.'
    else:
        detalhe = _form_errors_text(form)
        ok = False
        msg = f'Não foi possível registrar a distribuição. {detalhe}'
    return _ajax_or_redirect(
        request,
        ok=ok,
        message=msg,
        redirect_name='workflow:planejamento_detail',
        redirect_args=[processo.id],
    )


def planejamento_encaminhar(request, processo_id: int, destino: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    wf = _ensure_workflow(processo)
    tr = TRPlanejamento.objects.filter(processo=processo).first()
    if tr and tr.criterio_julgamento == TRPlanejamento.CriterioJulgamento.MENOR_PRECO_POR_LOTE:
        status_lotes = _status_lotes_tr(processo, tr)
        if status_lotes['duplicados'] or status_lotes['nao_alocados']:
            pendentes = len(status_lotes['nao_alocados'])
            duplicados = len(status_lotes['duplicados'])
            messages.error(
                request,
                f'Nao e possivel encaminhar: faltam {pendentes} item(ns) sem lote e ha {duplicados} item(ns) em lotes duplicados.',
            )
            return redirect('workflow:planejamento_detail', processo.id)

    destino = destino.upper()
    if destino not in [ModuloSistema.COMPRAS, ModuloSistema.LICITACAO]:
        messages.error(request, 'Destino inválido.')
        return redirect('workflow:planejamento_detail', processo.id)
    return redirect(f"{reverse('workflow:processo_comunicacao_interna', args=[processo.id])}?destino={destino}&origem=planejamento")


def planejamento_exportar(request, processo_id: int, doc: str, formato: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    doc = doc.lower(); formato = formato.lower()
    if doc == 'dfd':
        if formato == 'html':
            dfd, _ = PlanejamentoDFD.objects.get_or_create(processo=processo, defaults={'objeto_resumido': processo.objeto or ''})
            context = {
                'processo': processo,
                'dfd': dfd,
                'previsao_entrega_execucao_formatada': _format_date_label(dfd.previsao_entrega_execucao),
                'orgao': _orgao_ativo(),
                **_atendente_context(request.user),
            }
            return render(request, 'workflow/dfd_preview.html', context)
        if formato == 'docx': return export_dfd_docx(processo, request.user)
        if formato == 'pdf': return export_dfd_pdf(processo, request.user)
        if formato == 'xlsx': return export_dfd_xlsx(processo)
    if doc == 'mapa':
        if formato == 'html':
            return render(request, 'workflow/mapa_preview.html', {'processo': processo, 'orgao': _orgao_ativo(), **_mapa_preview_context(processo)})
        if formato == 'docx':
            return export_mapa_docx(processo)
        if formato == 'pdf':
            return export_mapa_pdf(processo)
        if formato == 'xlsx':
            return export_mapa_xlsx(processo)
    if doc == 'distribuicao':
        if formato == 'html':
            return render(
                request,
                'workflow/distribuicao_preview.html',
                {'processo': processo, 'orgao': _orgao_ativo(), **_distribuicao_preview_context(processo)},
            )
        if formato == 'docx':
            return export_distribuicao_docx(processo)
        if formato == 'pdf':
            return export_distribuicao_pdf(processo)
        if formato == 'xlsx':
            return export_distribuicao_xlsx(processo)
    if doc == 'tr' and formato == 'docx':
        return export_tr_docx(processo)
    messages.error(request, 'Exportação ainda não disponível para esta combinação.')
    return redirect('workflow:planejamento_detail', processo.id)


@xframe_options_exempt
def planejamento_preview_embed(request, processo_id: int, doc: str):
    processo = get_object_or_404(Processo, pk=processo_id)
    doc = (doc or '').lower()
    if doc == 'dfd':
        dfd, _ = PlanejamentoDFD.objects.get_or_create(processo=processo, defaults={'objeto_resumido': processo.objeto or ''})
        context = {
            'processo': processo,
            'dfd': dfd,
            'previsao_entrega_execucao_formatada': _format_date_label(dfd.previsao_entrega_execucao),
            'orgao': _orgao_ativo(),
            **_atendente_context(request.user),
        }
        return render(request, 'workflow/dfd_preview_embed.html', context)
    if doc == 'mapa':
        return render(request, 'workflow/mapa_preview_embed.html', {'processo': processo, 'orgao': _orgao_ativo(), **_mapa_preview_context(processo)})
    if doc == 'distribuicao':
        return render(
            request,
            'workflow/distribuicao_preview_embed.html',
            {'processo': processo, 'orgao': _orgao_ativo(), **_distribuicao_preview_context(processo)},
        )
    return HttpResponseNotFound('Pré-visualização não disponível para este documento.')


def planejamento_pessoas_autocomplete(request):
    termo = (request.GET.get('q') or '').strip()
    qs = Pessoa.objects.all().select_related('secretaria').order_by('nome')
    if termo:
        qs = qs.filter(nome__icontains=termo)
    pessoas = []
    for p in qs[:20]:
        pessoas.append({
            'id': p.id,
            'nome': p.nome,
            'cargo': p.cargo or '',
            'secretaria': p.secretaria.sigla if p.secretaria else '',
            'rotulo': f'{p.nome} - {p.cargo}' if p.cargo else p.nome,
        })
    return JsonResponse({'results': pessoas})


def planejamento_catalogo_autocomplete(request):
    termo = (request.GET.get('q') or '').strip()
    unidade = (request.GET.get('unidade') or '').strip()
    try:
        limit = max(1, min(int(request.GET.get('limit') or 40), 300))
    except Exception:
        limit = 40
    qs = DFDItemCatalogo.objects.all().order_by('descricao')
    if unidade:
        qs = qs.filter(unidade__iexact=unidade)
    if termo:
        filtro = Q(descricao__icontains=termo)
        if termo.isdigit():
            filtro |= Q(codigo=int(termo))
        qs = qs.filter(filtro)
    itens = []
    for item in qs[:limit]:
        itens.append({
            'id': item.id,
            'codigo': item.codigo,
            'descricao': item.descricao,
            'unidade': item.unidade or '',
            'rotulo': f'{item.codigo} - {item.descricao[:120]}',
        })
    return JsonResponse({'results': itens})





