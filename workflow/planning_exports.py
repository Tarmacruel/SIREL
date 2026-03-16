from datetime import datetime
from decimal import Decimal
from html import unescape as html_unescape
from io import BytesIO
import re
from xml.sax.saxutils import escape

from django.db import OperationalError, ProgrammingError
from django.db.models import Avg
from django.http import HttpResponse
from docx import Document
from docx.shared import Cm
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from core.models import OrgaoEntidade, Pessoa, ProcessoItem, Secretaria
from core.utils.formatters import fmt_brl, valor_por_extenso


def _resp(data: bytes, ctype: str, filename: str):
    response = HttpResponse(data, content_type=ctype)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _fmt_date(valor):
    if not valor:
        return "-"
    if hasattr(valor, "strftime"):
        return valor.strftime("%d/%m/%Y")
    try:
        return datetime.fromisoformat(str(valor)).strftime("%d/%m/%Y")
    except Exception:
        return str(valor)


def _pdf_text(value):
    return escape(str(value or "-")).replace("\n", "<br/>")


def _balance_reportlab_markup_tokens(texto):
    if not texto:
        return ""
    token_re = re.compile(r"(\[\[(?:/)?[BIU]\]\])")
    open_tokens = {"[[B]]": "B", "[[I]]": "I", "[[U]]": "U"}
    close_tokens = {"[[/B]]": "B", "[[/I]]": "I", "[[/U]]": "U"}
    close_for = {"B": "[[/B]]", "I": "[[/I]]", "U": "[[/U]]"}

    out = []
    stack = []
    for part in token_re.split(texto):
        if not part:
            continue
        if part in open_tokens:
            out.append(part)
            stack.append(open_tokens[part])
            continue
        if part in close_tokens:
            tag = close_tokens[part]
            if tag not in stack:
                continue
            while stack and stack[-1] != tag:
                out.append(close_for[stack.pop()])
            if stack and stack[-1] == tag:
                out.append(part)
                stack.pop()
            continue
        out.append(part)

    while stack:
        out.append(close_for[stack.pop()])
    return "".join(out)


def _pdf_rich_text(value):
    texto = str(value or "")
    if not texto:
        return "-"
    substituicoes = [
        (r"(?is)<\s*strong[^>]*>", "[[B]]"),
        (r"(?is)</\s*strong\s*>", "[[/B]]"),
        (r"(?is)<\s*b(?:\s[^>]*)?>", "[[B]]"),
        (r"(?is)</\s*b\s*>", "[[/B]]"),
        (r"(?is)<\s*em[^>]*>", "[[I]]"),
        (r"(?is)</\s*em\s*>", "[[/I]]"),
        (r"(?is)<\s*i(?:\s[^>]*)?>", "[[I]]"),
        (r"(?is)</\s*i\s*>", "[[/I]]"),
        (r"(?is)<\s*u(?:\s[^>]*)?>", "[[U]]"),
        (r"(?is)</\s*u\s*>", "[[/U]]"),
        (r"(?is)<\s*br\s*/?\s*>", "\n"),
        (r"(?is)</\s*p\s*>", "\n\n"),
        (r"(?is)</\s*div\s*>", "\n"),
        (r"(?is)<\s*li[^>]*>", "\n- "),
        (r"(?is)</\s*li\s*>", ""),
    ]
    for pattern, repl in substituicoes:
        texto = re.sub(pattern, repl, texto)

    texto = re.sub(r"(?is)<[^>]+>", "", texto)
    texto = html_unescape(texto)
    texto = escape(texto)
    texto = _balance_reportlab_markup_tokens(texto)
    texto = texto.replace("[[B]]", "<b>").replace("[[/B]]", "</b>")
    texto = texto.replace("[[I]]", "<i>").replace("[[/I]]", "</i>")
    texto = texto.replace("[[U]]", "<u>").replace("[[/U]]", "</u>")
    texto = texto.replace("\r\n", "\n").replace("\r", "\n")
    texto = re.sub(r"\n{3,}", "\n\n", texto).strip()
    if not texto:
        return "-"
    return texto.replace("\n", "<br/>")


def _user_identity(usuario):
    if not getattr(usuario, "is_authenticated", False):
        return "-", "-"
    nome = (usuario.get_full_name() or "").strip() or usuario.get_username()
    return nome, usuario.get_username()


def _orgao_payload():
    try:
        orgao = OrgaoEntidade.objects.order_by("-atualizado_em", "-id").first()
    except (OperationalError, ProgrammingError):
        orgao = None
    if not orgao:
        return {
            "obj": None,
            "nome": "Órgão não cadastrado",
            "razao_social": "",
            "cnpj": "",
            "endereco": "",
            "logo_path": "",
        }
    return {
        "obj": orgao,
        "nome": orgao.nome_fantasia or orgao.razao_social,
        "razao_social": orgao.razao_social or "",
        "cnpj": orgao.cnpj or "",
        "endereco": orgao.endereco_completo or "",
        "logo_path": orgao.logo.path if orgao.logo else "",
    }


def build_sd_payload(processo):
    try:
        dfd = processo.planejamento_dfd
    except Exception:
        dfd = None
    try:
        tr = processo.planejamento_tr
    except Exception:
        tr = None
    try:
        etp = processo.planejamento_etp
    except Exception:
        etp = None
    orgao = _orgao_payload()

    secretarias_vinculadas = []
    if dfd:
        secretarias_vinculadas = list(dfd.secretarias_vinculadas.select_related("secretaria").all())

    secretaria_assinatura = None
    if secretarias_vinculadas:
        if len(secretarias_vinculadas) > 1:
            secretaria_assinatura = Secretaria.objects.filter(sigla__iexact="ADM").first()
            if not secretaria_assinatura:
                secretaria_assinatura = Secretaria.objects.filter(nome__icontains="ADMINISTRA").first()
            if not secretaria_assinatura:
                secretaria_assinatura = secretarias_vinculadas[0].secretaria
        else:
            principal = next((x.secretaria for x in secretarias_vinculadas if x.principal), None)
            secretaria_assinatura = principal or secretarias_vinculadas[0].secretaria

    signatario = (
        Pessoa.objects.filter(secretaria=secretaria_assinatura).order_by("nome").first()
        if secretaria_assinatura else None
    )

    media_map = {}
    if etp:
        medias = (
            etp.cotacoes.filter(considerar_no_calculo=True)
            .values("item_id")
            .annotate(media=Avg("valor_unitario"))
        )
        media_map = {row["item_id"]: Decimal(str(row["media"])) for row in medias}

    items = []
    items_by_id = {}
    total_geral = Decimal("0.00")
    dfd_items = list(dfd.itens.all().order_by("codigo")) if dfd else []
    for item in dfd_items:
        quantidade = Decimal(str(item.quantidade or 0))
        valor_unitario = media_map.get(item.id)
        if valor_unitario is None:
            core_item = ProcessoItem.objects.filter(
                processo=processo,
                numero_item=item.codigo,
            ).order_by("-pncp_ultima_atualizacao", "-atualizado_em", "-id").first()
            if core_item:
                valor_unitario = Decimal(str(core_item.valor_referencia_unitario or core_item.valor_homologado_unitario or 0))
        valor_unitario = (valor_unitario or Decimal("0")).quantize(Decimal("0.01"))
        total_item = (valor_unitario * quantidade).quantize(Decimal("0.01"))
        row = {
            "item": item,
            "quantidade": quantidade,
            "valor_unitario": valor_unitario,
            "total_item": total_item,
        }
        items.append(row)
        items_by_id[item.id] = row
        total_geral += total_item
    total_geral = total_geral.quantize(Decimal("0.01"))

    lotes = []
    if tr and tr.lotes.exists():
        lotes_qs = tr.lotes.prefetch_related("itens").all().order_by("numero")
        for lote in lotes_qs:
            lote_items = []
            total_lote = Decimal("0.00")
            for item in lote.itens.all().order_by("codigo"):
                row = items_by_id.get(item.id)
                if not row:
                    continue
                lote_items.append(row)
                total_lote += row["total_item"]
            lotes.append({
                "numero": lote.numero,
                "titulo": lote.titulo or f"Lote {lote.numero}",
                "itens": lote_items,
                "total_lote": total_lote.quantize(Decimal("0.01")),
            })
    else:
        lotes.append({
            "numero": 1,
            "titulo": "Lote unico",
            "itens": items,
            "total_lote": total_geral,
        })

    dotacoes = []
    if tr:
        dotacoes = list(
            tr.dotacoes.select_related(
                "secretaria",
                "unidade_orcamentaria",
                "projeto_atividade",
                "elemento_despesa",
                "fonte_recurso",
            ).all()
        )

    return {
        "processo": processo,
        "dfd": dfd,
        "tr": tr,
        "orgao": orgao,
        "itens": items,
        "lotes": lotes,
        "dotacoes": dotacoes,
        "total_geral": total_geral,
        "secretarias_vinculadas": secretarias_vinculadas,
        "multiplas_secretarias": len(secretarias_vinculadas) > 1,
        "secretaria_assinatura": secretaria_assinatura,
        "signatario": signatario,
        "justificativa": (dfd.justificativa_contratacao if dfd else "") or "",
    }


def _dfd_payload(processo, usuario=None):
    dfd = processo.planejamento_dfd
    atendente_nome, atendente_identificacao = _user_identity(usuario)
    orgao = _orgao_payload()
    secretarias = [
        f"{s.secretaria.sigla} - {s.secretaria.nome}{' (principal)' if s.principal else ''}"
        for s in dfd.secretarias_vinculadas.select_related("secretaria").all()
    ]
    return {
        "dfd": dfd,
        "processo": processo,
        "atendente_nome": atendente_nome,
        "atendente_identificacao": atendente_identificacao,
        "orgao": orgao["obj"],
        "orgao_nome": orgao["nome"],
        "orgao_razao_social": orgao["razao_social"],
        "orgao_cnpj": orgao["cnpj"],
        "orgao_endereco": orgao["endereco"],
        "orgao_logo_path": orgao["logo_path"],
        "secretarias": secretarias,
        "itens": list(dfd.itens.all()),
        "previsao": _fmt_date(dfd.previsao_entrega_execucao),
        "data_demanda": _fmt_date(dfd.data_demanda),
        "responsavel_demanda": dfd.responsavel_demanda or "-",
        "cargo_funcao": dfd.cargo_funcao or "-",
    }


def export_dfd_docx(processo, usuario=None):
    data = _dfd_payload(processo, usuario)
    dfd = data["dfd"]

    doc = Document()
    if data["orgao_logo_path"]:
        try:
            doc.add_picture(data["orgao_logo_path"], width=Cm(2.8))
        except Exception:
            pass
    doc.add_paragraph(data["orgao_nome"])
    if data["orgao_razao_social"] and data["orgao_razao_social"] != data["orgao_nome"]:
        doc.add_paragraph(data["orgao_razao_social"])
    if data["orgao_cnpj"]:
        doc.add_paragraph(f"CNPJ: {data['orgao_cnpj']}")
    if data["orgao_endereco"]:
        doc.add_paragraph(data["orgao_endereco"])
    doc.add_paragraph("")

    doc.add_heading("Documento de Formalizacao da Demanda (DFD)", level=1)
    doc.add_paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}")

    cards = [
        ("Objeto resumido", dfd.objeto_resumido or "-"),
        ("Data da demanda", data["data_demanda"]),
        ("Modalidade pretendida", dfd.get_modalidade_pretendida_display()),
        ("Tipo de contratacao", dfd.get_tipo_contratacao_planejamento_display()),
        ("Especie da contratacao", dfd.get_especie_contratacao_display()),
        ("Cargo/Funcao", data["cargo_funcao"]),
        ("Atendente (registro)", data["atendente_nome"]),
        ("Identificacao do atendente", data["atendente_identificacao"]),
    ]
    info_table = doc.add_table(rows=(len(cards) + 1) // 2, cols=2)
    info_table.style = "Table Grid"
    for idx, (label, value) in enumerate(cards):
        row = idx // 2
        col = idx % 2
        cell = info_table.cell(row, col)
        cell.text = ""
        p1 = cell.add_paragraph(label.upper())
        p1.runs[0].bold = True
        p2 = cell.add_paragraph(str(value))
        p2.paragraph_format.space_after = 0

    def add_block(title, text):
        doc.add_paragraph("")
        t = doc.add_paragraph(title.upper())
        t.runs[0].bold = True
        doc.add_paragraph(text or "-")

    add_block("Descricao detalhada", dfd.descricao_detalhada)
    add_block("Justificativa da contratacao", dfd.justificativa_contratacao)
    add_block("Fundamento legal", dfd.fundamento_legal)
    add_block("Previsao de entrega/execucao", data["previsao"])

    doc.add_paragraph("")
    p = doc.add_paragraph("SECRETARIAS PARTICIPANTES")
    p.runs[0].bold = True
    if data["secretarias"]:
        for secretaria in data["secretarias"]:
            doc.add_paragraph(secretaria, style="List Bullet")
    else:
        doc.add_paragraph("-")

    doc.add_paragraph("")
    p = doc.add_paragraph("ITENS DA DEMANDA")
    p.runs[0].bold = True
    itens_table = doc.add_table(rows=1, cols=5)
    itens_table.style = "Table Grid"
    head = itens_table.rows[0].cells
    head[0].text = "No item"
    head[1].text = "Codigo"
    head[2].text = "Descricao"
    head[3].text = "Unidade"
    head[4].text = "Quantidade"
    for idx, item in enumerate(data["itens"], start=1):
        row = itens_table.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(item.codigo)
        row[2].text = item.descricao
        row[3].text = item.unidade
        row[4].text = f"{item.quantidade:.2f}"

    doc.add_paragraph("")
    p = doc.add_paragraph("ASSINATURA DO RESPONSAVEL PELA DEMANDA")
    p.runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph("________________________________________")
    doc.add_paragraph(data["responsavel_demanda"])
    doc.add_paragraph(data["cargo_funcao"])

    bio = BytesIO()
    doc.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"DFD_{processo.id}.docx",
    )


def export_dfd_xlsx(processo):
    dfd = processo.planejamento_dfd
    orgao = _orgao_payload()
    wb = Workbook()
    ws = wb.active
    ws.title = "DFD"
    ws.append(["Campo", "Valor"])
    ws.append(["Orgao", orgao["nome"]])
    ws.append(["Razao social", orgao["razao_social"] or "-"])
    ws.append(["CNPJ", orgao["cnpj"] or "-"])
    ws.append(["Endereco", orgao["endereco"] or "-"])
    ws.append(["Processo", f"{processo.numero_processo_adm}/{processo.ano_referencia}"])
    ws.append(["Objeto resumido", dfd.objeto_resumido])
    ws.append(["Descricao detalhada", dfd.descricao_detalhada])
    ws.append(["Data da demanda", _fmt_date(dfd.data_demanda)])
    ws.append(["Modalidade pretendida", dfd.get_modalidade_pretendida_display()])
    ws.append(["Tipo de contratacao", dfd.get_tipo_contratacao_planejamento_display()])
    ws.append(["Especie da contratacao", dfd.get_especie_contratacao_display()])
    ws.append(["Responsavel pela demanda", dfd.responsavel_demanda])
    ws.append(["Cargo/Funcao", dfd.cargo_funcao])

    ws2 = wb.create_sheet("Itens")
    ws2.append(["No item", "Codigo", "Descricao", "Unidade", "Quantidade"])
    for idx, item in enumerate(dfd.itens.all(), start=1):
        ws2.append([idx, item.codigo, item.descricao, item.unidade, float(item.quantidade)])

    bio = BytesIO()
    wb.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"DFD_{processo.id}.xlsx",
    )


def export_dfd_pdf(processo, usuario=None):
    data = _dfd_payload(processo, usuario)
    dfd = data["dfd"]

    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=28,
        bottomMargin=28,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleDFD",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#153153"),
    )
    subtitle_style = ParagraphStyle(
        "SubtitleDFD",
        parent=styles["BodyText"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    value_style = ParagraphStyle(
        "ValueDFD",
        parent=styles["BodyText"],
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#1d2b40"),
    )
    section_title_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading3"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#153153"),
    )

    elements = []
    if data["orgao_logo_path"]:
        try:
            elements.append(Image(data["orgao_logo_path"], width=70, height=70))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(data["orgao_nome"]), section_title_style))
    if data["orgao_razao_social"] and data["orgao_razao_social"] != data["orgao_nome"]:
        elements.append(Paragraph(_pdf_text(data["orgao_razao_social"]), subtitle_style))
    if data["orgao_cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ: {data['orgao_cnpj']}"), subtitle_style))
    if data["orgao_endereco"]:
        elements.append(Paragraph(_pdf_text(data["orgao_endereco"]), subtitle_style))
    elements.extend(
        [
            Spacer(1, 8),
            Paragraph("Documento de Formalizacao da Demanda (DFD)", title_style),
            Paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}", subtitle_style),
            Spacer(1, 10),
        ]
    )

    cards = [
        ("Objeto resumido", dfd.objeto_resumido or "-"),
        ("Data da demanda", data["data_demanda"]),
        ("Modalidade pretendida", dfd.get_modalidade_pretendida_display()),
        ("Tipo de contratacao", dfd.get_tipo_contratacao_planejamento_display()),
        ("Especie da contratacao", dfd.get_especie_contratacao_display()),
        ("Cargo/Funcao", data["cargo_funcao"]),
        ("Atendente (registro)", data["atendente_nome"]),
        ("Identificacao do atendente", data["atendente_identificacao"]),
    ]
    rows = []
    for i in range(0, len(cards), 2):
        left = cards[i]
        right = cards[i + 1] if i + 1 < len(cards) else ("", "")
        left_text = f"<b>{_pdf_text(left[0].upper())}</b><br/>{_pdf_text(left[1])}"
        right_text = f"<b>{_pdf_text(right[0].upper())}</b><br/>{_pdf_text(right[1])}" if right[0] else ""
        rows.append([Paragraph(left_text, value_style), Paragraph(right_text, value_style)])

    card_table = Table(rows, colWidths=[260, 260])
    card_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f6fb")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.extend([card_table, Spacer(1, 10)])

    def add_section(title, content):
        elements.append(Paragraph(_pdf_text(title.upper()), section_title_style))
        box = Table([[Paragraph(_pdf_text(content), value_style)]], colWidths=[520])
        box.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafd")),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        elements.append(box)
        elements.append(Spacer(1, 8))

    add_section("Descricao detalhada", dfd.descricao_detalhada)
    add_section("Justificativa da contratacao", dfd.justificativa_contratacao)
    add_section("Fundamento legal", dfd.fundamento_legal)
    add_section("Previsao de entrega/execucao", data["previsao"])

    secretarias_text = "\n".join(data["secretarias"]) if data["secretarias"] else "-"
    add_section("Secretarias participantes", secretarias_text)

    item_value_style = ParagraphStyle(
        "ItemValueDFD",
        parent=value_style,
        fontSize=9.5,
        leading=11.5,
        wordWrap="CJK",
    )
    item_value_style.splitLongWords = True

    elements.append(Paragraph("ITENS DA DEMANDA", section_title_style))
    itens_rows = [["No item", "Codigo", "Descricao", "Unidade", "Quantidade"]]
    for idx, item in enumerate(data["itens"], start=1):
        itens_rows.append(
            [
                Paragraph(str(idx), item_value_style),
                Paragraph(str(item.codigo), item_value_style),
                Paragraph(_pdf_text(item.descricao), item_value_style),
                Paragraph(_pdf_text(item.unidade), item_value_style),
                Paragraph(f"{item.quantidade:.2f}", item_value_style),
            ]
        )
    itens_table = Table(itens_rows, colWidths=[42, 50, 248, 70, 90], repeatRows=1)
    itens_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#153153")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 1), (1, -1), "CENTER"),
                ("ALIGN", (4, 1), (4, -1), "RIGHT"),
                ("WORDWRAP", (2, 1), (2, -1), "CJK"),
            ]
        )
    )
    elements.extend([itens_table, Spacer(1, 16)])

    elements.append(Paragraph("ASSINATURA DO RESPONSAVEL PELA DEMANDA", section_title_style))
    assinatura = Table(
        [[Paragraph(f"{_pdf_text(data['responsavel_demanda'])}<br/>{_pdf_text(data['cargo_funcao'])}", value_style)]],
        colWidths=[320],
    )
    assinatura.setStyle(
        TableStyle(
            [
                ("LINEABOVE", (0, 0), (0, 0), 1, colors.HexColor("#50617a")),
                ("TOPPADDING", (0, 0), (0, 0), 12),
            ]
        )
    )
    elements.append(Spacer(1, 18))
    elements.append(assinatura)

    pdf.build(elements)
    return _resp(bio.getvalue(), "application/pdf", f"DFD_{processo.id}.pdf")


def _mapa_payload(processo):
    orgao = _orgao_payload()
    try:
        etp = processo.planejamento_etp
    except Exception:
        etp = None
    try:
        dfd = processo.planejamento_dfd
    except Exception:
        dfd = None

    if not etp or not dfd:
        return {
            "processo": processo,
            "etp": None,
            "dfd": dfd,
            "orgao": orgao,
            "cotacoes": [],
            "itens_estimados": [],
            "total_estimado_processo": Decimal("0.00"),
        }

    cotacoes = list(etp.cotacoes.select_related("item", "fonte").order_by("item__codigo", "fonte__nome_fonte"))
    medias_qs = (
        etp.cotacoes.filter(considerar_no_calculo=True)
        .values("item_id")
        .annotate(media=Avg("valor_unitario"))
    )
    medias_map = {
        row["item_id"]: Decimal(str(row["media"]))
        for row in medias_qs
        if row.get("media") is not None
    }

    itens_estimados = []
    total_estimado_processo = Decimal("0.00")
    for item in dfd.itens.all().order_by("codigo"):
        quantidade = Decimal(str(item.quantidade or 0))
        media = medias_map.get(item.id)
        total_item = None
        if media is not None:
            total_item = (media * quantidade).quantize(Decimal("0.01"))
            total_estimado_processo += total_item
        itens_estimados.append(
            {
                "item": item,
                "quantidade": quantidade,
                "media": media,
                "total_item": total_item,
            }
        )

    return {
        "processo": processo,
        "etp": etp,
        "dfd": dfd,
        "orgao": orgao,
        "cotacoes": cotacoes,
        "itens_estimados": itens_estimados,
        "total_estimado_processo": total_estimado_processo.quantize(Decimal("0.01")),
    }


def export_mapa_xlsx(processo):
    data = _mapa_payload(processo)
    etp = data["etp"]
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumo"
    ws.append(["Processo", f"{processo.numero_processo_adm}/{processo.ano_referencia}"])
    ws.append(["Total estimado do processo", float(data["total_estimado_processo"])])
    ws.append(
        [
            "Total estimado do processo (formatado)",
            f"{fmt_brl(data['total_estimado_processo'])} ({valor_por_extenso(data['total_estimado_processo'])})",
        ]
    )
    ws.append([])
    ws.append(["Item", "Descricao", "Quantidade", "Media considerada", "Total estimado do item"])
    for linha in data["itens_estimados"]:
        ws.append(
            [
                linha["item"].codigo,
                linha["item"].descricao,
                float(linha["quantidade"]),
                float(linha["media"]) if linha["media"] is not None else None,
                float(linha["total_item"]) if linha["total_item"] is not None else None,
            ]
        )

    ws2 = wb.create_sheet("Cotacoes")
    ws2.append(["Item", "Descricao", "Fonte", "Valor unitario", "Considerar", "Inexequivel", "Sobrepreco"])
    if etp:
        for cot in etp.cotacoes.select_related("item", "fonte"):
            ws2.append(
                [
                    cot.item.codigo,
                    cot.item.descricao,
                    cot.fonte.nome_fonte,
                    float(cot.valor_unitario),
                    "SIM" if cot.considerar_no_calculo else "NAO",
                    "SIM" if cot.inexequivel else "NAO",
                    "SIM" if cot.sobrepreco else "NAO",
                ]
            )
    else:
        ws2.append(["-", "Sem ETP cadastrado", "-", "-", "-", "-", "-"])

    bio = BytesIO()
    wb.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"Mapa_ETP_{processo.id}.xlsx",
    )


def export_mapa_docx(processo):
    data = _mapa_payload(processo)
    doc = Document()
    orgao = data["orgao"]

    if orgao["logo_path"]:
        try:
            doc.add_picture(orgao["logo_path"], width=Cm(2.5))
        except Exception:
            pass
    doc.add_paragraph(orgao["nome"])
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        doc.add_paragraph(orgao["razao_social"])
    if orgao["cnpj"]:
        doc.add_paragraph(f"CNPJ: {orgao['cnpj']}")
    if orgao["endereco"]:
        doc.add_paragraph(orgao["endereco"])
    doc.add_paragraph("")

    doc.add_heading("Mapa Comparativo de Cotacoes (ETP)", level=1)
    doc.add_paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}")
    if not data["etp"]:
        doc.add_paragraph("Nenhum ETP cadastrado para este processo.")
    else:
        doc.add_paragraph(
            f"Total estimado do processo: {fmt_brl(data['total_estimado_processo'])} "
            f"({valor_por_extenso(data['total_estimado_processo'])})"
        )
        doc.add_paragraph("")
        p = doc.add_paragraph("ESTIMATIVAS POR ITEM")
        p.runs[0].bold = True
        tab_estimativa = doc.add_table(rows=1, cols=5)
        tab_estimativa.style = "Table Grid"
        head = tab_estimativa.rows[0].cells
        head[0].text = "Item"
        head[1].text = "Descricao"
        head[2].text = "Quantidade"
        head[3].text = "Media considerada"
        head[4].text = "Total estimado do item"
        for linha in data["itens_estimados"]:
            row = tab_estimativa.add_row().cells
            row[0].text = str(linha["item"].codigo)
            row[1].text = linha["item"].descricao
            row[2].text = f"{linha['quantidade']:.3f}"
            row[3].text = fmt_brl(linha["media"]) if linha["media"] is not None else "-"
            row[4].text = fmt_brl(linha["total_item"]) if linha["total_item"] is not None else "-"

        doc.add_paragraph("")
        p = doc.add_paragraph("COTACOES LANCADAS")
        p.runs[0].bold = True
        tab_cotacoes = doc.add_table(rows=1, cols=5)
        tab_cotacoes.style = "Table Grid"
        head2 = tab_cotacoes.rows[0].cells
        head2[0].text = "Item"
        head2[1].text = "Fonte"
        head2[2].text = "Valor"
        head2[3].text = "Situacao"
        head2[4].text = "Considerar"
        for cot in data["cotacoes"]:
            row = tab_cotacoes.add_row().cells
            row[0].text = f"{cot.item.codigo} - {cot.item.descricao[:110]}"
            row[1].text = cot.fonte.nome_fonte
            row[2].text = fmt_brl(Decimal(str(cot.valor_unitario)))
            if cot.inexequivel:
                situacao = "Abaixo de 50% da media"
            elif cot.sobrepreco:
                situacao = "Acima de 50% da media"
            else:
                situacao = "Faixa normal"
            row[3].text = situacao
            row[4].text = "Sim" if cot.considerar_no_calculo else "Nao"

    bio = BytesIO()
    doc.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"Mapa_ETP_{processo.id}.docx",
    )


def export_mapa_pdf(processo):
    data = _mapa_payload(processo)
    orgao = data["orgao"]
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=28,
        bottomMargin=28,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleMapa",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#153153"),
    )
    subtitle_style = ParagraphStyle(
        "SubtitleMapa",
        parent=styles["BodyText"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    value_style = ParagraphStyle(
        "ValueMapa",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=11.5,
        textColor=colors.HexColor("#1d2b40"),
    )
    section_title_style = ParagraphStyle(
        "SectionMapa",
        parent=styles["Heading3"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#153153"),
    )

    elements = []
    if orgao["logo_path"]:
        try:
            elements.append(Image(orgao["logo_path"], width=60, height=60))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(orgao["nome"]), section_title_style))
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        elements.append(Paragraph(_pdf_text(orgao["razao_social"]), subtitle_style))
    if orgao["cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ: {orgao['cnpj']}"), subtitle_style))
    if orgao["endereco"]:
        elements.append(Paragraph(_pdf_text(orgao["endereco"]), subtitle_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("Mapa Comparativo de Cotacoes (ETP)", title_style))
    elements.append(Paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}", subtitle_style))
    elements.append(Spacer(1, 8))

    if not data["etp"]:
        elements.append(Paragraph("Nenhum ETP cadastrado para este processo.", value_style))
    else:
        elements.append(
            Paragraph(
                f"<b>Total estimado do processo:</b> "
                f"{_pdf_text(fmt_brl(data['total_estimado_processo']))} "
                f"({_pdf_text(valor_por_extenso(data['total_estimado_processo']))})",
                value_style,
            )
        )
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("ESTIMATIVAS POR ITEM", section_title_style))

        resumo_rows = [[
            Paragraph("<b>Item</b>", value_style),
            Paragraph("<b>Descricao</b>", value_style),
            Paragraph("<b>Qtd.</b>", value_style),
            Paragraph("<b>Media</b>", value_style),
            Paragraph("<b>Total item</b>", value_style),
        ]]
        for linha in data["itens_estimados"]:
            resumo_rows.append(
                [
                    Paragraph(str(linha["item"].codigo), value_style),
                    Paragraph(_pdf_text(linha["item"].descricao), value_style),
                    Paragraph(f"{linha['quantidade']:.3f}", value_style),
                    Paragraph(_pdf_text(fmt_brl(linha["media"])) if linha["media"] is not None else "-", value_style),
                    Paragraph(_pdf_text(fmt_brl(linha["total_item"])) if linha["total_item"] is not None else "-", value_style),
                ]
            )
        resumo_table = Table(resumo_rows, colWidths=[40, 255, 60, 72, 93], repeatRows=1)
        resumo_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
                ]
            )
        )
        elements.extend([resumo_table, Spacer(1, 10)])
        elements.append(Paragraph("COTACOES LANCADAS", section_title_style))

        cot_rows = [[
            Paragraph("<b>Item</b>", value_style),
            Paragraph("<b>Fonte</b>", value_style),
            Paragraph("<b>Valor</b>", value_style),
            Paragraph("<b>Situacao</b>", value_style),
            Paragraph("<b>Considerar</b>", value_style),
        ]]
        for cot in data["cotacoes"]:
            if cot.inexequivel:
                situacao = "Abaixo de 50% da media"
            elif cot.sobrepreco:
                situacao = "Acima de 50% da media"
            else:
                situacao = "Faixa normal"
            cot_rows.append(
                [
                    Paragraph(_pdf_text(f"{cot.item.codigo} - {cot.item.descricao}"), value_style),
                    Paragraph(_pdf_text(cot.fonte.nome_fonte), value_style),
                    Paragraph(_pdf_text(fmt_brl(Decimal(str(cot.valor_unitario)))), value_style),
                    Paragraph(_pdf_text(situacao), value_style),
                    Paragraph("Sim" if cot.considerar_no_calculo else "Nao", value_style),
                ]
            )
        cot_table = Table(cot_rows, colWidths=[220, 120, 58, 90, 32], repeatRows=1)
        cot_styles = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
            ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
            ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (2, 1), (2, -1), "RIGHT"),
            ("ALIGN", (4, 1), (4, -1), "CENTER"),
        ]
        for idx, cot in enumerate(data["cotacoes"], start=1):
            if not cot.considerar_no_calculo:
                cot_styles.append(("TEXTCOLOR", (0, idx), (-1, idx), colors.HexColor("#7f8ca6")))
                cot_styles.append(("BACKGROUND", (0, idx), (-1, idx), colors.HexColor("#f0f2f6")))
        cot_table.setStyle(TableStyle(cot_styles))
        elements.append(cot_table)

    pdf.build(elements)
    return _resp(bio.getvalue(), "application/pdf", f"Mapa_ETP_{processo.id}.pdf")


def _distribuicao_payload(processo):
    orgao = _orgao_payload()
    try:
        tr = processo.planejamento_tr
    except Exception:
        tr = None
    try:
        dfd = processo.planejamento_dfd
    except Exception:
        dfd = None

    habilitada = bool(tr and not tr.nao_aplica_divisao_secretaria)
    linhas = []
    resumo_itens = []
    total_qtd_dfd = Decimal("0.000")
    total_qtd_distribuida = Decimal("0.000")

    if habilitada and dfd:
        linhas = list(
            tr.distribuicoes.select_related("secretaria", "item")
            .order_by("secretaria__sigla", "item__codigo")
        )
        totais_por_item = {}
        for linha in linhas:
            item_id = linha.item_id
            totais_por_item[item_id] = totais_por_item.get(item_id, Decimal("0")) + Decimal(
                str(linha.quantidade or 0)
            )

        for item in dfd.itens.all().order_by("codigo"):
            qtd_dfd = Decimal(str(item.quantidade or 0))
            qtd_distribuida = totais_por_item.get(item.id, Decimal("0"))
            saldo = qtd_dfd - qtd_distribuida
            resumo_itens.append(
                {
                    "item": item,
                    "quantidade_dfd": qtd_dfd,
                    "quantidade_distribuida": qtd_distribuida,
                    "saldo": saldo,
                }
            )
            total_qtd_dfd += qtd_dfd
            total_qtd_distribuida += qtd_distribuida

    return {
        "processo": processo,
        "orgao": orgao,
        "tr": tr,
        "dfd": dfd,
        "habilitada": habilitada,
        "linhas": linhas,
        "resumo_itens": resumo_itens,
        "total_qtd_dfd": total_qtd_dfd,
        "total_qtd_distribuida": total_qtd_distribuida,
    }


def export_distribuicao_xlsx(processo):
    data = _distribuicao_payload(processo)
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumo"
    ws.append(["Processo", f"{processo.numero_processo_adm}/{processo.ano_referencia}"])
    ws.append(
        [
            "Distribuicao por secretaria habilitada",
            "SIM" if data["habilitada"] else "NAO",
        ]
    )
    ws.append(["Total qtd. DFD", float(data["total_qtd_dfd"])])
    ws.append(["Total qtd. distribuida", float(data["total_qtd_distribuida"])])
    ws.append([])
    ws.append(["Item", "Descricao", "Unidade", "Qtd. DFD", "Qtd. distribuida", "Saldo"])
    for row in data["resumo_itens"]:
        ws.append(
            [
                row["item"].codigo,
                row["item"].descricao,
                row["item"].unidade,
                float(row["quantidade_dfd"]),
                float(row["quantidade_distribuida"]),
                float(row["saldo"]),
            ]
        )
    if not data["resumo_itens"]:
        ws.append(["-", "Sem itens para distribuicao", "-", 0, 0, 0])

    ws2 = wb.create_sheet("Distribuicoes")
    ws2.append(["Secretaria", "Item", "Descricao", "Unidade", "Quantidade"])
    for linha in data["linhas"]:
        ws2.append(
            [
                linha.secretaria.sigla,
                linha.item.codigo,
                linha.item.descricao,
                linha.item.unidade,
                float(linha.quantidade),
            ]
        )
    if not data["linhas"]:
        ws2.append(["-", "-", "Sem distribuicoes lancadas", "-", 0])

    bio = BytesIO()
    wb.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"Mapa_Distribuicao_{processo.id}.xlsx",
    )


def export_distribuicao_docx(processo):
    data = _distribuicao_payload(processo)
    doc = Document()
    orgao = data["orgao"]

    if orgao["logo_path"]:
        try:
            doc.add_picture(orgao["logo_path"], width=Cm(2.5))
        except Exception:
            pass
    doc.add_paragraph(orgao["nome"])
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        doc.add_paragraph(orgao["razao_social"])
    if orgao["cnpj"]:
        doc.add_paragraph(f"CNPJ: {orgao['cnpj']}")
    if orgao["endereco"]:
        doc.add_paragraph(orgao["endereco"])
    doc.add_paragraph("")

    doc.add_heading("Mapa de Distribuicao por Secretaria", level=1)
    doc.add_paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}")
    if not data["habilitada"]:
        doc.add_paragraph("Distribuicao por secretaria desabilitada no TR.")
    else:
        doc.add_paragraph(
            f"Total DFD: {data['total_qtd_dfd']:.3f} | "
            f"Total distribuido: {data['total_qtd_distribuida']:.3f}"
        )
        doc.add_paragraph("")

        p = doc.add_paragraph("RESUMO POR ITEM")
        p.runs[0].bold = True
        tab_resumo = doc.add_table(rows=1, cols=6)
        tab_resumo.style = "Table Grid"
        head = tab_resumo.rows[0].cells
        head[0].text = "Item"
        head[1].text = "Descricao"
        head[2].text = "Unidade"
        head[3].text = "Qtd. DFD"
        head[4].text = "Qtd. distribuida"
        head[5].text = "Saldo"
        for row in data["resumo_itens"]:
            cells = tab_resumo.add_row().cells
            cells[0].text = str(row["item"].codigo)
            cells[1].text = row["item"].descricao
            cells[2].text = row["item"].unidade
            cells[3].text = f"{row['quantidade_dfd']:.3f}"
            cells[4].text = f"{row['quantidade_distribuida']:.3f}"
            cells[5].text = f"{row['saldo']:.3f}"

        doc.add_paragraph("")
        p = doc.add_paragraph("LANCAMENTOS POR SECRETARIA")
        p.runs[0].bold = True
        tab_linhas = doc.add_table(rows=1, cols=5)
        tab_linhas.style = "Table Grid"
        head2 = tab_linhas.rows[0].cells
        head2[0].text = "Secretaria"
        head2[1].text = "Item"
        head2[2].text = "Descricao"
        head2[3].text = "Unidade"
        head2[4].text = "Quantidade"
        for linha in data["linhas"]:
            cells = tab_linhas.add_row().cells
            cells[0].text = linha.secretaria.sigla
            cells[1].text = str(linha.item.codigo)
            cells[2].text = linha.item.descricao
            cells[3].text = linha.item.unidade
            cells[4].text = f"{linha.quantidade:.3f}"

    bio = BytesIO()
    doc.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"Mapa_Distribuicao_{processo.id}.docx",
    )


def export_distribuicao_pdf(processo):
    data = _distribuicao_payload(processo)
    orgao = data["orgao"]
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=28,
        bottomMargin=28,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleDistribuicao",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#153153"),
    )
    subtitle_style = ParagraphStyle(
        "SubtitleDistribuicao",
        parent=styles["BodyText"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    value_style = ParagraphStyle(
        "ValueDistribuicao",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=11.5,
        textColor=colors.HexColor("#1d2b40"),
    )
    section_title_style = ParagraphStyle(
        "SectionDistribuicao",
        parent=styles["Heading3"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#153153"),
    )

    elements = []
    if orgao["logo_path"]:
        try:
            elements.append(Image(orgao["logo_path"], width=60, height=60))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(orgao["nome"]), section_title_style))
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        elements.append(Paragraph(_pdf_text(orgao["razao_social"]), subtitle_style))
    if orgao["cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ: {orgao['cnpj']}"), subtitle_style))
    if orgao["endereco"]:
        elements.append(Paragraph(_pdf_text(orgao["endereco"]), subtitle_style))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph("Mapa de Distribuicao por Secretaria", title_style))
    elements.append(Paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}", subtitle_style))
    elements.append(Spacer(1, 8))

    if not data["habilitada"]:
        elements.append(Paragraph("Distribuicao por secretaria desabilitada no TR.", value_style))
    else:
        elements.append(
            Paragraph(
                f"<b>Total DFD:</b> {data['total_qtd_dfd']:.3f} &nbsp;&nbsp; "
                f"<b>Total distribuido:</b> {data['total_qtd_distribuida']:.3f}",
                value_style,
            )
        )
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("RESUMO POR ITEM", section_title_style))

        resumo_rows = [[
            Paragraph("<b>Item</b>", value_style),
            Paragraph("<b>Descricao</b>", value_style),
            Paragraph("<b>Unid.</b>", value_style),
            Paragraph("<b>Qtd. DFD</b>", value_style),
            Paragraph("<b>Qtd. dist.</b>", value_style),
            Paragraph("<b>Saldo</b>", value_style),
        ]]
        for row in data["resumo_itens"]:
            resumo_rows.append(
                [
                    Paragraph(str(row["item"].codigo), value_style),
                    Paragraph(_pdf_text(row["item"].descricao), value_style),
                    Paragraph(_pdf_text(row["item"].unidade), value_style),
                    Paragraph(f"{row['quantidade_dfd']:.3f}", value_style),
                    Paragraph(f"{row['quantidade_distribuida']:.3f}", value_style),
                    Paragraph(f"{row['saldo']:.3f}", value_style),
                ]
            )
        resumo_table = Table(resumo_rows, colWidths=[34, 238, 52, 58, 68, 68], repeatRows=1)
        resumo_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (3, 1), (-1, -1), "RIGHT"),
                ]
            )
        )
        elements.extend([resumo_table, Spacer(1, 10)])

        elements.append(Paragraph("LANCAMENTOS POR SECRETARIA", section_title_style))
        linhas_rows = [[
            Paragraph("<b>Secretaria</b>", value_style),
            Paragraph("<b>Item</b>", value_style),
            Paragraph("<b>Descricao</b>", value_style),
            Paragraph("<b>Unid.</b>", value_style),
            Paragraph("<b>Quantidade</b>", value_style),
        ]]
        for linha in data["linhas"]:
            linhas_rows.append(
                [
                    Paragraph(_pdf_text(linha.secretaria.sigla), value_style),
                    Paragraph(str(linha.item.codigo), value_style),
                    Paragraph(_pdf_text(linha.item.descricao), value_style),
                    Paragraph(_pdf_text(linha.item.unidade), value_style),
                    Paragraph(f"{linha.quantidade:.3f}", value_style),
                ]
            )
        linhas_table = Table(linhas_rows, colWidths=[62, 34, 268, 52, 68], repeatRows=1)
        linhas_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                    ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                    ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (4, 1), (4, -1), "RIGHT"),
                ]
            )
        )
        elements.append(linhas_table)

    pdf.build(elements)
    return _resp(bio.getvalue(), "application/pdf", f"Mapa_Distribuicao_{processo.id}.pdf")


def export_ci_pdf_preview(
    *,
    processo,
    modulo_origem_label,
    modulo_destino_label,
    numero_formatado,
    data_comunicacao,
    assunto,
    mensagem,
    referencia,
    responsavel_envio='',
    signatarios=None,
    observacao='',
):
    orgao = _orgao_payload()
    signatarios_raw = list(signatarios or [])
    signatarios = []
    for sig in signatarios_raw:
        if isinstance(sig, dict):
            nome = str(sig.get('nome') or '').strip()
            cargo = str(sig.get('cargo') or '').strip()
        else:
            nome = str(sig or '').strip()
            cargo = ''
        if not nome:
            continue
        signatarios.append({'nome': nome, 'cargo': cargo})

    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=26,
        bottomMargin=26,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCI",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#153153"),
        alignment=1,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleCI",
        parent=styles["BodyText"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    section_title_style = ParagraphStyle(
        "SectionTitleCI",
        parent=styles["Heading3"],
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    value_style = ParagraphStyle(
        "ValueCI",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor("#1d2b40"),
    )

    elements = []
    if orgao["logo_path"]:
        try:
            elements.append(Image(orgao["logo_path"], width=56, height=56))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(orgao["nome"]), ParagraphStyle("OrgaoNome", parent=styles["Heading2"], fontSize=14, textColor=colors.HexColor("#153153"))))
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        elements.append(Paragraph(_pdf_text(orgao["razao_social"]), subtitle_style))
    if orgao["cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ {orgao['cnpj']}"), subtitle_style))
    if orgao["endereco"]:
        elements.append(Paragraph(_pdf_text(orgao["endereco"]), subtitle_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("COMUNICAÇÃO INTERNA", section_title_style))
    elements.append(Paragraph(f"CI n. {numero_formatado}", title_style))
    elements.append(Spacer(1, 8))

    meta_rows = [
        [Paragraph("<b>De:</b>", value_style), Paragraph(_pdf_text(modulo_origem_label), value_style)],
        [Paragraph("<b>Para:</b>", value_style), Paragraph(_pdf_text(modulo_destino_label), value_style)],
        [Paragraph("<b>Data:</b>", value_style), Paragraph(_pdf_text(_fmt_date(data_comunicacao)), value_style)],
        [Paragraph("<b>Ref.:</b>", value_style), Paragraph(_pdf_text(referencia), value_style)],
    ]
    meta_table = Table(meta_rows, colWidths=[120, 400])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#edf3fb")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("ASSUNTO", section_title_style))
    assunto_box = Table([[Paragraph(_pdf_text(assunto), value_style)]], colWidths=[520])
    assunto_box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f7fd")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    elements.append(assunto_box)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("MENSAGEM", section_title_style))
    elements.append(Paragraph(_pdf_rich_text(mensagem), value_style))
    if observacao:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph("OBSERVAÇÕES", section_title_style))
        elements.append(Paragraph(_pdf_rich_text(observacao), value_style))
    if responsavel_envio:
        elements.append(Spacer(1, 8))
        elements.append(Paragraph(f"<b>Responsável pelo envio:</b> {_pdf_text(responsavel_envio)}", value_style))

    elements.append(Spacer(1, 18))
    elements.append(Paragraph("SIGNATÁRIOS", section_title_style))
    if signatarios:
        sign_rows = []
        for sig in signatarios:
            nome = sig.get('nome') or '-'
            cargo = sig.get('cargo') or ''
            sign_rows.append([Paragraph("________________________________________", value_style)])
            sign_rows.append([Paragraph(_pdf_text(nome), value_style)])
            if cargo:
                sign_rows.append([Paragraph(_pdf_text(cargo), subtitle_style)])
            sign_rows.append([Spacer(1, 8)])
        sign_table = Table(sign_rows, colWidths=[520])
        sign_table.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")]))
        elements.append(sign_table)
    else:
        elements.append(Paragraph("Sem signatário selecionado.", value_style))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph("Documento gerado pelo SIREL Modular.", subtitle_style))
    elements.append(Paragraph("Controle de C.I. por departamento de origem e exercício.", subtitle_style))

    pdf.build(elements)
    return _resp(bio.getvalue(), "application/pdf", f"CI_{numero_formatado.replace('/', '_')}.pdf")


def export_sd_docx(processo):
    data = build_sd_payload(processo)
    doc = Document()
    orgao = data["orgao"]

    if orgao["logo_path"]:
        try:
            doc.add_picture(orgao["logo_path"], width=Cm(2.4))
        except Exception:
            pass
    doc.add_paragraph(orgao["nome"])
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        doc.add_paragraph(orgao["razao_social"])
    if orgao["cnpj"]:
        doc.add_paragraph(f"CNPJ: {orgao['cnpj']}")
    if orgao["endereco"]:
        doc.add_paragraph(orgao["endereco"])
    doc.add_paragraph("")

    doc.add_heading("Solicitacao de Despesa (SD)", level=1)
    doc.add_paragraph(f"Processo: {processo.numero_processo_adm}/{processo.ano_referencia}")
    doc.add_paragraph(f"Data: {_fmt_date(datetime.now())}")
    if data["secretaria_assinatura"]:
        doc.add_paragraph(f"Secretaria de assinatura: {data['secretaria_assinatura'].sigla} - {data['secretaria_assinatura'].nome}")

    doc.add_paragraph("")
    p = doc.add_paragraph("JUSTIFICATIVA DA CONTRATACAO")
    p.runs[0].bold = True
    doc.add_paragraph(data["justificativa"] or "-")

    doc.add_paragraph("")
    p = doc.add_paragraph("ITENS E VALORES ESTIMADOS")
    p.runs[0].bold = True
    itens_table = doc.add_table(rows=1, cols=8)
    itens_table.style = "Table Grid"
    head = itens_table.rows[0].cells
    head[0].text = "Lote"
    head[1].text = "No item"
    head[2].text = "Codigo"
    head[3].text = "Descricao"
    head[4].text = "Unidade"
    head[5].text = "Quantidade"
    head[6].text = "Valor unitario"
    head[7].text = "Total item"
    for lote in data["lotes"]:
        for row in lote["itens"]:
            item = row["item"]
            tr = itens_table.add_row().cells
            tr[0].text = str(lote["numero"])
            tr[1].text = str(item.codigo)
            tr[2].text = str(item.codigo)
            tr[3].text = item.descricao
            tr[4].text = item.unidade
            tr[5].text = f"{row['quantidade']:.3f}"
            tr[6].text = fmt_brl(row["valor_unitario"])
            tr[7].text = fmt_brl(row["total_item"])
    doc.add_paragraph(f"Total geral estimado: {fmt_brl(data['total_geral'])}")

    if data["lotes"]:
        doc.add_paragraph("")
        p = doc.add_paragraph("TOTAIS POR LOTE")
        p.runs[0].bold = True
        lotes_table = doc.add_table(rows=1, cols=3)
        lotes_table.style = "Table Grid"
        h = lotes_table.rows[0].cells
        h[0].text = "Lote"
        h[1].text = "Titulo"
        h[2].text = "Total"
        for lote in data["lotes"]:
            tr = lotes_table.add_row().cells
            tr[0].text = str(lote["numero"])
            tr[1].text = lote["titulo"]
            tr[2].text = fmt_brl(lote["total_lote"])

    doc.add_paragraph("")
    p = doc.add_paragraph("DOTACOES ORCAMENTARIAS")
    p.runs[0].bold = True
    dot_table = doc.add_table(rows=1, cols=6)
    dot_table.style = "Table Grid"
    h = dot_table.rows[0].cells
    h[0].text = "Secretaria"
    h[1].text = "Unidade"
    h[2].text = "Projeto"
    h[3].text = "Elemento"
    h[4].text = "Fonte"
    h[5].text = "Valor (estimado)"
    for d in data["dotacoes"]:
        tr = dot_table.add_row().cells
        tr[0].text = f"{d.secretaria.sigla} - {d.secretaria.nome}" if d.secretaria else "-"
        tr[1].text = f"{d.unidade_orcamentaria.sigla} - {d.unidade_orcamentaria.nome}" if d.unidade_orcamentaria else "-"
        tr[2].text = f"{d.projeto_atividade.codigo} - {d.projeto_atividade.descricao}" if d.projeto_atividade else "-"
        tr[3].text = f"{d.elemento_despesa.codigo} - {d.elemento_despesa.descricao}" if d.elemento_despesa else "-"
        tr[4].text = f"{d.fonte_recurso.codigo} - {d.fonte_recurso.descricao}" if d.fonte_recurso else "-"
        tr[5].text = "-"

    doc.add_paragraph("")
    p = doc.add_paragraph("ASSINATURA")
    p.runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph("________________________________________")
    if data["signatario"]:
        doc.add_paragraph(data["signatario"].nome)
        doc.add_paragraph(data["signatario"].cargo or "")
    else:
        doc.add_paragraph("Responsavel pela assinatura")
    if data["secretaria_assinatura"]:
        doc.add_paragraph(f"{data['secretaria_assinatura'].sigla} - {data['secretaria_assinatura'].nome}")

    bio = BytesIO()
    doc.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"SD_{processo.id}.docx",
    )


def export_sd_xlsx(processo):
    data = build_sd_payload(processo)
    wb = Workbook()
    ws = wb.active
    ws.title = "SD_Resumo"
    ws.append(["Campo", "Valor"])
    ws.append(["Processo", f"{processo.numero_processo_adm}/{processo.ano_referencia}"])
    ws.append(["Objeto", processo.objeto or "-"])
    ws.append(["Total estimado", float(data["total_geral"])])
    ws.append(["Total estimado (extenso)", valor_por_extenso(data["total_geral"])])
    ws.append([
        "Secretaria de assinatura",
        f"{data['secretaria_assinatura'].sigla} - {data['secretaria_assinatura'].nome}" if data["secretaria_assinatura"] else "-",
    ])
    ws.append(["Signatario", data["signatario"].nome if data["signatario"] else "-"])
    ws.append(["Justificativa", data["justificativa"] or "-"])

    ws_items = wb.create_sheet("Itens")
    ws_items.append(["Lote", "No item", "Codigo", "Descricao", "Unidade", "Quantidade", "Valor unitario", "Total item"])
    for lote in data["lotes"]:
        for row in lote["itens"]:
            item = row["item"]
            ws_items.append([
                lote["numero"],
                item.codigo,
                item.codigo,
                item.descricao,
                item.unidade,
                float(row["quantidade"]),
                float(row["valor_unitario"]),
                float(row["total_item"]),
            ])

    ws_lotes = wb.create_sheet("Totais_lote")
    ws_lotes.append(["Lote", "Titulo", "Total"])
    for lote in data["lotes"]:
        ws_lotes.append([lote["numero"], lote["titulo"], float(lote["total_lote"])])

    ws_dot = wb.create_sheet("Dotacoes")
    ws_dot.append(["Secretaria", "Unidade", "Projeto", "Elemento", "Fonte"])
    for d in data["dotacoes"]:
        ws_dot.append([
            f"{d.secretaria.sigla} - {d.secretaria.nome}" if d.secretaria else "-",
            f"{d.unidade_orcamentaria.sigla} - {d.unidade_orcamentaria.nome}" if d.unidade_orcamentaria else "-",
            f"{d.projeto_atividade.codigo} - {d.projeto_atividade.descricao}" if d.projeto_atividade else "-",
            f"{d.elemento_despesa.codigo} - {d.elemento_despesa.descricao}" if d.elemento_despesa else "-",
            f"{d.fonte_recurso.codigo} - {d.fonte_recurso.descricao}" if d.fonte_recurso else "-",
        ])

    bio = BytesIO()
    wb.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"SD_{processo.id}.xlsx",
    )


def export_sd_pdf(processo):
    data = build_sd_payload(processo)
    orgao = data["orgao"]

    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=28,
        rightMargin=28,
        topMargin=26,
        bottomMargin=26,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleSD",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#153153"),
    )
    subtitle_style = ParagraphStyle(
        "SubtitleSD",
        parent=styles["BodyText"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#5d6f8a"),
    )
    section_title_style = ParagraphStyle(
        "SectionTitleSD",
        parent=styles["Heading3"],
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#153153"),
    )
    value_style = ParagraphStyle(
        "ValueSD",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor("#1d2b40"),
    )

    elements = []
    if orgao["logo_path"]:
        try:
            elements.append(Image(orgao["logo_path"], width=58, height=58))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(orgao["nome"]), section_title_style))
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        elements.append(Paragraph(_pdf_text(orgao["razao_social"]), subtitle_style))
    if orgao["cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ: {orgao['cnpj']}"), subtitle_style))
    if orgao["endereco"]:
        elements.append(Paragraph(_pdf_text(orgao["endereco"]), subtitle_style))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph("Solicitacao de Despesa (SD)", title_style))
    elements.append(Paragraph(f"Processo {processo.numero_processo_adm}/{processo.ano_referencia}", subtitle_style))
    elements.append(Paragraph(f"Total estimado: {fmt_brl(data['total_geral'])} ({valor_por_extenso(data['total_geral'])})", subtitle_style))
    elements.append(Spacer(1, 10))

    meta_rows = [
        [Paragraph("<b>Secretaria de assinatura</b>", value_style), Paragraph(_pdf_text(
            f"{data['secretaria_assinatura'].sigla} - {data['secretaria_assinatura'].nome}" if data["secretaria_assinatura"] else "-"
        ), value_style)],
        [Paragraph("<b>Signatario</b>", value_style), Paragraph(_pdf_text(data["signatario"].nome if data["signatario"] else "-"), value_style)],
    ]
    meta_table = Table(meta_rows, colWidths=[180, 340])
    meta_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f3f6fb")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(meta_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("JUSTIFICATIVA DA CONTRATACAO", section_title_style))
    box = Table([[Paragraph(_pdf_text(data["justificativa"] or "-"), value_style)]], colWidths=[520])
    box.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafd")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    elements.append(box)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("ITENS E VALORES ESTIMADOS", section_title_style))
    itens_rows = [[
        Paragraph("<b>Lote</b>", value_style),
        Paragraph("<b>No item</b>", value_style),
        Paragraph("<b>Descricao</b>", value_style),
        Paragraph("<b>Unid.</b>", value_style),
        Paragraph("<b>Qtd.</b>", value_style),
        Paragraph("<b>Valor unit.</b>", value_style),
        Paragraph("<b>Total item</b>", value_style),
    ]]
    for lote in data["lotes"]:
        for row in lote["itens"]:
            item = row["item"]
            itens_rows.append([
                Paragraph(str(lote["numero"]), value_style),
                Paragraph(str(item.codigo), value_style),
                Paragraph(_pdf_text(item.descricao), value_style),
                Paragraph(_pdf_text(item.unidade), value_style),
                Paragraph(f"{row['quantidade']:.3f}", value_style),
                Paragraph(_pdf_text(fmt_brl(row["valor_unitario"])), value_style),
                Paragraph(_pdf_text(fmt_brl(row["total_item"])), value_style),
            ])
    itens_table = Table(itens_rows, colWidths=[34, 42, 210, 46, 48, 70, 70], repeatRows=1)
    itens_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (4, 1), (6, -1), "RIGHT"),
                ("WORDWRAP", (2, 1), (2, -1), "CJK"),
            ]
        )
    )
    elements.append(itens_table)
    elements.append(Spacer(1, 8))

    lotes_rows = [[Paragraph("<b>Lote</b>", value_style), Paragraph("<b>Titulo</b>", value_style), Paragraph("<b>Total</b>", value_style)]]
    for lote in data["lotes"]:
        lotes_rows.append([
            Paragraph(str(lote["numero"]), value_style),
            Paragraph(_pdf_text(lote["titulo"]), value_style),
            Paragraph(_pdf_text(fmt_brl(lote["total_lote"])), value_style),
        ])
    lotes_table = Table(lotes_rows, colWidths=[60, 340, 120], repeatRows=1)
    lotes_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("ALIGN", (2, 1), (2, -1), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    elements.append(Spacer(1, 8))
    elements.append(Paragraph("TOTAIS POR LOTE", section_title_style))
    elements.append(lotes_table)
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("DOTACOES ORCAMENTARIAS", section_title_style))
    dot_rows = [[
        Paragraph("<b>Secretaria</b>", value_style),
        Paragraph("<b>Unidade</b>", value_style),
        Paragraph("<b>Projeto</b>", value_style),
        Paragraph("<b>Elemento</b>", value_style),
        Paragraph("<b>Fonte</b>", value_style),
    ]]
    for d in data["dotacoes"]:
        dot_rows.append([
            Paragraph(_pdf_text(f"{d.secretaria.sigla} - {d.secretaria.nome}" if d.secretaria else "-"), value_style),
            Paragraph(_pdf_text(f"{d.unidade_orcamentaria.sigla} - {d.unidade_orcamentaria.nome}" if d.unidade_orcamentaria else "-"), value_style),
            Paragraph(_pdf_text(f"{d.projeto_atividade.codigo} - {d.projeto_atividade.descricao}" if d.projeto_atividade else "-"), value_style),
            Paragraph(_pdf_text(f"{d.elemento_despesa.codigo} - {d.elemento_despesa.descricao}" if d.elemento_despesa else "-"), value_style),
            Paragraph(_pdf_text(f"{d.fonte_recurso.codigo} - {d.fonte_recurso.descricao}" if d.fonte_recurso else "-"), value_style),
        ])
    dot_table = Table(dot_rows, colWidths=[102, 102, 102, 102, 112], repeatRows=1)
    dot_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    elements.append(dot_table)
    elements.append(Spacer(1, 16))

    assinatura_nome = data["signatario"].nome if data["signatario"] else "Responsavel pela assinatura"
    assinatura_cargo = data["signatario"].cargo if data["signatario"] else ""
    assinatura_secretaria = (
        f"{data['secretaria_assinatura'].sigla} - {data['secretaria_assinatura'].nome}"
        if data["secretaria_assinatura"] else ""
    )
    assinatura = Table(
        [[Paragraph(f"{_pdf_text(assinatura_nome)}<br/>{_pdf_text(assinatura_cargo)}<br/>{_pdf_text(assinatura_secretaria)}", value_style)]],
        colWidths=[320],
    )
    assinatura.setStyle(
        TableStyle(
            [
                ("LINEABOVE", (0, 0), (0, 0), 1, colors.HexColor("#50617a")),
                ("TOPPADDING", (0, 0), (0, 0), 12),
                ("ALIGN", (0, 0), (0, 0), "CENTER"),
            ]
        )
    )
    elements.append(Paragraph("ASSINATURA", section_title_style))
    elements.append(Spacer(1, 14))
    elements.append(assinatura)

    pdf.build(elements)
    return _resp(bio.getvalue(), "application/pdf", f"SD_{processo.id}.pdf")


def export_tr_docx(processo):
    tr = processo.planejamento_tr
    doc = Document()
    doc.add_heading("Resumo Estruturado do Termo de Referencia", level=1)
    doc.add_paragraph(f"Processo: {processo.numero_processo_adm}/{processo.ano_referencia}")
    doc.add_paragraph(f"Criterio de julgamento: {tr.get_criterio_julgamento_display()}")
    doc.add_paragraph(f"Exclusividade ME/EPP: {'SIM' if tr.permite_exclusividade_me_epp else 'NAO'}")
    doc.add_paragraph(f"Cota reservada: {'SIM' if tr.permite_cota_reservada else 'NAO'}")

    doc.add_heading("Lotes", level=2)
    for lote in tr.lotes.all():
        doc.add_paragraph(f"Lote {lote.numero} - {lote.titulo}")
        for item in lote.itens.all():
            doc.add_paragraph(f"Item {item.codigo} - {item.descricao}", style="List Bullet")

    doc.add_heading("Dotacoes", level=2)
    for d in tr.dotacoes.select_related(
        "secretaria",
        "unidade_orcamentaria",
        "projeto_atividade",
        "elemento_despesa",
        "fonte_recurso",
    ):
        doc.add_paragraph(
            f"{getattr(d.secretaria, 'sigla', '')} | "
            f"{getattr(d.unidade_orcamentaria, 'sigla', '')} | "
            f"{getattr(d.projeto_atividade, 'codigo', '')} | "
            f"{getattr(d.elemento_despesa, 'codigo', '')} | "
            f"{getattr(d.fonte_recurso, 'codigo', '')}"
        )

    bio = BytesIO()
    doc.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"TR_{processo.id}.docx",
    )


def export_frequencia_xlsx(payload: dict):
    usuario = payload["usuario_alvo"]
    mes_ref = payload["mes_referencia"].replace("-", "")
    wb = Workbook()
    ws = wb.active
    ws.title = "Folha mensal"
    ws.append(["Folha de Frequência Mensal"])
    ws.append(
        [
            f"Usuário: {(usuario.get_full_name() or '').strip() or usuario.get_username()}",
            f"Mês: {payload['mes_label']}",
        ]
    )
    ws.append([])
    ws.append(
        [
            "Data",
            "Dia",
            "Entrada",
            "Início intervalo",
            "Volta intervalo",
            "Saída",
            "Horas trabalhadas",
            "Horas extras",
            "Situação",
            "Justificativas",
        ]
    )
    for linha in payload["linhas"]:
        ws.append(
            [
                linha["data"].strftime("%d/%m/%Y"),
                linha["dia_semana"],
                linha["entrada"] or "-",
                linha["inicio_intervalo"] or "-",
                linha["fim_intervalo"] or "-",
                linha["saida"] or "-",
                linha["horas_total"],
                linha["horas_extras"],
                linha["status_label"],
                linha["justificativa_resumo"] or "-",
            ]
        )

    ws.append([])
    ws.append(["Dias úteis", payload["dias_uteis"]])
    ws.append(["Dias trabalhados", payload["dias_trabalhados"]])
    ws.append(["Justificados", payload["dias_justificados"]])
    ws.append(["Pendências", payload["pendencias"]])
    ws.append(["Total trabalhado", payload["total_horas"]])
    ws.append(["Total extras", payload["total_extras"]])

    ws2 = wb.create_sheet("Justificativas extras")
    ws2.append(["Data", "Dia", "Horas trabalhadas", "Horas extras", "Justificativa"])
    if payload["extras_justificativas"]:
        for extra in payload["extras_justificativas"]:
            ws2.append(
                [
                    extra["data"].strftime("%d/%m/%Y"),
                    extra["dia_semana"],
                    extra["horas_total"],
                    extra["horas_extras"],
                    extra["justificativa"] or "-",
                ]
            )
    else:
        ws2.append(["-", "-", "-", "-", "Sem horas extras no período."])

    for sheet in [ws, ws2]:
        for col in ["A", "B", "C", "D", "E", "F", "G", "H", "I", "J"]:
            sheet.column_dimensions[col].width = 20
        sheet.column_dimensions["A"].width = 14
        sheet.column_dimensions["B"].width = 10
        sheet.column_dimensions["J"].width = 42

    bio = BytesIO()
    wb.save(bio)
    return _resp(
        bio.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"folha_frequencia_{usuario.get_username()}_{mes_ref}.xlsx",
    )


def export_frequencia_pdf(payload: dict):
    usuario = payload["usuario_alvo"]
    orgao = _orgao_payload()
    mes_ref = payload["mes_referencia"].replace("-", "")
    bio = BytesIO()
    pdf = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=20,
        rightMargin=20,
        topMargin=24,
        bottomMargin=24,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "FreqTitle",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#153153"),
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "FreqSubtitle",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#4f6280"),
        spaceAfter=4,
    )
    section_style = ParagraphStyle(
        "FreqSection",
        parent=styles["Heading3"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1f4c86"),
        spaceBefore=8,
        spaceAfter=6,
    )
    value_style = ParagraphStyle(
        "FreqValue",
        parent=styles["Normal"],
        fontSize=7.8,
        leading=10,
        textColor=colors.HexColor("#1c2c42"),
    )

    elements = []
    if orgao["logo_path"]:
        try:
            elements.append(Image(orgao["logo_path"], width=56, height=56))
        except Exception:
            pass
    elements.append(Paragraph(_pdf_text(orgao["nome"]), section_style))
    if orgao["razao_social"] and orgao["razao_social"] != orgao["nome"]:
        elements.append(Paragraph(_pdf_text(orgao["razao_social"]), subtitle_style))
    if orgao["cnpj"]:
        elements.append(Paragraph(_pdf_text(f"CNPJ: {orgao['cnpj']}"), subtitle_style))
    if orgao["endereco"]:
        elements.append(Paragraph(_pdf_text(orgao["endereco"]), subtitle_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("Folha de Frequência Mensal", title_style))
    elements.append(
        Paragraph(
            _pdf_text(
                f"Usuário: {(usuario.get_full_name() or '').strip() or usuario.get_username()} | "
                f"Mês: {payload['mes_label']}"
            ),
            subtitle_style,
        )
    )
    elements.append(Spacer(1, 6))

    resumo = Table(
        [
            ["Dias úteis", str(payload["dias_uteis"]), "Dias trabalhados", str(payload["dias_trabalhados"])],
            ["Justificados", str(payload["dias_justificados"]), "Pendências", str(payload["pendencias"])],
            ["Total trabalhado", payload["total_horas"], "Total extras", payload["total_extras"]],
        ],
        colWidths=[120, 70, 120, 70],
    )
    resumo.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f2f6fc")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#cfd9e7")),
                ("INNERGRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#dbe3ef")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    elements.append(resumo)
    elements.append(Spacer(1, 8))

    rows = [[
        Paragraph("<b>Data</b>", value_style),
        Paragraph("<b>Dia</b>", value_style),
        Paragraph("<b>Entrada</b>", value_style),
        Paragraph("<b>Início</b>", value_style),
        Paragraph("<b>Volta</b>", value_style),
        Paragraph("<b>Saída</b>", value_style),
        Paragraph("<b>Trab.</b>", value_style),
        Paragraph("<b>Extras</b>", value_style),
        Paragraph("<b>Situação</b>", value_style),
        Paragraph("<b>Justificativas</b>", value_style),
    ]]
    for linha in payload["linhas"]:
        rows.append(
            [
                Paragraph(linha["data"].strftime("%d/%m/%Y"), value_style),
                Paragraph(linha["dia_semana"], value_style),
                Paragraph(_pdf_text(linha["entrada"] or "-"), value_style),
                Paragraph(_pdf_text(linha["inicio_intervalo"] or "-"), value_style),
                Paragraph(_pdf_text(linha["fim_intervalo"] or "-"), value_style),
                Paragraph(_pdf_text(linha["saida"] or "-"), value_style),
                Paragraph(_pdf_text(linha["horas_total"]), value_style),
                Paragraph(_pdf_text(linha["horas_extras"]), value_style),
                Paragraph(_pdf_text(linha["status_label"]), value_style),
                Paragraph(_pdf_text(linha["justificativa_resumo"] or "-"), value_style),
            ]
        )
    table = Table(rows, colWidths=[50, 24, 34, 34, 34, 34, 42, 40, 58, 170], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e7edf7")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d4dce8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    elements.append(table)

    elements.append(PageBreak())
    elements.append(Paragraph("Justificativas de Horas Extras", title_style))
    elements.append(
        Paragraph(
            "Página separada por dia com indicação das horas realizadas e justificativa.",
            subtitle_style,
        )
    )
    elements.append(Spacer(1, 4))

    extras = payload["extras_justificativas"]
    if not extras:
        elements.append(Paragraph("Nenhuma hora extra registrada no período.", value_style))
    else:
        for extra in extras:
            box = Table(
                [[
                    Paragraph(
                        _pdf_text(
                            f"Dia: {extra['data'].strftime('%d/%m/%Y')} ({extra['dia_semana']})\n"
                            f"Horas trabalhadas: {extra['horas_total']}\n"
                            f"Horas extras: {extra['horas_extras']}\n"
                            f"Justificativa: {extra['justificativa'] or '-'}"
                        ),
                        ParagraphStyle(
                            "FreqExtraValue",
                            parent=value_style,
                            fontSize=10,
                            leading=14,
                        ),
                    )
                ]],
                colWidths=[540],
            )
            box.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7faff")),
                        ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#d4dce8")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            elements.append(box)
            elements.append(Spacer(1, 6))

    pdf.build(elements)
    return _resp(
        bio.getvalue(),
        "application/pdf",
        f"folha_frequencia_{usuario.get_username()}_{mes_ref}.pdf",
    )
